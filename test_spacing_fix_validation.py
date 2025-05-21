#!/usr/bin/env python
"""
Test script to validate section header box height and spacing fixes.

This script generates a resume with our fixes and verifies:
1. Section header box heights are compact (using lineRule="auto" with line=276)
2. Spacing between sections is minimal (using space_after=0 on paragraphs before headers)
"""

import os
import sys
import json
import logging
import tempfile
import shutil
import uuid
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('spacing_fix_validation.log')
    ]
)
logger = logging.getLogger(__name__)

def create_test_resume_data(temp_dir):
    """Create test resume data for validation"""
    request_id = str(uuid.uuid4()).replace("-", "")[:16]
    logger.info(f"Creating test resume with request ID: {request_id}")
    
    # Contact data
    contact_data = {
        "name": "John Doe",
        "email": "john.doe@example.com",
        "phone": "+1 (555) 123-4567",
        "location": "New York, NY",
        "linkedin": "linkedin.com/in/johndoe"
    }
    
    with open(os.path.join(temp_dir, f"{request_id}_contact.json"), 'w') as f:
        json.dump(contact_data, f)
    
    # Summary data
    summary_data = {
        "summary": "Experienced software engineer with a strong background in Python, JavaScript, and cloud technologies. Proven track record of delivering high-quality software solutions and leading development teams."
    }
    
    with open(os.path.join(temp_dir, f"{request_id}_summary.json"), 'w') as f:
        json.dump(summary_data, f)
    
    # Experience data
    experience_data = {
        "experiences": [
            {
                "company": "Tech Company A",
                "location": "New York, NY",
                "position": "Senior Software Engineer",
                "dates": "Jan 2020 - Present",
                "role_description": "Lead developer for cloud-based applications and services.",
                "achievements": [
                    "Implemented CI/CD pipeline reducing deployment time by 50%",
                    "Developed microservice architecture for scalable backend systems",
                    "Led team of 5 engineers in successful product launches"
                ]
            },
            {
                "company": "Tech Company B",
                "location": "Boston, MA",
                "position": "Software Engineer",
                "dates": "Jun 2016 - Dec 2019",
                "role_description": "Full-stack developer for web applications.",
                "achievements": [
                    "Redesigned user interface improving user satisfaction by 35%",
                    "Optimized database queries reducing load times by 40%",
                    "Implemented automated testing framework with 85% code coverage"
                ]
            }
        ]
    }
    
    with open(os.path.join(temp_dir, f"{request_id}_experience.json"), 'w') as f:
        json.dump(experience_data, f)
    
    # Education data
    education_data = {
        "institutions": [
            {
                "institution": "University of Technology",
                "location": "New York, NY",
                "degree": "Master of Science in Computer Science",
                "dates": "2014 - 2016",
                "highlights": [
                    "GPA: 3.8/4.0",
                    "Specialization in Artificial Intelligence"
                ]
            }
        ]
    }
    
    with open(os.path.join(temp_dir, f"{request_id}_education.json"), 'w') as f:
        json.dump(education_data, f)
    
    # Skills data
    skills_data = {
        "Technical Skills": ["Python", "JavaScript", "React", "Node.js", "AWS", "Docker", "Kubernetes"],
        "Soft Skills": ["Leadership", "Communication", "Problem Solving", "Team Collaboration"]
    }
    
    with open(os.path.join(temp_dir, f"{request_id}_skills.json"), 'w') as f:
        json.dump(skills_data, f)
    
    return request_id

def validate_section_headers(docx_path):
    """Validate section header box heights and spacing"""
    # Open the document
    doc = Document(docx_path)
    
    # Find all section headers - use multiple methods to ensure we catch them
    section_headers = []
    for i, para in enumerate(doc.paragraphs):
        # Method 1: By style name
        if para.style and para.style.name == 'BoxedHeading2':
            section_headers.append((i, para))
            logger.info(f"Found section header by style name at index {i}: {para.text}")
            continue
            
        # Method 2: By text content (common section header names)
        section_keywords = ["PROFESSIONAL SUMMARY", "EXPERIENCE", "EDUCATION", "SKILLS", "PROJECTS"]
        if any(keyword in para.text.upper() for keyword in section_keywords):
            # Check if it has borders - a good indicator it's a section header
            p_pr = para._element.get_or_add_pPr()
            if p_pr.find(qn('w:pBdr')) is not None:
                section_headers.append((i, para))
                logger.info(f"Found section header by content+border at index {i}: {para.text}")
                continue
                
        # Method 3: By borders alone (most reliable for our box headers)
        p_pr = para._element.get_or_add_pPr()
        if p_pr.find(qn('w:pBdr')) is not None:
            borders = p_pr.find(qn('w:pBdr'))
            # Check if it has all four borders
            if (borders.find(qn('w:top')) is not None and
                borders.find(qn('w:left')) is not None and
                borders.find(qn('w:bottom')) is not None and
                borders.find(qn('w:right')) is not None):
                section_headers.append((i, para))
                logger.info(f"Found section header by borders at index {i}: {para.text}")
    
    logger.info(f"Found {len(section_headers)} section headers in generated document")
    
    # Check each section header's properties
    valid_line_rules = 0
    valid_spacing_before = 0
    valid_border_padding = 0
    
    for i, (idx, header) in enumerate(section_headers):
        # Check paragraph properties through XML
        p_pr = header._element.get_or_add_pPr()
        
        # Check line spacing
        spacing = p_pr.find(qn('w:spacing'))
        if spacing is not None:
            line_value = spacing.get(qn('w:line'))
            line_rule = spacing.get(qn('w:lineRule'))
            space_before = spacing.get(qn('w:before'))
            
            logger.info(f"Header {i+1} ({header.text}) line spacing: {line_value} ({line_rule})")
            logger.info(f"Header {i+1} ({header.text}) space before: {space_before}")
            
            # Validate line rule and value
            if line_rule == "auto" and line_value == "276":
                valid_line_rules += 1
                logger.info(f"✓ Header {i+1} has correct line rule and value")
            else:
                logger.warning(f"✗ Header {i+1} has incorrect line rule or value: rule={line_rule}, value={line_value}")
            
            # Validate space before
            if space_before == "0":
                valid_spacing_before += 1
                logger.info(f"✓ Header {i+1} has space_before=0")
            else:
                logger.warning(f"✗ Header {i+1} has incorrect space_before: {space_before}")
        
        # Check border padding
        borders = p_pr.find(qn('w:pBdr'))
        if borders is not None:
            top_border = borders.find(qn('w:top'))
            if top_border is not None:
                padding = top_border.get(qn('w:space'))
                logger.info(f"Header {i+1} ({header.text}) border padding: {padding}")
                
                # Validate padding is 20 (1pt * 20)
                if padding and padding == "20":
                    valid_border_padding += 1
                    logger.info(f"✓ Header {i+1} has correct border padding: 1pt (20 twips)")
                else:
                    logger.warning(f"✗ Header {i+1} has incorrect border padding: {padding} twips ({int(padding)/20 if padding else 'None'}pt)")
        else:
            logger.warning(f"✗ Header {i+1} has no borders defined")
    
    # Check spacing between sections
    valid_spacing_after = 0
    for i, (idx, header) in enumerate(section_headers[1:], 1):  # Skip the first header
        # Find the paragraph right before this header
        prev_idx = idx - 1
        if prev_idx >= 0 and prev_idx < len(doc.paragraphs):
            prev_para = doc.paragraphs[prev_idx]
            
            # Skip if it's another section header
            if prev_para.style and prev_para.style.name == 'BoxedHeading2':
                continue
                
            space_after = prev_para.paragraph_format.space_after
            
            # Check XML spacing directly
            p_pr = prev_para._element.get_or_add_pPr()
            spacing_xml = p_pr.find(qn('w:spacing'))
            
            space_after_xml = None
            if spacing_xml is not None:
                space_after_xml = spacing_xml.get(qn('w:after'))
            
            logger.info(f"Paragraph before header {i+1} ({header.text}): space_after={space_after.pt if space_after else 'None'}, XML w:after={space_after_xml}")
            
            # Validate space_after is 0 - check both API and XML values
            if (space_after and space_after.pt == 0) or space_after_xml == "0":
                valid_spacing_after += 1
                logger.info(f"✓ Paragraph before header {i+1} has space_after=0 (API or XML)")
            else:
                logger.warning(f"✗ Paragraph before header {i+1} has incorrect space_after: API={space_after.pt if space_after else 'None'}, XML={space_after_xml}")
    
    # Calculate success rates
    line_rule_success = (valid_line_rules / len(section_headers)) * 100 if section_headers else 0
    spacing_before_success = (valid_spacing_before / len(section_headers)) * 100 if section_headers else 0
    border_padding_success = (valid_border_padding / len(section_headers)) * 100 if section_headers else 0
    spacing_after_success = (valid_spacing_after / (len(section_headers) - 1)) * 100 if len(section_headers) > 1 else 0
    
    logger.info(f"Validation Results:")
    logger.info(f"- Line rule and value (auto/276): {valid_line_rules}/{len(section_headers)} ({line_rule_success:.1f}%)")
    logger.info(f"- Space before (0): {valid_spacing_before}/{len(section_headers)} ({spacing_before_success:.1f}%)")
    logger.info(f"- Border padding (1pt): {valid_border_padding}/{len(section_headers)} ({border_padding_success:.1f}%)")
    logger.info(f"- Space after paragraphs before headers (0): {valid_spacing_after}/{len(section_headers) - 1 if len(section_headers) > 1 else 0} ({spacing_after_success:.1f}%)")
    
    overall_success = (
        line_rule_success + 
        spacing_before_success + 
        border_padding_success + 
        spacing_after_success
    ) / 4
    
    logger.info(f"Overall validation success: {overall_success:.1f}%")
    
    return overall_success >= 90  # Consider 90% or above as successful

def test_spacing_fixes():
    """Test section header box height and spacing fixes"""
    try:
        from utils.docx_builder import build_docx
        
        # Create temporary directory for test data
        temp_dir = tempfile.mkdtemp()
        try:
            # Create test resume data
            request_id = create_test_resume_data(temp_dir)
            
            # Generate DOCX
            docx_bytes = build_docx(request_id, temp_dir)
            
            # Save for validation
            output_path = f"test_spacing_fix_{request_id}.docx"
            with open(output_path, 'wb') as f:
                f.write(docx_bytes.getvalue())
            
            logger.info(f"Generated test resume DOCX: {output_path}")
            
            # Validate the document
            validation_result = validate_section_headers(output_path)
            
            if validation_result:
                logger.info("✅ Spacing fixes validation PASSED")
                print(f"Spacing fixes validation PASSED! Generated file: {output_path}")
                print("Please open the generated file to visually inspect that:")
                print("1. Section header boxes have compact height around the text")
                print("2. There's minimal space between content and the next section header")
                return True
            else:
                logger.error("❌ Spacing fixes validation FAILED")
                print(f"Spacing fixes validation FAILED! See log for details.")
                return False
                
        finally:
            # Clean up
            shutil.rmtree(temp_dir)
            
    except Exception as e:
        logger.error(f"Error testing spacing fixes: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return False

if __name__ == "__main__":
    success = test_spacing_fixes()
    sys.exit(0 if success else 1) 