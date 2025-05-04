from yc_resume_generator import YCResumeGenerator
import os
import json
import logging
import time
import traceback
import re
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx2txt
from typing import Dict, List, Tuple, Optional, Any, Union
from datetime import datetime
from flask import current_app

# Third-party imports for Claude
from anthropic import Anthropic, RateLimitError

# Third-party imports for OpenAI
import openai
from openai import OpenAI
from claude_api_logger import api_logger

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Global variable to track the last LLM client used
last_llm_client = None

# Import our YC Resume Generator

# Import the new html_generator module
import html_generator

# Import utils
from utils.bullet_utils import strip_bullet_prefix


def clean_bullet_points(text: str) -> str:
    """
    Removes bullet point markers from text to prevent duplicate bullet points.
    
    Handles various bullet point formats:
    - Unicode bullets (•, ◦, ▪, ▫, etc.)
    - ASCII bullets (*, -, +, o)
    - Numbered bullets (1., 1), (1), etc.)
    
    Args:
        text (str): Text containing bullet points to clean
        
    Returns:
        str: Cleaned text with bullet point markers removed
    """
    if not text:
        return ""
    
    logger.info(f"Cleaning bullet points from text ({len(text)} chars)")
    
    # Split into lines to process each line individually
    lines = text.split('\n')
    cleaned_lines = []
    bullet_count = 0
    
    # Define comprehensive bullet point patterns
    # Unicode bullet symbols
    unicode_bullets = [
        '•', '◦', '▪', '▫', '■', '□', '▸', '►', '▹', '▻', '▷', '▶', '→', '⇒', '⟹', '⟶',
        '⇢', '⇨', '⟾', '➔', '➜', '➙', '➛', '➝', '➞', '➟', '➠', '➡', '➢', '➣', '➤', '➥', 
        '➦', '➧', '➨', '➩', '➪', '➫', '➬', '➭', '➮', '➯', '➱', '➲', '➳', '➵', '➸', '➼',
        '⦿', '⦾', '⧫', '⧮', '⧠', '⧔', '∙', '◆', '◇', '◈'
    ]
    
    # ASCII/common bullet symbols for regex
    ascii_bullets = r'[\*\-\+o~=#>]'
    
    # NEW: textual escape sequences for bullets (e.g., "u2022")
    bullet_escapes_pattern = re.compile(r'^(?:u2022|\\u2022|U\+2022|&#8226;|&bull;)\s+', re.IGNORECASE)

    for i, line in enumerate(lines):
        line = line.strip()
        
        # Skip empty lines
        if not line:
            cleaned_lines.append('')
            continue
        
        # Check for Unicode bullet symbols at the beginning
        if line and line[0] in unicode_bullets:
            # Find where the content starts after the bullet and any spaces
            content_start = 1
            while content_start < len(line) and line[content_start].isspace():
                content_start += 1
                
            # Extract just the content, removing the bullet and leading spaces
            if content_start < len(line):
                cleaned_line = line[content_start:]
                cleaned_lines.append(cleaned_line)
                bullet_count += 1
                continue
        
        # Check for ASCII bullet symbols (*, -, +, etc.) at the beginning
        ascii_bullet_match = re.match(r'^\s*[-*•]\s+', line)
        if ascii_bullet_match:
            cleaned_line = line[ascii_bullet_match.end():]
            cleaned_lines.append(cleaned_line)
            bullet_count += 1
            continue
            
        # Check for numbered bullets like "1.", "1)", "(1)", etc.
        numbered_match = re.match(
    r'^\s*(?:\(?\d+[\.\)\]]\s+|\d+[\.\)\]]\s+|\[\d+\]\s+)', line)
        if numbered_match:
            cleaned_line = line[numbered_match.end():]
            cleaned_lines.append(cleaned_line)
            bullet_count += 1
            continue
            
        # Check for textual bullet escape sequences (e.g., "u2022")
        escape_match = bullet_escapes_pattern.match(line)
        if escape_match:
            cleaned_line = line[escape_match.end():]
            cleaned_lines.append(cleaned_line)
            bullet_count += 1
            continue
        
        # No bullet point found, keep the line as is
        cleaned_lines.append(line)
    
    # Join the cleaned lines back into a single string
    cleaned_text = '\n'.join(cleaned_lines)
    
    # Validate the cleaning
    if bullet_count > 0:
        logger.info(f"Removed {bullet_count} bullet point markers from text")
    
    return cleaned_text


class LLMClient:
    """Base class for LLM API clients"""
    
    def __init__(self, api_key: str):
        self.api_key = api_key
        
    def tailor_resume_content(
    self,
    section_name: str,
    content: str,
     job_data: Dict) -> Union[Dict, List, str]:
        """Tailor resume content using LLM API - to be implemented by subclasses"""
        raise NotImplementedError("Subclasses must implement this method")


class ClaudeClient(LLMClient):
    """Client for interacting with Claude API"""
    
    def __init__(self, api_key: str, api_url: str = None):
        """Initialize the Claude API client"""
        super().__init__(api_key)
        self.api_key = api_key
        self.api_url = api_url
        self.client = None

        try:
            # Validate API key format for Claude
            if not api_key:
                raise ValueError("Claude API key is missing")

            # Import anthropic library (lazy import to handle missing
            # dependency)
            try:
                from anthropic import Anthropic
                logger.info("Using anthropic SDK for Claude API")
            except ImportError:
                logger.error(
                    "Failed to import anthropic SDK. Make sure it's installed with 'pip install anthropic'")
                raise

            # Initialize Claude client
            self.client = Anthropic(api_key=api_key)
            logger.info(f"Claude API client initialized successfully")

        except Exception as e:
            logger.error(f"Error initializing Claude API client: {str(e)}")
            logger.error(traceback.format_exc())
            raise ValueError(f"Failed to initialize Claude client: {str(e)}")
    
    def tailor_resume_content(
    self,
    section_name: str,
    content: str,
     job_data: Dict) -> Union[Dict, List, str]:
        """
        Tailor resume content using Claude API

        Args:
            section_name: Name of the section to tailor
            content: Content of the section
            job_data: Job data including requirements and skills

        Returns:
            Tailored content as structured data (dict/list) or string for simple sections
        """
        logger.info(f"Tailoring {section_name} with Claude API")
            
        # --- START FIX: Type-aware check for empty content ---
        # Check if content is None, empty string/list/dict, or whitespace-only string
        is_empty = False
        if not content:
            is_empty = True
        elif isinstance(content, str) and not content.strip():
            is_empty = True
            
        if is_empty:
        # --- END FIX ---
            logger.warning(f"Empty {section_name} content provided, skipping tailoring")
            return content

        if not self.client:
            logger.error("Claude client not initialized")
            return content

        try:
            # Extract job data
            job_title = job_data.get('job_title', 'the position')
            company = job_data.get('company', 'the company')
            requirements = job_data.get('requirements', [])
            skills = job_data.get('skills', [])

            # Prepare requirements and skills text
            requirements_text = "\n".join(
                [f"- {req}" for req in requirements]) if requirements else "Not specified"
            skills_text = ", ".join(skills) if skills else "Not specified"

            # Get job analysis if available
            analysis_prompt = ""
            if 'analysis' in job_data and isinstance(
                job_data['analysis'], dict):
                analysis = job_data['analysis']
                
                # Add candidate profile if available
                if 'candidate_profile' in analysis and analysis['candidate_profile']:
                    analysis_prompt += f"\n\nCANDIDATE PROFILE:\n{
    analysis['candidate_profile']}"
                
                # Add hard skills if available
                if 'hard_skills' in analysis and analysis['hard_skills']:
                    hard_skills = ", ".join(analysis['hard_skills'])
                    analysis_prompt += f"\n\nKEY HARD SKILLS:\n{hard_skills}"
                    
                # Add soft skills if available
                if 'soft_skills' in analysis and analysis['soft_skills']:
                    soft_skills = ", ".join(analysis['soft_skills'])
                    analysis_prompt += f"\n\nKEY SOFT SKILLS:\n{soft_skills}"
                
                # Add ideal candidate if available
                if 'ideal_candidate' in analysis and analysis['ideal_candidate']:
                    analysis_prompt += f"\n\nIDEAL CANDIDATE:\n{
    analysis['ideal_candidate']}"

            # Build section-specific prompts
            if section_name == "experience":
                # Convert the input list of job objects back to a JSON string for the prompt
                experience_json_input = json.dumps(content, indent=2) # Assumes content is the list of dicts

                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the work experience section (provided as a JSON list) to better match the requirements for a {job_title} position at {company}.

ORIGINAL EXPERIENCE SECTION (JSON):
```json
{experience_json_input}
```

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object containing ONLY the tailored experience list, matching the exact input structure but with tailored content AND the addition/modification of the "role_description" field:

```json
{{
  "experience": [
    {{
      "company": "Company Name",
      "location": "City, State",
      "position": "Job Title",
      "dates": "Time Period",
      "role_description": "A 1-2 sentence tailored or generated description of the role.",
      "achievements": [
        "Tailored Achievement 1",
        "Tailored Achievement 2"
      ]
    }}
    // ... other job entries
  ]
}}
```

Please process EACH job entry in the input JSON list and restructure it to better match the job requirements by:
1.  For the "role_description":
    *   If the original entry has a non-empty "role_description", TAILOR this description to highlight aspects relevant to the target job ({job_title}). Keep it concise (1-2 sentences).
    *   If the original entry has an empty, null, or missing "role_description", GENERATE a concise (1-2 sentences) description based on the "position" and "achievements" for that specific job entry. This description should summarize the core responsibilities or focus of the role in the context of the achievements listed.
2.  For the "achievements":
    • Rewrite EACH achievement as **one sentence ≤ 130 characters**.  
    • Sentence must follow the pattern  
      "**<Strong Verb>** + original core fact → **numbered outcome** → **who/what it influenced**".  
      – strong verb & influence must reference the job's requirements/ideal-candidate traits.  
      – numeric outcome (%, $, #) must be truthful to the underlying résumé fact; do **not** invent data.  
    • If the original bullet lacks a metric, infer a reasonable one **only if** a figure is present elsewhere in the entry; otherwise keep it qualitative but concise.  
    • Do not add new bullets, split bullets, or remove existing bullets.
3.  Maintain the original "company", "location", "position", and "dates" for each entry precisely.

IMPORTANT:
1. Do not include empty strings or whitespace-only strings in any arrays (like achievements).
2. Every achievement must contain meaningful content.
3. Ensure the "role_description" is present and populated for EVERY job entry in the output.
4. Each achievement must obey the 130-character limit and the action-impact-influence structure above.
5. Ensure the output is a valid JSON object containing ONLY the "experience" key with the list of tailored job objects.
"""

            elif section_name == "education":
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the education section to better match the requirements for a {job_title} position at {company}.

ORIGINAL EDUCATION SECTION:
            {content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "education": [
    {{
      "institution": "University Name",
      "location": "City, State",
      "degree": "Degree Name",
      "dates": "Time Period",
      "highlights": [
        "Highlight 1",
        "Highlight 2"
      ]
    }}
  ]
}}
```

Please rewrite the education section to better match the job requirements. Focus on:
1. Highlighting relevant coursework, projects, or achievements that match the job requirements
2. Emphasizing academic accomplishments that demonstrate skills needed for this position
3. Formatting in a way that emphasizes the most relevant educational experiences
4. Including any certifications or training that matches required skills

Keep the degree names, institutions, and dates exactly the same - only enhance descriptions to make them more relevant.
"""

            elif section_name == "skills":
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the skills section to better match the requirements for a {job_title} position at {company}.

ORIGINAL SKILLS SECTION:
{content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "skills": {{
    "technical": ["Skill 1", "Skill 2"],
    "soft": ["Skill 1", "Skill 2"],
    "other": ["Skill 1", "Skill 2"]
  }}
}}
```

Please rewrite the skills section to better match the job requirements. Focus on:
1. Reordering skills to prioritize those mentioned in the job description
2. Adding any missing skills that the candidate likely has based on their experience (must be reasonable to infer from other sections)
3. Grouping skills into relevant categories that align with the job posting
4. Rephrasing skills using the exact terminology from the job description
5. Removing skills that are irrelevant to this position if the list is very long

Only include skills that are authentic to the candidate based on their resume.
"""

            elif section_name == "projects":
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the projects section to better match the requirements for a {job_title} position at {company}.

ORIGINAL PROJECTS SECTION:
{content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "projects": [
    {{
      "title": "Project Name",
      "dates": "Time Period",
      "details": [
        "Detail 1",
        "Detail 2"
      ]
    }}
  ]
}}
```

Please rewrite the projects section to better match the job requirements. Focus on:
1. Highlighting projects that demonstrate skills required for this job
2. Emphasizing tools, technologies, and methodologies that match the job description
3. Quantifying project outcomes and impacts where possible
4. Rephrasing project descriptions to use terminology from the job posting
5. Focusing on the candidate's specific contributions and leadership roles

Keep the project titles and timelines the same, but enhance descriptions to better align with the job requirements.
"""

            else:
                # For other sections, use a generic prompt
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the {section_name} section to better match the requirements for a {job_title} position at {company}.

ORIGINAL SECTION:
{content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "{section_name}": "The tailored content goes here as a single string"
}}
```

Please rewrite this section to better match the job requirements while maintaining the same basic structure and information.
Focus on emphasizing elements most relevant to this job opportunity.
"""

            # Make the API call
            response = self.client.messages.create(
                model="claude-3-sonnet-20240229",
                max_tokens=4000,
                temperature=0.7,
                messages=[
                    {"role": "system", "content": "You are an expert resume tailor. Return only valid JSON responses in the specified format. Each \"achievements\" item must be a single sentence ≤ 130 characters."},
                    {"role": "user", "content": prompt}
                ]
            )

            # Get the response content
            response_content = response.content[0].text.strip()
            logger.info(
    f"Claude API response for {section_name}: {
        len(response_content)} chars")

            # Parse JSON response
            try:
                # First try to extract JSON if not properly formatted
                json_str = response_content
                if not json_str.startswith('{'):
                    # Try to extract JSON from text
                    json_start = response_content.find('{')
                    json_end = response_content.rfind('}') + 1
                    if json_start >= 0 and json_end > json_start:
                        json_str = response_content[json_start:json_end]

                json_response = json.loads(json_str)

                # Process JSON based on section type
                if section_name == "experience" and "experience" in json_response:
                    return json_response["experience"]
                elif section_name == "education" and "education" in json_response:
                    return json_response["education"]
                elif section_name == "skills" and "skills" in json_response:
                    return json_response["skills"]
                elif section_name == "projects" and "projects" in json_response:
                    return json_response["projects"]
                elif section_name in json_response:
                    # For other sections, just return the string content
                    formatted_text = json_response[section_name]
                    return formatted_text
                else:
                    logger.warning(f"JSON response missing expected '{section_name}' key")
                    return content

            except json.JSONDecodeError:
                logger.error(f"Failed to parse JSON from Claude response: {response_content[:100]}...")
                # Store as raw text for fallback
                return response_content
            
        except Exception as e:
            logger.error(f"Error in Claude API call: {str(e)}")
            return content


class OpenAIClient(LLMClient):
    """OpenAI API client for resume tailoring"""
    
    def __init__(self, api_key: str):
        """Initialize the OpenAI API client"""
        super().__init__(api_key)
        self.api_key = api_key
        self.client = None
        self.raw_responses = {}  # Store raw JSON responses from API
        self.initialize_client()

    def initialize_client(self):
        """Initialize the OpenAI client"""
        try:
            print(
                f"Initializing OpenAI client with API key starting with: {self.api_key[:8]}...")
            print(f"API key length: {len(self.api_key)} characters")
            print("Testing OpenAI API connection...")

            from openai import OpenAI
            self.client = OpenAI(api_key=self.api_key)

            # Test connection with a simple request - removed limit parameter
            self.client.models.list()
            print("OpenAI client initialized successfully")

        except Exception as e:
            print(f"Error initializing OpenAI client: {str(e)}")
            self.client = None
            raise

    def tailor_resume_content(
    self,
    section_name: str,
    content: str,
     job_data: Dict) -> Union[Dict, List, str]:
        """
        Tailor resume content using OpenAI API

        Args:
            section_name: Name of the section to tailor
            content: Content of the section
            job_data: Job data including requirements and skills

        Returns:
            Tailored content as structured data (dict/list) or string for simple sections
        """
        logger.info(f"Tailoring {section_name} with OpenAI API")
            
        # --- START FIX: Type-aware check for empty content ---
        # Check if content is None, empty string/list/dict, or whitespace-only string
        is_empty = False
        if not content:
            is_empty = True
        elif isinstance(content, str) and not content.strip():
            is_empty = True
            
        if is_empty:
        # --- END FIX ---
            logger.warning(f"Empty {section_name} content provided, skipping tailoring")
            return content

        if not self.client:
            logger.error("OpenAI client not initialized")
            return content

        try:
            # Extract job data
            job_title = job_data.get('job_title', 'the position')
            company = job_data.get('company', 'the company')
            requirements = job_data.get('requirements', [])
            skills = job_data.get('skills', [])

            # Prepare requirements and skills text
            requirements_text = "\n".join(
                [f"- {req}" for req in requirements]) if requirements else "Not specified"
            skills_text = ", ".join(skills) if skills else "Not specified"

            # Get job analysis if available
            analysis_prompt = ""
            if 'analysis' in job_data and isinstance(
                job_data['analysis'], dict):
                analysis = job_data['analysis']
                
                # Add candidate profile if available
                if 'candidate_profile' in analysis and analysis['candidate_profile']:
                    analysis_prompt += f"\n\nCANDIDATE PROFILE:\n{
    analysis['candidate_profile']}"
                
                # Add hard skills if available
                if 'hard_skills' in analysis and analysis['hard_skills']:
                    hard_skills = ", ".join(analysis['hard_skills'])
                    analysis_prompt += f"\n\nKEY HARD SKILLS:\n{hard_skills}"
                    
                # Add soft skills if available
                if 'soft_skills' in analysis and analysis['soft_skills']:
                    soft_skills = ", ".join(analysis['soft_skills'])
                    analysis_prompt += f"\n\nKEY SOFT SKILLS:\n{soft_skills}"
                
                # Add ideal candidate if available
                if 'ideal_candidate' in analysis and analysis['ideal_candidate']:
                    analysis_prompt += f"\n\nIDEAL CANDIDATE:\n{
    analysis['ideal_candidate']}"

            # Build section-specific prompts
            if section_name == "experience":
                # Convert the input list of job objects back to a JSON string for the prompt
                experience_json_input = json.dumps(content, indent=2) # Assumes content is the list of dicts

                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the work experience section (provided as a JSON list) to better match the requirements for a {job_title} position at {company}.

ORIGINAL EXPERIENCE SECTION (JSON):
```json
{experience_json_input}
```

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object containing ONLY the tailored experience list, matching the exact input structure but with tailored content AND the addition/modification of the "role_description" field:

```json
{{
  "experience": [
    {{
      "company": "Company Name",
      "location": "City, State",
      "position": "Job Title",
      "dates": "Time Period",
      "role_description": "A 1-2 sentence tailored or generated description of the role.",
      "achievements": [
        "Tailored Achievement 1",
        "Tailored Achievement 2"
      ]
    }}
    // ... other job entries
  ]
}}
```

Please process EACH job entry in the input JSON list and restructure it to better match the job requirements by:
1.  For the "role_description":
    *   If the original entry has a non-empty "role_description", TAILOR this description to highlight aspects relevant to the target job ({job_title}). Keep it concise (1-2 sentences).
    *   If the original entry has an empty, null, or missing "role_description", GENERATE a concise (1-2 sentences) description based on the "position" and "achievements" for that specific job entry. This description should summarize the core responsibilities or focus of the role in the context of the achievements listed.
2.  For the "achievements":
    • Rewrite EACH achievement as **one sentence ≤ 130 characters**.  
    • Sentence must follow the pattern  
      "**<Strong Verb>** + original core fact → **numbered outcome** → **who/what it influenced**".  
      – strong verb & influence must reference the job's requirements/ideal-candidate traits.  
      – numeric outcome (%, $, #) must be truthful to the underlying résumé fact; do **not** invent data.  
    • If the original bullet lacks a metric, infer a reasonable one **only if** a figure is present elsewhere in the entry; otherwise keep it qualitative but concise.  
    • Do not add new bullets, split bullets, or remove existing bullets.
3.  Maintain the original "company", "location", "position", and "dates" for each entry precisely.

IMPORTANT:
1. Do not include empty strings or whitespace-only strings in any arrays (like achievements).
2. Every achievement must contain meaningful content.
3. Ensure the "role_description" is present and populated for EVERY job entry in the output.
4. Each achievement must obey the 130-character limit and the action-impact-influence structure above.
5. Ensure the output is a valid JSON object containing ONLY the "experience" key with the list of tailored job objects.
"""

            elif section_name == "education":
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the education section to better match the requirements for a {job_title} position at {company}.

ORIGINAL EDUCATION SECTION:
            {content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "education": [
    {{
      "institution": "University Name",
      "location": "City, State",
      "degree": "Degree Name",
      "dates": "Time Period",
      "highlights": [
        "Highlight 1",
        "Highlight 2"
      ]
    }}
  ]
}}
```

Please rewrite the education section to better match the job requirements. Focus on:
1. Highlighting relevant coursework, projects, or achievements that match the job requirements
2. Emphasizing academic accomplishments that demonstrate skills needed for this position
3. Formatting in a way that emphasizes the most relevant educational experiences
4. Including any certifications or training that matches required skills

Keep the degree names, institutions, and dates exactly the same - only enhance descriptions to make them more relevant.
"""

            elif section_name == "skills":
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the skills section to better match the requirements for a {job_title} position at {company}.

ORIGINAL SKILLS SECTION:
{content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "skills": {{
    "technical": ["Skill 1", "Skill 2"],
    "soft": ["Skill 1", "Skill 2"],
    "other": ["Skill 1", "Skill 2"]
  }}
}}
```

Please rewrite the skills section to better match the job requirements. Focus on:
1. Reordering skills to prioritize those mentioned in the job description
2. Adding any missing skills that the candidate likely has based on their experience (must be reasonable to infer from other sections)
3. Grouping skills into relevant categories that align with the job posting
4. Rephrasing skills using the exact terminology from the job description
5. Removing skills that are irrelevant to this position if the list is very long

Only include skills that are authentic to the candidate based on their resume.
"""

            elif section_name == "projects":
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the projects section to better match the requirements for a {job_title} position at {company}.

ORIGINAL PROJECTS SECTION:
{content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "projects": [
    {{
      "title": "Project Name",
      "dates": "Time Period",
      "details": [
        "Detail 1",
        "Detail 2"
      ]
    }}
  ]
}}
```

Please rewrite the projects section to better match the job requirements. Focus on:
1. Highlighting projects that demonstrate skills relevant to this position
2. Emphasizing technical achievements and technologies that align with the job description
3. Quantifying impact and results wherever possible
4. Using terminology from the job description where appropriate

IMPORTANT:
1. Do not include empty strings or whitespace-only strings in any arrays
2. Every detail must contain meaningful content
3. Do not split single project descriptions into multiple entries
4. Make sure each bullet point contains complete, meaningful content
5. Ensure all project entries from the original resume are preserved

Maintain the original project names and dates - only enhance descriptions to make them more relevant.
"""

            else:
                # For other sections, use a generic prompt
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the {section_name} section to better match the requirements for a {job_title} position at {company}.

ORIGINAL SECTION:
{content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "{section_name}": "The tailored content goes here as a single string"
}}
```

Please rewrite this section to better match the job requirements while maintaining the same basic structure and information.
Focus on emphasizing elements most relevant to this job opportunity.
"""

            # Send the request to OpenAI API
            response = self.client.chat.completions.create(
                model="gpt-4o" if "4" in os.environ.get(
    'OPENAI_MODEL_NAME', 'gpt-4') else "gpt-3.5-turbo",
                messages=[
                    {"role": "system",
     "content": "You are an expert resume tailoring assistant. Return ONLY valid JSON in the specified format. Each \"achievements\" item must be a single sentence ≤ 130 characters."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=4096,
                top_p=1.0
            )

            # Get response text and extract JSON
            response_text = response.choices[0].message.content

            # Save raw response for debugging
            self.raw_responses[section_name] = response_text

            # Save to file
            try:
                # Create api_responses directory if it doesn't exist
                api_responses_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'api_responses')
                if not os.path.exists(api_responses_dir):
                    os.makedirs(api_responses_dir)

                # Generate filename with timestamp
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"{section_name}_response_{timestamp}.json"
                filepath = os.path.join(api_responses_dir, filename)

                # Write response to file
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(response_text)

                logger.info(
    f"Saved raw API response for {section_name} to {filepath}")
            except Exception as e:
                logger.error(f"Error saving raw API response: {str(e)}")

            logger.info(
    f"OpenAI API response for {section_name}: {
        len(response_text)} chars")
            
            # Log token usage
            prompt_tokens = response.usage.prompt_tokens
            completion_tokens = response.usage.completion_tokens
            logger.info(
    f"Completion tokens: {completion_tokens}, Prompt tokens: {prompt_tokens}")

            # Extract JSON from the response
            json_match = re.search(
    r'```json\s*(.*?)\s*```',
    response_text,
     re.DOTALL)

            if json_match:
                json_str = json_match.group(1)
            else:
                # If no json code block, try to find JSON object directly
                json_match = re.search(r'(\{.*\})', response_text, re.DOTALL)
                if json_match:
                    json_str = json_match.group(1)
                else:
                    # No JSON found, return the original content
                    logger.error(f"No JSON found in OpenAI response for {section_name}")
                    return content

            # Parse the JSON string
            try:
                json_response = json.loads(json_str)
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse JSON from OpenAI response: {e}")
                logger.error(f"JSON string: {json_str[:100]}...")
                # Store the raw response as fallback
                return response_text

            # Process JSON based on section type
            if section_name == "experience" and "experience" in json_response:
                return json_response["experience"]
            elif section_name == "education" and "education" in json_response:
                return json_response["education"]
            elif section_name == "skills" and "skills" in json_response:
                return json_response["skills"]
            elif section_name == "projects" and "projects" in json_response:
                return json_response["projects"]
            elif section_name in json_response:
                # For other sections, just return the string content
                formatted_text = json_response[section_name]
                return formatted_text
            else:
                logger.warning(f"JSON response missing expected '{section_name}' key")
                return content
                
        except Exception as e:
            logger.error(f"Error in OpenAI API call: {str(e)}")
            return content

    # Add a method to save all raw responses to a single file
    def save_all_raw_responses(self, resume_filename):
        """Save all raw API responses to a single file"""
        try:
            # Create api_responses directory if it doesn't exist
            api_responses_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'api_responses')
            if not os.path.exists(api_responses_dir):
                os.makedirs(api_responses_dir)

            # Generate filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_name = os.path.splitext(os.path.basename(resume_filename))[0]
            filename = f"{base_name}_all_responses_{timestamp}.json"
            filepath = os.path.join(api_responses_dir, filename)

            # Write all responses to file
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(self.raw_responses, f, indent=2)

            logger.info(f"Saved all raw API responses to {filepath}")
            return filepath
        except Exception as e:
            logger.error(f"Error saving all raw API responses: {str(e)}")
            return None

    def save_tailored_content_to_json(self):
        """DEPRECATED: This method saved formatted HTML to fixed filenames. Use the request_id based saving in tailor_resume_with_llm."""
        try:
            # Create api_responses directory if it doesn't exist
            api_responses_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'api_responses')
            if not os.path.exists(api_responses_dir):
                os.makedirs(api_responses_dir)
                
            # Log content of tailored_content for debugging
            sections_available = list(self.tailored_content.keys())
            logger.info(f"Sections available for saving: {sections_available}")
            
            # Validate that contact section exists
            if "contact" not in sections_available:
                logger.warning("Contact section not found in tailored content. This will result in missing contact information in the resume.")
                
            # Validate that summary section exists
            if "summary" not in sections_available:
                logger.warning("Summary section not found in tailored content. This will result in missing professional summary in the resume.")
            
            # Save each section to a separate JSON file without timestamp
            sections_saved = 0
            for section_name, content in self.tailored_content.items():
                # Skip empty content
                if not content:
                    logger.warning(f"Empty content for section {section_name}, skipping")
                    continue
                    
                # Create a simple JSON structure based on section type
                json_data = {}
                
                if section_name == "experience":
                    # Parse the HTML content to extract job entries
                    experience_entries = self._extract_experience_from_html(content)
                    json_data = experience_entries if experience_entries else []
                elif section_name in ["summary", "contact", "skills"]:
                    # For text sections, use simple content field
                    json_data = {"content": content}
                    logger.debug(f"Saving {section_name} with {len(content)} chars as simple content field")
                elif section_name == "education":
                    # For education, convert to proper JSON format (not raw string)
                    # Try to parse education data
                    json_data = {"content": content}
                    logger.debug(f"Saving {section_name} with {len(content)} chars in structured JSON format")
                elif section_name == "projects":
                    # For projects, convert to proper JSON format (not raw string)
                    json_data = {"content": content}
                    logger.debug(f"Saving {section_name} with {len(content)} chars in structured JSON format")
                else:
                    # For other sections, use simple content field
                    json_data = {"content": content}
                    logger.debug(f"Saving {section_name} with {len(content)} chars as simple content field")
                
                # Define the output file path (non-timestamped)
                filepath = os.path.join(api_responses_dir, f"{section_name}.json")
                
                # Write to file
                with open(filepath, 'w', encoding='utf-8') as f:
                    json.dump(json_data, f, indent=2)  # Always use json.dump, not raw strings
                
                sections_saved += 1
                logger.info(f"Saved {section_name} content to {filepath}")
            
            logger.info(f"Saved {sections_saved} sections to non-timestamped JSON files")
            
            # Verify contact.json was created
            contact_path = os.path.join(api_responses_dir, "contact.json")
            if os.path.exists(contact_path):
                logger.info(f"Verified contact.json exists at {contact_path}")
            else:
                logger.warning(f"Contact.json was not created at {contact_path}")
                
            # Verify summary.json was created
            summary_path = os.path.join(api_responses_dir, "summary.json")
            if os.path.exists(summary_path):
                logger.info(f"Verified summary.json exists at {summary_path}")
            else:
                logger.warning(f"Summary.json was not created at {summary_path}")
                
            return True
            
        except Exception as e:
            logger.error(f"Error saving tailored content to JSON: {str(e)}")
            logger.error(traceback.format_exc())
            return False
            
    def _extract_experience_from_html(self, html_content):
        """Extract structured experience data from HTML content"""
        try:
            # Initialize the list to hold job entries
            experience_entries = []
            
            # Simple pattern matching to extract job entries
            # This is a basic implementation - in a production environment, 
            # consider using a proper HTML parser like BeautifulSoup
            
            # Define regex patterns to identify parts of the experience entry
            company_pattern = r'<span class=[\'"]company[\'"]>(.*?)</span>'
            location_pattern = r'<span class=[\'"]location[\'"]>(.*?)</span>'
            position_pattern = r'<span class=[\'"]position[\'"]>(.*?)</span>'
            dates_pattern = r'<span class=[\'"]dates[\'"]>(.*?)</span>'
            
            # Split HTML into job blocks (roughly by top-level paragraphs)
            job_blocks = re.split(r'<p[^>]*>', html_content)
            
            current_job = {}
            content_items = []
            
            for block in job_blocks:
                if 'class=\'company\'' in block or 'class="company"' in block:
                    # If we already have a job being built, add it to our list
                    if current_job and 'company' in current_job:
                        if content_items:
                            current_job['content'] = content_items
                        experience_entries.append(current_job)
                        content_items = []
                    
                    # Start a new job
                    current_job = {}
                    
                    # Extract company
                    company_match = re.search(company_pattern, block)
                    if company_match:
                        current_job['company'] = company_match.group(1)
                    
                    # Extract location 
                    location_match = re.search(location_pattern, block)
                    if location_match:
                        current_job['location'] = location_match.group(1)
                
                elif 'class=\'position\'' in block or 'class="position"' in block:
                    # Extract position
                    position_match = re.search(position_pattern, block)
                    if position_match:
                        current_job['position'] = position_match.group(1)
                    
                    # Extract dates
                    dates_match = re.search(dates_pattern, block)
                    if dates_match:
                        current_job['dates'] = dates_match.group(1)
                
                elif '<li>' in block:
                    # Extract content items (bullet points)
                    items = re.findall(r'<li>(.*?)</li>', block)
                    content_items.extend(items)
                
                elif block.strip() and 'company' in current_job and not block.startswith(('</p>', '</ul>')):
                    # Add any other significant content
                    plain_text = re.sub(r'<[^>]+>', '', block).strip()
                    if plain_text:
                        content_items.append(plain_text)
            
            # Add the last job if we have one
            if current_job and 'company' in current_job:
                if content_items:
                    current_job['content'] = content_items
                experience_entries.append(current_job)
            
            return experience_entries
        except Exception as e:
            logger.error(f"Error extracting experience from HTML: {str(e)}")
            return []


def validate_bullet_point_cleaning(sections: Dict[str, str]) -> bool:
    """
    Validates that bullet points have been properly cleaned from resume sections.

    Args:
        sections (Dict[str, str]): Dictionary of resume sections

    Returns:
        bool: True if validation passed, False if bullet points were found
    """
    bullet_patterns = [
        # Unicode bullets
        r'^[•◦▪▫■□▸►▹▻▷▶→⇒⟹⟶⇢⇨⟾➔➜➙➛➝➞➟➠➡➢➣➤➥➦➧➨➩➪➫➬➭➮➯➱➲➳➵➸➼⦿⦾⧫⧮⧠⧔∙◆◇◈]',
        # ASCII bullets at beginning of line (after any whitespace)
        r'^\s*[\*\-\+o~=#>]\s+',
        # Textual escape forms for bullets (e.g., u2022, \u2022, U+2022, &#8226;, &bull;)
        r'^\s*(?:u2022|\\u2022|U\+2022|&#8226;|&bull;)\s+',
        # Numbered bullets
        r'^\s*(?:\(?\d+[\.\)\]]\s+|\d+[\.\)\]]\s+|\[\d+\]\s+)'
    ]

    for section, content in sections.items():
        if not content:
            continue
            
        lines = content.split('\n')
        for i, line in enumerate(lines):
            for pattern in bullet_patterns:
                if re.search(pattern, line):
                    logger.warning(f"Validation found bullet marker in {section} section, line {i + 1}: '{line[:30]}...'" )
                    return False

    return True


def validate_html_content(html_content: str) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.validate_html_content(html_content)


def format_section_content(content: str) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.format_section_content(content)


def format_job_entry(company: str, location: str, position: str, dates: str, content: List[str]) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.format_job_entry(company, location, position, dates, content)


def format_education_entry(institution: str, location: str, degree: str, dates: str, highlights: List[str]) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.format_education_entry(institution, location, degree, dates, highlights)


def format_project_entry(title: str, dates: str, details: List[str]) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.format_project_entry(title, dates, details)


def format_education_content(content: str) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.format_education_content(content)


def format_projects_content(content: str) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.format_projects_content(content)


def generate_preview_from_llm_responses(request_id: str, upload_folder: str, for_screen: bool = True) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.generate_preview_from_llm_responses(request_id, upload_folder, for_screen=for_screen)


def extract_resume_sections(doc_path: str) -> Dict[str, str]:
    """Extract sections from a resume document"""
    try:
        logger.info(f"Extracting sections from resume: {doc_path}")
        print(
    f"DEBUG: extract_resume_sections called with doc_path: {doc_path}")
        
        # Import config to check if LLM parsing is enabled
        from config import Config
        use_llm_parsing = Config.USE_LLM_RESUME_PARSING
        llm_provider_config = Config.LLM_RESUME_PARSER_PROVIDER
        
        # First try to parse with LLM if enabled and the module is available
        if use_llm_parsing:
            try:
                # Import the LLM parser module
                from llm_resume_parser import parse_resume_with_llm
                
                # Determine which LLM provider to use based on config or
                # available API keys
                llm_provider = llm_provider_config
                if llm_provider == "auto":
                    if os.environ.get("OPENAI_API_KEY"):
                        llm_provider = "openai"
                    elif os.environ.get("CLAUDE_API_KEY"):
                        llm_provider = "claude"
                    else:
                        logger.warning(
                            "No LLM API keys found for resume parsing. Will use traditional parsing.")
                        raise ImportError("No LLM API keys available")
                    
                # Try to parse with LLM
                logger.info(
    f"Attempting to parse resume with LLM ({llm_provider})...")
                llm_sections = parse_resume_with_llm(doc_path, llm_provider)
                
                # If LLM parsing succeeded, process sections, handling structured data correctly
                if llm_sections and any(
    content for content in llm_sections.values()):
                    logger.info(
                        "LLM parsing successful. Processing LLM-parsed sections.")
                    
                    # Process sections, handling structured data correctly
                    processed_sections = {}
                    for section, content in llm_sections.items():
                        if content:
                            # --- START FIX: Handle experience list --- 
                            if section == 'experience' and isinstance(content, list):
                                # Preserve the list structure for experience
                                processed_sections[section] = content
                                logger.info(f"LLM extracted structured {section} section: {len(content)} entries")
                            elif isinstance(content, str):
                                # Clean bullet points only for string content
                                cleaned_content = clean_bullet_points(content)
                                processed_sections[section] = cleaned_content
                                logger.info(
    f"LLM extracted and cleaned {section} section: {
        len(cleaned_content)} chars")
                            else:
                                # Handle unexpected content types if necessary
                                logger.warning(f"Unexpected content type for LLM section {section}: {type(content)}. Storing as string.")
                                processed_sections[section] = str(content)
                            # --- END FIX ---    
                        else:
                            processed_sections[section] = ""
                            logger.info(
    f"LLM found no content for {section} section")
                    
                    # Validate the cleaning was successful for STRING sections
                    string_sections_for_validation = { 
                        k: v for k, v in processed_sections.items() if isinstance(v, str) 
                    }
                    validation_success = validate_bullet_point_cleaning(
                        string_sections_for_validation)
                    if not validation_success:
                        logger.warning(
                            "Bullet point cleaning validation failed for some string sections.")
                    
                    return processed_sections # Return the processed sections
                else:
                    logger.warning(
                        "LLM parsing did not return usable results. Falling back to traditional parsing.")
            except (ImportError, Exception) as e:
                logger.warning(f"LLM parsing unavailable or failed: {str(e)}. Using traditional parsing.")
        else:
            logger.info(
                "LLM parsing is disabled by configuration. Using traditional parsing.")
        
        # If LLM parsing failed, is unavailable, or is disabled, fall back to
        # traditional parsing
        logger.info("Using traditional resume section extraction...")
        print(
    f"DEBUG: fallback to traditional parsing - will examine document: {doc_path}")
        
        # Extract plain text content from the DOCX file
        text = docx2txt.process(doc_path)
        
        # Parse the document into sections
        print(f"DEBUG: About to create Document object from: {doc_path}")
        doc = Document(doc_path)
        print(f"DEBUG: Document object created successfully - examining sections")
        
        # Initialize standard resume sections
        sections = {
            "contact": "",
            "summary": "",
            "experience": "",
            "education": "",
            "skills": "",
            "projects": "",
            "additional": ""
        }
        
        # Debug: Count total paragraphs with content
        total_paragraphs = sum(
    1 for para in doc.paragraphs if para.text.strip())
        logger.info(f"Total paragraphs with content: {total_paragraphs}")
        
        # Extract sections based on heading style
        current_section = "contact"  # Default to contact for initial content
        section_content = []
        
        # Log all potential section headers for debugging
        logger.info("Scanning document for potential section headers...")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # Debug potential headers
            if para.style.name.startswith('Heading') or any(
                p.bold for p in para.runs):
                logger.info(
    f"Potential header found at para {i}: '{text}' (Style: {
        para.style.name}, Bold: {
            any(
                p.bold for p in para.runs)})")
        
        # First pass - try to extract based on formatting
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # Check if this is a heading (section title)
            if para.style.name.startswith('Heading') or any(
                p.bold for p in para.runs):
                # Store previous section content
                if section_content:
                    sections[current_section] = "\n".join(section_content)
                    section_content = []
                
                # Determine new section type based on heading text
                if re.search(
    r'contact|info|email|phone|address',
     text.lower()):
                    current_section = "contact"
                elif re.search(r'summary|objective|profile|about', text.lower()):
                    current_section = "summary"
                elif re.search(r'experience|work|employment|job|career|professional', text.lower()):
                    current_section = "experience"
                elif re.search(r'education|degree|university|college|school|academic', text.lower()):
                    current_section = "education"
                elif re.search(r'skills|expertise|technologies|competencies|qualification|proficiencies', text.lower()):
                    current_section = "skills"
                elif re.search(r'projects|portfolio|works', text.lower()):
                    current_section = "projects"
                elif re.search(r'additional|interests|activities|volunteer|certification|awards|achievements', text.lower()):
                    current_section = "additional"
                else:
                    # Default to additional for unknown headings
                    current_section = "additional"
                
                logger.info(
    f"Section header detected: '{text}' -> categorized as '{current_section}'")
            else:
                # Add content to current section
                section_content.append(text)
        
        # Add the last section content
        if section_content:
            sections[current_section] = "\n".join(section_content)
        
        # Check if we actually found any sections beyond contact
        sections_found = sum(
    1 for section,
     content in sections.items() if content and section != "contact")
        logger.info(
    f"Sections found with standard detection: {sections_found}")

        # If we didn't find any sections or only found contact, try a simpler
        # approach
        if sections_found == 0:
            logger.info(
                "No sections detected with standard method. Using fallback approach...")
            
            # Fallback: Treat the entire document as experience section
            if total_paragraphs > 0:
                all_content = "\n".join(
    para.text for para in doc.paragraphs if para.text.strip())
                if len(all_content) > 0:
                    sections["experience"] = all_content
                    logger.info(
    f"Fallback: Added {
        len(all_content)} chars to experience section")
        
        # Clean bullet points from traditional parsing results
        cleaned_sections = {}
        for section, content in sections.items():
            if content:
                cleaned_content = clean_bullet_points(content)
                cleaned_sections[section] = cleaned_content
                logger.info(
    f"Traditional parsing: cleaned {section} section: {
        len(cleaned_content)} chars")
            else:
                cleaned_sections[section] = ""
        
        # Validate the cleaning was successful
        validation_success = validate_bullet_point_cleaning(cleaned_sections)
        if not validation_success:
            logger.warning(
                "Bullet point cleaning validation failed. There may still be bullet markers in the text.")
        
        return cleaned_sections
    
    except Exception as e:
        logger.error(f"Error extracting resume sections: {str(e)}")
        logger.error(traceback.format_exc())
        
        # Return at least an empty structure in case of error
        return {
            "contact": "",
            "summary": "",
            "experience": "Error extracting content. Please check your resume format.",
            "education": "",
            "skills": "",
            "projects": "",
            "additional": ""
        }


def tailor_resume_with_llm(
    resume_path: str,
    job_data: Dict,
    api_key: str,
    provider: str = "openai",
    api_url: str = None,
    request_id: str = None
) -> Tuple[Dict[str, Any], Union[ClaudeClient, OpenAIClient]]:
    """
    Main function to tailor resume using either Claude or OpenAI LLM
    
    Args:
        resume_path (str): Path to the original resume file (DOCX or PDF)
        job_data (Dict): Dictionary containing job requirements and skills
        api_key (str): API key for the selected provider
        provider (str): LLM provider ("claude" or "openai")
        api_url (str): API URL for Claude (optional)
        request_id (str): Unique identifier for this tailoring request
        
    Returns:
        Tuple[Dict[str, Any], Union[ClaudeClient, OpenAIClient]]: Dictionary of tailored sections, LLM client instance
    """
    global last_llm_client
    
    logger.info(f"Tailoring resume with {provider} LLM")
    
    # Extract resume sections
    resume_sections = extract_resume_sections(resume_path)
    
    # Initialize the appropriate LLM client
    llm_client = None
    if provider.lower() == "claude":
        llm_client = ClaudeClient(api_key, api_url)
    elif provider.lower() == "openai":
        llm_client = OpenAIClient(api_key)
    else:
        raise ValueError(f"Unsupported LLM provider: {provider}")
    
    # Tailor each section
    tailored_sections = {}
    section_order = ["contact", "summary", "experience", "education", "skills", "projects"]

    # Preserve contact section directly
    if "contact" in resume_sections:
        tailored_sections["contact"] = resume_sections["contact"].strip()
        # Store directly in client for potential later use (though we'll save from tailored_sections)
        logger.info(f"Preserving contact section for tailored resume: {len(tailored_sections['contact'])} chars")
    else:
        logger.warning("Contact section not found in resume")
        tailored_sections["contact"] = "" # Ensure key exists

    # Handle summary - generate if missing or tailor if present
    if "summary" not in resume_sections or not resume_sections["summary"].strip():
        logger.warning("No summary information found in resume sections, generating a new summary")
        try:
            tailored_sections["summary"] = generate_professional_summary(resume_sections, job_data, llm_client, provider)
            logger.info(f"Generated new professional summary: {len(tailored_sections['summary'])} chars")
        except Exception as summary_err:
            logger.error(f"Error generating professional summary: {summary_err}")
            tailored_sections["summary"] = "" # Default to empty if generation fails
    else:
        logger.info("Tailoring existing summary section")
        tailored_sections["summary"] = llm_client.tailor_resume_content("summary", resume_sections["summary"], job_data)

    # Tailor other sections
    for section_name in section_order:
        if section_name not in ["contact", "summary"]:
            # --- START FIX: Handle list type for experience check ---
            section_content = resume_sections.get(section_name) # Get content safely
            
            # Check if the section exists and has meaningful content (non-empty string or non-empty list)
            has_content_to_tailor = False
            if isinstance(section_content, str):
                if section_content.strip():
                    has_content_to_tailor = True
            elif isinstance(section_content, list):
                if section_content: # Check if list is not empty
                    has_content_to_tailor = True
            # Add checks for other types if necessary, e.g., dict for skills
            elif isinstance(section_content, dict):
                 if section_content: # Check if dict is not empty
                     has_content_to_tailor = True
            
            if has_content_to_tailor:
            # --- END FIX ---
                logger.info(f"Tailoring {section_name} section")
                # Pass the original content (string or list) directly to tailoring
                tailored_content_result = llm_client.tailor_resume_content(section_name, section_content, job_data)
                tailored_sections[section_name] = tailored_content_result
            else:
                logger.info(f"Skipping empty or missing section: {section_name}")
                # Ensure key exists even if skipped, potentially use the original empty value
                tailored_sections[section_name] = section_content if section_content is not None else ""

    # --- START: New saving logic (Step 3.3c) ---
    try:
        temp_data_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'temp_session_data')
        # Ensure directory exists (might be redundant if created in app.py, but safe)
        os.makedirs(temp_data_dir, exist_ok=True)

        sections_saved_count = 0
        logger.info(f"Attempting to save cleaned sections for request_id: {request_id}")
        # Use tailored_sections collected in this function scope
        for section_name, content in tailored_sections.items():
            if content is not None: # Save even if content is empty string, but not None
                # Determine the filename using request_id
                cleaned_filepath = os.path.join(temp_data_dir, f"{request_id}_{section_name}.json")
                # Save the actual content (which should be the final string or structured data)
                # We need to decide the format - let's save the direct content string/structure
                # For now, let's assume the values in tailored_sections are the final desired JSON-compatible structures
                # or strings. We'll save them directly.
                # TODO: Verify the format stored in tailored_sections matches what html_generator expects.
                # Let's wrap simple strings in a basic JSON structure for consistency.
                data_to_save = {}
                if isinstance(content, str):
                     # Attempt to parse if it looks like JSON, otherwise wrap it
                    try:
                        parsed_json = json.loads(content)
                        data_to_save = parsed_json # Save parsed JSON if valid
                    except json.JSONDecodeError:
                        # If not valid JSON, wrap the string content
                        data_to_save = { "content": content }
                elif isinstance(content, (dict, list)):
                    data_to_save = content # Save dicts/lists directly
                else:
                    # Handle other types if necessary, or log a warning
                    logger.warning(f"Unexpected data type for section {section_name}: {type(content)}. Saving as string.")
                    data_to_save = { "content": str(content) }

                try:
                    with open(cleaned_filepath, 'w', encoding='utf-8') as f:
                        json.dump(data_to_save, f, indent=2, ensure_ascii=False)
                    sections_saved_count += 1
                    logger.info(f"Saved cleaned {section_name} content for request {request_id} to {cleaned_filepath}")
                except Exception as save_err:
                    logger.error(f"Error saving cleaned section {section_name} for request {request_id} to {cleaned_filepath}: {save_err}")
            else:
                 logger.warning(f"Content for section {section_name} is None, skipping save for request {request_id}.")

        logger.info(f"Saved {sections_saved_count} cleaned sections for request {request_id} to {temp_data_dir}")

    except Exception as e:
        logger.error(f"Error saving cleaned session data for request {request_id}: {e}")
        logger.error(traceback.format_exc())
        # Decide if we should raise an error or just log
    # --- END: New saving logic ---

    # Return the collected tailored sections and the client instance
    last_llm_client = llm_client # Store client for potential reuse?
    return tailored_sections, llm_client

def generate_professional_summary(
    resume_sections: Dict[str, str],
    job_data: Dict,
    llm_client,
    provider: str
) -> str:
    """
    Generate a professional summary based on resume content and job data
    
    Args:
        resume_sections: Dictionary of resume sections
        job_data: Job data including requirements and skills
        llm_client: LLM client instance
        provider: LLM provider name
        
    Returns:
        Generated professional summary or empty string if failed
    """
    try:
        # Collect relevant information from the resume
        experience = resume_sections.get("experience", "")
        education = resume_sections.get("education", "")
        skills = resume_sections.get("skills", "")
        
        # Extract key information from job data
        job_title = job_data.get("title", "")
        job_requirements = job_data.get("requirements", "")
        candidate_profile = ""
        hard_skills = []
        soft_skills = []
        
        # Check if we have LLM-analyzed job data with more structured information
        job_analysis = job_data.get("analysis", {})
        if job_analysis:
            if "candidate_profile" in job_analysis:
                candidate_profile = job_analysis["candidate_profile"]
            if "hard_skills" in job_analysis:
                hard_skills = job_analysis["hard_skills"]
            if "soft_skills" in job_analysis:
                soft_skills = job_analysis["soft_skills"]
        
        # Create a combined string of skills for the prompt
        target_skills = ", ".join(hard_skills + soft_skills)
        
        # Create a prompt for summary generation
        prompt = f"""
Generate a professional summary section for a resume tailored to the following job:

JOB TITLE: {job_title}

CANDIDATE PROFILE:
{candidate_profile}

TARGET SKILLS:
{target_skills}

REQUIREMENTS:
{job_requirements}

Based on the following resume information:

EXPERIENCE:
{experience[:500]}...

EDUCATION:
{education[:300]}

SKILLS:
{skills[:300]}

Create a concise, powerful professional summary (3-5 sentences) that:
1. Positions the candidate as an ideal fit for this specific role
2. Highlights relevant experience and skills that match the job requirements
3. Demonstrates the candidate's unique value proposition
4. Is written in first person without using "I" statements
5. Uses strong action verbs and specific, quantifiable achievements where possible

Format the summary as a single paragraph without bullet points.
"""
        
        # Use the appropriate LLM to generate the summary
        if provider.lower() == "claude":
            summary = generate_summary_with_claude(prompt, llm_client)
        else:
            summary = generate_summary_with_openai(prompt, llm_client)
        
        return summary
        
    except Exception as e:
        logger.error(f"Error generating professional summary: {str(e)}")
        logger.error(traceback.format_exc())
        return ""

def generate_summary_with_claude(prompt: str, claude_client) -> str:
    """Generate a summary using Claude API"""
    try:
        # Use the Claude client to generate a summary
        response = claude_client.client.messages.create(
            model="claude-3-sonnet-20240229",
            max_tokens=400,
            temperature=0.5,
            system="""You are a professional resume writer specializing in creating concise, 
            impactful professional summaries. Your summaries highlight a candidate's strengths 
            and align with job requirements without hyperbole.""",
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        
        # Extract and clean the summary
        summary = response.content[0].text
        
        # Save the raw response
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        response_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'api_responses')
        if not os.path.exists(response_dir):
            os.makedirs(response_dir)
        
        with open(os.path.join(response_dir, f"summary_generation_response_{timestamp}.json"), "w") as f:
            json.dump({"provider": "claude", "prompt": prompt, "response": summary}, f, indent=2)
        
        logger.info(f"Generated summary with Claude: {len(summary)} chars")
        return summary
    
    except Exception as e:
        logger.error(f"Error generating summary with Claude: {str(e)}")
        return ""

def generate_summary_with_openai(prompt: str, openai_client) -> str:
    """Generate a summary using OpenAI API"""
    try:
        # Use the OpenAI client to generate a summary
        response = openai_client.client.chat.completions.create(
            model="gpt-4o",
            temperature=0.5,
            messages=[
                {"role": "system", "content": """You are a professional resume writer specializing in creating concise, 
                impactful professional summaries. Your summaries highlight a candidate's strengths 
                and align with job requirements without hyperbole."""},
                {"role": "user", "content": prompt}
            ]
        )
        
        # Extract the summary
        summary = response.choices[0].message.content
        
        # Save the raw response
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        response_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'api_responses')
        if not os.path.exists(response_dir):
            os.makedirs(response_dir)
        
        with open(os.path.join(response_dir, f"summary_generation_response_{timestamp}.json"), "w") as f:
            json.dump({
                "provider": "openai", 
                "prompt": prompt, 
                "response": summary,
                "prompt_tokens": response.usage.prompt_tokens,
                "completion_tokens": response.usage.completion_tokens,
                "total_tokens": response.usage.total_tokens
            }, f, indent=2)
        
        logger.info(f"Generated summary with OpenAI: {len(summary)} chars")
        return summary
    
    except Exception as e:
        logger.error(f"Error generating summary with OpenAI: {str(e)}")
        return ""

def generate_resume_preview(resume_path: str) -> str:
    """
    Generate HTML preview from resume
    
    Args:
        resume_path: Path to resume file
        
    Returns:
        HTML preview
    """
    logger.info(f"Generating resume preview for {resume_path}")
    
    # Check if we have direct LLM responses from last tailoring
    global last_llm_client
    if last_llm_client:
        logger.info("Using direct LLM responses for preview generation")
        preview_html = generate_preview_from_llm_responses(last_llm_client)
        if preview_html:
            return preview_html
    
    # Otherwise extract sections from file
    resume_sections = extract_resume_sections(resume_path)
    
    # Generate preview HTML
    preview_html = ""
    for section_name, content in resume_sections.items():
        if content.strip():
            formatted_content = format_section_content(content)
            preview_html += f'<div class="resume-section"><h2>{section_name.title()}</h2>{formatted_content}</div>'
    
    return preview_html
