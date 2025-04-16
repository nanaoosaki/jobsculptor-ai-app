import os
from typing import Dict, List, Optional, Union
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class YCEddieStyler:
    """
    YC-Eddie Style Resume Formatter
    
    Implements a consistent, professional resume style based on YC and Eddie's best practices.
    Creates a new document with standardized formatting rather than using the original as a template.
    """
    
    def __init__(self):
        """Initialize a new document with YC-Eddie styling"""
        print("DEBUG: Creating brand new DOCX document from scratch")
        self.document = Document()
        print("DEBUG: New document created, applying formatting")
        self._setup_document_properties()
        self._create_styles()
        print("DEBUG: Document properties and styles set up")
        
        # Ensure we have a clean document with no paragraphs
        if self.document.paragraphs:
            print(f"DEBUG: WARNING - New document has {len(self.document.paragraphs)} existing paragraphs, clearing them")
            for para in list(self.document.paragraphs):
                p = para._element
                if p.getparent() is not None:
                    p.getparent().remove(p)
            print("DEBUG: Document cleared of any existing content")
        
    def _setup_document_properties(self):
        """Set up document properties like margins and page size"""
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            
    def _create_styles(self):
        """Create all required paragraph styles for the document"""
        print("DEBUG: Creating custom document styles")
        
        # Default style (normal text)
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Contact info style
        contact_style = self.document.styles.add_style('Contact', WD_STYLE_TYPE.PARAGRAPH)
        font = contact_style.font
        font.name = 'Calibri'
        font.size = Pt(12)
        contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_style.paragraph_format.space_after = Pt(6)
        
        # Section header style - centered with box border
        header_style = self.document.styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
        font = header_style.font
        font.name = 'Calibri'
        font.size = Pt(14)
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 102)  # Dark blue TEXT ONLY
        header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_style.paragraph_format.space_after = Pt(6)
        
        # IMPORTANT: DO NOT set background shading for headers
        print("DEBUG: Created section header style with NO background color")
        
        # Bullet point style
        bullet_style = self.document.styles.add_style('BulletPoint', WD_STYLE_TYPE.PARAGRAPH)
        font = bullet_style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        bullet_style.paragraph_format.left_indent = Inches(0.25)
        bullet_style.paragraph_format.first_line_indent = Inches(-0.25)
        bullet_style.paragraph_format.space_after = Pt(6)
        
        # Company/role style
        role_style = self.document.styles.add_style('CompanyRole', WD_STYLE_TYPE.PARAGRAPH)
        font = role_style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.bold = True
        role_style.paragraph_format.space_after = Pt(3)
        
        # Date range style
        date_style = self.document.styles.add_style('DateRange', WD_STYLE_TYPE.PARAGRAPH)
        font = date_style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.italic = True
        date_style.paragraph_format.space_after = Pt(6)
        
        print("DEBUG: All document styles created successfully")
        
    def add_box_border(self, paragraph):
        """Add a box border (all four sides) to a paragraph"""
        p = paragraph._p  # Get the paragraph XML element
        pPr = p.get_or_add_pPr()  # Get the paragraph properties element
        pBdr = OxmlElement('w:pBdr')  # Create paragraph borders element
        
        # Add all four borders (top, right, bottom, left)
        for border_position in ['top', 'right', 'bottom', 'left']:
            border = OxmlElement(f'w:{border_position}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # 1pt
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Black border
            pBdr.append(border)
        
        pPr.append(pBdr)
        
    def add_bullet_point(self, paragraph, text):
        """Add a bullet point with an arrow character instead of a regular bullet"""
        run = paragraph.add_run('▸ ')  # Arrow character
        paragraph.add_run(text.strip())
        return paragraph
        
    def add_contact_section(self, contact_info: str):
        """Add the contact information section to the document"""
        contact_lines = contact_info.strip().split('\n')
        
        # First line is usually the name - make it bold and slightly larger
        if contact_lines:
            name_para = self.document.add_paragraph(style='Contact')
            name_run = name_para.add_run(contact_lines[0].strip())
            name_run.bold = True
            name_run.font.size = Pt(16)
            
            # Add remaining contact info
            for line in contact_lines[1:]:
                if line.strip():
                    contact_para = self.document.add_paragraph(style='Contact')
                    contact_para.add_run(line.strip())
        
        # Add a small space after contact section
        self.document.add_paragraph()
        
    def add_section_header(self, title: str):
        """Add a section header with proper styling"""
        print(f"DEBUG: Adding section header: {title}")
        
        # Add some spacing before section
        self.document.add_paragraph().add_run()
        
        # Create a new paragraph with the SectionHeader style
        header_para = self.document.add_paragraph(style='SectionHeader')
        
        # Add the section title in all caps
        run = header_para.add_run(title.upper())
        
        # Ensure the run has proper formatting (dark blue text, bold)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 102)  # Dark blue
        
        # Add the box border
        self.add_box_border(header_para)
        
        print(f"DEBUG: Finished adding section header with box border")
        
    def add_summary_section(self, summary_content: str):
        """Add the summary section to the document"""
        self.add_section_header("SUMMARY")
        
        if summary_content:
            summary_para = self.document.add_paragraph(style='Normal')
            summary_para.add_run(summary_content.strip())
    
    def add_experience_section(self, experience_content: str):
        """Add the experience section with proper formatting for companies, roles, dates and bullet points"""
        self.add_section_header("EXPERIENCE")
        
        if not experience_content:
            return
            
        # Split the experience section by companies/roles
        experiences = self._parse_experiences(experience_content)
        
        for exp in experiences:
            # Add company and role
            if exp.get('company_role'):
                company_para = self.document.add_paragraph(style='CompanyRole')
                company_para.add_run(exp['company_role'].strip())
                
            # Add date range if available
            if exp.get('date_range'):
                date_para = self.document.add_paragraph(style='DateRange')
                date_para.add_run(exp['date_range'].strip())
                
            # Add bullet points
            for bullet in exp.get('bullets', []):
                bullet_para = self.document.add_paragraph(style='BulletPoint')
                self.add_bullet_point(bullet_para, bullet.strip())
    
    def add_education_section(self, education_content: str):
        """Add the education section to the document"""
        self.add_section_header("EDUCATION")
        
        if not education_content:
            return
            
        # Split education entries and format each one
        education_entries = self._parse_education(education_content)
        
        for edu in education_entries:
            # Add institution and degree
            if edu.get('institution_degree'):
                edu_para = self.document.add_paragraph(style='CompanyRole')
                edu_para.add_run(edu['institution_degree'].strip())
                
            # Add date range if available
            if edu.get('date_range'):
                date_para = self.document.add_paragraph(style='DateRange')
                date_para.add_run(edu['date_range'].strip())
                
            # Add details or achievements
            for detail in edu.get('details', []):
                detail_para = self.document.add_paragraph(style='BulletPoint')
                self.add_bullet_point(detail_para, detail.strip())
                
    def add_skills_section(self, skills_content: str):
        """Add the skills section to the document"""
        self.add_section_header("SKILLS")
        
        if not skills_content:
            return
            
        # Handle different skill formats (paragraph or bullet list)
        lines = skills_content.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if line starts with a bullet point marker
            if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                skill_para = self.document.add_paragraph(style='BulletPoint')
                # Ensure consistent bullet style regardless of original marker
                clean_line = line.lstrip('•-* \t')
                self.add_bullet_point(skill_para, clean_line)
            else:
                # If not a bullet, check if it seems like a category with skills
                if ':' in line:
                    category, skills_list = line.split(':', 1)
                    skill_para = self.document.add_paragraph(style='Normal')
                    skill_para.add_run(category.strip() + ': ').bold = True
                    skill_para.add_run(skills_list.strip())
                else:
                    # Regular line
                    skill_para = self.document.add_paragraph(style='Normal')
                    skill_para.add_run(line)
                    
    def add_projects_section(self, projects_content: str):
        """Add the projects section to the document"""
        self.add_section_header("PROJECTS")
        
        if not projects_content:
            return
            
        # Split the projects section by project entries
        projects = self._parse_projects(projects_content)
        
        for project in projects:
            # Add project name
            if project.get('project_name'):
                project_para = self.document.add_paragraph(style='CompanyRole')
                project_para.add_run(project['project_name'].strip())
                
            # Add date range if available
            if project.get('date_range'):
                date_para = self.document.add_paragraph(style='DateRange')
                date_para.add_run(project['date_range'].strip())
                
            # Add bullet points
            for bullet in project.get('bullets', []):
                bullet_para = self.document.add_paragraph(style='BulletPoint')
                self.add_bullet_point(bullet_para, bullet.strip())
                
    def add_additional_section(self, additional_content: str, section_title: str = "ADDITIONAL INFORMATION"):
        """Add an additional information section to the document"""
        self.add_section_header(section_title)
        
        if not additional_content:
            return
            
        # Split the content by lines
        lines = additional_content.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if line starts with a bullet point marker
            if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                para = self.document.add_paragraph(style='BulletPoint')
                clean_line = line.lstrip('•-* \t')
                self.add_bullet_point(para, clean_line)
            else:
                # Regular paragraph
                para = self.document.add_paragraph(style='Normal')
                para.add_run(line)
    
    def _parse_experiences(self, experience_content: str) -> List[Dict[str, Union[str, List[str]]]]:
        """Parse experience content into structured data with company/role, date range, and bullets"""
        result = []
        current_experience = {}
        current_bullets = []
        
        lines = experience_content.strip().split('\n')
        
        # Simple state machine for parsing
        state = "company_role"  # Start with company/role
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if state == "company_role":
                current_experience = {'company_role': line}
                state = "date_range"
                
            elif state == "date_range":
                current_experience['date_range'] = line
                current_bullets = []
                state = "bullets"
                
            elif state == "bullets":
                # Check if this line looks like a new company
                if not (line.startswith('•') or line.startswith('-') or line.startswith('*')):
                    # Save the previous experience entry
                    if current_experience:
                        current_experience['bullets'] = current_bullets
                        result.append(current_experience)
                    
                    # Start a new experience entry
                    current_experience = {'company_role': line}
                    state = "date_range"
                else:
                    # Add bullet point
                    clean_line = line.lstrip('•-* \t')
                    current_bullets.append(clean_line)
        
        # Add the last experience entry
        if current_experience:
            current_experience['bullets'] = current_bullets
            result.append(current_experience)
            
        return result
    
    def _parse_education(self, education_content: str) -> List[Dict[str, Union[str, List[str]]]]:
        """Parse education content into structured data with institution/degree, date range, and details"""
        result = []
        current_education = {}
        current_details = []
        
        lines = education_content.strip().split('\n')
        
        # Simple state machine for parsing
        state = "institution_degree"  # Start with institution/degree
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if state == "institution_degree":
                current_education = {'institution_degree': line}
                state = "date_range"
                
            elif state == "date_range":
                current_education['date_range'] = line
                current_details = []
                state = "details"
                
            elif state == "details":
                # Check if this line looks like a new institution
                if not (line.startswith('•') or line.startswith('-') or line.startswith('*')):
                    # Save the previous education entry
                    if current_education:
                        current_education['details'] = current_details
                        result.append(current_education)
                    
                    # Start a new education entry
                    current_education = {'institution_degree': line}
                    state = "date_range"
                else:
                    # Add detail
                    clean_line = line.lstrip('•-* \t')
                    current_details.append(clean_line)
        
        # Add the last education entry
        if current_education:
            current_education['details'] = current_details
            result.append(current_education)
            
        return result
    
    def _parse_projects(self, projects_content: str) -> List[Dict[str, Union[str, List[str]]]]:
        """Parse projects content into structured data with project name, date range, and bullet points"""
        result = []
        current_project = {}
        current_bullets = []
        
        lines = projects_content.strip().split('\n')
        
        # Simple state machine for parsing
        state = "project_name"  # Start with project name
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if state == "project_name":
                current_project = {'project_name': line}
                state = "date_range"
                
            elif state == "date_range":
                current_project['date_range'] = line
                current_bullets = []
                state = "bullets"
                
            elif state == "bullets":
                # Check if this line looks like a new project
                if not (line.startswith('•') or line.startswith('-') or line.startswith('*')):
                    # Save the previous project entry
                    if current_project:
                        current_project['bullets'] = current_bullets
                        result.append(current_project)
                    
                    # Start a new project entry
                    current_project = {'project_name': line}
                    state = "date_range"
                else:
                    # Add bullet point
                    clean_line = line.lstrip('•-* \t')
                    current_bullets.append(clean_line)
        
        # Add the last project entry
        if current_project:
            current_project['bullets'] = current_bullets
            result.append(current_project)
            
        return result
    
    def save(self, output_path: str):
        """Save the document to the specified path"""
        self.document.save(output_path)


def create_resume_document(
    contact: str = "", 
    summary: str = "", 
    experience: str = "", 
    education: str = "", 
    skills: str = "", 
    projects: str = "", 
    additional: str = "",
    output_path: str = "resume_yc_eddie_style.docx"
) -> str:
    """
    Create a complete resume document with YC-Eddie styling and save to the specified path
    
    Args:
        contact: Contact information text
        summary: Summary/Profile text
        experience: Experience section text
        education: Education section text
        skills: Skills section text
        projects: Projects section text
        additional: Additional information text
        output_path: Path to save the document
        
    Returns:
        The path where the document was saved
    """
    print(f"DEBUG: Creating new resume document from scratch using YCEddieStyler")
    print(f"DEBUG: Output path will be: {output_path}")
    print(f"DEBUG: Contact info length: {len(contact)} chars")
    print(f"DEBUG: Summary length: {len(summary)} chars")
    print(f"DEBUG: Experience length: {len(experience)} chars")
    print(f"DEBUG: Education length: {len(education)} chars")
    print(f"DEBUG: Skills length: {len(skills)} chars")
    print(f"DEBUG: Projects length: {len(projects)} chars")
    
    # Create a completely new styler instance with a fresh document
    styler = YCEddieStyler()
    
    print(f"DEBUG: Created new YCEddieStyler instance with clean document")
    
    # Add each section to the document
    if contact:
        styler.add_contact_section(contact)
        print(f"DEBUG: Added contact section")
    
    if summary:
        styler.add_summary_section(summary)
        print(f"DEBUG: Added summary section")
    
    if experience:
        styler.add_experience_section(experience)
        print(f"DEBUG: Added experience section")
    
    if education:
        styler.add_education_section(education)
        print(f"DEBUG: Added education section")
    
    if skills:
        styler.add_skills_section(skills)
        print(f"DEBUG: Added skills section")
    
    if projects:
        styler.add_projects_section(projects)
        print(f"DEBUG: Added projects section")
    
    if additional:
        styler.add_additional_section(additional)
        print(f"DEBUG: Added additional section")
    
    # Save the document
    print(f"DEBUG: About to save document to: {output_path}")
    styler.save(output_path)
    print(f"DEBUG: Document saved successfully")
    
    return output_path 