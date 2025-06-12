import os
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def analyze_template(template_path):
    """Analyze template document structure and styles"""
    template_doc = docx.Document(template_path)
    
    # Extract styles
    styles = {
        'paragraph_styles': [],
        'character_styles': [],
        'table_styles': []
    }
    
    for style in template_doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            styles['paragraph_styles'].append(style.name)
        elif style.type == WD_STYLE_TYPE.CHARACTER:
            styles['character_styles'].append(style.name)
        elif style.type == WD_STYLE_TYPE.TABLE:
            styles['table_styles'].append(style.name)
    
    # Extract sections and structure
    sections = []
    current_section = {'heading': '', 'content': []}
    
    for para in template_doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Check if this is a heading
        if any(run.bold for run in para.runs) or para.style.name.startswith('Heading'):
            # Save previous section if it exists
            if current_section['heading']:
                sections.append(current_section)
            
            # Start new section
            current_section = {'heading': text, 'content': []}
        else:
            current_section['content'].append({
                'text': text,
                'style': para.style.name,
                'alignment': para.alignment
            })
    
    # Add the last section
    if current_section['heading']:
        sections.append(current_section)
    
    # Extract document properties
    properties = {
        'margins': {
            'top': template_doc.sections[0].top_margin.inches,
            'bottom': template_doc.sections[0].bottom_margin.inches,
            'left': template_doc.sections[0].left_margin.inches,
            'right': template_doc.sections[0].right_margin.inches
        },
        'page_size': {
            'width': template_doc.sections[0].page_width.inches,
            'height': template_doc.sections[0].page_height.inches
        }
    }
    
    return {
        'styles': styles,
        'sections': sections,
        'properties': properties
    }

def format_resume_to_template(user_resume_path, template_path, output_path):
    """Format user resume according to template structure"""
    # Analyze template
    template_analysis = analyze_template(template_path)
    
    # Load user resume
    user_doc = docx.Document(user_resume_path)
    
    # Create new document based on template
    new_doc = docx.Document(template_path)
    
    # Clear template content while preserving styles
    for i in range(len(new_doc.paragraphs)-1, -1, -1):
        p = new_doc.paragraphs[i]
        p_element = p._element
        p_element.getparent().remove(p_element)
    
    # Extract content from user resume
    user_content = {
        'paragraphs': [],
        'tables': []
    }
    
    for para in user_doc.paragraphs:
        if para.text.strip():
            user_content['paragraphs'].append({
                'text': para.text,
                'style': para.style.name,
                'alignment': para.alignment,
                'bold': any(run.bold for run in para.runs),
                'italic': any(run.italic for run in para.runs)
            })
    
    for table in user_doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data.append(row_data)
        user_content['tables'].append(table_data)
    
    # Categorize user content
    sections = {
        'contact_info': [],
        'summary': [],
        'education': [],
        'experience': [],
        'skills': [],
        'projects': [],
        'other': []
    }
    
    current_section = 'other'
    
    for para in user_content['paragraphs']:
        text = para['text'].lower()
        
        # Determine section based on heading text
        if para['bold'] or para['style'].startswith('Heading'):
            if any(term in text for term in ['contact', 'email', 'phone', 'address']):
                current_section = 'contact_info'
            elif any(term in text for term in ['summary', 'objective', 'profile']):
                current_section = 'summary'
            elif any(term in text for term in ['education', 'academic']):
                current_section = 'education'
            elif any(term in text for term in ['experience', 'employment', 'work']):
                current_section = 'experience'
            elif any(term in text for term in ['skill', 'technology', 'competenc']):
                current_section = 'skills'
            elif any(term in text for term in ['project', 'portfolio']):
                current_section = 'projects'
            else:
                current_section = 'other'
        
        # Add paragraph to current section
        sections[current_section].append(para)
    
    # Add content to new document following template structure
    # Contact Information
    if sections['contact_info']:
        p = new_doc.add_paragraph()
        p.style = 'Heading 1'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("CONTACT INFORMATION")
        run.bold = True
        
        for para in sections['contact_info']:
            if not para['bold']:  # Skip the heading
                p = new_doc.add_paragraph()
                p.style = 'Normal'
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(para['text'])
    
    # Summary
    if sections['summary']:
        p = new_doc.add_paragraph()
        p.style = 'Heading 1'
        run = p.add_run("PROFESSIONAL SUMMARY")
        run.bold = True
        
        for para in sections['summary']:
            if not para['bold']:  # Skip the heading
                p = new_doc.add_paragraph()
                p.style = 'Normal'
                p.add_run(para['text'])
    
    # Experience
    if sections['experience']:
        p = new_doc.add_paragraph()
        p.style = 'Heading 1'
        run = p.add_run("WORK EXPERIENCE")
        run.bold = True
        
        current_job = None
        for para in sections['experience']:
            if para['bold'] and not current_job:  # Job title
                current_job = para['text']
                p = new_doc.add_paragraph()
                p.style = 'Heading 2'
                run = p.add_run(para['text'])
                run.bold = True
            elif para['bold'] and current_job:  # New job title
                current_job = para['text']
                p = new_doc.add_paragraph()
                p.style = 'Heading 2'
                run = p.add_run(para['text'])
                run.bold = True
            else:  # Job details
                p = new_doc.add_paragraph()
                p.style = 'Normal'
                p.add_run(para['text'])
    
    # Education
    if sections['education']:
        p = new_doc.add_paragraph()
        p.style = 'Heading 1'
        run = p.add_run("EDUCATION")
        run.bold = True
        
        for para in sections['education']:
            if para['bold']:  # Institution
                p = new_doc.add_paragraph()
                p.style = 'Heading 2'
                run = p.add_run(para['text'])
                run.bold = True
            else:  # Degree details
                p = new_doc.add_paragraph()
                p.style = 'Normal'
                p.add_run(para['text'])
    
    # Skills
    if sections['skills']:
        p = new_doc.add_paragraph()
        p.style = 'Heading 1'
        run = p.add_run("SKILLS")
        run.bold = True
        
        for para in sections['skills']:
            if not para['bold']:  # Skip the heading
                p = new_doc.add_paragraph()
                p.style = 'Normal'
                p.add_run(para['text'])
    
    # Projects
    if sections['projects']:
        p = new_doc.add_paragraph()
        p.style = 'Heading 1'
        run = p.add_run("PROJECTS")
        run.bold = True
        
        for para in sections['projects']:
            if para['bold']:  # Project title
                p = new_doc.add_paragraph()
                p.style = 'Heading 2'
                run = p.add_run(para['text'])
                run.bold = True
            else:  # Project details
                p = new_doc.add_paragraph()
                p.style = 'Normal'
                p.add_run(para['text'])
    
    # Other sections
    if sections['other']:
        p = new_doc.add_paragraph()
        p.style = 'Heading 1'
        run = p.add_run("ADDITIONAL INFORMATION")
        run.bold = True
        
        for para in sections['other']:
            if not para['bold']:  # Skip headings
                p = new_doc.add_paragraph()
                p.style = 'Normal'
                p.add_run(para['text'])
    
    # Save the formatted resume
    new_doc.save(output_path)
    
    return output_path

def create_formatted_resume(user_resume_filename, upload_folder):
    """Create a formatted resume based on the template"""
    user_resume_path = os.path.join(upload_folder, user_resume_filename)
    template_path = os.path.join(upload_folder, 'template_resume.docx')
    
    # Generate output filename
    output_filename = user_resume_filename.replace('.docx', '_formatted.docx')
    output_path = os.path.join(upload_folder, output_filename)
    
    # Format the resume
    format_resume_to_template(user_resume_path, template_path, output_path)
    
    return output_filename
