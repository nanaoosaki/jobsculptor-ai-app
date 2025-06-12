import os
import json
import time
import docx
import shutil
import traceback
from claude_integration import tailor_resume_with_claude, generate_resume_preview
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

def test_full_workflow():
    """
    Run a complete workflow with a fresh test resume
    """
    # Create a test resume DOCX
    print("Creating test resume...")
    
    test_resume_path = "test_resume.docx"
    doc = docx.Document()
    
    # Add name and contact
    p = doc.add_paragraph()
    p.alignment = 1  # Center
    run = p.add_run("John Smith")
    run.bold = True
    run.font.size = docx.shared.Pt(16)
    
    p = doc.add_paragraph()
    p.alignment = 1  # Center
    p.add_run("555-123-4567 | john.smith@example.com | www.linkedin.com/in/johnsmith")
    
    # Add summary
    p = doc.add_paragraph()
    run = p.add_run("SUMMARY")
    run.bold = True
    
    p = doc.add_paragraph()
    p.add_run("Data scientist with 5 years of experience in machine learning and analytics. Proficient in Python, R, and SQL with a focus on predictive modeling and data visualization.")
    
    # Add experience
    p = doc.add_paragraph()
    run = p.add_run("EXPERIENCE")
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("SENIOR DATA SCIENTIST")
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("XYZ Company, New York, NY | 2020-Present")
    run.italic = True
    
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Developed machine learning models for customer segmentation, increasing marketing ROI by 25%")
    
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Created Python scripts to automate data cleaning and preprocessing, saving 10 hours per week")
    
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Presented findings to executive team and stakeholders, influencing product roadmap")
    
    p = doc.add_paragraph()
    run = p.add_run("DATA ANALYST")
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("ABC Corp, Boston, MA | 2018-2020")
    run.italic = True
    
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Built dashboards and reports using Tableau and Power BI for executive review")
    
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Used SQL queries to extract and analyze customer data for product recommendations")
    
    # Add skills
    p = doc.add_paragraph()
    run = p.add_run("SKILLS")
    run.bold = True
    
    p = doc.add_paragraph()
    p.add_run("Programming: Python, R, SQL, JavaScript")
    
    p = doc.add_paragraph()
    p.add_run("Data Science: Machine Learning, Statistics, A/B Testing, Data Visualization")
    
    p = doc.add_paragraph()
    p.add_run("Tools: Tableau, Power BI, Excel, Git, AWS")
    
    # Add education
    p = doc.add_paragraph()
    run = p.add_run("EDUCATION")
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("MASTER OF SCIENCE IN DATA SCIENCE")
    run.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run("University of Data Science, Boston, MA | 2018")
    run.italic = True
    
    # Save the document
    doc.save(test_resume_path)
    print(f"Test resume created at {test_resume_path}")
    
    # Define job requirements
    job_data = {
        "job_title": "Principal Data Scientist - AI Foundations",
        "company": "Capital One",
        "requirements": [
            "Experience with Python, AWS, PyTorch, Machine Learning",
            "Experience working with cross-functional teams to deliver data science solutions"
        ],
        "skills": [
            "Python",
            "AWS",
            "PyTorch",
            "Machine Learning",
            "Deep Learning",
            "AI",
            "NLP",
            "Data Science",
            "Big Data",
            "Spark"
        ]
    }
    
    # Get API key from environment
    api_key = os.environ.get("CLAUDE_API_KEY") or os.environ.get("ANTHROPIC_API_KEY")
    
    if not api_key:
        print("Error: No Claude API key found in environment variables")
        return
    
    print("Using Claude API key:", api_key[:5] + "...")
    
    # Copy to uploads folder
    upload_folder = "static/uploads"
    os.makedirs(upload_folder, exist_ok=True)
    
    test_resume_upload_path = os.path.join(upload_folder, "test_resume.docx")
    
    shutil.copy(test_resume_path, test_resume_upload_path)
    
    print(f"Copied resume to {test_resume_upload_path}")
    
    # Tailor the resume
    print("\nTailoring resume...")
    
    try:
        output_filename, output_path = tailor_resume_with_claude(
            test_resume_upload_path,
            job_data,
            api_key,
            None
        )
        
        print(f"Resume tailored successfully!")
        print(f"Output file: {output_path}")
        
        # Generate HTML preview
        preview_html = generate_resume_preview(output_path)
        
        # Save HTML to a file for inspection
        with open("test_tailored_preview.html", "w") as f:
            f.write("""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Test Tailored Resume Preview</title>
                <style>
                    body { font-family: Arial, sans-serif; line-height: 1.6; }
                    .resume-preview { max-width: 800px; margin: 0 auto; padding: 20px; }
                    .resume-heading-1 { color: #2c3e50; margin-top: 20px; }
                    .resume-heading-2 { color: #3498db; margin-top: 15px; }
                    .resume-paragraph { margin: 8px 0; }
                    .resume-bullet-list { margin: 8px 0; padding-left: 20px; }
                    .resume-bullet-item { margin: 5px 0; }
                </style>
            </head>
            <body>
                <h1>Test Tailored Resume Preview</h1>
                """ + preview_html + """
            </body>
            </html>
            """)
        
        print(f"\nPreview HTML saved to test_tailored_preview.html")
        print("Open this file in a browser to view the tailored resume")
        
        # Check logs
        print("\nChecking logs for Claude API responses...")
        time.sleep(3)  # Give time for logs to be written
        
        log_dir = "logs"
        log_files = [f for f in os.listdir(log_dir) if f.startswith("claude_api_log_")]
        
        if log_files:
            print(f"\nLog files found:")
            for log_file in log_files:
                log_path = os.path.join(log_dir, log_file)
                print(f" - {log_path}")
                
                # Read and display log content
                try:
                    with open(log_path, 'r', encoding='utf-8') as f:
                        log_content = f.read().strip()
                        if log_content and log_content != "[]":
                            logs = json.loads(log_content)
                            print(f"   Log entries: {len(logs)}")
                            
                            # Look for logs created in the last minute
                            recent_time = time.time() - 300  # Last 5 minutes
                            recent_logs = []
                            
                            for log in logs:
                                log_time = log.get('timestamp', '')
                                if log_time:
                                    try:
                                        # Convert ISO timestamp to seconds since epoch
                                        log_time = time.mktime(time.strptime(log_time.split('.')[0], '%Y-%m-%dT%H:%M:%S'))
                                        if log_time > recent_time:
                                            recent_logs.append(log)
                                    except:
                                        pass
                            
                            print(f"   Recent log entries (last 5 minutes): {len(recent_logs)}")
                            
                            for i, log in enumerate(recent_logs):
                                print(f"\n   Recent log entry #{i+1}:")
                                print(f"   - Timestamp: {log.get('timestamp')}")
                                print(f"   - Section: {log.get('section')}")
                                print(f"   - Resume ID: {log.get('resume_id')}")
                                
                                # Check if we have tailored content
                                response_data = log.get('response', {})
                                if 'tailored_content' in response_data:
                                    content = response_data['tailored_content']
                                    print(f"\n   TAILORED CONTENT PREVIEW:\n   {content[:300]}...")
                                    print(f"\n   Token usage: {log.get('token_usage', {})}")
                        else:
                            print("   Empty log file")
                except Exception as e:
                    print(f"   Error reading log: {str(e)}")
        else:
            print("\nNo log files were found.")
        
    except Exception as e:
        print(f"Error during tailoring: {str(e)}")
        print(traceback.format_exc())
    
    print("\nTest completed!")
    
    # Open the preview in browser
    os.system("start test_tailored_preview.html")

if __name__ == "__main__":
    test_full_workflow() 