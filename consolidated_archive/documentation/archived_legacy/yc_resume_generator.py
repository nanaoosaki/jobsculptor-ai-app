import os
import re
import logging
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.shared import RGBColor
from docx.oxml.ns import qn
from yc_eddie_styler import YCEddieStyler

logger = logging.getLogger(__name__)

class YCResumeGenerator:
    """
    Generates a resume in the YC/Eddie style format.
    
    This class creates a clean, minimalist resume with consistent formatting following
    the YC/Eddie style guidelines, which emphasize:
    - Clean typography with a single font
    - Minimal formatting (bold for headers only)
    - Consistent spacing and alignment
    - No templates or fancy styling
    """
    
    def __init__(self):
        """Initialize the YC Resume Generator with default settings"""
        # Font settings
        self.font_name = "Calibri"
        self.name_size = Pt(16)
        self.header_size = Pt(12)
        self.body_size = Pt(11)
        
        # Spacing settings
        self.paragraph_spacing = Pt(6)
        self.section_spacing = Pt(12)
        
        # Margins (in inches)
        self.margin_top = 0.5
        self.margin_bottom = 0.5
        self.margin_left = 0.75
        self.margin_right = 0.75
    
    def _set_document_properties(self, doc):
        """Set basic document properties like margins"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(self.margin_top)
            section.bottom_margin = Inches(self.margin_bottom)
            section.left_margin = Inches(self.margin_left)
            section.right_margin = Inches(self.margin_right)
    
    def _add_contact_section(self, doc, content):
        """Add the contact information section to the document"""
        if not content:
            logger.warning("No contact information provided")
            return
        
        # Name paragraph with larger font
        name_lines = content.strip().split('\n', 1)
        if name_lines:
            name = name_lines[0].strip()
            name_paragraph = doc.add_paragraph()
            name_run = name_paragraph.add_run(name)
            name_run.font.name = self.font_name
            name_run.font.size = self.name_size
            name_run.font.bold = True
            name_paragraph.space_after = self.paragraph_spacing
        
        # Contact details
        if len(name_lines) > 1:
            contact_details = name_lines[1].strip()
            contact_paragraph = doc.add_paragraph()
            
            # Process contact lines
            for line in contact_details.split('\n'):
                if line.strip():
                    contact_run = contact_paragraph.add_run(line.strip())
                    contact_run.font.name = self.font_name
                    contact_run.font.size = self.body_size
                    contact_paragraph.add_run('\n')
            
            contact_paragraph.space_after = self.section_spacing
    
    def _add_section(self, doc, title, content):
        """Add a general section with title and content to the document"""
        if not content:
            logger.warning(f"No content provided for {title} section")
            return
        
        # Add section header
        header = doc.add_paragraph()
        header_run = header.add_run(title.upper())
        header_run.font.name = self.font_name
        header_run.font.size = self.header_size
        header_run.font.bold = True
        header.space_after = self.paragraph_spacing
        
        # Process content
        current_paragraph = None
        
        # Split content into lines
        lines = content.split('\n')
        
        for line in lines:
            stripped_line = line.strip()
            
            # Skip empty lines
            if not stripped_line:
                if current_paragraph:
                    current_paragraph.space_after = self.paragraph_spacing
                    current_paragraph = None
                continue
            
            # Check if it's a bullet point
            is_bullet = stripped_line.startswith(('•', '-', '*', '·', '○', '▪', '▫', '◦'))
            
            # For bullet points, ensure we have a new paragraph
            if is_bullet:
                current_paragraph = doc.add_paragraph()
                # Remove the bullet character
                if is_bullet:
                    stripped_line = stripped_line[1:].strip()
                
                # Add content
                content_run = current_paragraph.add_run(stripped_line)
                content_run.font.name = self.font_name
                content_run.font.size = self.body_size
                
                # Format as bullet point
                current_paragraph.style = 'List Bullet'
                
            # For non-bullet text
            else:
                # Check if this might be a sub-header (shorter line that's not a continuation)
                if len(stripped_line) < 100 and (not current_paragraph or stripped_line.endswith((':'))):
                    if current_paragraph:
                        current_paragraph.space_after = self.paragraph_spacing
                    
                    current_paragraph = doc.add_paragraph()
                    content_run = current_paragraph.add_run(stripped_line)
                    content_run.font.name = self.font_name
                    content_run.font.size = self.body_size
                    content_run.font.bold = True
                else:
                    # Continue existing paragraph or create new one
                    if not current_paragraph:
                        current_paragraph = doc.add_paragraph()
                    
                    # Add content
                    if current_paragraph.text:
                        current_paragraph.add_run(' ')
                    
                    content_run = current_paragraph.add_run(stripped_line)
                    content_run.font.name = self.font_name
                    content_run.font.size = self.body_size
        
        # Add space after the section
        if current_paragraph:
            current_paragraph.space_after = self.section_spacing
        else:
            # Add an empty paragraph for spacing if no content was added
            empty = doc.add_paragraph()
            empty.space_after = self.section_spacing
    
    def generate_tailored_resume(self, sections, output_path):
        """
        Generate a tailored resume document in YC/Eddie style
        
        Args:
            sections (dict): Dictionary containing the resume sections
            output_path (str): Path where to save the generated document
            
        Returns:
            str: Path to the generated document
        """
        # Create a new document from scratch
        doc = Document()
        
        # Set document properties
        self._set_document_properties(doc)
        
        # Add contact information
        if "contact" in sections:
            self._add_contact_section(doc, sections["contact"])
        
        # Add summary section
        if "summary" in sections and sections["summary"]:
            self._add_section(doc, "SUMMARY", sections["summary"])
        
        # Add experience section
        if "experience" in sections and sections["experience"]:
            self._add_section(doc, "EXPERIENCE", sections["experience"])
        
        # Add education section
        if "education" in sections and sections["education"]:
            self._add_section(doc, "EDUCATION", sections["education"])
        
        # Add skills section
        if "skills" in sections and sections["skills"]:
            self._add_section(doc, "SKILLS", sections["skills"])
        
        # Add projects section
        if "projects" in sections and sections["projects"]:
            self._add_section(doc, "PROJECTS", sections["projects"])
        
        # Add additional section
        if "additional" in sections and sections["additional"]:
            self._add_section(doc, "ADDITIONAL INFORMATION", sections["additional"])
        
        # Save the document
        try:
            doc.save(output_path)
            logger.info(f"Successfully saved YC-style resume to {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            return None 