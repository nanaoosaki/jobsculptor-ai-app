#!/usr/bin/env python3
"""
Test script to verify design token loading in NumberingEngine
"""

import sys
from pathlib import Path

# Add the project root to Python path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

from word_styles.numbering_engine import NumberingEngine
from docx import Document

def test_design_token_loading():
    """Test that NumberingEngine loads and uses design tokens correctly."""
    print("🧪 Testing design token loading in NumberingEngine...")
    
    # Create a test document and engine
    doc = Document()
    engine = NumberingEngine()
    
    # Create a paragraph with content
    para = doc.add_paragraph()
    para.add_run("Test bullet point")
    
    # Apply bullet and capture the log output
    print("\n📋 Applying bullet with design token loading...")
    engine.apply_native_bullet(para, num_id=1, level=0)
    
    # Inspect the actual XML to see what values were applied
    pPr = para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
    ind = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ind') if pPr is not None else None
    
    if ind is not None:
        left_twips = ind.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left')
        hanging_twips = ind.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hanging')
        
        print(f"\n✅ ACTUAL VALUES APPLIED:")
        print(f"   Left indent: {left_twips} twips ({int(left_twips)/1440:.3f} inches)")
        print(f"   Hanging indent: {hanging_twips} twips ({int(hanging_twips)/1440:.3f} inches)")
        
        # Convert back to cm to verify
        left_cm = int(left_twips) / 567
        hanging_cm = int(hanging_twips) / 567
        print(f"   Left indent: {left_cm:.3f} cm")
        print(f"   Hanging indent: {hanging_cm:.3f} cm")
        
        # Check if it matches our design tokens
        expected_cm = 0.33
        if abs(left_cm - expected_cm) < 0.01:
            print(f"✅ SUCCESS: Design tokens are being used! (Expected: {expected_cm} cm)")
        else:
            print(f"❌ FAILURE: Design tokens not used. Expected: {expected_cm} cm, Got: {left_cm:.3f} cm")
    else:
        print("❌ FAILURE: No indentation XML found")

def test_direct_token_loading():
    """Test loading design tokens directly."""
    print("\n🧪 Testing direct design token loading...")
    
    try:
        from style_engine import StyleEngine
        tokens = StyleEngine.load_tokens()
        
        bullet_left = tokens.get("docx-bullet-left-indent-cm", "NOT_FOUND")
        bullet_hanging = tokens.get("docx-bullet-hanging-indent-cm", "NOT_FOUND")
        
        print(f"✅ Design tokens loaded:")
        print(f"   docx-bullet-left-indent-cm: {bullet_left}")
        print(f"   docx-bullet-hanging-indent-cm: {bullet_hanging}")
        
        return True
    except Exception as e:
        print(f"❌ Error loading design tokens: {e}")
        return False

if __name__ == "__main__":
    print("🚀 Design Token Integration Test")
    print("=" * 50)
    
    # Test direct token loading first
    tokens_ok = test_direct_token_loading()
    
    if tokens_ok:
        # Test actual usage in NumberingEngine
        test_design_token_loading()
    else:
        print("❌ Skipping NumberingEngine test due to token loading failure")
    
    print("\n" + "=" * 50)
    print("�� Test complete!") 