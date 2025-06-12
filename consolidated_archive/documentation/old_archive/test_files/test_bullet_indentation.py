"""
Test script to verify bullet indentation formatting with design tokens.

This script tests the new professional bullet formatting:
- Bullet aligned at: 0.13" (0.33 cm)
- Text indented at: 0.38" (0.965 cm)
"""

import os
import logging
from docx import Document
from docx.shared import Pt

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Import our components
from style_engine import StyleEngine
from word_styles.numbering_engine import NumberingEngine
from utils.docx_builder import add_bullet_point_native, _create_document_styles

def test_bullet_indentation():
    """Test the new bullet indentation formatting."""
    
    logger.info("🎯 Testing professional bullet indentation formatting...")
    
    # Create a new document
    doc = Document()
    
    # Load design tokens
    design_tokens = StyleEngine.load_tokens()
    logger.info(f"📋 Loaded design tokens")
    
    # Check the bullet indentation values
    left_indent_cm = design_tokens.get("docx-bullet-left-indent-cm")
    hanging_indent_cm = design_tokens.get("docx-bullet-hanging-indent-cm")
    
    logger.info(f"📏 Design token values:")
    logger.info(f"   Left indent: {left_indent_cm} cm = {float(left_indent_cm)*0.394:.3f} inches")
    logger.info(f"   Hanging indent: {hanging_indent_cm} cm = {float(hanging_indent_cm)*0.394:.3f} inches")
    
    # Verify we have the target values
    expected_left = "0.33"  # 0.13 inches - paragraph left indent
    expected_hanging = "0"  # 0 inches - no hanging, bullet and text align
    
    if left_indent_cm == expected_left:
        logger.info("✅ Left indent matches target (0.13 inches)")
    else:
        logger.error(f"❌ Left indent mismatch: got {left_indent_cm}, expected {expected_left}")
    
    if hanging_indent_cm == expected_hanging:
        logger.info("✅ Hanging indent matches target (0 inches - no hanging)")
    else:
        logger.error(f"❌ Hanging indent mismatch: got {hanging_indent_cm}, expected {expected_hanging}")
    
    # Create document styles
    docx_styles = {"design_tokens": design_tokens}
    try:
        _create_document_styles(doc, docx_styles)
        logger.info("✅ Created document styles")
    except Exception as e:
        logger.warning(f"Could not create document styles: {e}")
    
    # Create numbering engine
    numbering_engine = NumberingEngine()
    
    # Add test content
    header_para = doc.add_paragraph()
    header_para.add_run("Professional Bullet Formatting Test")
    header_para.style = 'Heading 1'
    
    doc.add_paragraph("This document tests the new professional bullet indentation:")
    
    # Test bullets with design token integration
    test_bullets = [
        "First bullet point with professional indentation",
        "Second bullet showing consistent alignment",
        "Third bullet demonstrating design token control",
        "Fourth bullet verifying Word native numbering system"
    ]
    
    for bullet_text in test_bullets:
        try:
            para = add_bullet_point_native(
                doc=doc,
                text=bullet_text,
                numbering_engine=numbering_engine,
                docx_styles=docx_styles
            )
            logger.info(f"✅ Created bullet: {bullet_text[:30]}...")
        except Exception as e:
            logger.error(f"❌ Failed to create bullet: {e}")
    
    # Add verification section
    doc.add_paragraph("\nExpected measurements in Word:")
    doc.add_paragraph("• Bullet aligned at: 0.13 inches")
    doc.add_paragraph("• Text aligned at: 0.13 inches (same as bullet)")
    doc.add_paragraph("• Paragraph left: 0.13 inches")
    doc.add_paragraph("• Hanging indent: 0 inches (no hanging)")
    
    doc.add_paragraph("\nTo verify:")
    doc.add_paragraph("1. Select a bullet paragraph")
    doc.add_paragraph("2. Right-click → Paragraph")
    doc.add_paragraph("3. Check 'Bullets and Numbering' settings")
    
    # Save the document
    output_path = "test_bullet_indentation_corrected.docx"
    doc.save(output_path)
    
    logger.info(f"✅ Test document saved: {output_path}")
    logger.info("🎯 Open the document in Word to verify bullet alignment!")
    
    return output_path

def test_conversion_functions():
    """Test the conversion utility functions."""
    
    logger.info("🔧 Testing conversion functions...")
    
    # Test NumberingEngine conversion methods
    engine = NumberingEngine()
    
    # Test cm to twips conversion
    test_cm = 0.965  # 0.38 inches
    twips = engine.cm_to_twips(test_cm)
    expected_twips = int(0.965 * 567)  # Should be ~547
    
    logger.info(f"📏 Conversion test: {test_cm} cm = {twips} twips (expected ~{expected_twips})")
    
    if abs(twips - expected_twips) <= 1:  # Allow 1 twip rounding difference
        logger.info("✅ cm_to_twips conversion correct")
    else:
        logger.error(f"❌ cm_to_twips conversion failed")
    
    # Test twips to inches conversion
    inches = engine.twips_to_inches(twips)
    expected_inches = 0.38
    
    logger.info(f"📏 Conversion test: {twips} twips = {inches:.3f} inches (expected ~{expected_inches})")
    
    if abs(inches - expected_inches) <= 0.01:  # Allow small floating point difference
        logger.info("✅ twips_to_inches conversion correct")
    else:
        logger.error(f"❌ twips_to_inches conversion failed")

if __name__ == "__main__":
    logger.info("🚀 Starting bullet indentation tests...")
    
    # Test conversion functions
    test_conversion_functions()
    
    # Test bullet formatting
    output_file = test_bullet_indentation()
    
    logger.info("🎉 Testing complete!")
    logger.info(f"📄 Open {output_file} in Microsoft Word to verify formatting")
    logger.info("📋 Expected: Bullets and text both aligned at 0.13\"") 