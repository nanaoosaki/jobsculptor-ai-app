import os
import sys
import json
import logging
import uuid
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.style import WD_STYLE_TYPE

# Add the current directory to the path so we can import our modules
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('section_header_box_test.log')
    ]
)
logger = logging.getLogger(__name__)

def test_section_header_box():
    """Test that DOCX section headers have the correct border properties"""
    try:
        # Import our modules
        from style_engine import StyleEngine, TokenAccessor
        
        # Load tokens
        tokens = StyleEngine.load_tokens()
        tokens_access = TokenAccessor(tokens)
        
        # Create a test document
        doc = Document()
        
        # Create the BoxedHeading2 style
        style_name = StyleEngine.create_boxed_heading_style(doc, tokens)
        logger.info(f"Created style: {style_name}")
        
        # Add a section header with the style
        para = doc.add_paragraph("PERSONAL SUMMARY")
        StyleEngine.apply_boxed_section_header_style(doc, para, tokens)
        
        # Add some regular content
        doc.add_paragraph("This is regular content text that follows the section header.")
        
        # Add another section header
        para = doc.add_paragraph("PROFESSIONAL EXPERIENCE")
        StyleEngine.apply_boxed_section_header_style(doc, para, tokens)
        
        # Add some regular content
        doc.add_paragraph("This is more regular content text.")
        
        # Generate a unique filename
        test_id = str(uuid.uuid4())[:8]
        output_file = f"test_section_header_box_{test_id}.docx"
        
        # Save the document
        doc.save(output_file)
        logger.info(f"Saved test document to {output_file}")
        
        # Validate the document
        validate_section_header_box(output_file, tokens_access)
        
    except Exception as e:
        logger.error(f"Test failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return False
    
    return True

def validate_section_header_box(docx_file, tokens_access):
    """Validate that the section headers in the document have proper box styling"""
    try:
        # Open the document with python-docx
        doc = Document(docx_file)
        
        # Find paragraphs with BoxedHeading2 style
        boxed_headers = [p for p in doc.paragraphs if p.style.name == 'BoxedHeading2']
        
        if not boxed_headers:
            logger.error("No BoxedHeading2 style paragraphs found in the document")
            return False
        
        logger.info(f"Found {len(boxed_headers)} boxed headers")
        
        # Check the first boxed header
        header = boxed_headers[0]
        
        # Access the paragraph's XML to check border properties
        p_pr = header._element.get_or_add_pPr()
        p_borders = p_pr.get_or_add_pBdr()
        
        # Check if all four borders are present
        borders = {
            'top': p_borders.top if hasattr(p_borders, 'top') else None,
            'left': p_borders.left if hasattr(p_borders, 'left') else None,
            'bottom': p_borders.bottom if hasattr(p_borders, 'bottom') else None,
            'right': p_borders.right if hasattr(p_borders, 'right') else None
        }
        
        # Log border properties for debugging
        for side, border in borders.items():
            if border is not None:
                logger.info(f"{side} border: val={border.val}, sz={border.sz}, color={border.color}")
            else:
                logger.error(f"No {side} border found")
                
        # Expected values from tokens
        border_width_pt = tokens_access.get("sectionHeader.border.widthPt", 1)
        border_color = tokens_access.get("sectionHeader.border.color", "#000000").lstrip('#')
        border_style = tokens_access.get("sectionHeader.border.style", "single")
        
        # Validate border properties
        for side, border in borders.items():
            if border is None:
                logger.error(f"Missing {side} border")
                return False
                
            # Check border style
            if border.val != border_style:
                logger.error(f"{side} border style incorrect. Expected: {border_style}, Got: {border.val}")
                return False
                
            # Check border width (in 8ths of a point)
            expected_width = str(int(border_width_pt * 8))
            if border.sz != expected_width:
                logger.error(f"{side} border width incorrect. Expected: {expected_width}, Got: {border.sz}")
                return False
                
            # Check border color
            if border.color.lower() != border_color.lower():
                logger.error(f"{side} border color incorrect. Expected: {border_color}, Got: {border.color}")
                return False
        
        # Check outline level for accessibility (should be 1 for Heading 2)
        outline_level = None
        if hasattr(p_pr, 'outlineLvl') and p_pr.outlineLvl is not None:
            outline_level = p_pr.outlineLvl.val
            
        if outline_level != '1':  # 0 = Heading 1, 1 = Heading 2
            logger.error(f"Outline level incorrect. Expected: 1, Got: {outline_level}")
            return False
            
        logger.info("Section header box validation passed!")
        return True
        
    except Exception as e:
        logger.error(f"Validation failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return False

if __name__ == "__main__":
    test_section_header_box() 