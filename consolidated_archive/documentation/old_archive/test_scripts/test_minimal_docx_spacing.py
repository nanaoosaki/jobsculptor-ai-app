#!/usr/bin/env python
"""
Minimal test script to diagnose section header box height and spacing issues in DOCX.
This creates a simple document with normal paragraphs and section headers to isolate styling issues.
"""

import os
import logging
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from style_engine import StyleEngine

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def create_minimal_test_document(output_filename="test_minimal_spacing.docx"):
    """
    Creates a minimal test document with section headers and paragraphs
    to isolate and diagnose spacing issues.
    """
    doc = Document()
    
    # Load tokens
    tokens = StyleEngine.load_tokens()
    
    # Create a function to apply section header box styling
    def apply_section_header_style(paragraph, line_rule="exact", line_value=240):
        """Apply section header styling with specified line rule and value"""
        # Make text bold
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(14)
        
        # Set paragraph formatting
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.0
        
        # Border properties
        border_width_pt = 1
        border_color = "000000"  # Black
        border_style = "single"
        padding_pt = 1
        
        # Apply direct XML for spacing with specified line rule and value
        spacing_xml = f'''
        <w:spacing {nsdecls("w")} 
            w:before="0" 
            w:after="80" 
            w:line="{line_value}" 
            w:lineRule="{line_rule}" 
            w:beforeAutospacing="0" 
            w:afterAutospacing="0"/>
        '''
        
        # Apply borders
        border_xml = f'''
        <w:pBdr {nsdecls("w")}>
            <w:top w:val="{border_style}" w:sz="{int(border_width_pt * 8)}" w:space="{int(padding_pt * 20)}" w:color="{border_color}"/>
            <w:left w:val="{border_style}" w:sz="{int(border_width_pt * 8)}" w:space="{int(padding_pt * 20)}" w:color="{border_color}"/>
            <w:bottom w:val="{border_style}" w:sz="{int(border_width_pt * 8)}" w:space="{int(padding_pt * 20)}" w:color="{border_color}"/>
            <w:right w:val="{border_style}" w:sz="{int(border_width_pt * 8)}" w:space="{int(padding_pt * 20)}" w:color="{border_color}"/>
        </w:pBdr>
        '''
        
        # Get paragraph properties
        p_pr = paragraph._element.get_or_add_pPr()
        
        # Remove any existing spacing to prevent conflicts
        for existing in p_pr.xpath('./w:spacing'):
            p_pr.remove(existing)
        
        # Add new spacing
        p_pr.append(parse_xml(spacing_xml))
        
        # Remove any existing borders to prevent conflicts
        for existing in p_pr.xpath('./w:pBdr'):
            p_pr.remove(existing)
        
        # Add new borders
        p_pr.append(parse_xml(border_xml))
        
        logger.info(f"Applied section header style with lineRule={line_rule}, line={line_value}")
    
    # Create a function to apply zero spacing to a paragraph
    def set_zero_spacing_after(paragraph):
        """Set space_after to 0 on a paragraph"""
        paragraph.paragraph_format.space_after = Pt(0)
        
        # Also apply through XML for maximum control
        p_pr = paragraph._element.get_or_add_pPr()
        spacing = p_pr.find(qn('w:spacing'))
        if spacing is not None:
            spacing.set(qn('w:after'), '0')
        else:
            spacing_xml = f'<w:spacing {nsdecls("w")} w:after="0"/>'
            p_pr.append(parse_xml(spacing_xml))
        
        logger.info(f"Set space_after=0 on paragraph: '{paragraph.text[:30]}...'")
    
    # Create test document structure
    
    # First paragraph with zero space after
    para1 = doc.add_paragraph("First normal paragraph with zero space after.")
    set_zero_spacing_after(para1)
    
    # First section header - using "auto" line rule with 276 value
    header1 = doc.add_paragraph("FIRST SECTION HEADER")
    apply_section_header_style(header1, line_rule="auto", line_value=276)
    
    # Content paragraph with normal spacing
    para2 = doc.add_paragraph("Content paragraph in first section with normal spacing.")
    
    # Last paragraph in section with zero space after
    para3 = doc.add_paragraph("Last paragraph in first section with zero space after.")
    set_zero_spacing_after(para3)
    
    # Second section header - using "exact" line rule with 280 value
    header2 = doc.add_paragraph("SECOND SECTION HEADER")
    apply_section_header_style(header2, line_rule="exact", line_value=280)
    
    # Content paragraph
    para4 = doc.add_paragraph("Content paragraph in second section.")
    
    # Third section header - using "exact" line rule with 240 value
    header3 = doc.add_paragraph("THIRD SECTION HEADER")
    apply_section_header_style(header3, line_rule="exact", line_value=240)
    
    # Content paragraph
    para5 = doc.add_paragraph("Content paragraph in third section.")
    
    # Fourth section header - using "exact" line rule with 320 value
    header4 = doc.add_paragraph("FOURTH SECTION HEADER")
    apply_section_header_style(header4, line_rule="exact", line_value=320)
    
    # Final paragraph
    para6 = doc.add_paragraph("Final paragraph in the document.")
    
    # Log all paragraph properties before saving
    logger.info("Paragraph properties before saving:")
    for idx, para in enumerate(doc.paragraphs):
        space_after = para.paragraph_format.space_after
        space_after_pt = space_after.pt if space_after else "None"
        logger.info(f"Para {idx}: '{para.text[:30]}...', space_after={space_after_pt} pt")
    
    # Save the document
    doc.save(output_filename)
    logger.info(f"Saved test document to {output_filename}")
    
    return output_filename

if __name__ == "__main__":
    filename = create_minimal_test_document()
    print(f"Created test document: {filename}")
    print("Please open this document in Word and:")
    print("1. Check the height of each section header box")
    print("2. Note which line rule and value creates the most compact box")
    print("3. Verify spacing between sections")
    print("4. Use Shift+F1 on each header to reveal formatting details") 