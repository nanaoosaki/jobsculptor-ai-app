"""
Universal Role Box Renderer

Single source of truth for role/position box styling across HTML, PDF, and DOCX formats.
Eliminates content duplication by providing canonical rendering.
"""

from typing import Any, Dict, Optional
import logging
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class RoleBox:
    """Universal role/position box rendering - prevents duplication"""
    
    REQUIRED_TOKEN_SCHEMA = {
        "typography": {
            "fontSize": {"roleBox": "required"},
            "fontColor": {"roleBox": {"hex": "required"}},
            "fontFamily": {"primary": "required"}
        },
        "roleBox": {
            "spacing": {"rightMarginPt": "required"},
            "dates": {
                "fontSize": "required",
                "fontColor": {"hex": "required"}
            },
            "borderColor": "required",
            "borderWidth": "required",
            "borderRadius": "required",
            "padding": "required"
        }
    }
    
    def __init__(self, tokens: Dict, role: str, dates: Optional[str] = None):
        """
        Initialize role box with design tokens
        
        Args:
            tokens: Design tokens dictionary
            role: Role/position title
            dates: Optional date range
        """
        self.tokens = tokens
        self.role = role
        self.dates = dates or ""
        self._validate_required_tokens()
        logger.debug(f"Created RoleBox for '{role}' with dates '{dates}'")
    
    def _validate_required_tokens(self):
        """Validate that all required design tokens are present"""
        missing_tokens = []
        
        def check_nested_path(tokens_dict, path_dict, current_path=""):
            for key, value in path_dict.items():
                full_path = f"{current_path}.{key}" if current_path else key
                
                if key not in tokens_dict:
                    missing_tokens.append(full_path)
                    continue
                
                if isinstance(value, dict):
                    if isinstance(tokens_dict[key], dict):
                        check_nested_path(tokens_dict[key], value, full_path)
                    else:
                        missing_tokens.append(full_path)
                elif value == "required" and not tokens_dict[key]:
                    missing_tokens.append(full_path)
        
        check_nested_path(self.tokens, self.REQUIRED_TOKEN_SCHEMA)
        
        if missing_tokens:
            raise ValueError(f"Required design tokens missing: {', '.join(missing_tokens)}")
    
    def to_html(self) -> str:
        """Single role box HTML - no duplication possible"""
        styles = self._build_css_styles()
        
        # Create single, canonical role box structure
        html_parts = [
            f'<div class="role-box" style="{styles["container"]}" role="presentation" aria-label="{self.role} {self.dates}">',
            f'<span class="role" style="{styles["role"]}">{self.role}</span>'
        ]
        
        if self.dates:
            html_parts.append(f'&nbsp;<span class="dates" style="{styles["dates"]}">{self.dates}</span>')
        
        html_parts.append('</div>')
        
        result = ''.join(html_parts)
        logger.info(f"Generated single HTML role box: '{self.role}' -> no duplication")
        return result
    
    def _build_css_styles(self) -> Dict[str, str]:
        """Build CSS styles from design tokens"""
        role_font_size = self.tokens["typography"]["fontSize"]["roleBox"]
        role_color = self.tokens["typography"]["fontColor"]["roleBox"]["hex"]
        font_family = self.tokens["typography"]["fontFamily"]["primary"]
        dates_font_size = self.tokens["roleBox"]["dates"]["fontSize"]
        dates_color = self.tokens["roleBox"]["dates"]["fontColor"]["hex"]
        
        # Get border styling tokens
        border_color = self.tokens["roleBox"]["borderColor"]
        border_width = self.tokens["roleBox"]["borderWidth"] 
        border_radius = self.tokens["roleBox"]["borderRadius"]
        padding = self.tokens["roleBox"]["padding"]
        
        return {
            "container": "; ".join([
                "display: flex",
                "justify-content: space-between",
                "align-items: baseline",
                "margin: 8pt 0 4pt 0",
                "line-height: 1.2",
                f"border: {border_width}pt solid {border_color}",
                f"border-radius: {border_radius}pt",
                f"padding: {padding}pt 8pt",
                "background-color: transparent"
            ]),
            "role": "; ".join([
                f"font-size: {role_font_size}",
                f"color: {role_color}",
                f"font-family: {font_family}",
                "font-weight: bold",
                "flex: 1"
            ]),
            "dates": "; ".join([
                f"font-size: {dates_font_size}",
                f"color: {dates_color}",
                f"font-family: {font_family}",
                "font-weight: normal",
                "text-align: right"
            ])
        }
    
    def to_docx(self, doc) -> Any:
        """Table-based role box for DOCX with token-driven styling"""
        # Create two-column table for role and dates
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        # Set column widths - role takes most space, dates are right-aligned
        total_width = Cm(17.0)
        table.columns[0].width = Cm(12.0)  # Role column
        table.columns[1].width = Cm(5.0)   # Dates column
        
        # Role cell
        role_cell = table.rows[0].cells[0]
        role_para = role_cell.paragraphs[0]
        role_para.text = self.role
        self._apply_docx_role_styling(role_para)
        
        # Dates cell (if dates provided)
        dates_cell = table.rows[0].cells[1]
        if self.dates:
            dates_para = dates_cell.paragraphs[0]
            dates_para.text = self.dates
            self._apply_docx_dates_styling(dates_para)
            dates_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Remove table borders for clean appearance
        self._remove_table_borders(table)
        
        logger.info(f"Generated DOCX role box: '{self.role}' with dates '{self.dates}'")
        return table
    
    def _apply_docx_role_styling(self, paragraph):
        """Apply role text styling from design tokens"""
        role_font_size = self.tokens["typography"]["fontSize"]["roleBox"]
        role_color = self.tokens["typography"]["fontColor"]["roleBox"]["hex"]
        font_family = self.tokens["typography"]["fontFamily"]["primary"]
        
        # Extract numeric font size
        if isinstance(role_font_size, str) and role_font_size.endswith('pt'):
            size_value = float(role_font_size[:-2])
        else:
            size_value = float(role_font_size)
        
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = Pt(size_value)
        run.font.name = font_family
        run.font.bold = True
        
        # Apply color
        if role_color.startswith('#'):
            from docx.shared import RGBColor
            hex_color = role_color[1:]
            rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
            run.font.color.rgb = RGBColor(*rgb)
    
    def _apply_docx_dates_styling(self, paragraph):
        """Apply dates text styling from design tokens"""
        dates_font_size = self.tokens["roleBox"]["dates"]["fontSize"]
        dates_color = self.tokens["roleBox"]["dates"]["fontColor"]["hex"]
        font_family = self.tokens["typography"]["fontFamily"]["primary"]
        
        # Extract numeric font size
        if isinstance(dates_font_size, str) and dates_font_size.endswith('pt'):
            size_value = float(dates_font_size[:-2])
        else:
            size_value = float(dates_font_size)
        
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = Pt(size_value)
        run.font.name = font_family
        run.font.bold = False
        
        # Apply color
        if dates_color.startswith('#'):
            from docx.shared import RGBColor
            hex_color = dates_color[1:]
            rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
            run.font.color.rgb = RGBColor(*rgb)
    
    def _remove_table_borders(self, table):
        """Remove borders from table for clean role box appearance"""
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn
        
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                
                for side in ['top', 'bottom', 'left', 'right']:
                    border = OxmlElement(f'w:{side}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                
                tcPr.append(tcBorders)
    
    def get_debug_info(self) -> Dict:
        """Get debug information about this role box"""
        return {
            "role": self.role,
            "dates": self.dates,
            "role_font_size": self.tokens["typography"]["fontSize"]["roleBox"],
            "role_color": self.tokens["typography"]["fontColor"]["roleBox"]["hex"],
            "dates_font_size": self.tokens["roleBox"]["dates"]["fontSize"],
            "dates_color": self.tokens["roleBox"]["dates"]["fontColor"]["hex"]
        } 