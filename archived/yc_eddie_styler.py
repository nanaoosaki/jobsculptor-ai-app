import os
import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

class YCEddieStyler:
    """
    Creates a resume with the YC/Eddie style formatting:
    - Clean, minimalist design
    - No templates or fancy formatting
    - Consistent section headers
    - Bold company names, regular position titles
    - Clear date ranges
    - Bullet points for accomplishments
    """
    
    def __init__(self):
        """Initialize the YC Eddie Styler"""
        self.doc = None
        
    def create_document(self):
        """Create a new document with proper styles"""
        self.doc = Document()
        
        # Remove default styles and set up document
        self._setup_document()
        
        return self.doc
    
    def _setup_document(self):
        """Set up document styles and formatting"""
        # Set narrow margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Create styles
        self._create_styles()
    
    def _create_styles(self):
        """Create the document styles"""
        # Create name style
        name_style = self.doc.styles.add_style('Name', WD_STYLE_TYPE.PARAGRAPH)
        name_style.font.name = 'Calibri'
        name_style.font.size = Pt(16)
        name_style.font.bold = True
        
        # Create contact style
        contact_style = self.doc.styles.add_style('Contact', WD_STYLE_TYPE.PARAGRAPH)
        contact_style.font.name = 'Calibri'
        contact_style.font.size = Pt(10)
        
        # Create section heading style
        heading_style = self.doc.styles.add_style('Section', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(12)
        heading_style.font.bold = True
        heading_style.font.all_caps = True
        
        # Create company name style
        company_style = self.doc.styles.add_style('Company', WD_STYLE_TYPE.PARAGRAPH)
        company_style.font.name = 'Calibri'
        company_style.font.size = Pt(11)
        company_style.font.bold = True
        
        # Create job title style
        title_style = self.doc.styles.add_style('JobTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(11)
        title_style.font.italic = True
        
        # Create bullet style
        bullet_style = self.doc.styles.add_style('Bullet', WD_STYLE_TYPE.PARAGRAPH)
        bullet_style.font.name = 'Calibri'
        bullet_style.font.size = Pt(10)
        
        # Create regular paragraph style
        para_style = self.doc.styles.add_style('Regular', WD_STYLE_TYPE.PARAGRAPH)
        para_style.font.name = 'Calibri'
        para_style.font.size = Pt(10)
    
    def add_name(self, name):
        """Add the candidate's name at the top"""
        p = self.doc.add_paragraph(name, style='Name')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p
    
    def add_contact_info(self, contact_text):
        """Add contact information"""
        lines = contact_text.split('\n')
        p = self.doc.add_paragraph(style='Contact')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run(' | ')
            p.add_run(line.strip())
        
        return p
    
    def add_section_heading(self, heading_text):
        """Add a section heading"""
        p = self.doc.add_paragraph(heading_text, style='Section')
        return p
    
    def add_company(self, company_text):
        """Add a company name"""
        p = self.doc.add_paragraph(company_text, style='Company')
        return p
    
    def add_job_title(self, title_text):
        """Add a job title and date"""
        p = self.doc.add_paragraph(title_text, style='JobTitle')
        return p
    
    def add_bullet_point(self, text):
        """Add a bullet point"""
        p = self.doc.add_paragraph(style='Bullet')
        
        # Use actual bullet character
        p.style.paragraph_format.left_indent = Inches(0.25)
        p.style.paragraph_format.first_line_indent = Inches(-0.25)
        
        run = p.add_run('• ')
        run.font.name = 'Symbol'
        
        p.add_run(text)
        return p
    
    def add_paragraph(self, text):
        """Add a regular paragraph"""
        p = self.doc.add_paragraph(text, style='Regular')
        return p
    
    def add_skills(self, skills_text):
        """Add skills with special formatting"""
        lines = skills_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                self.add_bullet_point(line[1:].strip())
            else:
                if ':' in line:
                    category, skills = line.split(':', 1)
                    p = self.doc.add_paragraph(style='Regular')
                    
                    # Bold category name
                    category_run = p.add_run(category.strip() + ': ')
                    category_run.bold = True
                    
                    # Regular skills
                    p.add_run(skills.strip())
                else:
                    self.add_paragraph(line)
    
    def add_education(self, education_text):
        """Add education with special formatting"""
        lines = education_text.split('\n')
        current_school = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # If line is bold in the original, treat as school name
            if line.isupper() or 'UNIVERSITY' in line.upper() or 'COLLEGE' in line.upper() or 'SCHOOL' in line.upper():
                p = self.doc.add_paragraph(style='Company')
                p.add_run(line)
                current_school = True
            elif '|' in line and current_school:
                # This is likely a "Degree | Date" line
                p = self.doc.add_paragraph(style='JobTitle')
                p.add_run(line)
                current_school = False
            else:
                self.add_paragraph(line)
    
    def save(self, output_path):
        """Save the document"""
        self.doc.save(output_path)
        return output_path 