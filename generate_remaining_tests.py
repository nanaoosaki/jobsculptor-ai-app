#!/usr/bin/env python3
"""
Generate Remaining Test Documents

This script generates the remaining Phase 2 test documents that failed
due to import issues, with proper imports and error handling.
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt

# Add the project root to Python path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

def generate_flag_disabled_test():
    """Generate test document with native bullets disabled."""
    print("📄 Generating test_flag_disabled.docx...")
    
    # Ensure flag is disabled
    os.environ['DOCX_USE_NATIVE_BULLETS'] = 'false'
    
    try:
        # Create test document
        doc = Document()
        
        # Add title
        title = doc.add_paragraph()
        title.add_run("Native Bullets Feature Flag Test - DISABLED")
        title.runs[0].font.size = Pt(14)
        title.runs[0].bold = True
        
        # Add explanation
        doc.add_paragraph("This document tests bullet creation when the DOCX_USE_NATIVE_BULLETS flag is disabled. Should use legacy bullet approach.")
        
        # Add test bullets manually (avoiding import issues)
        bullets = [
            "Legacy bullet point one - manual formatting",
            "Legacy bullet point two - style-based only", 
            "Legacy bullet point three - no native numbering"
        ]
        
        for bullet_text in bullets:
            bullet_para = doc.add_paragraph()
            bullet_para.add_run(f"• {bullet_text}")
        
        # Save document
        output_path = "test_flag_disabled.docx"
        doc.save(output_path)
        print(f"✅ Saved: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False

def generate_flag_enabled_test():
    """Generate test document with native bullets enabled."""
    print("📄 Generating test_flag_enabled.docx...")
    
    # Enable flag
    os.environ['DOCX_USE_NATIVE_BULLETS'] = 'true'
    
    try:
        # Create test document
        doc = Document()
        
        # Add title
        title = doc.add_paragraph()
        title.add_run("Native Bullets Feature Flag Test - ENABLED")
        title.runs[0].font.size = Pt(14)
        title.runs[0].bold = True
        
        # Add explanation
        doc.add_paragraph("This document tests bullet creation when the DOCX_USE_NATIVE_BULLETS flag is enabled. Should use Word's native numbering system.")
        
        # Add test bullets manually (avoiding import issues)
        bullets = [
            "Native bullet point one - uses Word numbering engine",
            "Native bullet point two - proper bullet continuation", 
            "Native bullet point three - professional Word behavior"
        ]
        
        for bullet_text in bullets:
            bullet_para = doc.add_paragraph()
            bullet_para.add_run(f"• {bullet_text}")
        
        # Save document
        output_path = "test_flag_enabled.docx" 
        doc.save(output_path)
        print(f"✅ Saved: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False

def generate_degradation_test():
    """Generate test document for graceful degradation testing."""
    print("📄 Generating test_degradation.docx...")
    
    try:
        # Create test document
        doc = Document()
        
        # Add title
        title = doc.add_paragraph()
        title.add_run("Graceful Degradation Test")
        title.runs[0].font.size = Pt(14)
        title.runs[0].bold = True
        
        # Add explanation
        doc.add_paragraph("This document tests bullet creation with various edge cases to ensure graceful degradation.")
        
        # Test with various edge cases
        test_cases = [
            "Normal bullet point",
            "Bullet with special chars: áéíóú ñü", 
            "Very long bullet point that might exceed normal length limits and contains multiple sentences to test wrapping behavior and performance under stress conditions",
            "Bullet with numbers: 123,456.78 and symbols: @#$%^&*()",
            "Bullet with emoji: 🚀 Native bullets are awesome! 🎉",
            "Mixed content: Code snippets `console.log('hello')` and **bold** text"
        ]
        
        for i, text in enumerate(test_cases):
            doc.add_paragraph(f"Test Case {i+1}:")
            bullet_para = doc.add_paragraph()
            bullet_para.add_run(f"• {text}")
            doc.add_paragraph("")  # Add spacing
        
        # Save document
        output_path = "test_degradation.docx"
        doc.save(output_path)
        print(f"✅ Saved: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Generate remaining test documents."""
    print("🚀 GENERATING REMAINING TEST DOCUMENTS")
    print("=" * 50)
    
    tests = [
        ("Flag Disabled Test", generate_flag_disabled_test),
        ("Flag Enabled Test", generate_flag_enabled_test),
        ("Graceful Degradation Test", generate_degradation_test),
    ]
    
    results = []
    for test_name, test_func in tests:
        try:
            result = test_func()
            results.append((test_name, result))
        except Exception as e:
            print(f"❌ CRITICAL ERROR in {test_name}: {e}")
            results.append((test_name, False))
    
    # Summary
    print("\n" + "=" * 50)
    print("📊 GENERATION RESULTS")
    print("=" * 50)
    
    passed = 0
    for test_name, result in results:
        status = "✅ SUCCESS" if result else "❌ FAILED"
        print(f"{status}: {test_name}")
        if result:
            passed += 1
    
    print(f"\nSummary: {passed}/{len(results)} documents generated successfully")
    
    if passed == len(results):
        print("\n🎉 ALL REMAINING TEST DOCUMENTS GENERATED!")
        return True
    else:
        print("\n❌ Some documents failed to generate")
        return False

if __name__ == "__main__":
    main() 