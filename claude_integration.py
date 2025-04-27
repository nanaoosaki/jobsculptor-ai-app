from yc_resume_generator import YCResumeGenerator
import os
import json
import logging
import time
import traceback
import re
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx2txt
from typing import Dict, List, Tuple, Optional, Any, Union
from datetime import datetime
from flask import current_app
import uuid  # Import UUID module

# Third-party imports for Claude
from anthropic import Anthropic, RateLimitError

# Third-party imports for OpenAI
import openai
from openai import OpenAI
from claude_api_logger import api_logger

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Global variable to track the last LLM client used
last_llm_client = None

# Import the new html_generator module
import html_generator

# Shared bullet utility
from utils.bullet_utils import strip_bullet_prefix, BULLET_ESCAPE_RE


def clean_bullet_points(text: str) -> str:
    """
    Removes bullet point markers from text to prevent duplicate bullet points.
    
    Handles various bullet point formats:
    - Unicode bullets (•, ◦, ▪, ▫, etc.)
    - ASCII bullets (*, -, +, o)
    - Numbered bullets (1., 1), (1), etc.)
    
    Args:
        text (str): Text containing bullet points to clean
        
    Returns:
        str: Cleaned text with bullet point markers removed
    """
    if not text:
        return ""
    
    logger.info(f"Cleaning bullet points from text ({len(text)} chars)")
    
    # Simple wrapper using shared util – ensures single source of truth
    lines = text.split('\n')
    cleaned_lines = [strip_bullet_prefix(l) for l in lines]

    cleaned_text = '\n'.join(cleaned_lines)

    # Optional: quick validation log
    stray = sum(1 for l in cleaned_lines if BULLET_ESCAPE_RE.match(l))
    if stray:
        logger.warning(f"clean_bullet_points(): {stray} bullet markers remained after cleaning")
    return cleaned_text


class LLMClient:
    """Base class for LLM API clients"""
    
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.tailored_content = {}  # Store tailored responses (cleaned Python structures)
        self.request_id = str(uuid.uuid4()) # Generate unique ID for this instance/request
        self.temp_data_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'temp_session_data')
        os.makedirs(self.temp_data_dir, exist_ok=True) # Ensure temp dir exists
        
    def tailor_resume_content(
    self,
    section_name: str,
    content: str,
     job_data: Dict) -> Optional[Union[Dict, List, str]]: # Return type can vary
        """Tailor resume content using LLM API - to be implemented by subclasses"""
        raise NotImplementedError("Subclasses must implement this method")

    def _save_cleaned_section_to_temp_file(self, section_name: str, data: Union[Dict, List, str]):
        """Saves the cleaned Python data structure to a temporary session file."""
        if data is None:
            logger.warning(f"Attempted to save None data for section {section_name} with request ID {self.request_id}. Skipping.")
            return

        try:
            # Handle simple text sections (summary, contact) - wrap in dict for consistency
            if section_name in ["summary", "contact"] and isinstance(data, str):
                 save_data = {"content": data}
                 logger.debug(f"Wrapping string content for {section_name} in dict for saving.")
            else:
                 save_data = data # Assumes data is already list/dict for other sections

            temp_filename = f"{self.request_id}_{section_name}.json"
            temp_filepath = os.path.join(self.temp_data_dir, temp_filename)

            with open(temp_filepath, 'w', encoding='utf-8') as f:
                json.dump(save_data, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Stored cleaned, structured data for request {self.request_id} for section {section_name}. Method: temp_file, Location: {temp_filepath}")

        except TypeError as te:
            logger.error(f"TypeError saving {section_name} (Req ID: {self.request_id}) to JSON: {te}. Content type: {type(data)}")
            logger.error(f"Problematic content snippet: {str(data)[:200]}")
        except Exception as e:
            logger.error(f"Error saving {section_name} (Req ID: {self.request_id}) to temp file: {e}")
            logger.error(traceback.format_exc())

    def get_request_id(self) -> str:
         """Returns the unique request ID for this tailoring operation."""
         return self.request_id


class ClaudeClient(LLMClient):
    """Client for interacting with Claude API"""
    
    def __init__(self, api_key: str, api_url: str = None):
        """Initialize the Claude API client"""
        super().__init__(api_key)
        self.api_key = api_key
        self.api_url = api_url
        self.client = None
        self.tailored_content = {}

        try:
            # Validate API key format for Claude
            if not api_key:
                raise ValueError("Claude API key is missing")

            # Import anthropic library (lazy import to handle missing
            # dependency)
            try:
                from anthropic import Anthropic
                logger.info("Using anthropic SDK for Claude API")
            except ImportError:
                logger.error(
                    "Failed to import anthropic SDK. Make sure it's installed with 'pip install anthropic'")
                raise

            # Initialize Claude client
            self.client = Anthropic(api_key=api_key)
            logger.info(f"Claude API client initialized successfully (Req ID: {self.request_id})")

        except Exception as e:
            logger.error(f"Error initializing Claude API client: {str(e)}")
            logger.error(traceback.format_exc())
            raise ValueError(f"Failed to initialize Claude client: {str(e)}")
    
    def tailor_resume_content(
    self,
    section_name: str,
    content: str,
     job_data: Dict) -> Optional[Union[Dict, List, str]]: # Return type can vary
        """
        Tailor resume content using Claude API.
        Parses response, cleans structured data, stores it in self.tailored_content,
        saves cleaned data to temp file, and returns cleaned data.
        """
        logger.info(f"Tailoring {section_name} with Claude API (Req ID: {self.request_id})")
            
        if not content or not content.strip():
            logger.warning(f"Empty {section_name} content provided, skipping tailoring")
            return None

        if not self.client:
            logger.error("Claude client not initialized")
            return content

        try:
            # Extract job data
            job_title = job_data.get('job_title', 'the position')
            company = job_data.get('company', 'the company')
            requirements = job_data.get('requirements', [])
            skills = job_data.get('skills', [])

            # Prepare requirements and skills text
            requirements_text = "\n".join(
                [f"- {req}" for req in requirements]) if requirements else "Not specified"
            skills_text = ", ".join(skills) if skills else "Not specified"

            # Get job analysis if available
            analysis_prompt = ""
            if 'analysis' in job_data and isinstance(
                job_data['analysis'], dict):
                analysis = job_data['analysis']
                
                # Add candidate profile if available
                if 'candidate_profile' in analysis and analysis['candidate_profile']:
                    analysis_prompt += f"\n\nCANDIDATE PROFILE:\n{
    analysis['candidate_profile']}"
                
                # Add hard skills if available
                if 'hard_skills' in analysis and analysis['hard_skills']:
                    hard_skills = ", ".join(analysis['hard_skills'])
                    analysis_prompt += f"\n\nKEY HARD SKILLS:\n{hard_skills}"
                    
                # Add soft skills if available
                if 'soft_skills' in analysis and analysis['soft_skills']:
                    soft_skills = ", ".join(analysis['soft_skills'])
                    analysis_prompt += f"\n\nKEY SOFT SKILLS:\n{soft_skills}"
                
                # Add ideal candidate if available
                if 'ideal_candidate' in analysis and analysis['ideal_candidate']:
                    analysis_prompt += f"\n\nIDEAL CANDIDATE:\n{
    analysis['ideal_candidate']}"

            # Build section-specific prompts
            if section_name == "experience":
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the work experience section to better match the requirements for a {job_title} position at {company}.

ORIGINAL EXPERIENCE SECTION:
{content}

JOB REQUIREMENTS:
            {requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "experience": [
    {{
      "company": "Company Name",
      "location": "City, State",
      "position": "Job Title",
      "dates": "Time Period",
      "achievements": [
        "Achievement 1",
        "Achievement 2"
      ]
    }}
  ]
}}
```

Please restructure the experience section to better match the job requirements by:
1. Highlighting relevant experience that demonstrates skills required for this job
2. Focusing on achievements and results rather than just responsibilities
3. Quantifying accomplishments with metrics where possible
4. Using terminology from the job description where appropriate
5. Maintaining the original company names, job titles, and dates

Do not add any fictional experiences or embellish beyond what is reasonable based on the original content.
"""

            elif section_name == "education":
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the education section to better match the requirements for a {job_title} position at {company}.

ORIGINAL EDUCATION SECTION:
            {content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "education": [
    {{
      "institution": "University Name",
      "location": "City, State",
      "degree": "Degree Name",
      "dates": "Time Period",
      "highlights": [
        "Highlight 1",
        "Highlight 2"
      ]
    }}
  ]
}}
```

Please rewrite the education section to better match the job requirements. Focus on:
1. Highlighting relevant coursework, projects, or achievements that match the job requirements
2. Emphasizing academic accomplishments that demonstrate skills needed for this position
3. Formatting in a way that emphasizes the most relevant educational experiences
4. Including any certifications or training that matches required skills

Keep the degree names, institutions, and dates exactly the same - only enhance descriptions to make them more relevant.
"""

            elif section_name == "skills":
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the skills section to better match the requirements for a {job_title} position at {company}.

ORIGINAL SKILLS SECTION:
{content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "skills": {{
    "technical": ["Skill 1", "Skill 2"],
    "soft": ["Skill 1", "Skill 2"],
    "other": ["Skill 1", "Skill 2"]
  }}
}}
```

Please rewrite the skills section to better match the job requirements. Focus on:
1. Reordering skills to prioritize those mentioned in the job description
2. Adding any missing skills that the candidate likely has based on their experience (must be reasonable to infer from other sections)
3. Grouping skills into relevant categories that align with the job posting
4. Rephrasing skills using the exact terminology from the job description
5. Removing skills that are irrelevant to this position if the list is very long

Only include skills that are authentic to the candidate based on their resume.
"""

            elif section_name == "projects":
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the projects section to better match the requirements for a {job_title} position at {company}.

ORIGINAL PROJECTS SECTION:
{content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "projects": [
    {{
      "title": "Project Name",
      "dates": "Time Period",
      "details": [
        "Detail 1",
        "Detail 2"
      ]
    }}
  ]
}}
```

Please rewrite the projects section to better match the job requirements. Focus on:
1. Highlighting projects that demonstrate skills required for this job
2. Emphasizing tools, technologies, and methodologies that match the job description
3. Quantifying project outcomes and impacts where possible
4. Rephrasing project descriptions to use terminology from the job posting
5. Focusing on the candidate's specific contributions and leadership roles

Keep the project titles and timelines the same, but enhance descriptions to better align with the job requirements.
"""

            else:
                # For other sections, use a generic prompt
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the {section_name} section to better match the requirements for a {job_title} position at {company}.

ORIGINAL SECTION:
{content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "{section_name}": "The tailored content goes here as a single string"
}}
```

Please rewrite this section to better match the job requirements while maintaining the same basic structure and information.
Focus on emphasizing elements most relevant to this job opportunity.
"""

            # Make the API call
            response = self.client.messages.create(
                model="claude-3-sonnet-20240229",
                max_tokens=4000,
                temperature=0.7,
                messages=[
                    {"role": "system", "content": "You are an expert resume tailor. Return only valid JSON responses in the specified format."},
                    {"role": "user", "content": prompt}
                ]
            )

            # Get the response content
            response_content = response.content[0].text.strip()
            logger.info(
    f"Claude API response for {section_name}: {
        len(response_content)} chars")

            # Parse JSON response
            cleaned_data = None # Initialize variable to store cleaned data
            try:
                # First try to extract JSON if not properly formatted
                json_str = response_content
                if not json_str.startswith('{'):
                    # Try to extract JSON from text
                    json_start = response_content.find('{')
                    json_end = response_content.rfind('}') + 1
                    if json_start >= 0 and json_end > json_start:
                        json_str = response_content[json_start:json_end]

                json_response = json.loads(json_str)
                logger.debug(f"Parsed LLM response for {section_name} (Req ID: {self.request_id}) before cleaning: {str(json_response)[:200]}...")

                # Process JSON based on section type
                if section_name == "experience" and "experience" in json_response:
                    # Clean the structured data directly after parsing
                    experience_data = json_response["experience"]
                    cleaned_experience_data = []
                    if isinstance(experience_data, list):
                        for job in experience_data:
                            if isinstance(job, dict):
                                # Clean achievements list within the job dictionary
                                achievements = job.get('achievements', [])
                                if isinstance(achievements, list):
                                    job['achievements'] = [strip_bullet_prefix(a) for a in achievements if isinstance(a, str) and a.strip()]
                                cleaned_experience_data.append(job)
                            else:
                                logger.warning(f"Skipping non-dictionary item in experience data: {job}")
                        cleaned_data = cleaned_experience_data # Store cleaned data
                    else:
                        logger.warning(f"Experience data is not a list: {experience_data}")
                        cleaned_data = response_content # Use raw as fallback
                elif section_name == "education" and "education" in json_response:
                    # Clean education highlights
                    education_data = json_response["education"]
                    cleaned_education_data = []
                    if isinstance(education_data, list):
                        for edu in education_data:
                             if isinstance(edu, dict):
                                highlights = edu.get('highlights', [])
                                if isinstance(highlights, list):
                                    edu['highlights'] = [strip_bullet_prefix(h) for h in highlights if isinstance(h, str) and h.strip()]
                                cleaned_education_data.append(edu)
                             else:
                                 logger.warning(f"Skipping non-dictionary item in education data: {edu}")
                    cleaned_data = cleaned_education_data
                elif section_name == "skills" and "skills" in json_response:
                     # Clean skills lists
                    skills_data = json_response["skills"]
                    if isinstance(skills_data, dict):
                        for key in ['technical', 'soft', 'other']:
                            if key in skills_data and isinstance(skills_data[key], list):
                                skills_data[key] = [strip_bullet_prefix(s) for s in skills_data[key] if isinstance(s, str) and s.strip()]
                        cleaned_data = skills_data
                    else:
                        logger.warning(f"Skills data is not a dictionary: {skills_data}")
                        cleaned_data = response_content
                elif section_name == "projects" and "projects" in json_response:
                     # Clean project details
                    projects_data = json_response["projects"]
                    cleaned_projects_data = []
                    if isinstance(projects_data, list):
                        for proj in projects_data:
                            if isinstance(proj, dict):
                                details = proj.get('details', [])
                                if isinstance(details, list):
                                    proj['details'] = [strip_bullet_prefix(d) for d in details if isinstance(d, str) and d.strip()]
                                cleaned_projects_data.append(proj)
                            else:
                               logger.warning(f"Skipping non-dictionary item in projects data: {proj}")
                    cleaned_data = cleaned_projects_data
                elif section_name in json_response:
                    # For other sections, just return the string content
                    formatted_text = json_response[section_name]
                    cleaned_data = clean_bullet_points(formatted_text) # Apply general cleaning
                else:
                    logger.warning(f"JSON response missing expected '{section_name}' key")
                    cleaned_data = content # Fallback to original content

                # Log after cleaning
                logger.debug(f"Parsed LLM response for {section_name} (Req ID: {self.request_id}) after cleaning: {str(cleaned_data)[:200]}...")

            except json.JSONDecodeError:
                logger.error(f"Failed to parse JSON from Claude response: {response_content[:100]}...")
                cleaned_data = response_content # Store raw response as the data to save

            # Store the final cleaned data (structure or raw fallback) in the instance
            self.tailored_content[section_name] = cleaned_data
            
            # Save the cleaned data to a temporary file
            self._save_cleaned_section_to_temp_file(section_name, cleaned_data)

            # Return the cleaned data structure (or raw text fallback)
            return cleaned_data
            
        except Exception as e:
            logger.error(f"Error in Claude API call: {str(e)}")
            cleaned_data = content # Fallback to original content
            self.tailored_content[section_name] = cleaned_data
            self._save_cleaned_section_to_temp_file(section_name, cleaned_data)
            return cleaned_data


class OpenAIClient(LLMClient):
    """OpenAI API client for resume tailoring"""
    
    def __init__(self, api_key: str):
        """Initialize the OpenAI API client"""
        super().__init__(api_key)
        self.api_key = api_key
        self.client = None
        self.raw_responses = {}  # Store raw JSON responses from API
        self.initialize_client()

    def initialize_client(self):
        """Initialize the OpenAI client"""
        try:
            print(
                f"Initializing OpenAI client with API key starting with: {self.api_key[:8]}...")
            print(f"API key length: {len(self.api_key)} characters")
            print("Testing OpenAI API connection...")

            from openai import OpenAI
            self.client = OpenAI(api_key=self.api_key)

            # Test connection with a simple request - removed limit parameter
            self.client.models.list()
            print("OpenAI client initialized successfully")

        except Exception as e:
            print(f"Error initializing OpenAI client: {str(e)}")
            self.client = None
            raise

    def tailor_resume_content(
    self,
    section_name: str,
    content: str,
     job_data: Dict) -> Optional[Union[Dict, List, str]]: # Return type can vary
        """
        Tailor resume content using OpenAI API.
        Parses response, cleans structured data, stores it in self.tailored_content,
        saves cleaned data to temp file, and returns cleaned data.
        """
        logger.info(f"Tailoring {section_name} with OpenAI API (Req ID: {self.request_id})")
            
        if not content or not content.strip():
            logger.warning(f"Empty {section_name} content provided, skipping tailoring")
            return None

        if not self.client:
            logger.error("OpenAI client not initialized")
            return content

        try:
            # Extract job data
            job_title = job_data.get('job_title', 'the position')
            company = job_data.get('company', 'the company')
            requirements = job_data.get('requirements', [])
            skills = job_data.get('skills', [])

            # Prepare requirements and skills text
            requirements_text = "\n".join(
                [f"- {req}" for req in requirements]) if requirements else "Not specified"
            skills_text = ", ".join(skills) if skills else "Not specified"

            # Get job analysis if available
            analysis_prompt = ""
            if 'analysis' in job_data and isinstance(
                job_data['analysis'], dict):
                analysis = job_data['analysis']
                
                # Add candidate profile if available
                if 'candidate_profile' in analysis and analysis['candidate_profile']:
                    analysis_prompt += f"\n\nCANDIDATE PROFILE:\n{
    analysis['candidate_profile']}"
                
                # Add hard skills if available
                if 'hard_skills' in analysis and analysis['hard_skills']:
                    hard_skills = ", ".join(analysis['hard_skills'])
                    analysis_prompt += f"\n\nKEY HARD SKILLS:\n{hard_skills}"
                    
                # Add soft skills if available
                if 'soft_skills' in analysis and analysis['soft_skills']:
                    soft_skills = ", ".join(analysis['soft_skills'])
                    analysis_prompt += f"\n\nKEY SOFT SKILLS:\n{soft_skills}"
                
                # Add ideal candidate if available
                if 'ideal_candidate' in analysis and analysis['ideal_candidate']:
                    analysis_prompt += f"\n\nIDEAL CANDIDATE:\n{
    analysis['ideal_candidate']}"

            # Build section-specific prompts
            if section_name == "experience":
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the work experience section to better match the requirements for a {job_title} position at {company}.

ORIGINAL EXPERIENCE SECTION:
{content}

JOB REQUIREMENTS:
            {requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "experience": [
    {{
      "company": "Company Name",
      "location": "City, State",
      "position": "Job Title",
      "dates": "Time Period",
      "achievements": [
        "Achievement 1",
        "Achievement 2"
      ]
    }}
  ]
}}
```

Please restructure the experience section to better match the job requirements by:
1. Highlighting relevant experience that demonstrates skills required for this job
2. Focusing on achievements and results rather than just responsibilities
3. Quantifying accomplishments with metrics where possible
4. Using terminology from the job description where appropriate
5. Maintaining the original company names, job titles, and dates

IMPORTANT:
1. Do not include empty strings or whitespace-only strings in any arrays
2. Every achievement must contain meaningful content
3. Do not split single achievements into multiple entries
4. Make sure each bullet point contains complete, meaningful content

Do not add any fictional experiences or embellish beyond what is reasonable based on the original content.
"""

            elif section_name == "education":
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the education section to better match the requirements for a {job_title} position at {company}.

ORIGINAL EDUCATION SECTION:
            {content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "education": [
    {{
      "institution": "University Name",
      "location": "City, State",
      "degree": "Degree Name",
      "dates": "Time Period",
      "highlights": [
        "Highlight 1",
        "Highlight 2"
      ]
    }}
  ]
}}
```

Please rewrite the education section to better match the job requirements. Focus on:
1. Highlighting relevant coursework, projects, or achievements that match the job requirements
2. Emphasizing academic accomplishments that demonstrate skills needed for this position
3. Formatting in a way that emphasizes the most relevant educational experiences
4. Including any certifications or training that matches required skills

IMPORTANT:
1. Do not include empty strings or whitespace-only strings in any arrays
2. Every highlight must contain meaningful content
3. Do not split single highlights into multiple entries
4. Make sure each bullet point contains complete, meaningful content
5. Ensure all education entries from the original resume are preserved

Keep the degree names, institutions, and dates exactly the same - only enhance descriptions to make them more relevant.
"""

            elif section_name == "skills":
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the skills section to better match the requirements for a {job_title} position at {company}.

ORIGINAL SKILLS SECTION:
{content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "skills": {{
    "technical": ["Skill 1", "Skill 2"],
    "soft": ["Skill 1", "Skill 2"],
    "other": ["Skill 1", "Skill 2"]
  }}
}}
```

Please rewrite the skills section to better match the job requirements. Focus on:
1. Reordering skills to prioritize those mentioned in the job description
2. Adding any missing skills that the candidate likely has based on their experience (must be reasonable to infer from other sections)
3. Grouping skills into relevant categories that align with the job posting
4. Rephrasing skills using the exact terminology from the job description
5. Removing skills that are irrelevant to this position if the list is very long

IMPORTANT:
1. Do not include empty strings or whitespace-only strings in any arrays
2. Every skill must contain meaningful content
3. Do not create duplicate skills
4. Skills should be concise - typically one to three words each
"""

            elif section_name == "projects":
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the projects section to better match the requirements for a {job_title} position at {company}.

ORIGINAL PROJECTS SECTION:
{content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "projects": [
    {{
      "title": "Project Name",
      "dates": "Time Period (optional)",
      "details": [
        "Detail 1",
        "Detail 2"
      ]
    }}
  ]
}}
```

Please rewrite the projects section to better match the job requirements. Focus on:
1. Highlighting projects that demonstrate skills relevant to this position
2. Emphasizing technical achievements and technologies that align with the job description
3. Quantifying impact and results wherever possible
4. Using terminology from the job description where appropriate

IMPORTANT:
1. Do not include empty strings or whitespace-only strings in any arrays
2. Every detail must contain meaningful content
3. Do not split single project descriptions into multiple entries
4. Make sure each bullet point contains complete, meaningful content
5. Ensure all project entries from the original resume are preserved

Maintain the original project names and dates - only enhance descriptions to make them more relevant.
"""

            else:
                # For other sections, use a generic prompt
                prompt = f"""
You are an expert resume tailoring assistant. Your task is to tailor the {section_name} section to better match the requirements for a {job_title} position at {company}.

ORIGINAL SECTION:
{content}

JOB REQUIREMENTS:
{requirements_text}

REQUIRED SKILLS:
{skills_text}{analysis_prompt}

Return your response as a structured JSON object with the following format:

```json
{{
  "{section_name}": "The tailored content goes here as a single string"
}}
```

Please rewrite this section to better match the job requirements while maintaining the same basic structure and information.
Focus on emphasizing elements most relevant to this job opportunity.
"""

            # Send the request to OpenAI API
            response = self.client.chat.completions.create(
                model="gpt-4o" if "4" in os.environ.get(
    'OPENAI_MODEL_NAME', 'gpt-4') else "gpt-3.5-turbo",
                messages=[
                    {"role": "system",
     "content": "You are an expert resume tailoring assistant."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=4096,
                top_p=1.0
            )

            # Get response text and extract JSON
            response_text = response.choices[0].message.content

            # Save raw response for debugging
            self.raw_responses[section_name] = response_text

            # Save to file
            try:
                # Create api_responses directory if it doesn't exist
                api_responses_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'api_responses')
                if not os.path.exists(api_responses_dir):
                    os.makedirs(api_responses_dir)

                # Generate filename with timestamp
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"{section_name}_response_{timestamp}.json"
                filepath = os.path.join(api_responses_dir, filename)

                # Write response to file
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(response_text)

                logger.info(
    f"Saved raw API response for {section_name} to {filepath}")
            except Exception as e:
                logger.error(f"Error saving raw API response: {str(e)}")

            logger.info(
    f"OpenAI API response for {section_name}: {
        len(response_text)} chars")
            
            # Log token usage
            prompt_tokens = response.usage.prompt_tokens
            completion_tokens = response.usage.completion_tokens
            logger.info(
    f"Completion tokens: {completion_tokens}, Prompt tokens: {prompt_tokens}")

            # Extract JSON from the response
            json_match = re.search(
    r'```json\s*(.*?)\s*```',
    response_text,
     re.DOTALL)

            if json_match:
                json_str = json_match.group(1)
            else:
                # If no json code block, try to find JSON object directly
                json_match = re.search(r'(\{.*\})', response_text, re.DOTALL)
                if json_match:
                    json_str = json_match.group(1)
                else:
                    # No JSON found, return the original content
                    logger.error(f"No JSON found in OpenAI response for {section_name}")
                    return content

            # Parse the JSON string
            cleaned_data = None # Initialize variable
            try:
                json_response = json.loads(json_str)
                logger.debug(f"Parsed LLM response for {section_name} (Req ID: {self.request_id}) before cleaning: {str(json_response)[:200]}...")

                # Process JSON based on section type - Apply cleaning and store structure
            if section_name == "experience" and "experience" in json_response:
                    # Clean the structured data directly after parsing
                    experience_data = json_response["experience"]
                    cleaned_experience_data = []
                    if isinstance(experience_data, list):
                        for job in experience_data:
                            if isinstance(job, dict):
                                # Clean achievements list within the job dictionary
                                achievements = job.get('achievements', [])
                                if isinstance(achievements, list):
                                    job['achievements'] = [strip_bullet_prefix(a) for a in achievements if isinstance(a, str) and a.strip()]
                                cleaned_experience_data.append(job)
                            else:
                                logger.warning(f"Skipping non-dictionary item in experience data: {job}")
                        cleaned_data = cleaned_experience_data # Store cleaned data
                    else:
                        logger.warning(f"Experience data is not a list: {experience_data}")
                        cleaned_data = response_text # Use raw as fallback
            elif section_name == "education" and "education" in json_response:
                    # Clean education highlights
                    education_data = json_response["education"]
                    cleaned_education_data = []
                    if isinstance(education_data, list):
                        for edu in education_data:
                             if isinstance(edu, dict):
                                highlights = edu.get('highlights', [])
                                if isinstance(highlights, list):
                                    edu['highlights'] = [strip_bullet_prefix(h) for h in highlights if isinstance(h, str) and h.strip()]
                                cleaned_education_data.append(edu)
                             else:
                                 logger.warning(f"Skipping non-dictionary item in education data: {edu}")
                    cleaned_data = cleaned_education_data
            elif section_name == "skills" and "skills" in json_response:
                     # Clean skills lists
                    skills_data = json_response["skills"]
                    if isinstance(skills_data, dict):
                        for key in ['technical', 'soft', 'other']:
                            if key in skills_data and isinstance(skills_data[key], list):
                                skills_data[key] = [strip_bullet_prefix(s) for s in skills_data[key] if isinstance(s, str) and s.strip()]
                        cleaned_data = skills_data
                    else:
                        logger.warning(f"Skills data is not a dictionary: {skills_data}")
                        cleaned_data = response_text
            elif section_name == "projects" and "projects" in json_response:
                     # Clean project details
                    projects_data = json_response["projects"]
                    cleaned_projects_data = []
                    if isinstance(projects_data, list):
                        for proj in projects_data:
                            if isinstance(proj, dict):
                                details = proj.get('details', [])
                                if isinstance(details, list):
                                    proj['details'] = [strip_bullet_prefix(d) for d in details if isinstance(d, str) and d.strip()]
                                cleaned_projects_data.append(proj)
                            else:
                               logger.warning(f"Skipping non-dictionary item in projects data: {proj}")
                    cleaned_data = cleaned_projects_data
                elif section_name in json_response and isinstance(json_response[section_name], str):
                    # Generic section (summary, contact): Clean the single string
                    raw_text = json_response[section_name]
                    cleaned_data = clean_bullet_points(raw_text) # Apply general cleaning
            else:
                logger.warning(f"JSON response missing expected '{section_name}' key")
                    cleaned_data = content # Fallback to original content

                # Log after cleaning
                logger.debug(f"Parsed LLM response for {section_name} (Req ID: {self.request_id}) after cleaning: {str(cleaned_data)[:200]}...")

            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse JSON from OpenAI response: {e}")
                logger.error(f"JSON string: {json_str[:100]}...")
                cleaned_data = response_text # Store raw response as the data to save

            # Store the final cleaned data (structure or raw fallback) in the instance
            self.tailored_content[section_name] = cleaned_data
            
            # Save the cleaned data to a temporary file
            self._save_cleaned_section_to_temp_file(section_name, cleaned_data)

            # Return the cleaned data structure (or raw text fallback)
            return cleaned_data
                
        except Exception as e:
            logger.error(f"Error in OpenAI API call: {str(e)}")
            cleaned_data = content # Fallback to original content
            self.tailored_content[section_name] = cleaned_data
            self._save_cleaned_section_to_temp_file(section_name, cleaned_data)
            return cleaned_data

    def save_all_raw_responses(self, resume_filename):
        """Save all raw API responses to a single file"""
        try:
            # Create api_responses directory if it doesn't exist
            api_responses_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'api_responses')
            if not os.path.exists(api_responses_dir):
                os.makedirs(api_responses_dir)

            # Generate filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_name = os.path.splitext(os.path.basename(resume_filename))[0]
            filename = f"{base_name}_all_responses_{timestamp}.json"
            filepath = os.path.join(api_responses_dir, filename)

            # Write all responses to file
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(self.raw_responses, f, indent=2)

            logger.info(f"Saved all raw API responses to {filepath}")
            return filepath
        except Exception as e:
            logger.error(f"Error saving all raw API responses: {str(e)}")
            return None

    def save_tailored_content_to_json(self):
        """Save tailored content (Python structures) to non-timestamped JSON files"""
        try:
            # Create api_responses directory if it doesn't exist
            api_responses_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'api_responses')
            if not os.path.exists(api_responses_dir):
                os.makedirs(api_responses_dir)
                
            # Log content of tailored_content for debugging
            sections_available = list(self.tailored_content.keys())
            logger.info(f"Sections available for saving: {sections_available}")
            
            # Validate sections exist
            if "contact" not in sections_available:
                logger.warning("Contact section not found in tailored content.")
            if "summary" not in sections_available:
                logger.warning("Summary section not found in tailored content.")
            
            # Save each section to a separate JSON file without timestamp
            sections_saved = 0
            for section_name, content in self.tailored_content.items():
                # Skip empty content (allow empty lists/dicts for structure)
                if content is None: 
                    logger.warning(f"None content for section {section_name}, skipping")
                    continue
                
                # Define the output file path (non-timestamped)
                filepath = os.path.join(api_responses_dir, f"{section_name}.json")
                
                # Determine the data to save (content should already be correct Python structure)
                json_data = content

                # Handle simple text sections (summary, contact) - wrap in dict
                if section_name in ["summary", "contact"] and isinstance(content, str):
                     json_data = {"content": content}
                     logger.debug(f"Saving {section_name} string wrapped in dict")
                elif section_name == "experience" and not isinstance(content, list):
                     logger.warning(f"Experience content is not a list, attempting to save raw: {content}")
                     json_data = {"raw_content": str(content)} # Save raw if not list
                elif section_name == "education" and not isinstance(content, list):
                     logger.warning(f"Education content is not a list, attempting to save raw: {content}")
                     json_data = {"raw_content": str(content)} # Save raw if not list
                elif section_name == "skills" and not isinstance(content, dict):
                     logger.warning(f"Skills content is not a dict, attempting to save raw: {content}")
                     json_data = {"raw_content": str(content)} # Save raw if not dict
                elif section_name == "projects" and not isinstance(content, list):
                     logger.warning(f"Projects content is not a list, attempting to save raw: {content}")
                     json_data = {"raw_content": str(content)} # Save raw if not list
                
                # Write the Python structure directly to JSON file
                try:
                    with open(filepath, 'w', encoding='utf-8') as f:
                        # Save with ensure_ascii=False to preserve unicode characters
                        json.dump(json_data, f, indent=2, ensure_ascii=False) 
                sections_saved += 1
                    logger.info(f"Saved {section_name} structure to {filepath}")
                except TypeError as te:
                     logger.error(f"TypeError saving {section_name} to JSON: {te}. Content type: {type(json_data)}")
                     logger.error(f"Problematic content snippet: {str(json_data)[:200]}")
                except Exception as e:
                     logger.error(f"Error saving {section_name} to JSON: {e}")
            
            logger.info(f"Saved {sections_saved} sections to non-timestamped JSON files")
            
            # Verify key files were created
            contact_path = os.path.join(api_responses_dir, "contact.json")
            summary_path = os.path.join(api_responses_dir, "summary.json")
            experience_path = os.path.join(api_responses_dir, "experience.json")
            
            if not os.path.exists(contact_path): logger.warning(f"Contact.json was not created at {contact_path}")
            if not os.path.exists(summary_path): logger.warning(f"Summary.json was not created at {summary_path}")
            if not os.path.exists(experience_path): logger.warning(f"Experience.json was not created at {experience_path}")
                
            return True
            
        except Exception as e:
            logger.error(f"Error saving tailored content to JSON: {str(e)}")
            logger.error(traceback.format_exc())
            return False


def validate_bullet_point_cleaning(sections: Dict[str, str]) -> bool:
    """
    Validates that bullet points have been properly cleaned from resume sections.

    Args:
        sections (Dict[str, str]): Dictionary of resume sections

    Returns:
        bool: True if validation passed, False if bullet points were found
    """
    bullet_patterns = [BULLET_ESCAPE_RE]

    for section, content in sections.items():
        if not content:
            continue
            
        lines = content.split('\n')
        for i, line in enumerate(lines):
            for pattern in bullet_patterns:
                if re.search(pattern, line):
                    logger.warning(f"Validation found bullet marker in {section} section, line {i + 1}: '{line[:30]}...'" )
                    return False

    return True


def validate_html_content(html_content: str) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.validate_html_content(html_content)


def format_section_content(content: str) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.format_section_content(content)


def format_job_entry(company: str, location: str, position: str, dates: str, content: List[str]) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.format_job_entry(company, location, position, dates, content)


def format_education_entry(institution: str, location: str, degree: str, dates: str, highlights: List[str]) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.format_education_entry(institution, location, degree, dates, highlights)


def format_project_entry(title: str, dates: str, details: List[str]) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.format_project_entry(title, dates, details)


def format_education_content(content: str) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.format_education_content(content)


def format_projects_content(content: str) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.format_projects_content(content)


def generate_preview_from_llm_responses(llm_client, for_screen: bool = True) -> str:
    """Proxy function that calls the implementation in html_generator"""
    return html_generator.generate_preview_from_llm_responses(llm_client, for_screen=for_screen)


def extract_resume_sections(doc_path: str) -> Dict[str, str]:
    """Extract sections from a resume document"""
    try:
        logger.info(f"Extracting sections from resume: {doc_path}")
        print(
    f"DEBUG: extract_resume_sections called with doc_path: {doc_path}")
        
        # Import config to check if LLM parsing is enabled
        from config import Config
        use_llm_parsing = Config.USE_LLM_RESUME_PARSING
        llm_provider_config = Config.LLM_RESUME_PARSER_PROVIDER
        
        # First try to parse with LLM if enabled and the module is available
        if use_llm_parsing:
            try:
                # Import the LLM parser module
                from llm_resume_parser import parse_resume_with_llm
                
                # Determine which LLM provider to use based on config or
                # available API keys
                llm_provider = llm_provider_config
                if llm_provider == "auto":
                    if os.environ.get("OPENAI_API_KEY"):
                        llm_provider = "openai"
                    elif os.environ.get("CLAUDE_API_KEY"):
                        llm_provider = "claude"
                    else:
                        logger.warning(
                            "No LLM API keys found for resume parsing. Will use traditional parsing.")
                        raise ImportError("No LLM API keys available")
                    
                # Try to parse with LLM
                logger.info(
    f"Attempting to parse resume with LLM ({llm_provider})...")
                llm_sections = parse_resume_with_llm(doc_path, llm_provider)
                
                # If LLM parsing succeeded, clean bullet points from each
                # section
                if llm_sections and any(
    content for content in llm_sections.values()):
                    logger.info(
                        "LLM parsing successful. Cleaning bullet points from LLM-parsed sections.")
                    
                    # Clean bullet points from each section
                    cleaned_sections = {}
                    for section, content in llm_sections.items():
                        if content:
                            cleaned_content = clean_bullet_points(content)
                            cleaned_sections[section] = cleaned_content
                            logger.info(
    f"LLM extracted and cleaned {section} section: {
        len(cleaned_content)} chars")
                        else:
                            cleaned_sections[section] = ""
                            logger.info(
    f"LLM found no content for {section} section")
                    
                    # Validate the cleaning was successful
                    validation_success = validate_bullet_point_cleaning(
                        cleaned_sections)
                    if not validation_success:
                        logger.warning(
                            "Bullet point cleaning validation failed. There may still be bullet markers in the text.")
                    
                    return cleaned_sections
                else:
                    logger.warning(
                        "LLM parsing did not return usable results. Falling back to traditional parsing.")
            except (ImportError, Exception) as e:
                logger.warning(f"LLM parsing unavailable or failed: {str(e)}. Using traditional parsing.")
        else:
            logger.info(
                "LLM parsing is disabled by configuration. Using traditional parsing.")
        
        # If LLM parsing failed, is unavailable, or is disabled, fall back to
        # traditional parsing
        logger.info("Using traditional resume section extraction...")
        print(
    f"DEBUG: fallback to traditional parsing - will examine document: {doc_path}")
        
        # Extract plain text content from the DOCX file
        text = docx2txt.process(doc_path)
        
        # Parse the document into sections
        print(f"DEBUG: About to create Document object from: {doc_path}")
        doc = Document(doc_path)
        print(f"DEBUG: Document object created successfully - examining sections")
        
        # Initialize standard resume sections
        sections = {
            "contact": "",
            "summary": "",
            "experience": "",
            "education": "",
            "skills": "",
            "projects": "",
            "additional": ""
        }
        
        # Debug: Count total paragraphs with content
        total_paragraphs = sum(
    1 for para in doc.paragraphs if para.text.strip())
        logger.info(f"Total paragraphs with content: {total_paragraphs}")
        
        # Extract sections based on heading style
        current_section = "contact"  # Default to contact for initial content
        section_content = []
        
        # Log all potential section headers for debugging
        logger.info("Scanning document for potential section headers...")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # Debug potential headers
            if para.style.name.startswith('Heading') or any(
                p.bold for p in para.runs):
                logger.info(
    f"Potential header found at para {i}: '{text}' (Style: {
        para.style.name}, Bold: {
            any(
                p.bold for p in para.runs)})")
        
        # First pass - try to extract based on formatting
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # Skip empty paragraphs
            if not text:
                continue
            
            # Check if this is a heading (section title)
            if para.style.name.startswith('Heading') or any(
                p.bold for p in para.runs):
                # Store previous section content
                if section_content:
                    sections[current_section] = "\n".join(section_content)
                    section_content = []
                
                # Determine new section type based on heading text
                if re.search(
    r'contact|info|email|phone|address',
     text.lower()):
                    current_section = "contact"
                elif re.search(r'summary|objective|profile|about', text.lower()):
                    current_section = "summary"
                elif re.search(r'experience|work|employment|job|career|professional', text.lower()):
                    current_section = "experience"
                elif re.search(r'education|degree|university|college|school|academic', text.lower()):
                    current_section = "education"
                elif re.search(r'skills|expertise|technologies|competencies|qualification|proficiencies', text.lower()):
                    current_section = "skills"
                elif re.search(r'projects|portfolio|works', text.lower()):
                    current_section = "projects"
                elif re.search(r'additional|interests|activities|volunteer|certification|awards|achievements', text.lower()):
                    current_section = "additional"
                else:
                    # Default to additional for unknown headings
                    current_section = "additional"
                
                logger.info(
    f"Section header detected: '{text}' -> categorized as '{current_section}'")
            else:
                # Add content to current section
                section_content.append(text)
        
        # Add the last section content
        if section_content:
            sections[current_section] = "\n".join(section_content)
        
        # Check if we actually found any sections beyond contact
        sections_found = sum(
    1 for section,
     content in sections.items() if content and section != "contact")
        logger.info(
    f"Sections found with standard detection: {sections_found}")

        # If we didn't find any sections or only found contact, try a simpler
        # approach
        if sections_found == 0:
            logger.info(
                "No sections detected with standard method. Using fallback approach...")
            
            # Fallback: Treat the entire document as experience section
            if total_paragraphs > 0:
                all_content = "\n".join(
    para.text for para in doc.paragraphs if para.text.strip())
                if len(all_content) > 0:
                    sections["experience"] = all_content
                    logger.info(
    f"Fallback: Added {
        len(all_content)} chars to experience section")
        
        # Clean bullet points from traditional parsing results
        cleaned_sections = {}
        for section, content in sections.items():
            if content:
                cleaned_content = clean_bullet_points(content)
                cleaned_sections[section] = cleaned_content
                logger.info(
    f"Traditional parsing: cleaned {section} section: {
        len(cleaned_content)} chars")
            else:
                cleaned_sections[section] = ""
        
        # Validate the cleaning was successful
        validation_success = validate_bullet_point_cleaning(cleaned_sections)
        if not validation_success:
            logger.warning(
                "Bullet point cleaning validation failed. There may still be bullet markers in the text.")
        
        return cleaned_sections
    
    except Exception as e:
        logger.error(f"Error extracting resume sections: {str(e)}")
        logger.error(traceback.format_exc())
        
        # Return at least an empty structure in case of error
        return {
            "contact": "",
            "summary": "",
            "experience": "Error extracting content. Please check your resume format.",
            "education": "",
            "skills": "",
            "projects": "",
            "additional": ""
        }


def tailor_resume_with_llm(resume_path: str, job_data: Dict, api_key: str, provider: str = "openai", api_url: str = None) -> Tuple[str, str, LLMClient, str]: # Added request_id to return tuple
    """
    Tailor resume with LLM (Claude or OpenAI).
    Returns output PDF info, the LLM client instance, and the unique request ID.
    """
    global last_llm_client
    
    logger.info(f"Tailoring resume with {provider} LLM")
    
    # Extract resume sections
    resume_sections = extract_resume_sections(resume_path)
    
    # Initialize the appropriate LLM client (generates request_id)
    llm_client = None
    if provider.lower() == "claude":
        llm_client = ClaudeClient(api_key, api_url)
    elif provider.lower() == "openai":
        llm_client = OpenAIClient(api_key)
    else:
        raise ValueError(f"Unsupported LLM provider: {provider}")
    
    request_id = llm_client.get_request_id() # Get the generated ID
    logger.info(f"Starting tailoring process with Request ID: {request_id}")
    
    # Add contact section directly to tailored_content without tailoring it
    if resume_sections.get("contact", "").strip():
        logger.info(f"Preserving contact section for tailored resume: {len(resume_sections['contact'])} chars")
        llm_client.tailored_content["contact"] = resume_sections["contact"]
    else:
        logger.warning("No contact information found in resume sections")
        
    # Handle summary section - either preserve existing or generate a new one
    if resume_sections.get("summary", "").strip():
        logger.info(f"Preserving existing summary section for tailored resume: {len(resume_sections['summary'])} chars")
        llm_client.tailored_content["summary"] = resume_sections["summary"]
    else:
        logger.warning("No summary information found in resume sections, generating a new summary")
        generated_summary = generate_professional_summary(resume_sections, job_data, llm_client, provider)
        if generated_summary:
            logger.info(f"Generated new professional summary: {len(generated_summary)} chars")
            llm_client.tailored_content["summary"] = generated_summary
            llm_client._save_cleaned_section_to_temp_file("summary", generated_summary)
        else:
            logger.error("Failed to generate professional summary")
    
    # Tailor each section
    for section_name, content in resume_sections.items():
        if content.strip() and section_name in ["experience", "education", "skills", "projects"]:
            logger.info(f"Tailoring {section_name} section")
            tailored_content = llm_client.tailor_resume_content(section_name, content, job_data)
            resume_sections[section_name] = tailored_content
    
    # Generate output filename
    resume_filename = os.path.basename(resume_path)
    base_name = os.path.splitext(resume_filename)[0]
    output_filename = f"{base_name}_tailored_{provider}.pdf"
    output_dir = os.path.dirname(resume_path)
    output_path = os.path.join(output_dir, output_filename)
    
    # Store LLM client for preview generation
    last_llm_client = llm_client
    
    # Return info including the request_id
    return output_filename, output_path, llm_client, request_id

def generate_professional_summary(resume_sections: Dict[str, str], job_data: Dict, llm_client, provider: str) -> str:
    """
    Generate a professional summary based on resume content and job data
    
    Args:
        resume_sections: Dictionary of resume sections
        job_data: Job data including requirements and skills
        llm_client: LLM client instance
        provider: LLM provider name
        
    Returns:
        Generated professional summary or empty string if failed
    """
    try:
        # Collect relevant information from the resume
        experience = resume_sections.get("experience", "")
        education = resume_sections.get("education", "")
        skills = resume_sections.get("skills", "")
        
        # Extract key information from job data
        job_title = job_data.get("title", "")
        job_requirements = job_data.get("requirements", "")
        candidate_profile = ""
        hard_skills = []
        soft_skills = []
        
        # Check if we have LLM-analyzed job data with more structured information
        job_analysis = job_data.get("analysis", {})
        if job_analysis:
            if "candidate_profile" in job_analysis:
                candidate_profile = job_analysis["candidate_profile"]
            if "hard_skills" in job_analysis:
                hard_skills = job_analysis["hard_skills"]
            if "soft_skills" in job_analysis:
                soft_skills = job_analysis["soft_skills"]
        
        # Create a combined string of skills for the prompt
        target_skills = ", ".join(hard_skills + soft_skills)
        
        # Create a prompt for summary generation
        prompt = f"""
Generate a professional summary section for a resume tailored to the following job:

JOB TITLE: {job_title}

CANDIDATE PROFILE:
{candidate_profile}

TARGET SKILLS:
{target_skills}

REQUIREMENTS:
{job_requirements}

Based on the following resume information:

EXPERIENCE:
{experience[:500]}...

EDUCATION:
{education[:300]}

SKILLS:
{skills[:300]}

Create a concise, powerful professional summary (3-5 sentences) that:
1. Positions the candidate as an ideal fit for this specific role
2. Highlights relevant experience and skills that match the job requirements
3. Demonstrates the candidate's unique value proposition
4. Is written in first person without using "I" statements
5. Uses strong action verbs and specific, quantifiable achievements where possible

Format the summary as a single paragraph without bullet points.
"""
        
        # Use the appropriate LLM to generate the summary
        if provider.lower() == "claude":
            summary = generate_summary_with_claude(prompt, llm_client)
        else:
            summary = generate_summary_with_openai(prompt, llm_client)
        
        return summary
        
    except Exception as e:
        logger.error(f"Error generating professional summary: {str(e)}")
        logger.error(traceback.format_exc())
        return ""

def generate_summary_with_claude(prompt: str, claude_client) -> str:
    """Generate a summary using Claude API"""
    try:
        # Use the Claude client to generate a summary
        response = claude_client.client.messages.create(
            model="claude-3-sonnet-20240229",
            max_tokens=400,
            temperature=0.5,
            system="""You are a professional resume writer specializing in creating concise, 
            impactful professional summaries. Your summaries highlight a candidate's strengths 
            and align with job requirements without hyperbole.""",
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        
        # Extract and clean the summary
        summary = response.content[0].text
        
        # Save the raw response
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        response_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'api_responses')
        if not os.path.exists(response_dir):
            os.makedirs(response_dir)
        
        with open(os.path.join(response_dir, f"summary_generation_response_{timestamp}.json"), "w") as f:
            json.dump({"provider": "claude", "prompt": prompt, "response": summary}, f, indent=2)
        
        logger.info(f"Generated summary with Claude: {len(summary)} chars")
        return summary
    
    except Exception as e:
        logger.error(f"Error generating summary with Claude: {str(e)}")
        return ""

def generate_summary_with_openai(prompt: str, openai_client) -> str:
    """Generate a summary using OpenAI API"""
    try:
        # Use the OpenAI client to generate a summary
        response = openai_client.client.chat.completions.create(
            model="gpt-4o",
            temperature=0.5,
            messages=[
                {"role": "system", "content": """You are a professional resume writer specializing in creating concise, 
                impactful professional summaries. Your summaries highlight a candidate's strengths 
                and align with job requirements without hyperbole."""},
                {"role": "user", "content": prompt}
            ]
        )
        
        # Extract the summary
        summary = response.choices[0].message.content
        
        # Save the raw response
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        response_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'api_responses')
        if not os.path.exists(response_dir):
            os.makedirs(response_dir)
        
        with open(os.path.join(response_dir, f"summary_generation_response_{timestamp}.json"), "w") as f:
            json.dump({
                "provider": "openai", 
                "prompt": prompt, 
                "response": summary,
                "prompt_tokens": response.usage.prompt_tokens,
                "completion_tokens": response.usage.completion_tokens,
                "total_tokens": response.usage.total_tokens
            }, f, indent=2)
        
        logger.info(f"Generated summary with OpenAI: {len(summary)} chars")
        return summary
    
    except Exception as e:
        logger.error(f"Error generating summary with OpenAI: {str(e)}")
        return ""

def generate_resume_preview(resume_path: str) -> str:
    """
    Generate HTML preview from resume
    
    Args:
        resume_path: Path to resume file
        
    Returns:
        HTML preview
    """
    logger.info(f"Generating resume preview for {resume_path}")
    
    # Check if we have direct LLM responses from last tailoring
    global last_llm_client
    if last_llm_client:
        logger.info("Using direct LLM responses for preview generation")
        preview_html = generate_preview_from_llm_responses(last_llm_client)
        if preview_html:
            return preview_html
    
    # Otherwise extract sections from file
    resume_sections = extract_resume_sections(resume_path)
    
    # Generate preview HTML
    preview_html = ""
    for section_name, content in resume_sections.items():
        if content.strip():
            formatted_content = format_section_content(content)
            preview_html += f'<div class="resume-section"><h2>{section_name.title()}</h2>{formatted_content}</div>'
    
    return preview_html
