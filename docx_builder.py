def _fix_spacing_between_sections(doc):
    """
    Fix spacing between sections by setting space_after=0 on paragraphs
    that appear right before a section header.
    
    This addresses the issue of excessive space between the last content
    paragraph of one section and the next section header.
    
    Args:
        doc: The DOCX document to fix
    """
    try:
        from docx.shared import Pt
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml
        
        logger.info("Starting section spacing fix process...")
        
        # Find all section headers by style
        section_headers = []
        for i, para in enumerate(doc.paragraphs):
            if para.style and para.style.name == 'BoxedHeading2':
                section_headers.append(i)
                logger.info(f"Found section header at index {i}: '{para.text[:30]}'")
        
        logger.info(f"Found {len(section_headers)} section headers in document")
        
        # No section headers or only one (no "between" to fix)
        if len(section_headers) < 2:
            logger.info("Less than 2 section headers found, no spacing to fix")
            return
        
        # For each section header (except the first), fix the paragraph right before it
        fixed_count = 0
        for header_idx in section_headers[1:]:  # Skip the first header
            # Find the paragraph right before this header
            prev_para_idx = header_idx - 1
            
            # Skip if it's an invalid index
            if prev_para_idx < 0 or prev_para_idx >= len(doc.paragraphs):
                logger.warning(f"Invalid previous paragraph index: {prev_para_idx}")
                continue
                
            # Get the paragraph before the section header
            prev_para = doc.paragraphs[prev_para_idx]
            prev_text = prev_para.text[:30] + "..." if len(prev_para.text) > 30 else prev_para.text
            
            # Skip if it's another section header (would be unusual but possible)
            if prev_para.style and prev_para.style.name == 'BoxedHeading2':
                logger.info(f"Skipping paragraph at index {prev_para_idx} - it's another section header")
                continue
            
            # Get current spacing for logging
            current_space_after = prev_para.paragraph_format.space_after
            current_space_pt = current_space_after.pt if current_space_after else "None"
            logger.info(f"Fixing spacing on paragraph at index {prev_para_idx}: '{prev_text}' (current space_after={current_space_pt}pt)")
                
            # Set space_after to 0
            prev_para.paragraph_format.space_after = Pt(0)
            
            # Apply the change through XML as well for maximum control
            if hasattr(prev_para._element, 'get_or_add_pPr'):
                p_pr = prev_para._element.get_or_add_pPr()
                
                # Find existing spacing element or create it
                spacing = p_pr.find(qn('w:spacing'))
                
                if spacing is not None:
                    # Set the w:after attribute to 0
                    spacing.set(qn('w:after'), '0')
                    logger.info(f"Updated existing spacing XML for paragraph {prev_para_idx}")
                else:
                    # Create a new spacing element with all necessary attributes to override defaults
                    spacing_xml = f'<w:spacing {nsdecls("w")} w:after="0" w:afterAutospacing="0"/>'
                    p_pr.append(parse_xml(spacing_xml))
                    logger.info(f"Added new spacing XML for paragraph {prev_para_idx}")
            
            fixed_count += 1
            logger.info(f"Fixed spacing on paragraph {prev_para_idx}")
        
        logger.info(f"Completed section spacing fix: modified {fixed_count} paragraphs")
        
        # Verify fix was applied
        if fixed_count > 0:
            verification_count = 0
            for header_idx in section_headers[1:]:
                prev_para_idx = header_idx - 1
                if prev_para_idx >= 0 and prev_para_idx < len(doc.paragraphs):
                    prev_para = doc.paragraphs[prev_para_idx]
                    if prev_para.style and prev_para.style.name != 'BoxedHeading2':
                        space_after = prev_para.paragraph_format.space_after
                        if space_after and space_after.pt == 0:
                            verification_count += 1
            
            logger.info(f"Verification: {verification_count} of {fixed_count} paragraphs confirmed with space_after=0")
    
    except Exception as e:
        logger.error(f"Error fixing spacing between sections: {e}")
        import traceback
        logger.error(traceback.format_exc())
        # Continue execution - this is a fix attempt, not critical functionality 

def tighten_before_headers(doc):
    """
    Finds 'BoxedHeading2' paragraphs and ensures that any preceding empty paragraphs
    have their spacing (before and after) set to zero. This prevents unwanted gaps
    caused by empty paragraphs that might have default spacing.
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    logger.info("Starting tighten_before_headers process...")
    fixed_instances_count = 0
    for p_idx, p_obj in enumerate(doc.paragraphs):
        if p_obj.style and p_obj.style.name == 'BoxedHeading2':
            logger.debug(f"Found BoxedHeading2 at paragraph index {p_idx}: '{p_obj.text[:30]}'")
            # Get the underlying lxml element for the current paragraph
            current_xml_p = p_obj._element
            prev_xml_p = current_xml_p.getprevious()
            
            processed_prev_paras = 0
            while prev_xml_p is not None and prev_xml_p.tag.endswith('}p'):
                # Check if the previous paragraph is empty (has no w:t elements)
                # An empty paragraph in XML might look like <w:p><w:pPr>...</w:pPr></w:p> or <w:p/>
                # A paragraph with text will have <w:r><w:t>text</w:t></w:r>
                # So, we check for the absence of any w:t (text) elements within it.
                if not prev_xml_p.xpath('.//w:t', namespaces=prev_xml_p.nsmap):
                    logger.debug(f"  Found preceding empty paragraph (index ~{p_idx - 1 - processed_prev_paras}). Applying zero spacing.")
                    # Get or add pPr for the empty paragraph
                    pPr = prev_xml_p.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
                        prev_xml_p.insert(0, pPr)
                    
                    # Remove any existing spacing element first to avoid conflicts
                    for existing_spacing in pPr.xpath('./w:spacing'):
                        pPr.remove(existing_spacing)
                    
                    # Add new comprehensive zero spacing XML
                    # Using a basic lineRule="auto" and a common line value like 240 (12pt) for safety,
                    # though for an empty paragraph, it might not matter as much.
                    # The critical parts are w:before="0" and w:after="0".
                    spacing_xml_fix = f'<w:spacing {nsdecls("w")} w:before="0" w:after="0" w:line="240" w:lineRule="auto" w:beforeAutospacing="0" w:afterAutospacing="0"/>'
                    pPr.append(parse_xml(spacing_xml_fix))
                    fixed_instances_count +=1
                else:
                    # Found a non-empty paragraph, stop searching backwards for this header
                    logger.debug(f"  Found preceding non-empty paragraph (index ~{p_idx - 1 - processed_prev_paras}). Stopping backward search for this header.")
                    break 
                
                prev_xml_p = prev_xml_p.getprevious()
                processed_prev_paras += 1
                if processed_prev_paras >= 3: # Safety break: only check up to 3 preceding paragraphs
                    logger.debug("  Reached safety limit of 3 preceding paragraphs checked.")
                    break
        
    logger.info(f"Completed tighten_before_headers process. Applied zero spacing to {fixed_instances_count} empty preceding paragraphs.")

def build_docx(request_id: str, temp_dir: str, debug: bool = False) -> BytesIO:
    """
    Build a DOCX file from the resume data for the given request ID.
    
    Args:
        request_id: The unique request ID for the resume
        temp_dir: Directory containing the temp session data files
        debug: Whether to enable debugging output
        
    Returns:
        BytesIO object containing the DOCX file data
    """
    try:
        logger.info(f"Building DOCX for request ID: {request_id}")
        logger.info(f"Temp directory path: {temp_dir}")
        logger.info(f"Debug mode: {debug}")
        
        # Check for files related to this request_id
        try:
            file_list = os.listdir(temp_dir)
            matching_files = [f for f in file_list if request_id in f]
            logger.info(f"Files containing request ID: {matching_files}")
        except Exception as e:
            logger.error(f"Error listing directory: {e}")
        
        # Load DOCX styles from StyleManager
        docx_styles = StyleManager.load_docx_styles()
        if not docx_styles:
            logger.error("No DOCX styles found. Using defaults.")
            docx_styles = {
                "page": {"marginTopCm": 1.0, "marginBottomCm": 1.0, "marginLeftCm": 2.0, "marginRightCm": 2.0},
                "heading1": {"fontFamily": "Calibri", "fontSizePt": 16},
                "heading2": {"fontFamily": "Calibri", "fontSizePt": 14, "color": [0, 0, 102], "spaceAfterPt": 6},
                "heading3": {"fontFamily": "Calibri", "fontSizePt": 11, "bold": True, "spaceAfterPt": 4},
                "body": {"fontFamily": "Calibri", "fontSizePt": 11}
            }
        
        # Create a new Document
        doc = Document()
        
        # Create custom document styles
        custom_styles = _create_document_styles(doc, docx_styles)
        
        # Configure page margins using global margin tokens
        section = doc.sections[0]
        page_config = docx_styles.get("page", {})
        
        # Get margin values from global tokens
        structured_tokens = StyleEngine.get_structured_tokens()
        global_margins = structured_tokens.get("global", {}).get("margins", {})
        global_left_margin_cm = float(global_margins.get("leftCm", "2.0"))
        global_right_margin_cm = float(global_margins.get("rightCm", "2.0"))
        
        # Set margins from tokens and style config
        section.top_margin = Cm(page_config.get("marginTopCm", 1.0))
        section.bottom_margin = Cm(page_config.get("marginBottomCm", 1.0))
        section.left_margin = Cm(global_left_margin_cm)  # Use global token
        section.right_margin = Cm(global_right_margin_cm)  # Use global token
        
        logger.info(f"Applied document margins from global tokens: Left={global_left_margin_cm}cm, Right={global_right_margin_cm}cm")
        
        # ------ CONTACT SECTION ------
        logger.info("Processing Contact section...")
        
        # Open and inspect the contact file directly to debug structure issues
        contact_file_path = os.path.join(temp_dir, f"{request_id}_contact.json")
        if os.path.exists(contact_file_path):
            try:
                with open(contact_file_path, 'r', encoding='utf-8') as f:
                    raw_content = f.read()
                    logger.info(f"Raw contact file content (first 200 chars): {raw_content[:200]}")
            except Exception as e:
                logger.error(f"Error reading raw contact file: {e}")
        
        contact = load_section_json(request_id, "contact", temp_dir)
        logger.info(f"Contact data loaded: {bool(contact)}")
        logger.info(f"Contact data type: {type(contact)}")
        
        if isinstance(contact, dict):
            logger.info(f"Contact keys: {list(contact.keys())}")
        
        if contact:
            # Verify contact is a dictionary
            if not isinstance(contact, dict):
                logger.warning(f"Contact section is not a dictionary: {type(contact)}")
                contact = {}
            
            # Check if contact has a 'content' key (alternate structure)
            if 'content' in contact:
                logger.info("Found 'content' key in contact section")
                # Extract actual contact data from content
                try:
                    # Try to access the contact data
                    contact_data = contact['content']
                    if isinstance(contact_data, dict):
                        logger.info(f"Contact content is a dictionary with keys: {list(contact_data.keys())}")
                        contact = contact_data
                    elif isinstance(contact_data, list) and len(contact_data) > 0 and isinstance(contact_data[0], dict):
                        logger.info(f"Contact content is a list of dictionaries, using first item")
                        contact = contact_data[0]
                    elif isinstance(contact_data, str):
                        # Handle string content - parse it into structured data
                        logger.info("Contact content is a string, attempting to parse")
                        
                        # Parse the contact information from the string
                        lines = contact_data.strip().split('\n')
                        
                        # Extract name from the first line
                        name = lines[0].strip() if lines else ""
                        logger.info(f"Extracted name from string: {name}")
                        
                        # Extract contact details from subsequent lines
                        contact_details = {}
                        if len(lines) > 1:
                            details_line = lines[1].strip()
                            # Split by common separators
                            details_parts = [p.strip() for p in details_line.replace('|', '|').split('|')]
                            
                            logger.info(f"Extracted contact details: {details_parts}")
                            
                            for part in details_parts:
                                part = part.strip()
                                # Try to identify the type of contact detail
                                if '@' in part:
                                    contact_details['email'] = part
                                elif 'P:' in part or 'Phone:' in part or any(c.isdigit() for c in part):
                                    contact_details['phone'] = part
                                elif 'LinkedIn' in part or 'linkedin.com' in part:
                                    contact_details['linkedin'] = part
                                elif 'Github' in part or 'github.com' in part:
                                    contact_details['github'] = part
                                elif any(loc in part.lower() for loc in ['street', 'ave', 'road', 'blvd', 'city', 'state']):
                                    contact_details['location'] = part
                        
                        # Create a new contact object
                        contact = {'name': name, **contact_details}
                        logger.info(f"Created structured contact data: {contact}")
                    else:
                        logger.info(f"Contact content is type: {type(contact_data)}")
                except Exception as e:
                    logger.warning(f"Error accessing contact content: {e}")
            
            # Print full contact data for debugging
            logger.info(f"Final contact data structure to use: {contact}")
            
            # Name - handle potential different structures
            name = ""
            if "name" in contact:
                name = contact["name"]
            elif "full_name" in contact:
                name = contact["full_name"]
            
            if name:
                name_para = doc.add_paragraph(name)
                _apply_paragraph_style(name_para, "heading1", docx_styles)
                
                # Contact details
                contact_parts = []
                if "location" in contact and contact["location"]:
                    contact_parts.append(contact["location"])
                if "phone" in contact and contact["phone"]:
                    contact_parts.append(contact["phone"])
                if "email" in contact and contact["email"]:
                    contact_parts.append(contact["email"])
                if "linkedin" in contact and contact["linkedin"]:
                    contact_parts.append(contact["linkedin"])
                
                # Additional possible contact fields
                if "website" in contact and contact["website"]:
                    contact_parts.append(contact["website"])
                if "github" in contact and contact["github"]:
                    contact_parts.append(contact["github"])
                
                contact_text = " | ".join(contact_parts)
                contact_para = doc.add_paragraph(contact_text)
                _apply_paragraph_style(contact_para, "body", docx_styles)
                contact_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add a horizontal line
                doc.add_paragraph("").paragraph_format.space_after = Pt(6)
            else:
                logger.warning("No name found in contact data, skipping contact section")
        else:
            logger.warning("No contact data found")
            
            # Try a fallback approach - look for the file directly
            fallback_files = [f for f in os.listdir(temp_dir) if f.endswith("_contact.json") and request_id in f]
            if fallback_files:
                logger.info(f"Found fallback contact files: {fallback_files}")
                try:
                    with open(os.path.join(temp_dir, fallback_files[0]), 'r', encoding='utf-8') as f:
                        fallback_contact = json.load(f)
                        logger.info(f"Loaded fallback contact data: {fallback_contact}")
                        
                        # Extract name and add to document
                        if isinstance(fallback_contact, dict):
                            name = fallback_contact.get("name", "")
                            if name:
                                name_para = doc.add_paragraph(name)
                                _apply_paragraph_style(name_para, "heading1", docx_styles)
                                logger.info(f"Added name from fallback: {name}")
                except Exception as e:
                    logger.error(f"Error processing fallback contact data: {e}")
        
        # ------ SUMMARY SECTION ------
        logger.info("Processing Summary section...")
        summary = load_section_json(request_id, "summary", temp_dir)
        logger.info(f"Summary data loaded: {bool(summary)}")
        
        # Handle both direct summary and summary with 'content' key
        summary_text = ""
        if summary:
            if isinstance(summary, dict):
                if "summary" in summary:
                    summary_text = summary.get("summary", "")
                elif "content" in summary:
                    # Alternative structure with content key
                    content = summary.get("content", "")
                    if isinstance(content, dict) and "summary" in content:
                        summary_text = content.get("summary", "")
                    else:
                        summary_text = str(content)
            else:
                summary_text = str(summary)
                
            if summary_text:
                # Add section header with helper function
                summary_header = add_section_header(doc, "PROFESSIONAL SUMMARY", docx_styles)
                
                # Add summary content
                summary_para = doc.add_paragraph(summary_text, style='MR_SummaryText')
                
                # Space after section
                doc.add_paragraph("").paragraph_format.space_after = Pt(6)
        
        # ------ EXPERIENCE SECTION ------
        logger.info("Processing Experience section...")
        experience = load_section_json(request_id, "experience", temp_dir)
        logger.info(f"Experience data loaded: {bool(experience)}")
        logger.info(f"Experience contains 'experiences' key: {isinstance(experience, dict) and 'experiences' in experience}")
        
        # Handle both dictionary with 'experiences' key and direct list of experiences
        experiences_list = []
        if experience:
            if isinstance(experience, dict) and "experiences" in experience:
                experiences_list = experience.get("experiences", [])
            elif isinstance(experience, list):
                # Direct list of experiences
                experiences_list = experience
            else:
                logger.warning(f"Unexpected experience format: {type(experience)}")
                
            if experiences_list:
                # Add section header with helper function
                exp_header = add_section_header(doc, "EXPERIENCE", docx_styles)
                
                # Verify experiences is a list
                if not isinstance(experiences_list, list):
                    logger.warning(f"Experiences is not a list: {type(experiences_list)}")
                    experiences_list = []
                    
                # Add each job
                for job in experiences_list:
                    # Verify job is a dictionary
                    if not isinstance(job, dict):
                        logger.warning(f"Job is not a dictionary: {type(job)}")
                        continue
                        
                    # Company and location - use the helper function for consistent formatting
                    company = job.get('company', '')
                    location = job.get('location', '')
                    
                    if company or location:
                        company_para = format_right_aligned_pair(
                            doc,
                            company,
                            location,
                            "heading3",
                            "heading3",
                            docx_styles
                        )
                    
                    # Position and dates - use role box for consistent formatting
                    position = job.get('position', '')
                    if not position and job.get('title'):  # Fallback to 'title' if 'position' is not available
                        position = job.get('title', '')
                    dates = job.get('dates', '')
                    
                    if position:
                        # Import the role box function
                        from word_styles.section_builder import add_role_box
                        
                        # Add role box with position and dates
                        role_box_table = add_role_box(doc, position, dates if dates else None)
                    
                    logger.info(f"Formatted experience entry: '{company} - {position}'")
                    
                    # Role description - use the helper function for consistent formatting
                    if job.get('role_description'):
                        role_para = add_role_description(doc, job.get('role_description'), docx_styles)
                    
                    # Achievements/bullets - use the helper function for consistent formatting
                    for achievement in job.get('achievements', []):
                        bullet_para = create_bullet_point(doc, achievement, docx_styles)
                
        # ------ EDUCATION SECTION ------
        logger.info("Processing Education section...")
        education = load_section_json(request_id, "education", temp_dir)
        logger.info(f"Education data loaded: {bool(education)}")
        logger.info(f"Education contains 'institutions' key: {isinstance(education, dict) and 'institutions' in education}")
        
        # Handle both dictionary with 'institutions' key and direct list of institutions
        institutions_list = []
        if education:
            if isinstance(education, dict) and "institutions" in education:
                institutions_list = education.get("institutions", [])
            elif isinstance(education, list):
                # Direct list of institutions
                institutions_list = education
            else:
                logger.warning(f"Unexpected education format: {type(education)}")
                
            if institutions_list:
                # Add section header with helper function
                edu_header = add_section_header(doc, "EDUCATION", docx_styles)
                
                # Verify institutions is a list
                if not isinstance(institutions_list, list):
                    logger.warning(f"Institutions is not a list: {type(institutions_list)}")
                    institutions_list = []
                    
                # Add each institution
                for school in institutions_list:
                    # Verify school is a dictionary
                    if not isinstance(school, dict):
                        logger.warning(f"School is not a dictionary: {type(school)}")
                        continue
                        
                    # Institution and location - use the helper function for consistent formatting
                    institution = school.get('institution', '')
                    location = school.get('location', '')
                    
                    if institution or location:
                        school_para = format_right_aligned_pair(
                            doc,
                            institution,
                            location,
                            "heading3",
                            "heading3",
                            docx_styles
                        )
                    
                    # Degree and dates - use the helper function for consistent formatting
                    degree = school.get('degree', '')
                    dates = school.get('dates', '')
                    
                    if degree or dates:
                        degree_para = format_right_aligned_pair(
                            doc,
                            degree,
                            dates,
                            "body",
                            "body",
                            docx_styles
                        )
                    
                    logger.info(f"Formatted education entry: '{institution} - {degree}'")
                    
                    # Highlights/bullets - use the helper function for consistent formatting
                    for highlight in school.get('highlights', []):
                        bullet_para = create_bullet_point(doc, highlight, docx_styles)
        
        # ------ SKILLS SECTION ------
        logger.info("Processing Skills section...")
        skills = load_section_json(request_id, "skills", temp_dir)
        logger.info(f"Skills data loaded: {bool(skills)}")
        logger.info(f"Skills data type: {type(skills)}")
        logger.info(f"Skills content sample: {str(skills)[:100]}")
        
        if skills:
            # Add section header with helper function 
            skills_header = add_section_header(doc, "SKILLS", docx_styles)
            
            # Add skills content - handle different possible formats
            if isinstance(skills, dict) and "skills" in skills:
                skills_content = skills.get("skills", "")
                logger.info(f"Skills content type: {type(skills_content)}")
                
                # Check if skills content is a list
                if isinstance(skills_content, list):
                    logger.info("Processing skills as inline list with commas")
                    # Handle skills as a comma-separated list on a single line
                    skills_text = ", ".join([str(skill) for skill in skills_content])
                    skills_para = doc.add_paragraph(skills_text, style='MR_SkillList')
                    logger.info(f"Applied MR_SkillList style to skills list")
                elif isinstance(skills_content, dict):
                    logger.info("Processing skills as dict")
                    # Handle skills as dictionary with categories
                    for category, skill_list in skills_content.items():
                        # Add category as subheading
                        category_para = doc.add_paragraph(category.upper(), style='MR_SkillCategory')
                        logger.info(f"Applied MR_SkillCategory style to category: {category}")
                        
                        # Add skills in this category as a comma-separated list
                        if isinstance(skill_list, list):
                            skills_text = ", ".join([str(skill) for skill in skill_list])
                            skills_para = doc.add_paragraph(skills_text, style='MR_SkillList')
                            logger.info(f"Applied MR_SkillList style to skills in category: {category}")
                        else:
                            # Not a list, just add as text
                            skill_para = doc.add_paragraph(str(skill_list), style='MR_SkillList')
                            logger.info(f"Applied MR_SkillList style to non-list skills in category: {category}")
                else:
                    logger.info("Processing skills as string")
                    # Handle skills as string or any other type
                    skills_para = doc.add_paragraph(str(skills_content), style='MR_SkillList')
                    logger.info(f"Applied MR_SkillList style to skills string")
            elif isinstance(skills, dict):
                logger.info("Processing skills dict directly")
                # Direct display of skills object if it doesn't have "skills" key
                for category, skill_list in skills.items():
                    # Add category as subheading
                    category_para = doc.add_paragraph(category.upper(), style='MR_SkillCategory')
                    logger.info(f"Applied MR_SkillCategory style to category: {category}")
                    
                    # Add skills in this category as a comma-separated list
                    if isinstance(skill_list, list):
                        skills_text = ", ".join([str(skill) for skill in skill_list])
                        skills_para = doc.add_paragraph(skills_text, style='MR_SkillList')
                        logger.info(f"Applied MR_SkillList style to skills in category: {category}")
                    else:
                        # Not a list, just add as text
                        skill_para = doc.add_paragraph(str(skill_list), style='MR_SkillList')
                        logger.info(f"Applied MR_SkillList style to non-list skills in category: {category}")
            elif isinstance(skills, list):
                logger.info("Processing skills as top-level list")
                # Handle the case where skills is a list directly
                skills_text = ", ".join([str(skill) for skill in skills])
                skills_para = doc.add_paragraph(skills_text, style='MR_SkillList')
                logger.info(f"Applied MR_SkillList style to top-level skills list")
            else:
                logger.info(f"Processing skills as fallback type: {type(skills)}")
                # Fallback for any other format
                skills_para = doc.add_paragraph(str(skills), style='MR_SkillList')
                logger.info(f"Applied MR_SkillList style to fallback skills content")
        
        # ------ PROJECTS SECTION ------
        logger.info("Processing Projects section...")
        projects = load_section_json(request_id, "projects", temp_dir)
        logger.info(f"Projects data loaded: {bool(projects)}")
        logger.info(f"Projects contains 'projects' key: {isinstance(projects, dict) and 'projects' in projects}")
        
        # Handle both dictionary with 'projects' key, direct list of projects, and content key
        projects_list = []
        if projects:
            if isinstance(projects, dict) and "projects" in projects:
                projects_list = projects.get("projects", [])
            elif isinstance(projects, dict) and "content" in projects:
                # Handle content key similar to other sections
                logger.info("Found 'content' key in projects section")
                content = projects.get("content", [])
                if isinstance(content, list):
                    projects_list = content
                elif isinstance(content, dict) and "projects" in content:
                    projects_list = content.get("projects", [])
                elif isinstance(content, str):
                    # In some cases, content might be a string
                    logger.info("Projects content is a string, adding as single project")
                    # Add projects section header using helper function
                    projects_header = add_section_header(doc, "PROJECTS", docx_styles)
                    
                    # Add the string content directly as paragraph
                    projects_para = doc.add_paragraph(content)
                    _apply_paragraph_style(projects_para, "body", docx_styles)
                    
                    # Skip the rest of the projects processing
                    projects_list = []
                else:
                    logger.warning(f"Unexpected projects content format: {type(content)}")
            elif isinstance(projects, list):
                # Direct list of projects
                projects_list = projects
            else:
                logger.warning(f"Unexpected projects format: {type(projects)}")
                
            if projects_list:
                # Add section header with helper function
                projects_header = add_section_header(doc, "PROJECTS", docx_styles)
                
                # Verify projects is a list
                if not isinstance(projects_list, list):
                    logger.warning(f"Projects is not a list: {type(projects_list)}")
                    projects_list = []
                    
                # Add each project
                for project in projects_list:
                    # Verify project is a dictionary
                    if not isinstance(project, dict):
                        logger.warning(f"Project is not a dictionary: {type(project)}")
                        continue
                        
                    # Project title and dates - use the helper function for consistent formatting
                    title = project.get('title', '')
                    if not title and project.get('name'):  # Fallback to 'name' if 'title' is not available
                        title = project.get('name', '')
                    dates = project.get('dates', '')
                    if not dates and project.get('date'):  # Fallback to 'date' if 'dates' is not available
                        dates = project.get('date', '')
                    
                    if title or dates:
                        title_para = format_right_aligned_pair(
                            doc,
                            title,
                            dates,
                            "heading3",
                            "body",
                            docx_styles
                        )
                    
                    logger.info(f"Formatted project entry: '{title}'")
                    
                    # Project details - use the helper function for consistent formatting
                    for detail in project.get('details', []):
                        bullet_para = create_bullet_point(doc, detail, docx_styles)
        
        # Save DOCX to BytesIO
        logger.info("Saving DOCX to BytesIO...")
        
        # IMPORTANT FIX: Fix the spacing between sections by finding the last paragraph 
        # before each section header and setting its space_after to 0
        logger.info("Applying section spacing fix before saving document...")
        tighten_before_headers(doc)
        
        # Save the document
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Generate debug report if requested
        if debug:
            try:
                from utils.docx_debug import generate_debug_report
                import json
                
                # Generate debug report
                debug_report = generate_debug_report(doc)
                
                # Save debug report to file
                debug_path = Path(__file__).parent.parent / f'debug_{request_id}.json'
                with open(debug_path, 'w') as f:
                    json.dump(debug_report, f, indent=2)
                
                logger.info(f"Debug report saved to {debug_path}")
                
                # Also save a copy of the DOCX for inspection
                debug_docx_path = Path(__file__).parent.parent / f'debug_{request_id}.docx'
                with open(debug_docx_path, 'wb') as f:
                    f.write(output.getvalue())
                
                logger.info(f"Debug DOCX saved to {debug_docx_path}")
                
                # Reset the output buffer position
                output.seek(0)
            except Exception as e:
                logger.error(f"Error generating debug report: {e}")
        
        logger.info(f"Successfully built DOCX for request ID: {request_id}")
        return output
    
    except Exception as e:
        # Enhanced error reporting
        logger.error(f"Error building DOCX for request ID {request_id}: {e}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error details: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise 