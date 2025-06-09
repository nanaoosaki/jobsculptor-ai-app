"""
O3 Core Bullet Consistency Engine

This module implements O3's final bullet consistency fixes based on the 
comprehensive analysis and "build-then-reconcile" architecture.

Key Features:
- Document-level bullet state management
- Atomic bullet operations with validation
- Comprehensive post-generation reconciliation
- Integration with B-series edge case handling
- Production-ready bullet consistency guarantee

Author: O3 Expert Analysis + Resume Tailor Team
Status: Phase 4 Implementation - Production Ready
"""

import logging
import time
from typing import Dict, List, Optional, Tuple, Set, Any
from dataclasses import dataclass
from enum import Enum
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml

# Integration with B-series modules
from utils.unicode_bullet_sanitizer import sanitize_bullet_text
from utils.numid_collision_manager import allocate_safe_numid
from utils.xml_repair_system import analyze_docx_xml_issues
from utils.style_collision_handler import validate_style_for_bullets

logger = logging.getLogger(__name__)


class BulletState(Enum):
    """States of bullet point processing."""
    PENDING = "pending"           # Bullet created but not validated
    VALIDATED = "validated"       # Bullet validated and confirmed
    FAILED = "failed"            # Bullet failed validation
    RECONCILED = "reconciled"    # Bullet fixed during reconciliation
    STABLE = "stable"            # Bullet confirmed stable in final document


@dataclass
class BulletMetadata:
    """Comprehensive metadata for a bullet point."""
    paragraph_id: str
    text_content: str
    num_id: int
    abstract_num_id: int
    level: int
    style_name: str
    state: BulletState
    created_at: datetime
    validated_at: Optional[datetime] = None
    reconciled_at: Optional[datetime] = None
    error_count: int = 0
    last_error: Optional[str] = None


class O3BulletCoreEngine:
    """
    O3's core bullet consistency engine implementing the final architecture.
    
    This engine implements O3's key insights:
    - Build bullets without immediate verification (trust phase)
    - Validate entire document state after completion (verify phase)
    - Atomic reconciliation with guaranteed consistency
    - Comprehensive state tracking and error recovery
    """
    
    def __init__(self, document_id: str):
        """Initialize the O3 core engine for a specific document."""
        
        self.document_id = document_id
        
        # O3: Document-level state tracking
        self.bullet_registry: Dict[str, BulletMetadata] = {}
        self.numbering_allocations: Dict[int, str] = {}  # numId -> section_name
        self.document_state = "building"  # building, validating, reconciling, complete
        
        # O3: Performance and error tracking
        self.stats = {
            'bullets_created': 0,
            'bullets_validated': 0,
            'bullets_reconciled': 0,
            'bullets_failed': 0,
            'reconciliation_time_ms': 0,
            'total_processing_time_ms': 0
        }
        
        # O3: Configuration
        self.config = {
            'max_reconciliation_attempts': 3,
            'validation_timeout_ms': 5000,
            'enable_nuclear_cleanup': True,
            'enable_b_series_integration': True
        }
        
        logger.info(f"O3: Core bullet engine initialized for document {document_id}")
    
    def create_bullet_trusted(self, doc: Document, text: str, section_name: str,
                            numbering_engine: Any, docx_styles: Dict[str, Any]) -> Tuple[Any, str]:
        """
        Create bullet point using O3's "trust" approach - no immediate validation.
        
        Args:
            doc: Document object
            text: Bullet text content
            section_name: Section name for numId allocation
            numbering_engine: Numbering engine instance
            docx_styles: Style configuration
            
        Returns:
            Tuple of (paragraph, bullet_id)
        """
        start_time = time.time()
        
        try:
            # O3: B-series integration - sanitize text first
            if self.config['enable_b_series_integration']:
                text = sanitize_bullet_text(text)
            
            # O3: Generate unique bullet ID
            bullet_id = f"bullet_{self.document_id}_{len(self.bullet_registry)}"
            
            # O3: Allocate safe numId using B9
            if self.config['enable_b_series_integration']:
                num_id, abstract_num_id = allocate_safe_numid(self.document_id, section_name, "MR_BulletPoint")
            else:
                # Fallback to simple allocation
                num_id = 100
                abstract_num_id = 100
            
            # O3: Create paragraph with trust-based approach
            para = doc.add_paragraph()
            para.add_run(text.strip())
            
            # O3: Apply style first (more reliable)
            try:
                style = doc.styles['MR_BulletPoint']
                para.style = style
                logger.debug(f"O3: Applied MR_BulletPoint style to bullet '{text[:30]}...'")
            except KeyError:
                logger.warning("O3: MR_BulletPoint style not found, creating bullets may fail")
            
            # O3: Apply numbering (trust that it will work)
            try:
                numbering_engine.apply_native_bullet(para, num_id=num_id, level=0)
                logger.debug(f"O3: Applied numbering numId={num_id} to bullet '{text[:30]}...'")
            except Exception as e:
                logger.warning(f"O3: Numbering application failed for bullet '{text[:30]}...': {e}")
            
            # O3: Register bullet metadata
            metadata = BulletMetadata(
                paragraph_id=bullet_id,
                text_content=text,
                num_id=num_id,
                abstract_num_id=abstract_num_id,
                level=0,
                style_name="MR_BulletPoint",
                state=BulletState.PENDING,
                created_at=datetime.now()
            )
            
            self.bullet_registry[bullet_id] = metadata
            self.numbering_allocations[num_id] = section_name
            self.stats['bullets_created'] += 1
            
            processing_time = (time.time() - start_time) * 1000
            logger.debug(f"O3: Created trusted bullet '{text[:30]}...' in {processing_time:.1f}ms")
            
            return para, bullet_id
            
        except Exception as e:
            logger.error(f"O3: Failed to create trusted bullet '{text[:30]}...': {e}")
            self.stats['bullets_failed'] += 1
            raise
    
    def validate_document_bullets(self, doc: Document) -> Dict[str, Any]:
        """
        Validate all bullets in the document after construction is complete.
        
        Args:
            doc: Complete document to validate
            
        Returns:
            Validation results summary
        """
        start_time = time.time()
        self.document_state = "validating"
        
        logger.info(f"O3: Starting document-wide bullet validation for {len(self.bullet_registry)} bullets")
        
        validation_results = {
            'total_bullets': len(self.bullet_registry),
            'validated_bullets': 0,
            'failed_bullets': 0,
            'validation_issues': [],
            'needs_reconciliation': False
        }
        
        # O3: Iterate through all document paragraphs
        for i, para in enumerate(doc.paragraphs):
            # Check if this paragraph should be a bullet
            if para.style and para.style.name == 'MR_BulletPoint':
                bullet_id = self._find_bullet_by_paragraph(para)
                
                if bullet_id:
                    metadata = self.bullet_registry[bullet_id]
                    is_valid = self._validate_single_bullet(para, metadata)
                    
                    if is_valid:
                        metadata.state = BulletState.VALIDATED
                        metadata.validated_at = datetime.now()
                        validation_results['validated_bullets'] += 1
                        self.stats['bullets_validated'] += 1
                    else:
                        metadata.state = BulletState.FAILED
                        validation_results['failed_bullets'] += 1
                        validation_results['needs_reconciliation'] = True
                        
                        issue = {
                            'bullet_id': bullet_id,
                            'paragraph_index': i,
                            'text': para.text[:50],
                            'issue': metadata.last_error or 'Unknown validation failure'
                        }
                        validation_results['validation_issues'].append(issue)
                        
                        logger.warning(f"O3: Bullet validation failed: {issue}")
        
        # O3: Check for orphaned bullets (registered but not found in document)
        found_bullets = set()
        for para in doc.paragraphs:
            if para.style and para.style.name == 'MR_BulletPoint':
                bullet_id = self._find_bullet_by_paragraph(para)
                if bullet_id:
                    found_bullets.add(bullet_id)
        
        orphaned_bullets = set(self.bullet_registry.keys()) - found_bullets
        if orphaned_bullets:
            logger.warning(f"O3: Found {len(orphaned_bullets)} orphaned bullets")
            validation_results['validation_issues'].extend([
                {
                    'bullet_id': bid,
                    'paragraph_index': -1,
                    'text': self.bullet_registry[bid].text_content[:50],
                    'issue': 'Orphaned bullet - registered but not found in document'
                }
                for bid in orphaned_bullets
            ])
            validation_results['needs_reconciliation'] = True
        
        validation_time = (time.time() - start_time) * 1000
        logger.info(f"O3: Document validation complete in {validation_time:.1f}ms - "
                   f"{validation_results['validated_bullets']}/{validation_results['total_bullets']} valid")
        
        return validation_results
    
    def reconcile_document_bullets(self, doc: Document, numbering_engine: Any) -> Dict[str, Any]:
        """
        Perform comprehensive bullet reconciliation using O3's approach.
        
        Args:
            doc: Document to reconcile
            numbering_engine: Numbering engine for repairs
            
        Returns:
            Reconciliation results summary
        """
        start_time = time.time()
        self.document_state = "reconciling"
        
        logger.info("O3: Starting comprehensive bullet reconciliation")
        
        reconciliation_results = {
            'bullets_processed': 0,
            'bullets_repaired': 0,
            'bullets_stable': 0,
            'repair_attempts': 0,
            'success_rate': 0.0
        }
        
        # O3: Multi-pass reconciliation for maximum reliability
        for attempt in range(self.config['max_reconciliation_attempts']):
            logger.info(f"O3: Reconciliation attempt {attempt + 1}/{self.config['max_reconciliation_attempts']}")
            
            repaired_this_pass = 0
            
            # O3: Scan all paragraphs for bullet issues
            for i, para in enumerate(doc.paragraphs):
                if para.style and para.style.name == 'MR_BulletPoint':
                    reconciliation_results['bullets_processed'] += 1
                    
                    # O3: Check if bullet has proper numbering
                    has_numbering = self._check_paragraph_numbering(para)
                    
                    if not has_numbering:
                        logger.info(f"O3: Repairing bullet at paragraph {i}: '{para.text[:30]}...'")
                        
                        # O3: Attempt repair
                        repair_success = self._repair_bullet_numbering(para, numbering_engine)
                        reconciliation_results['repair_attempts'] += 1
                        
                        if repair_success:
                            repaired_this_pass += 1
                            reconciliation_results['bullets_repaired'] += 1
                            self.stats['bullets_reconciled'] += 1
                        else:
                            logger.warning(f"O3: Failed to repair bullet at paragraph {i}")
                    else:
                        # O3: Bullet is already good
                        reconciliation_results['bullets_stable'] += 1
            
            logger.info(f"O3: Reconciliation pass {attempt + 1} repaired {repaired_this_pass} bullets")
            
            # O3: If no repairs needed, we're done
            if repaired_this_pass == 0:
                logger.info("O3: No repairs needed, reconciliation complete")
                break
        
        # O3: Calculate success rate
        total_bullets = reconciliation_results['bullets_processed']
        stable_bullets = reconciliation_results['bullets_stable'] + reconciliation_results['bullets_repaired']
        
        if total_bullets > 0:
            reconciliation_results['success_rate'] = (stable_bullets / total_bullets) * 100
        
        reconciliation_time = (time.time() - start_time) * 1000
        self.stats['reconciliation_time_ms'] = reconciliation_time
        
        logger.info(f"O3: Reconciliation complete in {reconciliation_time:.1f}ms - "
                   f"{reconciliation_results['success_rate']:.1f}% success rate")
        
        self.document_state = "complete"
        return reconciliation_results
    
    def _find_bullet_by_paragraph(self, para: Any) -> Optional[str]:
        """Find bullet ID by matching paragraph text."""
        para_text = para.text.strip()
        
        for bullet_id, metadata in self.bullet_registry.items():
            if metadata.text_content.strip() == para_text:
                return bullet_id
        
        return None
    
    def _validate_single_bullet(self, para: Any, metadata: BulletMetadata) -> bool:
        """Validate a single bullet paragraph."""
        
        try:
            # O3: Check for numPr element
            if not self._check_paragraph_numbering(para):
                metadata.last_error = "Missing numPr element"
                metadata.error_count += 1
                return False
            
            # O3: Check style consistency
            if not para.style or para.style.name != metadata.style_name:
                metadata.last_error = f"Style mismatch: expected {metadata.style_name}, got {para.style.name if para.style else 'None'}"
                metadata.error_count += 1
                return False
            
            # O3: Check text content
            if para.text.strip() != metadata.text_content.strip():
                metadata.last_error = "Text content mismatch"
                metadata.error_count += 1
                return False
            
            return True
            
        except Exception as e:
            metadata.last_error = f"Validation exception: {e}"
            metadata.error_count += 1
            return False
    
    def _check_paragraph_numbering(self, para: Any) -> bool:
        """Check if paragraph has proper bullet numbering."""
        
        try:
            # O3: Look for numPr element in paragraph XML
            pPr = para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
            if pPr is None:
                return False
            
            numPr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
            if numPr is None:
                return False
            
            # O3: Check for required numId and ilvl elements
            numId = numPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
            ilvl = numPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
            
            if numId is None or ilvl is None:
                return False
            
            # O3: Validate numId value
            num_id_val = numId.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            if not num_id_val or not num_id_val.isdigit():
                return False
            
            return True
            
        except Exception as e:
            logger.debug(f"O3: Error checking paragraph numbering: {e}")
            return False
    
    def _repair_bullet_numbering(self, para: Any, numbering_engine: Any) -> bool:
        """Repair bullet numbering for a paragraph."""
        
        try:
            # O3: Find appropriate numId for this bullet
            num_id = 100  # Default fallback
            
            # Try to find the original numId from our registry
            bullet_id = self._find_bullet_by_paragraph(para)
            if bullet_id and bullet_id in self.bullet_registry:
                num_id = self.bullet_registry[bullet_id].num_id
            
            # O3: Apply numbering repair
            numbering_engine.apply_native_bullet(para, num_id=num_id, level=0)
            
            # O3: Verify repair worked
            time.sleep(0.01)  # Brief pause for XML state to settle
            return self._check_paragraph_numbering(para)
            
        except Exception as e:
            logger.error(f"O3: Bullet repair failed: {e}")
            return False
    
    def get_engine_summary(self) -> Dict[str, Any]:
        """Get comprehensive engine status summary."""
        
        return {
            'document_id': self.document_id,
            'document_state': self.document_state,
            'bullet_count': len(self.bullet_registry),
            'statistics': self.stats.copy(),
            'configuration': self.config.copy(),
            'success_metrics': {
                'creation_success_rate': ((self.stats['bullets_created'] - self.stats['bullets_failed']) / 
                                        max(self.stats['bullets_created'], 1)) * 100,
                'reconciliation_success_rate': (self.stats['bullets_reconciled'] / 
                                               max(self.stats['bullets_failed'], 1)) * 100
            }
        }


# Global engine registry for document-specific engines
_engine_registry: Dict[str, O3BulletCoreEngine] = {}


def get_o3_engine(document_id: str) -> O3BulletCoreEngine:
    """Get or create O3 engine for a document."""
    
    if document_id not in _engine_registry:
        _engine_registry[document_id] = O3BulletCoreEngine(document_id)
    
    return _engine_registry[document_id]


def cleanup_o3_engine(document_id: str):
    """Clean up O3 engine for a document."""
    
    if document_id in _engine_registry:
        del _engine_registry[document_id]
        logger.info(f"O3: Cleaned up engine for document {document_id}")


def get_all_engines_summary() -> Dict[str, Any]:
    """Get summary of all active O3 engines."""
    
    return {
        'active_engines': len(_engine_registry),
        'engines': {
            doc_id: engine.get_engine_summary()
            for doc_id, engine in _engine_registry.items()
        }
    } 