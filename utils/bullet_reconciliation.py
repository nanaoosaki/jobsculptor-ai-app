"""
Bullet Reconciliation Engine

This module implements O3's "Build-Then-Reconcile" architecture for achieving
100% bullet consistency in DOCX documents. It replaces the unreliable 
"verify-in-loop" approach with a single, deterministic reconciliation pass
after document construction is complete.

Key Features:
- Full document tree scanning (includes tables, headers, footers)
- Multi-level list support with level preservation
- Performance monitoring and guard-rails
- Comprehensive error handling for malformed XML
- Memory monitoring for large documents
- Request ID propagation for logging

Author: Resume Tailor Team + O3 Expert Review
Status: Production-Ready Implementation
"""

import logging
import time
import tracemalloc
from typing import Dict, List, Tuple, Optional, Any
from pathlib import Path

try:
    from lxml import etree
except ImportError:
    import xml.etree.ElementTree as etree

from docx import Document
from docx.oxml import parse_xml
from flask import current_app

# A12: XML Namespace Helper - Extract WordprocessingML namespace constant
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

logger = logging.getLogger(__name__)


class BulletReconciliationEngine:
    """
    Handles post-build verification and repair of bullet numbering.
    
    This engine implements O3's "Build-Then-Reconcile" architecture:
    1. Scan the entire document tree for bullet paragraphs
    2. Verify each paragraph has proper <w:numPr> numbering
    3. Repair any missing numbering while preserving original levels
    4. Monitor performance and memory usage
    5. Provide comprehensive logging with request context
    """
    
    def __init__(self, request_id: Optional[str] = None):
        """
        Initialize the reconciliation engine.
        
        Args:
            request_id: Optional request ID for logging context (B8)
        """
        self.request_id = request_id or "unknown"
        self.logger = logger
        
        # B8: Request-ID Propagation for multi-tenant log search
        if hasattr(logger, 'bind'):
            self.logger = logger.bind(request_id=self.request_id)
        
        # Performance monitoring state
        self.start_time = None
        self.memory_start = None
        
    def reconcile_bullet_styles(self, doc: Document, numbering_engine, num_id: int) -> Dict[str, Any]:
        """
        Single pass to ensure all bullets have native numbering.
        
        This is the core reconciliation method that:
        1. Scans the full document tree (A3: includes tables, headers, footers)
        2. Finds all paragraphs with MR_BulletPoint style
        3. Verifies each has proper <w:numPr> numbering
        4. Repairs missing numbering while preserving levels (A2)
        5. Monitors performance and memory usage (A6, B5)
        
        Args:
            doc: The Document object to reconcile
            numbering_engine: NumberingEngine instance for applying bullets
            num_id: The numbering ID to use for repairs
            
        Returns:
            Dict with reconciliation statistics and performance metrics
        """
        self._start_performance_monitoring()
        
        try:
            # A9: Logging Noise Control - detailed logs only in debug mode
            if current_app and current_app.debug:
                self.logger.debug(f"🛡️ Starting bullet reconciliation for request {self.request_id}")
            else:
                self.logger.info(f"🛡️ Starting bullet reconciliation for request {self.request_id}")
            
            # B1: Handle user-supplied style name collision
            self._handle_style_collision(doc)
            
            # A3 + B10: Scan full document tree including tables, headers, footers, text-boxes
            bullet_paragraphs = self._scan_bullet_paragraphs(doc)
            
            repaired_count = 0
            total_bullets = len(bullet_paragraphs)
            errors = []
            
            for para_info in bullet_paragraphs:
                para = para_info['paragraph']
                location = para_info['location']
                original_level = para_info['level']
                
                try:
                    # B6: Pre-existing broken XML handling
                    if not self._verify_numbering(para):
                        if current_app and current_app.debug:
                            self.logger.debug(f"🔧 Repairing bullet in {location}: '{para.text[:40]}...'")
                        
                        # A2: Preserve original bullet level
                        level_to_use = original_level if original_level is not None else 0
                        
                        # Apply native bullet formatting
                        numbering_engine.apply_native_bullet(para, num_id=num_id, level=level_to_use)
                        repaired_count += 1
                        
                        if current_app and current_app.debug:
                            self.logger.debug(f"  ✅ Repaired with level {level_to_use}")
                
                except Exception as e:
                    error_msg = f"Failed to repair bullet in {location}: {type(e).__name__}: {e}"
                    self.logger.error(error_msg)
                    errors.append(error_msg)
            
            # Performance and memory monitoring
            duration_ms, memory_diff_mb = self._end_performance_monitoring()
            
            # A6: Performance Guard-Rail - warn if reconciliation is slow
            if duration_ms > 200:
                self.logger.warning(f"🐌 Reconcile slow: {duration_ms:.1f}ms on {total_bullets} paragraphs")
            
            # B5: Memory Guard Rail for large documents
            if memory_diff_mb > 30 or total_bullets > 5000:
                self.logger.warning(f"📊 Memory impact: {memory_diff_mb:.1f}MB, {total_bullets} paragraphs")
            
            # A9: Production logging - summary format
            if current_app and current_app.debug:
                self.logger.debug(f"🛡️ Reconciliation complete. Repaired: {repaired_count}/{total_bullets}")
            else:
                self.logger.info(f"🛡️ Reconciled {repaired_count}/{total_bullets} bullets ({duration_ms:.1f}ms)")
            
            return {
                'total_bullets': total_bullets,
                'repaired_count': repaired_count,
                'duration_ms': duration_ms,
                'memory_diff_mb': memory_diff_mb,
                'errors': errors,
                'request_id': self.request_id
            }
            
        except Exception as e:
            self.logger.error(f"🚨 Reconciliation failed: {type(e).__name__}: {e}")
            self.logger.error(f"🚨 Traceback: {traceback.format_exc()}")
            raise
    
    def _handle_style_collision(self, doc: Document) -> None:
        """
        B1: Handle user-supplied style name collision.
        
        If uploaded résumé already contains a style called 'MR_BulletPoint',
        rename it to avoid our style definition being overridden.
        """
        try:
            existing_style = None
            for style in doc.styles:
                if style.name == 'MR_BulletPoint':
                    # Check if this is our style (has our custom properties) or user's
                    if not hasattr(style, '_custom_manus_style') and style.base_style is None:
                        existing_style = style
                        break
            
            if existing_style:
                # Rename user's style to avoid collision
                new_name = 'MR_BulletPoint__orig'
                counter = 1
                while any(s.name == new_name for s in doc.styles):
                    new_name = f'MR_BulletPoint__orig_{counter}'
                    counter += 1
                
                existing_style.name = new_name
                self.logger.info(f"🔄 Renamed existing MR_BulletPoint style to {new_name}")
                
        except Exception as e:
            self.logger.warning(f"⚠️ Could not check for style collision: {e}")
    
    def _scan_bullet_paragraphs(self, doc: Document) -> List[Dict[str, Any]]:
        """
        A3 + B10: Find all paragraphs with MR_BulletPoint style in the full document.
        
        This scans:
        - Main document body (including tables)
        - Headers and footers (B10)
        - Text boxes and drawing canvases (B10)
        
        Returns:
            List of dicts with paragraph info including location and original level
        """
        bullet_paragraphs = []
        
        try:
            # A3: Scan main body using full document tree (includes tables)
            body_paras = self._scan_body_paragraphs(doc)
            bullet_paragraphs.extend(body_paras)
            
            # B10: Scan headers and footers
            header_footer_paras = self._scan_headers_footers(doc)
            bullet_paragraphs.extend(header_footer_paras)
            
            # B10: Scan text boxes and drawing canvases
            drawing_paras = self._scan_drawing_elements(doc)
            bullet_paragraphs.extend(drawing_paras)
            
        except Exception as e:
            self.logger.error(f"Error scanning bullet paragraphs: {e}")
            # Fallback to basic doc.paragraphs scan
            for i, para in enumerate(doc.paragraphs):
                if para.style and para.style.name == 'MR_BulletPoint':
                    bullet_paragraphs.append({
                        'paragraph': para,
                        'location': f'main_body_fallback[{i}]',
                        'level': self._get_original_level(para)
                    })
        
        return bullet_paragraphs
    
    def _scan_body_paragraphs(self, doc: Document) -> List[Dict[str, Any]]:
        """Scan main document body using XPath for full tree traversal."""
        paragraphs = []
        
        try:
            # A3: Use XPath to get all paragraphs in document body (includes tables)
            body_element = doc._body._element
            
            # Fix: Use proper namespace prefix registration for XPath
            nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            para_elements = body_element.xpath('.//w:p', namespaces=nsmap)
            
            for i, para_element in enumerate(para_elements):
                # Find corresponding paragraph object
                para_obj = self._element_to_paragraph(doc, para_element)
                if para_obj and para_obj.style and para_obj.style.name == 'MR_BulletPoint':
                    paragraphs.append({
                        'paragraph': para_obj,
                        'location': f'main_body[{i}]',
                        'level': self._get_original_level(para_obj)
                    })
                    
        except Exception as e:
            self.logger.warning(f"XPath body scan failed, using fallback: {e}")
            # Fallback to standard iteration
            for i, para in enumerate(doc.paragraphs):
                if para.style and para.style.name == 'MR_BulletPoint':
                    paragraphs.append({
                        'paragraph': para,
                        'location': f'main_body_fallback[{i}]',
                        'level': self._get_original_level(para)
                    })
        
        return paragraphs
    
    def _scan_headers_footers(self, doc: Document) -> List[Dict[str, Any]]:
        """B10: Scan headers and footers for bullet paragraphs."""
        paragraphs = []
        
        try:
            # Get header parts
            header_parts = getattr(doc.part, '_header_parts', [])
            for i, header_part in enumerate(header_parts):
                para_elements = header_part.element.xpath(f'.//{W}p')
                for j, para_element in enumerate(para_elements):
                    para_obj = self._element_to_paragraph_generic(para_element)
                    if para_obj and self._has_bullet_style(para_element):
                        paragraphs.append({
                            'paragraph': para_obj,
                            'location': f'header[{i}][{j}]',
                            'level': self._get_original_level_from_element(para_element)
                        })
            
            # Get footer parts
            footer_parts = getattr(doc.part, '_footer_parts', [])
            for i, footer_part in enumerate(footer_parts):
                para_elements = footer_part.element.xpath(f'.//{W}p')
                for j, para_element in enumerate(para_elements):
                    para_obj = self._element_to_paragraph_generic(para_element)
                    if para_obj and self._has_bullet_style(para_element):
                        paragraphs.append({
                            'paragraph': para_obj,
                            'location': f'footer[{i}][{j}]',
                            'level': self._get_original_level_from_element(para_element)
                        })
                        
        except Exception as e:
            self.logger.warning(f"Header/footer scan failed: {e}")
        
        return paragraphs
    
    def _scan_drawing_elements(self, doc: Document) -> List[Dict[str, Any]]:
        """B10: Scan text boxes and drawing canvases for bullet paragraphs."""
        paragraphs = []
        
        try:
            # Scan for drawing elements that might contain text boxes
            drawing_elements = doc._body._element.xpath('.//w:drawing')
            for i, drawing in enumerate(drawing_elements):
                para_elements = drawing.xpath(f'.//{W}p')
                for j, para_element in enumerate(para_elements):
                    para_obj = self._element_to_paragraph_generic(para_element)
                    if para_obj and self._has_bullet_style(para_element):
                        paragraphs.append({
                            'paragraph': para_obj,
                            'location': f'drawing[{i}][{j}]',
                            'level': self._get_original_level_from_element(para_element)
                        })
                        
        except Exception as e:
            self.logger.warning(f"Drawing elements scan failed: {e}")
        
        return paragraphs
    
    def _element_to_paragraph(self, doc: Document, para_element) -> Optional[Any]:
        """Convert XML element to paragraph object."""
        try:
            # Find matching paragraph in doc.paragraphs
            for para in doc.paragraphs:
                if para._element == para_element:
                    return para
        except Exception:
            pass
        return None
    
    def _element_to_paragraph_generic(self, para_element) -> Optional[Any]:
        """Create a generic paragraph object from XML element."""
        try:
            # For headers/footers/drawings, create a minimal paragraph wrapper
            from docx.text.paragraph import Paragraph
            return Paragraph(para_element, None)
        except Exception:
            pass
        return None
    
    def _has_bullet_style(self, para_element) -> bool:
        """Check if paragraph element has MR_BulletPoint style."""
        try:
            # Look for style reference
            pPr = para_element.find(f'{W}pPr')
            if pPr is not None:
                pStyle = pPr.find(f'{W}pStyle')
                if pStyle is not None:
                    style_val = pStyle.get(f'{W}val')
                    return style_val == 'MR_BulletPoint'
        except Exception:
            pass
        return False
    
    def _verify_numbering(self, para) -> bool:
        """
        Check if paragraph has <w:numPr> element.
        
        B6: Wrap in try/except to handle pre-existing broken XML.
        """
        try:
            pPr = para._element.find(f'{W}pPr')
            if pPr is not None:
                numPr = pPr.find(f'{W}numPr')
                return numPr is not None
        except (AttributeError, etree.Error) as e:
            # B6: Treat XML errors as "needs repair"
            self.logger.warning(f"XML error during numbering verification: {e}")
            return False
        
        return False
    
    def _get_original_level(self, para) -> Optional[int]:
        """A2: Get original bullet level from paragraph."""
        try:
            return self._get_original_level_from_element(para._element)
        except Exception:
            return None
    
    def _get_original_level_from_element(self, para_element) -> Optional[int]:
        """A2: Get original bullet level from XML element."""
        try:
            pPr = para_element.find(f'{W}pPr')
            if pPr is not None:
                numPr = pPr.find(f'{W}numPr')
                if numPr is not None:
                    ilvl = numPr.find(f'{W}ilvl')
                    if ilvl is not None:
                        return int(ilvl.get(f'{W}val', '0'))
        except (AttributeError, ValueError, etree.Error):
            pass
        return None
    
    def _start_performance_monitoring(self) -> None:
        """A6 + B5: Start performance and memory monitoring."""
        self.start_time = time.time()
        
        # B5: Start memory monitoring for large document guard-rail
        try:
            tracemalloc.start()
            self.memory_start = tracemalloc.get_traced_memory()[0]
        except Exception:
            self.memory_start = None
    
    def _end_performance_monitoring(self) -> Tuple[float, float]:
        """A6 + B5: End monitoring and return duration and memory diff."""
        # Calculate duration
        duration_ms = 0.0
        if self.start_time:
            duration_ms = (time.time() - self.start_time) * 1000
        
        # Calculate memory difference
        memory_diff_mb = 0.0
        if self.memory_start is not None:
            try:
                current_memory = tracemalloc.get_traced_memory()[0]
                memory_diff_mb = (current_memory - self.memory_start) / (1024 * 1024)
                tracemalloc.stop()
            except Exception:
                pass
        
        return duration_ms, memory_diff_mb 