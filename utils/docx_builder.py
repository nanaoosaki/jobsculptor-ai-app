"""
DOCX Builder for Resume Tailor Application

Generates Microsoft Word (.docx) files with consistent styling based on design tokens.
"""

import os
import logging
import json
import traceback
import re
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional, Any, Union

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

# Enhanced architecture imports
try:
    from style_manager import StyleManager
    USE_STYLE_MANAGER = True
except ImportError:
    USE_STYLE_MANAGER = False

try:
    from style_engine import StyleEngine
    USE_STYLE_ENGINE = True
except ImportError:
    USE_STYLE_ENGINE = False

# Import our style registry and section builder
try:
    from word_styles.registry import USE_STYLE_REGISTRY, get_or_create_style, apply_direct_paragraph_formatting
    from word_styles.section_builder import add_section_header as registry_add_section_header
    from word_styles.section_builder import add_content_paragraph, add_bullet_point, remove_empty_paragraphs
    from word_styles.section_builder import add_role_box
except ImportError:
    USE_STYLE_REGISTRY = False
    logger.warning("⚠️ Style registry not available - using fallback mode")

# Import native numbering engine
try:
    from word_styles.numbering_engine import NumberingEngine
    USE_NATIVE_NUMBERING = True
except ImportError:
    USE_NATIVE_NUMBERING = False
    logger.warning("⚠️ NumberingEngine not available - native bullets disabled")

# Import bullet reconciliation engine
try:
    from utils.bullet_reconciliation import BulletReconciliationEngine
    USE_BULLET_RECONCILIATION = True
except ImportError:
    USE_BULLET_RECONCILIATION = False
    logger.warning("⚠️ BulletReconciliationEngine not available - reconciliation disabled")

# Simple stub for rendering tracer to avoid import errors
def trace(name):
    """Simple stub decorator for tracing function calls"""
    def decorator(func):
        return func
    return decorator

# Import our new universal renderers
try:
    from rendering.components import SectionHeader, RoleBox
    USE_UNIVERSAL_RENDERERS = True
except ImportError:
    USE_UNIVERSAL_RENDERERS = False

# FEATURE FLAGS FOR GRADUAL ROLLOUT
DOCX_USE_NATIVE_BULLETS = os.getenv('DOCX_USE_NATIVE_BULLETS', 'true').lower() == 'true'

# Only enable native bullets if both the flag is set AND the engine is available
NATIVE_BULLETS_ENABLED = DOCX_USE_NATIVE_BULLETS and USE_NATIVE_NUMBERING

logger = logging.getLogger(__name__)

logger.info(f"🎯 DOCX Feature Flags: NATIVE_BULLETS={DOCX_USE_NATIVE_BULLETS}, ENGINE_AVAILABLE={USE_NATIVE_NUMBERING}, ENABLED={NATIVE_BULLETS_ENABLED}")

def load_section_json(request_id: str, section_name: str, temp_dir: str) -> Dict[str, Any]:
    """Load a section's JSON data from the temporary session directory."""
    try:
        file_path = os.path.join(temp_dir, f"{request_id}_{section_name}.json")
        logger.info(f"Looking for section file: {file_path}")
        
        if not os.path.exists(file_path):
            logger.warning(f"File not found: {file_path}")
            return {}
            
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            logger.info(f"Successfully loaded '{section_name}' section: {type(data)}")
            if isinstance(data, dict):
                logger.info(f"Dict keys for {section_name}: {list(data.keys())}")
            elif isinstance(data, list):
                logger.info(f"List length for {section_name}: {len(data)}")
            return data
    except FileNotFoundError:
        logger.warning(f"Section file not found: {file_path}")
        return {}
    except json.JSONDecodeError as e:
        logger.error(f"Error decoding JSON from {file_path}: {e}")
        # Try to get raw content for debugging
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                logger.error(f"Raw content from {file_path} (first 100 chars): {content[:100]}...")
        except Exception:
            pass
        return {}
    except Exception as e:
        logger.error(f"Error loading section JSON: {e}")
        return {}

def _apply_paragraph_style(doc: Document, p: Paragraph, style_name: str, docx_styles: Dict[str, Any]):
    """Enhanced style application for paragraphs in DOCX."""
    if not p.runs:
        logger.warning(f"🔧 SKIPPING style application to empty paragraph (no runs)")
        return  # Skip empty paragraphs
    
    logger.info(f"🔧 ATTEMPTING to apply style '{style_name}' to paragraph with text: '{p.text[:50]}...'")
    
    try:
        # Use the passed-in doc object's styles collection
        available_styles = [s.name for s in doc.styles] 
        logger.info(f"🔍 Available styles in document: {len(available_styles)} total")
        logger.info(f"🔍 Checking if '{style_name}' is in available styles...")
        
        if style_name in available_styles:
            logger.info(f"✅ Style '{style_name}' found in doc.styles, attempting assignment...")
            
            # Store the original style for comparison
            original_style = p.style.name if p.style else "None"
            logger.info(f"📝 Original paragraph style: '{original_style}'")
            
            # Get a reference to the style object to verify it exists
            try:
                style_obj = doc.styles[style_name]
                logger.info(f"✅ Successfully retrieved style object for '{style_name}': {style_obj}")
                logger.info(f"📋 Style object properties: type={style_obj.type}, base_style={style_obj.base_style.name if style_obj.base_style else 'None'}")
                
                # Check spacing properties of the style
                if hasattr(style_obj, 'paragraph_format'):
                    pf = style_obj.paragraph_format
                    logger.info(f"📋 Style spacing: space_before={pf.space_before}, space_after={pf.space_after}")
                
            except Exception as style_obj_error:
                logger.error(f"❌ Failed to retrieve style object for '{style_name}': {style_obj_error}")
                return
            
            # Attempt the style assignment
            try:
                p.style = style_name
                logger.info(f"🎯 Style assignment call completed for '{style_name}'")
            except Exception as assign_error:
                logger.error(f"❌ Style assignment threw exception: {assign_error}")
                logger.error(f"❌ Assignment exception type: {type(assign_error).__name__}")
                return
            
            # Verify the assignment worked
            new_style = p.style.name if p.style else "None"
            logger.info(f"📝 Style after assignment: '{new_style}'")
            
            if new_style == style_name:
                logger.info(f"✅ Successfully set paragraph style to: {style_name} using doc.styles")
            else:
                logger.error(f"❌ Style assignment failed! Expected '{style_name}', got '{new_style}'")
                logger.error(f"❌ This indicates the assignment was silently ignored or overridden")
        else:
            logger.error(f"❌ Style '{style_name}' NOT found in doc.styles.")
            logger.error(f"❌ Available styles: {available_styles[:10]}... (showing first 10)")
            # Attempt to use by name directly, python-docx might find it if it's a built-in like 'Normal'
            try:
                p.style = style_name
                logger.info(f"✅ Fallback: Successfully set style to '{style_name}' by direct assignment (might be built-in).")
            except Exception as e_fb:
                logger.error(f"❌ Fallback direct assignment for style '{style_name}' also failed: {e_fb}")
    except Exception as e:
        logger.error(f"❌ Failed to set paragraph style to {style_name}: {e}")
        logger.error(f"❌ Exception type: {type(e).__name__}")
        # import traceback # Already imported at top
        logger.error(f"❌ Full traceback: {traceback.format_exc()}")
    
    # Get style configuration from the new specification
    styles_section = docx_styles.get("styles", {})
    style_config = styles_section.get(style_name, {})
    if not style_config:
        logger.warning(f"⚠️ Style config not found for: {style_name}")
        logger.warning(f"⚠️ Available style configs: {list(styles_section.keys())}")
        return
    
    logger.info(f"🔧 Applying direct formatting from style config for '{style_name}'...")
    logger.info(f"📋 Style config keys: {list(style_config.keys())}")
    
    # Apply font properties to all runs in the paragraph for consistency
    for run in p.runs:
        font = run.font
        if "fontSizePt" in style_config:
            font.size = Pt(style_config["fontSizePt"])
            logger.info(f"🔧 Applied font size: {style_config['fontSizePt']}pt")
        if "fontFamily" in style_config:
            font.name = style_config["fontFamily"]
            logger.info(f"🔧 Applied font family: {style_config['fontFamily']}")
        if "color" in style_config:
            r, g, b = style_config["color"]
            font.color.rgb = RGBColor(r, g, b)
            logger.info(f"🔧 Applied color: RGB({r}, {g}, {b})")
        if style_config.get("bold", False):
            font.bold = True
            logger.info(f"🔧 Applied bold formatting")
        if style_config.get("italic", False):
            font.italic = True
            logger.info(f"🔧 Applied italic formatting")
        if style_config.get("allCaps", False):
            font.all_caps = True
            logger.info(f"🔧 Applied all caps formatting")
    
    # Apply paragraph formatting (REMOVED spaceAfterPt and spaceBeforePt to prevent style override)
    # NOTE: Per O3's advice, removing direct spacing formatting to let the style handle it
    logger.info(f"🔧 Applying paragraph formatting...")
    
    # ✅ o3's FIX 1: Remove conflicting direct indentation for bullet points
    # Only apply direct indentation for NON-BULLET styles to prevent L-0/L-1 conflicts
    if style_name != "MR_BulletPoint" and "indentCm" in style_config:
        p.paragraph_format.left_indent = Cm(style_config["indentCm"])
        logger.info(f"🔧 Applied left indent: {style_config['indentCm']}cm")
    elif style_name == "MR_BulletPoint":
        logger.info(f"🚫 Skipped direct left indent for {style_name} - letting XML numbering (L-1) control it")
        
    if style_name != "MR_BulletPoint" and "hangingIndentCm" in style_config:
        p.paragraph_format.first_line_indent = Cm(-style_config["hangingIndentCm"])
        logger.info(f"🔧 Applied hanging indent: {style_config['hangingIndentCm']}cm")
    elif style_name == "MR_BulletPoint":
        logger.info(f"🚫 Skipped direct hanging indent for {style_name} - letting XML numbering (L-1) control it")
        
    if "lineHeight" in style_config:
        p.paragraph_format.line_spacing = style_config["lineHeight"]
        logger.info(f"🔧 Applied line spacing: {style_config['lineHeight']}")
    
    # Apply alignment
    alignment = style_config.get("alignment", "left")
    if alignment == "center":
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logger.info(f"🔧 Applied center alignment")
    elif alignment == "right":
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        logger.info(f"🔧 Applied right alignment")
    else:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logger.info(f"🔧 Applied left alignment")
    
    # Apply right tab for company lines if specified
    if "rightTabCm" in style_config:
        from docx.enum.table import WD_TABLE_ALIGNMENT
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(style_config["rightTabCm"]))
        logger.info(f"🔧 Applied right tab stop at: {style_config['rightTabCm']}cm")
    
    # Apply background shading for role boxes and section headers
    if "backgroundColor" in style_config:
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        r, g, b = style_config["backgroundColor"]
        hex_color = f"{r:02x}{g:02x}{b:02x}"
        
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
        p._element.get_or_add_pPr().append(parse_xml(shading_xml))
        logger.info(f"🔧 Applied background color: #{hex_color}")
    
    # Apply borders for section headers and role boxes
    if "borderColor" in style_config and "borderWidthPt" in style_config:
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        r, g, b = style_config["borderColor"]
        border_hex = f"{r:02x}{g:02x}{b:02x}"
        border_width_8th_pt = int(style_config["borderWidthPt"] * 8)  # Convert to 1/8th points
        
        # For section headers, apply full border; for role boxes, apply as needed
        if style_name == "MR_SectionHeader":
            # Section headers get a border on all sides
            border_xml = f'''
                <w:pBdr {nsdecls("w")}>
                    <w:top w:val="single" w:sz="{border_width_8th_pt}" w:space="0" w:color="{border_hex}"/>
                    <w:left w:val="single" w:sz="{border_width_8th_pt}" w:space="0" w:color="{border_hex}"/>
                    <w:bottom w:val="single" w:sz="{border_width_8th_pt}" w:space="0" w:color="{border_hex}"/>
                    <w:right w:val="single" w:sz="{border_width_8th_pt}" w:space="0" w:color="{border_hex}"/>
                </w:pBdr>
            '''
        else:
            # Role boxes get a border on all sides too
            border_xml = f'''
                <w:pBdr {nsdecls("w")}>
                    <w:top w:val="single" w:sz="{border_width_8th_pt}" w:space="0" w:color="{border_hex}"/>
                    <w:left w:val="single" w:sz="{border_width_8th_pt}" w:space="0" w:color="{border_hex}"/>
                    <w:bottom w:val="single" w:sz="{border_width_8th_pt}" w:space="0" w:color="{border_hex}"/>
                    <w:right w:val="single" w:sz="{border_width_8th_pt}" w:space="0" w:color="{border_hex}"/>
                </w:pBdr>
            '''
        
        p._element.get_or_add_pPr().append(parse_xml(border_xml))
        logger.info(f"🔧 Applied border: {style_config['borderWidthPt']}pt, color #{border_hex}")

    # Check final paragraph spacing to understand the inconsistency
    final_space_before = p.paragraph_format.space_before
    final_space_after = p.paragraph_format.space_after
    logger.info(f"📋 Final paragraph spacing: before={final_space_before}, after={final_space_after}")

    # FINAL DIAGNOSTIC: Check the actual final style of the paragraph
    final_style_name = p.style.name if p.style else "None"
    logger.info(f"🔍 FINAL DIAGNOSTIC: Paragraph style after all processing = '{final_style_name}' (Expected: '{style_name}')")
    if final_style_name != style_name:
        logger.warning(f"⚠️ FINAL DIAGNOSTIC FAILED: Paragraph using '{final_style_name}' instead of '{style_name}'!")
        
        # Additional diagnostics for style assignment failure
        if final_style_name == "Normal":
            logger.warning(f"⚠️ Paragraph reverted to 'Normal' style - this suggests style assignment was overridden")
            logger.warning(f"⚠️ This is the root cause of the spacing inconsistency issue")

def _create_document_styles(doc, docx_styles):
    """Create custom styles in the document for consistent formatting."""
    # Create all custom styles using the StyleEngine
    try:
        # Create all custom styles at once
        custom_styles = StyleEngine.create_docx_custom_styles(doc)
        logger.info(f"Successfully created custom DOCX styles: {list(custom_styles.keys())}")
        return custom_styles
    except Exception as e:
        logger.warning(f"Error creating custom styles: {e}")
        logger.warning("Falling back to legacy style creation")
        
        # Fallback to original implementation
        styles = {}
        
        # Create section header style
        try:
            section_style = StyleEngine.create_docx_section_header_style(doc)
            styles["SectionHeader"] = section_style
            logger.info("Successfully created custom SectionHeader style with fallback method")
        except Exception as e:
            logger.warning(f"Error creating SectionHeader style with fallback method: {e}")
        
        # Create bullet style
        try:
            bullet_style = StyleEngine.create_docx_bullet_style(doc)
            styles["CustomBullet"] = bullet_style
            logger.info("Successfully created custom CustomBullet style with fallback method")
        except Exception as e:
            logger.warning(f"Error creating CustomBullet style with fallback method: {e}")
            
        return styles

def format_right_aligned_pair(doc: Document, left_text: str, right_text: str, left_style: str, right_style: str, docx_styles: Dict[str, Any]):
    """Creates a paragraph with left-aligned and right-aligned text using tab stops."""
    para = doc.add_paragraph()
    
    logger.info(f"➡️ COMPANY/INSTITUTION ENTRY: '{left_text}' | '{right_text}' using style '{left_style}'")
    
    # CRITICAL FIX: Add text content FIRST, then apply style
    # Add left-aligned text
    if left_text:
        left_run = para.add_run(left_text)
        # Apply styling for the left text (always bold)
        left_run.bold = True
    
    # NOW apply the paragraph style (paragraph has runs, so it won't be skipped)
    _apply_paragraph_style(doc, para, left_style, docx_styles) 
    
    # SPACING: Company/institution spacing is now handled by design tokens
    # via the MR_Company style definition (paragraph-spacing-company-before)
    
    # Verify the final style assignment specifically for company entries
    final_style_name = para.style.name if para.style else "None"
    logger.info(f"🔍 COMPANY/INSTITUTION STYLE CHECK: Final style = '{final_style_name}' (Expected: '{left_style}')")
    if final_style_name != left_style:
        logger.error(f"❌ COMPANY/INSTITUTION STYLE FAILED: '{left_text}' using '{final_style_name}' instead of '{left_style}'!")
    else:
        logger.info(f"✅ COMPANY/INSTITUTION STYLE SUCCESS: '{left_text}' correctly using '{final_style_name}'")
    
    # Get page margins from design tokens directly
    from style_engine import StyleEngine
    design_tokens = StyleEngine.load_tokens()
    
    # Get page margins in cm from design tokens
    global_left_margin_cm = float(design_tokens.get("docx-page-margin-left-cm", "1.5"))
    global_right_margin_cm = float(design_tokens.get("docx-page-margin-right-cm", "1.5"))
    
    # Get section properties for page width
    section = doc.sections[0]
    page_width_emu = section.page_width if section.page_width is not None else 0
    
    # Convert page width from EMU to cm
    if page_width_emu > 0:
        page_width_cm = page_width_emu / 360000.0  # 1 cm = 360000 EMUs
    else:
        # Fallback to A4 width if page_width not available
        page_width_cm = 21.0  # Standard A4 width
    
    # Calculate content width in cm using design token margins
    content_width_cm = page_width_cm - global_left_margin_cm - global_right_margin_cm
    
    # Adjust for any content indentation from the paragraph style
    # The MR_Company style has 0 indentation, so tab position should be at content width
    tab_position_cm = content_width_cm
    
    # Set tab position at the right edge of content area
    if content_width_cm <= 0:  # Fallback if dimensions are invalid
        logger.warning(f"Calculated non-positive content_width_cm ({content_width_cm}). Defaulting tab stop to 16cm.")
        tab_position_cm = 16.0  # A reasonable default
    
    logger.info(f"Right-align calculation: L={global_left_margin_cm}cm, R={global_right_margin_cm}cm, PW={page_width_cm:.2f}cm, Content={content_width_cm:.2f}cm, TabPos={tab_position_cm:.2f}cm")
    
    # Remove any existing tab stops to prevent conflicts
    para.paragraph_format.tab_stops.clear_all()
    
    # Add the new tab stop at the calculated position
    if tab_position_cm > 0:
        para.paragraph_format.tab_stops.add_tab_stop(Cm(tab_position_cm), WD_TAB_ALIGNMENT.RIGHT)
    else:
        logger.warning(f"Calculated tab position {tab_position_cm:.2f}cm is not positive. Skipping tab stop addition.")
    
    # Add right-aligned text with tab
    if right_text:
        para.add_run('\t')  # Add tab
        right_run = para.add_run(right_text)
    
    # Spacing and indentation are now handled by the 'MR_Company' style plus direct formatting
    return para

def add_bullet_point_native(doc: Document, text: str, numbering_engine: NumberingEngine = None, 
                           num_id: int = None, level: int = 0, docx_styles: Dict[str, Any] = None) -> Paragraph:
    """
    Create native Word bullet point using built-in numbering system.
    
    This is the enhanced bullet creation that produces professional Word behavior
    instead of manual spacing that can be overridden.
    
    Args:
        doc: Document to add bullet to
        text: Bullet text content
        numbering_engine: Pre-configured NumberingEngine (optional)
        num_id: Override numbering ID (if None, uses safe allocation)
        level: Indentation level (default: 0)
        docx_styles: Style definitions for design token system
        
    Returns:
        The created paragraph with native bullets
    """
    logger.debug(f"Creating native bullet: '{text[:40]}...'")
    
    # Ensure we have a numbering engine
    if numbering_engine is None:
        numbering_engine = NumberingEngine()

    # C1/C2: Use safe ID allocation if no specific numId provided
    if num_id is None:
        # O3's Guard Rail: In production, num_id should be provided by safe allocation
        logger.warning("add_bullet_point_native: num_id=None - using emergency safe allocation")
        safe_num_id, safe_abstract_num_id = NumberingEngine._allocate_safe_ids(doc)
        num_id = safe_num_id
        abstract_num_id = safe_abstract_num_id
        logger.debug(f"C1/C2: Emergency safe allocation - numId: {num_id}, abstractNumId: {abstract_num_id}")
    else:
        # Backward compatibility - if num_id is provided, use it for abstract too
        abstract_num_id = num_id

    # Only create numbering definition once per document
    if not hasattr(doc, '_mr_numbering_definitions_created'):
        doc._mr_numbering_definitions_created = set()
    
    if num_id not in doc._mr_numbering_definitions_created:
        numbering_engine.get_or_create_numbering_definition(doc, num_id=num_id, abstract_num_id=abstract_num_id)
        doc._mr_numbering_definitions_created.add(num_id)
        logger.info(f"✅ C1/C2: Created numbering definition - numId: {num_id}, abstractNumId: {abstract_num_id}")
    else:
        logger.debug(f"✅ C1/C2: Using cached numbering definition - numId: {num_id}")
    
    # Create paragraph with content FIRST
    para = doc.add_paragraph()
    para.add_run(text.strip())
    
    # Apply MR_BulletPoint style
    if docx_styles:
        try:
            _apply_paragraph_style(doc, para, "MR_BulletPoint", docx_styles)
        except Exception as e:
            logger.warning(f"Could not apply MR_BulletPoint via design tokens: {e}")
            # Try fallback style application
            try:
                para.style = 'MR_BulletPoint'
            except Exception as e2:
                logger.warning(f"MR_BulletPoint style not available, using default: {e2}")
    else:
        # Handle case where docx_styles is None (test environment)
        try:
            para.style = 'MR_BulletPoint'
        except Exception as e:
            logger.warning(f"MR_BulletPoint style not available in test environment, using default: {e}")
            # In test environment, just use default paragraph style
    
    # Apply native numbering - TRUST that it works
    numbering_engine.apply_native_bullet(para, num_id=num_id, level=level)
    
    return para

def add_bullet_point_legacy(doc: Document, text: str, docx_styles: Dict[str, Any] = None) -> Paragraph:
    """
    LEGACY: Create bullet point with manual formatting (deprecated).
    
    ⚠️ G-3 FIX: All direct indent overrides removed to prevent spacing bug resurrection.
    This function should only be used as fallback when native numbering fails.
    
    🚨 o3's WARNING: This function creates L-0 direct formatting that can override
    native numbering indentation. Use should be eliminated in production.
    
    Args:
        doc: Document to add bullet to  
        text: Bullet text content
        docx_styles: Style definitions for design token system
        
    Returns:
        The created paragraph with manual bullet
    """
    import warnings
    warnings.warn(
        "Legacy bullet path deprecated; may produce wrong spacing. "
        "Use add_bullet_point_native() for professional Word behavior.",
        DeprecationWarning,
        stacklevel=2
    )
    
    # o3's ENHANCED WARNING
    logger.error(f"🚨 LEGACY BULLET PATH ACTIVE - This creates L-0 direct formatting!")
    logger.error(f"🚨 Text: '{text[:50]}...'")
    logger.error(f"🚨 This may override native numbering indentation!")
    
    # Optional: Completely disable legacy path in production
    if os.getenv('DOCX_DISABLE_LEGACY_BULLETS', 'false').lower() == 'true':
        raise RuntimeError(f"Legacy bullet path disabled by DOCX_DISABLE_LEGACY_BULLETS. Fix native path instead.")
    
    logger.debug(f"Creating legacy bullet: '{text[:50]}...'")
    
    # Create paragraph with content FIRST (content-first architecture)
    bullet_para = doc.add_paragraph()
    
    # ✅ CRITICAL FIX: Ensure bullet character is always added
    clean_text = text.strip()
    if not clean_text.startswith('•'):
        bullet_text = f"• {clean_text}"
    else:
        bullet_text = clean_text
    
    bullet_para.add_run(bullet_text)
    logger.info(f"🔫 LEGACY: Added bullet text: '{bullet_text[:50]}...'")
    
    # FIXED: Apply MR_BulletPoint style through design token system
    # This ensures the zero spacing from design tokens is properly applied
    if docx_styles:
        try:
            _apply_paragraph_style(doc, bullet_para, "MR_BulletPoint", docx_styles)
            logger.debug("✅ Applied MR_BulletPoint style via design token system")
        except Exception as e:
            logger.warning(f"Could not apply MR_BulletPoint via design tokens: {e}")
            # Fallback to direct style assignment
            try:
                bullet_para.style = 'MR_BulletPoint'
                logger.debug("✅ Applied MR_BulletPoint style via direct assignment")
            except Exception as e2:
                logger.warning(f"Could not apply MR_BulletPoint style at all: {e2}")
    else:
        # Fallback when no docx_styles provided
        try:
            bullet_para.style = 'MR_BulletPoint'
            logger.debug("✅ Applied MR_BulletPoint style via direct assignment (no docx_styles)")
        except Exception as e:
            logger.warning(f"Could not apply MR_BulletPoint style: {e}")
    
    # ❌ REMOVED: XML override that was fighting the design token system
    # The MR_BulletPoint style from design tokens already has spaceAfterPt: 0
    # Adding XML spacing was OVERRIDING the design token values!
    
    # ❌ G-3 FIX: These lines REMOVED to prevent spacing override bug:
    # bullet_para.paragraph_format.left_indent = Pt(18)
    # bullet_para.paragraph_format.first_line_indent = Pt(-18)
    
    # o3's WARNING: Check if style application secretly added direct formatting
    if bullet_para.paragraph_format.left_indent or bullet_para.paragraph_format.first_line_indent:
        left_info = f"{int(bullet_para.paragraph_format.left_indent.twips)} twips" if bullet_para.paragraph_format.left_indent else "None"
        first_info = f"{int(bullet_para.paragraph_format.first_line_indent.twips)} twips" if bullet_para.paragraph_format.first_line_indent else "None"
        logger.error(f"🚨 STYLE APPLICATION ADDED DIRECT FORMATTING!")
        logger.error(f"🚨 Left: {left_info}, First Line: {first_info}")
        logger.error(f"🚨 This is the source of L-0 override bug!")
    
    logger.debug(f"✅ Legacy bullet created with design token zero spacing")
    return bullet_para

def create_bullet_point(doc: Document, text: str, docx_styles: Dict[str, Any] = None, 
                       numbering_engine: NumberingEngine = None, num_id: int = 100,
                       o3_engine: Any = None, section_name: str = "unknown") -> Paragraph:
    """
    O3 Enhanced bullet creation using "Build-Then-Reconcile" architecture.
    
    This function implements O3's TRUST approach:
    1. Create content using O3 engine if available
    2. Apply style  
    3. Apply numbering
    4. TRUST that it works (no immediate verification)
    5. Reconciliation will fix any issues later
    
    Args:
        doc: Document to add bullet to
        text: Bullet text content  
        docx_styles: Style definitions
        numbering_engine: NumberingEngine instance for applying bullets
        num_id: Numbering ID (default: 100)
        o3_engine: O3 core engine for enhanced bullet handling
        section_name: Section name for O3 engine tracking
        
    Returns:
        The created paragraph with bullets
    """
    if not text or not text.strip():
        logger.warning("Empty bullet text provided, skipping")
        return None

    # O3 Enhanced: Use O3 engine if available for comprehensive bullet management
    if o3_engine is not None:
        try:
            logger.debug(f"O3: Creating trusted bullet via O3 engine: '{text[:40]}...'")
            para, bullet_id = o3_engine.create_bullet_trusted(
                doc=doc,
                text=text,
                section_name=section_name,
                numbering_engine=numbering_engine,
                docx_styles=docx_styles,
                num_id=num_id  # C1/C2: Pass the safe numId from caller
            )
            logger.debug(f"O3: Bullet created successfully with ID {bullet_id}")
            return para
        except Exception as e:
            logger.warning(f"O3: Engine failed, falling back to standard path: {e}")
            # Continue to standard path

    # B3: Locale-specific bullet glyph sanitization
    clean_text = text.strip()
    # Strip leading bullet glyphs (handles Western •, Chinese ·, Japanese ・, etc.)
    import re
    # Unicode General Punctuation bullet characters
    bullet_pattern = r'^[\u2022\u2023\u2043\u204C\u204D\u2219\u22C5\u2010-\u2015\u2043\uFF65\u30FB\u2024\u2025\u2026]*\s*'
    clean_text = re.sub(bullet_pattern, '', clean_text).strip()

    # Phase 1: Try native numbering if enabled
    if NATIVE_BULLETS_ENABLED:
        try:
            logger.debug(f"Creating native bullet: '{clean_text[:40]}...'")
            para = add_bullet_point_native(doc, clean_text, numbering_engine, num_id=num_id, docx_styles=docx_styles)
            logger.debug(f"Native bullet created successfully")
            return para
                
        except Exception as e:
            # Log error but don't retry - reconciliation will fix it
            logger.warning(f"Native bullet creation failed for '{clean_text[:40]}...': {e}")
            # Continue to fallback
    
    # Phase 2: Fallback to legacy approach
    try:
        logger.debug(f"Using legacy bullet path for: '{clean_text[:40]}...'")
        para = add_bullet_point_legacy(doc, clean_text, docx_styles)
        logger.debug("Legacy bullet created successfully")  
        return para
    except Exception as e:
        logger.error(f"Both native and legacy bullet creation failed: {e}")
        # Emergency fallback - basic paragraph with bullet
        para = doc.add_paragraph()
        para.add_run(f"• {clean_text}")
        logger.warning("Used emergency fallback paragraph")
        return para

@trace("docx.section_header")
def add_section_header(doc: Document, section_name: str) -> Any:
    """
    Add a section header using the new comprehensive specification.
    
    This implementation uses the MR_SectionHeader style with proper borders,
    fonts, and spacing as defined in the comprehensive DOCX specification.
    
    Args:
        doc: The document to add the section header to
        section_name: The section header text
        
    Returns:
        The rendered section header paragraph
    """
    # Use the new comprehensive specification approach
    try:
        # Get the DOCX styles
        docx_styles = StyleManager.load_docx_styles()
        
        # Create the section header paragraph
        header_para = doc.add_paragraph(section_name.upper())
        
        # Apply the MR_SectionHeader style using our enhanced style application
        _apply_paragraph_style(doc, header_para, "MR_SectionHeader", docx_styles)
        
        logger.info(f"Generated section header: '{section_name}' using MR_SectionHeader from comprehensive specification")
        return header_para
    except Exception as e:
        logger.warning(f"Comprehensive specification failed for '{section_name}': {e}")
        logger.warning("Falling back to legacy registry approach")
        
        # Fallback to legacy approach
        if USE_STYLE_REGISTRY:
            section_para = registry_add_section_header(doc, section_name)
            logger.info(f"Applied BoxedHeading2Table style to section header: {section_name}")
            return section_para
        else:
            # Emergency fallback
            para = doc.add_paragraph(section_name.upper())
            para.style = 'Heading 2'
            logger.warning(f"Used emergency fallback for section header: {section_name}")
            return para

def add_role_description(doc, text, docx_styles):
    """Adds a consistently formatted role description paragraph."""
    if not text:
        return None
    
    # Use our new custom style
    role_para = doc.add_paragraph(text, style='MR_RoleDescription')
    
    # USE DESIGN TOKENS instead of hardcoded values
    if docx_styles and 'MR_RoleDescription' in docx_styles:
        # Get spacing from the style configuration (which comes from design tokens)
        style_config = docx_styles['MR_RoleDescription']
        space_before = style_config.get('spaceBeforePt', 0)
        space_after = style_config.get('spaceAfterPt', 0)
        
        role_para.paragraph_format.space_before = Pt(space_before)
        role_para.paragraph_format.space_after = Pt(space_after)
        
        logger.info(f"Applied design token spacing: before={space_before}pt, after={space_after}pt")
    else:
        # Fallback to design tokens directly if style config not available
        from style_engine import StyleEngine
        design_tokens = StyleEngine.load_tokens()
        space_before = int(design_tokens.get("paragraph-spacing-roledesc-before", "0"))
        space_after = int(design_tokens.get("paragraph-spacing-roledesc-after", "0"))
        
        role_para.paragraph_format.space_before = Pt(space_before)
        role_para.paragraph_format.space_after = Pt(space_after)
    
    logger.info(f"Applied MR_RoleDescription with design token spacing to: {str(text)[:30]}...")
    return role_para

def tighten_before_headers(doc):
    """
    Finds paragraphs before section headers and sets spacing to zero.
    
    This enhanced implementation addresses the issues identified:
    1. Properly detects empty paragraphs that might cause unwanted spacing
    2. Applies spacing changes directly to XML for maximum compatibility
    3. Handles various paragraph types correctly
    
    Args:
        doc: The document to process
        
    Returns:
        Number of paragraphs fixed
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    logger.info("Starting tighten_before_headers process...")
    fixed_instances_count = 0
    
    # If using style registry, first remove any unwanted empty paragraphs
    if USE_STYLE_REGISTRY:
        # This doesn't delete intentional EmptyParagraph style paragraphs
        remove_count = remove_empty_paragraphs(doc)
        logger.info(f"Removed {remove_count} unwanted empty paragraphs")
    
    # Find all section headers
    header_indices = []
    for i, para in enumerate(doc.paragraphs):
        # Check by style name first (most reliable)
        if para.style and para.style.name in ['BoxedHeading2', 'Heading 2']:
            header_indices.append(i)
            continue
            
        # Check by borders (also reliable)
        p_pr = para._element.get_or_add_pPr()
        if p_pr.find(qn('w:pBdr')) is not None:
            header_indices.append(i)
            continue
            
        # Check by common section header names as fallback
        section_keywords = ["PROFESSIONAL SUMMARY", "SUMMARY", "EXPERIENCE", 
                            "EDUCATION", "SKILLS", "PROJECTS"]
        if any(keyword in para.text.upper() for keyword in section_keywords):
            header_indices.append(i)
            
    logger.info(f"Found {len(header_indices)} section headers")
    
    # Process paragraphs before each header (skip first header)
    for header_idx in header_indices[1:]:  # Skip first header (no content before it)
        if header_idx > 0:  # Safety check
            prev_para = doc.paragraphs[header_idx - 1]
            
            # Apply spacing change
            p_pr = prev_para._element.get_or_add_pPr()
            
            # Create spacing node with zero space after
            spacing_xml = f'<w:spacing {nsdecls("w")} w:after="0"/>'
            
            # Remove existing spacing element to avoid conflicts
            for existing in p_pr.xpath('./w:spacing'):
                p_pr.remove(existing)
                
            # Add new spacing element
            p_pr.append(parse_xml(spacing_xml))
            
            # Also set via API for maximum compatibility
            prev_para.paragraph_format.space_after = Pt(0)
            
            logger.info(f"Fixed spacing: Set space_after=0 on paragraph before section header at index {header_idx}")
            fixed_instances_count += 1
    
    # DIAGNOSTIC: Check if MR_Company style was actually created (O3 Checklist #2)
    logger.info("🔍 DIAGNOSTIC #2: Listing all document styles before save...")
    all_styles = [s.name for s in doc.styles]
    logger.info(f"📝 All document styles: {all_styles}")
    if "MR_Company" in all_styles:
        logger.info("✅ DIAGNOSTIC #2 PASSED: MR_Company style found in document")
    else:
        logger.warning("⚠️ DIAGNOSTIC #2 FAILED: MR_Company style NOT found in document!")
        logger.info("🛠️ Applying O3's backup solution: Creating robust MR_Company style...")
        _create_robust_company_style(doc)
    
    return fixed_instances_count

# Replace old _fix_spacing_between_sections with new implementation
_fix_spacing_between_sections = tighten_before_headers

def parse_contact_string(contact_text: str) -> Dict[str, str]:
    """
    Parse contact information from a string format.
    Handles both pipe-separated and space-separated formats with Unicode symbols.
    
    Args:
        contact_text: Raw contact string from JSON
        
    Returns:
        Dictionary with parsed contact fields
    """
    contact_data = {}
    
    # Split by lines and filter out empty lines
    lines = [line.strip() for line in contact_text.strip().split('\n') if line.strip()]
    
    if not lines:
        return contact_data
    
    # First line is the name
    contact_data['name'] = lines[0].strip()
    logger.info(f"Extracted name: {contact_data['name']}")
    
    # Process remaining lines for contact details
    for line in lines[1:]:
        line = line.strip()
        if not line:
            continue
            
        logger.info(f"Processing contact line: {line}")
        
        # Split by common separators (spaces, pipes, tabs)
        # Handle both pipe-separated and space-separated formats
        if '|' in line:
            parts = [p.strip() for p in line.split('|') if p.strip()]
        else:
            # For space-separated format, split by multiple spaces or Unicode symbols
            # Use regex to split on multiple spaces or symbol boundaries
            parts = re.split(r'\s{2,}|(?<=\S)\s+(?=[☎✉📧📞🌐])|(?<=[☎✉📧📞🌐])\s+', line)
            parts = [p.strip() for p in parts if p.strip()]
        
        logger.info(f"Split contact parts: {parts}")
        
        # Parse each part to identify contact type
        for part in parts:
            part = part.strip()
            if not part:
                continue
                
            # Remove Unicode symbols for easier parsing
            clean_part = re.sub(r'[☎✉📧📞🌐]', '', part).strip()
            
            # Email detection
            if '@' in clean_part and re.match(r'^[^@]+@[^@]+\.[^@]+$', clean_part):
                contact_data['email'] = clean_part
                logger.info(f"Found email: {clean_part}")
            
            # Phone number detection (various formats)
            elif re.search(r'[\d\-\.\(\)\+\s]{7,}', clean_part) and any(c.isdigit() for c in clean_part):
                # Additional validation for phone numbers
                digit_count = sum(1 for c in clean_part if c.isdigit())
                if digit_count >= 7:  # Minimum digits for a phone number
                    contact_data['phone'] = clean_part
                    logger.info(f"Found phone: {clean_part}")
            
            # URL detection (GitHub, LinkedIn, websites)
            elif any(domain in clean_part.lower() for domain in ['github.com', 'linkedin.com', 'www.', 'http']):
                if 'github' in clean_part.lower():
                    contact_data['github'] = clean_part
                    logger.info(f"Found GitHub: {clean_part}")
                elif 'linkedin' in clean_part.lower():
                    contact_data['linkedin'] = clean_part
                    logger.info(f"Found LinkedIn: {clean_part}")
                else:
                    contact_data['website'] = clean_part
                    logger.info(f"Found website: {clean_part}")
            
            # Location detection (contains common location indicators)
            elif any(indicator in clean_part.lower() for indicator in ['street', 'ave', 'road', 'blvd', 'city', 'state', ', ']):
                if len(clean_part) > 3:  # Avoid single letters
                    contact_data['location'] = clean_part
                    logger.info(f"Found location: {clean_part}")
            
            # If none of the above, and it's substantial text, treat as additional info
            elif len(clean_part) > 2 and not contact_data.get('location'):
                # Could be location or other contact info
                contact_data['other'] = clean_part
                logger.info(f"Found other contact info: {clean_part}")
    
    logger.info(f"Final parsed contact data: {contact_data}")
    return contact_data

def _cleanup_bullet_direct_formatting(doc: Document) -> int:
    """
    o3's Nuclear Option: Remove all direct indentation from bullet paragraphs.
    
    This addresses the rogue L-0 direct formatting that causes Word to show
    "Left: 0" instead of the proper bullet indentation from L-1 XML numbering.
    
    Returns:
        Number of bullet paragraphs cleaned
    """
    cleaned_count = 0
    
    for para in doc.paragraphs:
        if para.style and para.style.name == "MR_BulletPoint":
            # Check if paragraph has rogue direct formatting
            before_left = para.paragraph_format.left_indent
            before_first = para.paragraph_format.first_line_indent
            
            if before_left or before_first:
                # o3's nuclear cleanup - remove ALL direct formatting
                para.paragraph_format.left_indent = None
                para.paragraph_format.first_line_indent = None
                cleaned_count += 1
                
                if before_left:
                    left_twips = int(before_left.twips)
                    logger.debug(f"🧹 Removed rogue left indent: {left_twips} twips from '{para.text[:30]}...'")
                if before_first:
                    first_twips = int(before_first.twips)
                    logger.debug(f"🧹 Removed rogue first line indent: {first_twips} twips from '{para.text[:30]}...'")
    
    if cleaned_count > 0:
        logger.info(f"🧹 o3's Nuclear Cleanup: Removed direct formatting from {cleaned_count} bullet paragraphs")
    else:
        logger.debug("🧹 o3's Nuclear Cleanup: No direct formatting found to remove")
    
    return cleaned_count

def _detect_rogue_bullet_formatting(doc: Document, checkpoint_name: str) -> int:
    """
    o3's DRAMATIC DIAGNOSTIC: Detect rogue direct formatting on bullet paragraphs.
    
    This runs after every major operation to catch exactly when and where
    direct formatting gets added to bullet paragraphs.
    
    Args:
        doc: Document to scan
        checkpoint_name: Name of the checkpoint for logging
        
    Returns:
        Number of paragraphs with rogue formatting
    """
    rogue_count = 0
    
    logger.info(f"🔍 CHECKPOINT '{checkpoint_name}': Scanning for rogue bullet formatting...")
    
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == "MR_BulletPoint":
            # Check for rogue direct formatting
            left_indent = para.paragraph_format.left_indent
            first_line_indent = para.paragraph_format.first_line_indent
            
            if left_indent or first_line_indent:
                rogue_count += 1
                
                left_info = f"{int(left_indent.twips)} twips" if left_indent else "None"
                first_info = f"{int(first_line_indent.twips)} twips" if first_line_indent else "None"
                
                logger.error(f"🚨 ROGUE FORMATTING DETECTED at checkpoint '{checkpoint_name}':")
                logger.error(f"   Paragraph {i}: '{para.text[:50]}...'")
                logger.error(f"   Left indent: {left_info}")
                logger.error(f"   First line indent: {first_info}")
                
                # Check if it has numbering properties
                pPr = para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                if pPr is not None:
                    numPr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                    if numPr is not None:
                        logger.error(f"   Has numbering: YES (this is the L-0 override bug!)")
                    else:
                        logger.error(f"   Has numbering: NO")
    
    if rogue_count == 0:
        logger.info(f"✅ CHECKPOINT '{checkpoint_name}': No rogue formatting detected")
    else:
        logger.error(f"❌ CHECKPOINT '{checkpoint_name}': {rogue_count} paragraphs have rogue formatting")
    
    return rogue_count

def build_docx(request_id: str, temp_dir: str, debug: bool = False) -> BytesIO:
    """
    Build a DOCX file from the resume data for the given request ID.
    
    Args:
        request_id: The unique request ID for the resume
        temp_dir: Directory containing the temp session data files
        debug: Whether to enable debugging output
        
    Returns:
        BytesIO object containing the DOCX file data
    """
    try:
        logger.info(f"Building DOCX for request ID: {request_id}")
        logger.info(f"Temp directory path: {temp_dir}")
        logger.info(f"Debug mode: {debug}")
        
        # Check for files related to this request_id
        try:
            file_list = os.listdir(temp_dir)
            matching_files = [f for f in file_list if request_id in f]
            logger.info(f"Files containing request ID: {matching_files}")
        except Exception as e:
            logger.error(f"Error listing directory: {e}")
        
        # Load DOCX styles from StyleManager
        docx_styles = StyleManager.load_docx_styles()
        if not docx_styles:
            logger.error("No DOCX styles found. Using defaults.")
            docx_styles = {
                "page": {"marginTopCm": 1.0, "marginBottomCm": 1.0, "marginLeftCm": 2.0, "marginRightCm": 2.0},
                "heading1": {"fontFamily": "Calibri", "fontSizePt": 16},
                "heading2": {"fontFamily": "Calibri", "fontSizePt": 14, "color": [0, 0, 102], "spaceAfterPt": 6},
                "heading3": {"fontFamily": "Calibri", "fontSizePt": 11, "bold": True, "spaceAfterPt": 4},
                "body": {"fontFamily": "Calibri", "fontSizePt": 11}
            }
        
        # Create a new Document
        doc = Document()
        
        # Initialize NumberingEngine for optimal performance (O3 recommendation)
        # A4: Use singleton with request ID for isolation
        numbering_engine = None
        custom_num_id = None
        o3_engine = None
        
        if NATIVE_BULLETS_ENABLED:
            try:
                # A4: Get singleton instance with request isolation
                numbering_engine = NumberingEngine.get_instance(request_id)
                
                # C1/C2: Use safe allocation instead of global counter
                custom_num_id, custom_abstract_num_id = NumberingEngine._allocate_safe_ids(doc)
                
                logger.info(f"🔧 C1/C2: Creating safe numbering definition (numId={custom_num_id}, abstractNumId={custom_abstract_num_id})...")
                numbering_engine.get_or_create_numbering_definition(doc, num_id=custom_num_id, abstract_num_id=custom_abstract_num_id)
                
                # O3: Initialize O3 core engine for enhanced bullet management
                from utils.o3_bullet_core_engine import get_o3_engine
                o3_engine = get_o3_engine(request_id)
                logger.info("🚀 O3: Core bullet engine initialized")
                
                logger.info("✅ NumberingEngine initialized with safe bullet definition")
            except Exception as e:
                logger.warning(f"Failed to initialize NumberingEngine: {e}")
                numbering_engine = None
                custom_num_id = 100  # Fallback ID
        
        # Create custom document styles
        custom_styles = _create_document_styles(doc, docx_styles)
        
        # o3's CHECKPOINT 1: After style creation
        _detect_rogue_bullet_formatting(doc, "AFTER_STYLE_CREATION")
        
        # **FIX: Ensure all custom styles are actually available in the document**
        logger.info("🔧 VERIFYING: Checking if all custom styles are available...")
        expected_styles = ['MR_SectionHeader', 'MR_Content', 'MR_RoleDescription', 'MR_BulletPoint', 
                          'MR_SummaryText', 'MR_SkillCategory', 'MR_SkillList', 'MR_Company']
        available_styles = [s.name for s in doc.styles]
        
        missing_styles = [style for style in expected_styles if style not in available_styles]
        if missing_styles:
            logger.error(f"❌ MISSING STYLES: {missing_styles}")
            logger.error(f"❌ Available styles: {available_styles}")
            # Force recreation of missing styles
            logger.info("🔧 FORCING recreation of missing styles...")
            try:
                # Try to force recreation
                from style_engine import StyleEngine
                StyleEngine.create_docx_custom_styles(doc)
                logger.info("✅ Successfully forced style recreation")
            except Exception as e:
                logger.error(f"❌ Failed to force style recreation: {e}")
        else:
            logger.info(f"✅ ALL STYLES AVAILABLE: {expected_styles}")
        
        # Get page configuration from new comprehensive specification
        page_config = docx_styles.get("page", {})
        
        # Set all margins from the new specification (1.5cm each)
        section = doc.sections[0]
        section.top_margin = Cm(page_config.get("marginTopCm", 1.5))
        section.bottom_margin = Cm(page_config.get("marginBottomCm", 1.5))
        section.left_margin = Cm(page_config.get("marginLeftCm", 1.5))
        section.right_margin = Cm(page_config.get("marginRightCm", 1.5))
        
        logger.info(f"Applied document margins from specification: Top={page_config.get('marginTopCm', 1.5)}cm, Bottom={page_config.get('marginBottomCm', 1.5)}cm, Left={page_config.get('marginLeftCm', 1.5)}cm, Right={page_config.get('marginRightCm', 1.5)}cm")
        
        # ------ CONTACT SECTION ------
        logger.info("Processing Contact section...")
        
        # Open and inspect the contact file directly to debug structure issues
        contact_file_path = os.path.join(temp_dir, f"{request_id}_contact.json")
        if os.path.exists(contact_file_path):
            try:
                with open(contact_file_path, 'r', encoding='utf-8') as f:
                    raw_content = f.read()
                    logger.info(f"Raw contact file content (first 200 chars): {raw_content[:200]}")
            except Exception as e:
                logger.error(f"Error reading raw contact file: {e}")
        
        contact = load_section_json(request_id, "contact", temp_dir)
        logger.info(f"Contact data loaded: {bool(contact)}")
        logger.info(f"Contact data type: {type(contact)}")
        
        if isinstance(contact, dict):
            logger.info(f"Contact keys: {list(contact.keys())}")
        
        if contact:
            # Verify contact is a dictionary
            if not isinstance(contact, dict):
                logger.warning(f"Contact section is not a dictionary: {type(contact)}")
                contact = {}
            
            # Check if contact has a 'content' key (alternate structure)
            if 'content' in contact:
                logger.info("Found 'content' key in contact section")
                # Extract actual contact data from content
                try:
                    # Try to access the contact data
                    contact_data = contact['content']
                    if isinstance(contact_data, dict):
                        logger.info(f"Contact content is a dictionary with keys: {list(contact_data.keys())}")
                        contact = contact_data
                    elif isinstance(contact_data, list) and len(contact_data) > 0 and isinstance(contact_data[0], dict):
                        logger.info(f"Contact content is a list of dictionaries, using first item")
                        contact = contact_data[0]
                    elif isinstance(contact_data, str):
                        # Handle string content - parse it into structured data
                        logger.info("Contact content is a string, attempting to parse")
                        
                        # Parse the contact information from the string
                        contact = parse_contact_string(contact_data)
                        logger.info(f"Created structured contact data: {contact}")
                    else:
                        logger.info(f"Contact content is type: {type(contact_data)}")
                except Exception as e:
                    logger.warning(f"Error accessing contact content: {e}")
            
            # Print full contact data for debugging
            logger.info(f"Final contact data structure to use: {contact}")
            
            # Name - handle potential different structures
            name = ""
            if "name" in contact:
                name = contact["name"]
            elif "full_name" in contact:
                name = contact["full_name"]
            
            if name:
                name_para = doc.add_paragraph(name)
                _apply_paragraph_style(doc, name_para, "MR_Name", docx_styles)
                
                # Contact details
                contact_parts = []
                if "location" in contact and contact["location"]:
                    contact_parts.append(contact["location"])
                if "phone" in contact and contact["phone"]:
                    contact_parts.append(contact["phone"])
                if "email" in contact and contact["email"]:
                    contact_parts.append(contact["email"])
                if "linkedin" in contact and contact["linkedin"]:
                    contact_parts.append(contact["linkedin"])
                
                # Additional possible contact fields
                if "website" in contact and contact["website"]:
                    contact_parts.append(contact["website"])
                if "github" in contact and contact["github"]:
                    contact_parts.append(contact["github"])
                
                contact_text = " • ".join(contact_parts)  # Use bullet separator as per specification
                contact_para = doc.add_paragraph(contact_text)
                _apply_paragraph_style(doc, contact_para, "MR_Contact", docx_styles)
            else:
                logger.warning("No name found in contact data, skipping contact section")
        else:
            logger.warning("No contact data found")
            
            # Try a fallback approach - look for the file directly
            fallback_files = [f for f in os.listdir(temp_dir) if f.endswith("_contact.json") and request_id in f]
            if fallback_files:
                logger.info(f"Found fallback contact files: {fallback_files}")
                try:
                    with open(os.path.join(temp_dir, fallback_files[0]), 'r', encoding='utf-8') as f:
                        fallback_contact = json.load(f)
                        logger.info(f"Loaded fallback contact data: {fallback_contact}")
                        
                        # Extract name and add to document
                        if isinstance(fallback_contact, dict):
                            name = fallback_contact.get("name", "")
                            if name:
                                name_para = doc.add_paragraph(name)
                                _apply_paragraph_style(doc, name_para, "MR_Name", docx_styles)
                                logger.info(f"Added name from fallback: {name}")
                except Exception as e:
                    logger.error(f"Error processing fallback contact data: {e}")
        
        # ------ SUMMARY SECTION ------
        logger.info("Processing Summary section...")
        summary = load_section_json(request_id, "summary", temp_dir)
        logger.info(f"Summary data loaded: {bool(summary)}")
        
        # Handle both direct summary and summary with 'content' key
        summary_text = ""
        if summary:
            if isinstance(summary, dict):
                if "summary" in summary:
                    summary_text = summary.get("summary", "")
                elif "content" in summary:
                    # Alternative structure with content key
                    content = summary.get("content", "")
                    if isinstance(content, dict) and "summary" in content:
                        summary_text = content.get("summary", "")
                    else:
                        summary_text = str(content)
            else:
                summary_text = str(summary)
                
            if summary_text:
                # Add section header with helper function
                summary_header = add_section_header(doc, "PROFESSIONAL SUMMARY")
                
                # Add summary content
                summary_para = doc.add_paragraph(summary_text, style='MR_SummaryText')
                
                # Space after section
                doc.add_paragraph("").paragraph_format.space_after = Pt(6)
        
        # ------ EXPERIENCE SECTION ------
        logger.info("Processing Experience section...")
        experience = load_section_json(request_id, "experience", temp_dir)
        logger.info(f"Experience data loaded: {bool(experience)}")
        logger.info(f"Experience contains 'experiences' key: {isinstance(experience, dict) and 'experiences' in experience}")
        
        # Handle both dictionary with 'experiences' key and direct list of experiences
        experiences_list = []
        if experience:
            if isinstance(experience, dict) and "experiences" in experience:
                experiences_list = experience.get("experiences", [])
            elif isinstance(experience, list):
                # Direct list of experiences
                experiences_list = experience
            else:
                logger.warning(f"Unexpected experience format: {type(experience)}")
                
            if experiences_list:
                # Add section header with helper function
                exp_header = add_section_header(doc, "EXPERIENCE")
                
                # Verify experiences is a list
                if not isinstance(experiences_list, list):
                    logger.warning(f"Experiences is not a list: {type(experiences_list)}")
                    experiences_list = []
                    
                # Add each job
                for job in experiences_list:
                    # Verify job is a dictionary
                    if not isinstance(job, dict):
                        logger.warning(f"Job is not a dictionary: {type(job)}")
                        continue
                        
                    # Company and location - use the helper function for consistent formatting
                    company = job.get('company', '')
                    location = job.get('location', '')
                    
                    if company or location:
                        company_para = format_right_aligned_pair(
                            doc,
                            company,
                            location,
                            "MR_Company",
                            "MR_Company",
                            docx_styles
                        )
                    
                    # Position and dates - use role box for consistent styling with HTML/PDF
                    position = job.get('position', '')
                    if not position and job.get('title'):  # Fallback to 'title' if 'position' is not available
                        position = job.get('title', '')
                    dates = job.get('dates', '')
                    
                    if position or dates:
                        # Use role box instead of format_right_aligned_pair for consistent styling
                        if USE_STYLE_REGISTRY:
                            role_box_table = add_role_box(doc, position, dates)
                            logger.info(f"Added role box for position: '{position}' with dates: '{dates}'")
                        else:
                            # Fallback to original approach if style registry is not available
                            position_para = format_right_aligned_pair(
                                doc,
                                position,
                                dates,
                                "body",
                                "body",
                                docx_styles
                            )
                            logger.info(f"Used fallback formatting for position: '{position}' with dates: '{dates}'")
                    
                    logger.info(f"Formatted experience entry: '{company} - {position}'")
                    
                    # Role description - use the helper function for consistent formatting
                    if job.get('role_description'):
                        role_para = add_role_description(doc, job.get('role_description'), docx_styles)
                    
                    # Achievements/bullets - use the helper function for consistent formatting
                    for achievement in job.get('achievements', []):
                        bullet_para = create_bullet_point(
                            doc, achievement, docx_styles, numbering_engine, 
                            num_id=custom_num_id, o3_engine=o3_engine, section_name="experience"
                        )
                        
                        # o3's CHECKPOINT: After each bullet creation
                        _detect_rogue_bullet_formatting(doc, f"AFTER_BULLET_{achievement[:20]}")
                
                # o3's CHECKPOINT 2: After all experience bullets
                _detect_rogue_bullet_formatting(doc, "AFTER_ALL_EXPERIENCE_BULLETS")
        
        # ------ EDUCATION SECTION ------
        logger.info("Processing Education section...")
        education = load_section_json(request_id, "education", temp_dir)
        logger.info(f"Education data loaded: {bool(education)}")
        logger.info(f"Education contains 'institutions' key: {isinstance(education, dict) and 'institutions' in education}")
        
        # Handle both dictionary with 'institutions' key and direct list of institutions
        institutions_list = []
        if education:
            if isinstance(education, dict) and "institutions" in education:
                institutions_list = education.get("institutions", [])
            elif isinstance(education, list):
                # Direct list of institutions
                institutions_list = education
            else:
                logger.warning(f"Unexpected education format: {type(education)}")
                
            if institutions_list:
                # Add section header with helper function
                edu_header = add_section_header(doc, "EDUCATION")
                
                # Verify institutions is a list
                if not isinstance(institutions_list, list):
                    logger.warning(f"Institutions is not a list: {type(institutions_list)}")
                    institutions_list = []
                    
                # Add each institution
                for school in institutions_list:
                    # Verify school is a dictionary
                    if not isinstance(school, dict):
                        logger.warning(f"School is not a dictionary: {type(school)}")
                        continue
                        
                    # Institution and location - use the helper function for consistent formatting
                    institution = school.get('institution', '')
                    location = school.get('location', '')
                    
                    if institution or location:
                        school_para = format_right_aligned_pair(
                            doc,
                            institution,
                            location,
                            "MR_Company",
                            "MR_Company",
                            docx_styles
                        )
                    
                    # Degree and dates - use role box for consistent styling with HTML/PDF
                    degree = school.get('degree', '')
                    dates = school.get('dates', '')
                    
                    if degree or dates:
                        # Use role box instead of format_right_aligned_pair for consistent styling
                        if USE_STYLE_REGISTRY:
                            role_box_table = add_role_box(doc, degree, dates)
                            logger.info(f"Added role box for degree: '{degree}' with dates: '{dates}'")
                        else:
                            # Fallback to original approach if style registry is not available
                            degree_para = format_right_aligned_pair(
                                doc,
                                degree,
                                dates,
                                "body",
                                "body",
                                docx_styles
                            )
                            logger.info(f"Used fallback formatting for degree: '{degree}' with dates: '{dates}'")
                    
                    logger.info(f"Formatted education entry: '{institution} - {degree}'")
                    
                    # Highlights/bullets - use the helper function for consistent formatting
                    for highlight in school.get('highlights', []):
                        bullet_para = create_bullet_point(
                            doc, highlight, docx_styles, numbering_engine, 
                            num_id=custom_num_id, o3_engine=o3_engine, section_name="education"
                        )
        
        # ------ SKILLS SECTION ------
        logger.info("Processing Skills section...")
        skills = load_section_json(request_id, "skills", temp_dir)
        logger.info(f"Skills data loaded: {bool(skills)}")
        logger.info(f"Skills data type: {type(skills)}")
        logger.info(f"Skills content sample: {str(skills)[:100]}")
        
        if skills:
            # Add section header with helper function 
            skills_header = add_section_header(doc, "SKILLS")
            
            # Add skills content - handle different possible formats
            if isinstance(skills, dict) and "skills" in skills:
                skills_content = skills.get("skills", "")
                logger.info(f"Skills content type: {type(skills_content)}")
                
                # Check if skills content is a list
                if isinstance(skills_content, list):
                    logger.info("Processing skills as inline list with commas")
                    # Handle skills as a comma-separated list on a single line
                    skills_text = ", ".join([str(skill) for skill in skills_content])
                    skills_para = doc.add_paragraph(skills_text, style='MR_SkillList')
                    logger.info(f"Applied MR_SkillList style to skills list")
                elif isinstance(skills_content, dict):
                    logger.info("Processing skills as dict")
                    # Handle skills as dictionary with categories
                    for category, skill_list in skills_content.items():
                        # Add category as subheading
                        category_para = doc.add_paragraph(category.upper(), style='MR_SkillCategory')
                        logger.info(f"Applied MR_SkillCategory style to category: {category}")
                        
                        # Add skills in this category as a comma-separated list
                        if isinstance(skill_list, list):
                            skills_text = ", ".join([str(skill) for skill in skill_list])
                            skills_para = doc.add_paragraph(skills_text, style='MR_SkillList')
                            logger.info(f"Applied MR_SkillList style to skills in category: {category}")
                        else:
                            # Not a list, just add as text
                            skill_para = doc.add_paragraph(str(skill_list), style='MR_SkillList')
                            logger.info(f"Applied MR_SkillList style to non-list skills in category: {category}")
                else:
                    logger.info("Processing skills as string")
                    # Handle skills as string or any other type
                    skills_para = doc.add_paragraph(str(skills_content), style='MR_SkillList')
                    logger.info(f"Applied MR_SkillList style to skills string")
            elif isinstance(skills, dict):
                logger.info("Processing skills dict directly")
                # Direct display of skills object if it doesn't have "skills" key
                for category, skill_list in skills.items():
                    # Add category as subheading
                    category_para = doc.add_paragraph(category.upper(), style='MR_SkillCategory')
                    logger.info(f"Applied MR_SkillCategory style to category: {category}")
                    
                    # Add skills in this category as a comma-separated list
                    if isinstance(skill_list, list):
                        skills_text = ", ".join([str(skill) for skill in skill_list])
                        skills_para = doc.add_paragraph(skills_text, style='MR_SkillList')
                        logger.info(f"Applied MR_SkillList style to skills in category: {category}")
                    else:
                        # Not a list, just add as text
                        skill_para = doc.add_paragraph(str(skill_list), style='MR_SkillList')
                        logger.info(f"Applied MR_SkillList style to non-list skills in category: {category}")
            elif isinstance(skills, list):
                logger.info("Processing skills as top-level list")
                # Handle the case where skills is a list directly
                skills_text = ", ".join([str(skill) for skill in skills])
                skills_para = doc.add_paragraph(skills_text, style='MR_SkillList')
                logger.info(f"Applied MR_SkillList style to top-level skills list")
            else:
                logger.info(f"Processing skills as fallback type: {type(skills)}")
                # Fallback for any other format
                skills_para = doc.add_paragraph(str(skills), style='MR_SkillList')
                logger.info(f"Applied MR_SkillList style to fallback skills content")
        
        # ------ PROJECTS SECTION ------
        logger.info("Processing Projects section...")
        projects = load_section_json(request_id, "projects", temp_dir)
        logger.info(f"Projects data loaded: {bool(projects)}")
        logger.info(f"Projects contains 'projects' key: {isinstance(projects, dict) and 'projects' in projects}")
        
        # Handle both dictionary with 'projects' key, direct list of projects, and content key
        projects_list = []
        if projects:
            if isinstance(projects, dict) and "projects" in projects:
                projects_list = projects.get("projects", [])
            elif isinstance(projects, dict) and "content" in projects:
                # Handle content key similar to other sections
                logger.info("Found 'content' key in projects section")
                content = projects.get("content", [])
                if isinstance(content, list):
                    projects_list = content
                elif isinstance(content, dict) and "projects" in content:
                    projects_list = content.get("projects", [])
                elif isinstance(content, str):
                    # In some cases, content might be a string
                    logger.info("Projects content is a string, adding as single project")
                    # Add projects section header using helper function
                    projects_header = add_section_header(doc, "PROJECTS")
                    
                    # Add the string content directly as paragraph
                    projects_para = doc.add_paragraph(content)
                    _apply_paragraph_style(doc, projects_para, "body", docx_styles)
                    
                    # Skip the rest of the projects processing
                    projects_list = []
                else:
                    logger.warning(f"Unexpected projects content format: {type(content)}")
            elif isinstance(projects, list):
                # Direct list of projects
                projects_list = projects
            else:
                logger.warning(f"Unexpected projects format: {type(projects)}")
                
            if projects_list:
                # Add section header with helper function
                projects_header = add_section_header(doc, "PROJECTS")
                
                # Verify projects is a list
                if not isinstance(projects_list, list):
                    logger.warning(f"Projects is not a list: {type(projects_list)}")
                    projects_list = []
                    
                # Add each project
                for project in projects_list:
                    # Verify project is a dictionary
                    if not isinstance(project, dict):
                        logger.warning(f"Project is not a dictionary: {type(project)}")
                        continue
                        
                    # Project title and dates - use role box for consistent styling with HTML/PDF
                    title = project.get('title', '')
                    if not title and project.get('name'):  # Fallback to 'name' if 'title' is not available
                        title = project.get('name', '')
                    dates = project.get('dates', '')
                    if not dates and project.get('date'):  # Fallback to 'date' if 'dates' is not available
                        dates = project.get('date', '')
                    
                    if title or dates:
                        # Use role box instead of format_right_aligned_pair for consistent styling
                        if USE_STYLE_REGISTRY:
                            role_box_table = add_role_box(doc, title, dates)
                            logger.info(f"Added role box for project title: '{title}' with dates: '{dates}'")
                        else:
                            # Fallback to original approach if style registry is not available
                            title_para = format_right_aligned_pair(
                                doc,
                                title,
                                dates,
                                "MR_Company",
                                "body",
                                docx_styles
                            )
                            logger.info(f"Used fallback formatting for project title: '{title}' with dates: '{dates}'")
                    
                    logger.info(f"Formatted project entry: '{title}'")
                    
                    # Project details - use the helper function for consistent formatting
                    for detail in project.get('details', []):
                        bullet_para = create_bullet_point(
                            doc, detail, docx_styles, numbering_engine, 
                            num_id=custom_num_id, o3_engine=o3_engine, section_name="projects"
                        )
        
        # Fix spacing between sections - use our enhanced implementation
        if USE_STYLE_REGISTRY:
            tighten_before_headers(doc)
        else:
            _fix_spacing_between_sections(doc)
        
        # DIAGNOSTIC: Check if MR_Company style was actually created (O3 Checklist #2)
        logger.info("🔍 DIAGNOSTIC #2: Listing all document styles before save...")
        all_styles = [s.name for s in doc.styles]
        logger.info(f"📝 All document styles: {all_styles}")
        if "MR_Company" in all_styles:
            logger.info("✅ DIAGNOSTIC #2 PASSED: MR_Company style found in document")
        else:
            logger.warning("⚠️ DIAGNOSTIC #2 FAILED: MR_Company style NOT found in document!")
            logger.info("🛠️ Applying O3's backup solution: Creating robust MR_Company style...")
            _create_robust_company_style(doc)
        
        # 🚀 O3's Enhanced "Build-Then-Reconcile" Architecture: Final Bullet Consistency Pass
        if NATIVE_BULLETS_ENABLED and numbering_engine and custom_num_id:
            try:
                logger.info("🚀 O3: Starting enhanced bullet reconciliation...")
                
                # O3: Use enhanced reconciliation if O3 engine is available
                if o3_engine is not None:
                    # O3 Enhanced reconciliation
                    reconciliation_stats = o3_engine.reconcile_document_bullets(doc, numbering_engine)
                    
                    # Log O3 results
                    bullets_processed = reconciliation_stats.get('bullets_processed', 0)
                    bullets_repaired = reconciliation_stats.get('bullets_repaired', 0)
                    bullets_stable = reconciliation_stats.get('bullets_stable', 0)
                    success_rate = reconciliation_stats.get('success_rate', 0)
                    
                    logger.info(f"🚀 O3: Processed {bullets_processed} bullets - "
                               f"{bullets_repaired} repaired, {bullets_stable} stable ({success_rate:.1f}% success)")
                    
                    # Get O3 engine summary
                    engine_summary = o3_engine.get_engine_summary()
                    logger.info(f"🚀 O3: Engine state: {engine_summary['document_state']}, "
                               f"Total bullets: {engine_summary['bullet_count']}")
                    
                else:
                    # Fallback to legacy reconciliation if O3 engine not available
                    if USE_BULLET_RECONCILIATION:
                        logger.info("🛡️ Starting legacy bullet reconciliation pass...")
                        reconciler = BulletReconciliationEngine(request_id)
                        reconciliation_stats = reconciler.reconcile_bullet_styles(doc, numbering_engine, custom_num_id)
                        
                        # Log legacy results
                        total_bullets = reconciliation_stats.get('total_bullets', 0)
                        repaired_count = reconciliation_stats.get('repaired_count', 0)
                        duration_ms = reconciliation_stats.get('duration_ms', 0)
                        errors = reconciliation_stats.get('errors', [])
                        
                        if repaired_count > 0:
                            logger.info(f"🛡️ Legacy reconciliation fixed {repaired_count}/{total_bullets} bullets ({duration_ms:.1f}ms)")
                        else:
                            logger.info(f"🛡️ All {total_bullets} bullets already consistent ({duration_ms:.1f}ms)")
                        
                        if errors:
                            logger.warning(f"🛡️ {len(errors)} reconciliation errors occurred")
                            for error in errors[:3]:  # Log first 3 errors
                                logger.warning(f"  - {error}")
                
            except Exception as e:
                logger.error(f"🚀 O3: Enhanced reconciliation failed: {e}")
                logger.warning("🚀 O3: Proceeding with document save despite reconciliation failure")
        
        logger.info("Saving DOCX to BytesIO...")
        
        # 🚨 Legacy O3 Cleanup: Remove direct formatting issues
        # This preserves existing functionality while the reconciliation handles the main logic
        cleaned_count = _cleanup_bullet_direct_formatting(doc)
        if cleaned_count > 0:
            logger.info(f"🧹 Cleaned {cleaned_count} bullet paragraphs with direct formatting issues")
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Generate debug report if requested
        if debug:
            try:
                from utils.docx_debug import generate_debug_report
                import json
                
                # Generate debug report
                debug_report = generate_debug_report(doc)
                
                # Save debug report to file
                debug_path = Path(__file__).parent.parent / f'debug_{request_id}.json'
                with open(debug_path, 'w') as f:
                    json.dump(debug_report, f, indent=2)
                
                logger.info(f"Debug report saved to {debug_path}")
                
                # Also save a copy of the DOCX for inspection
                debug_docx_path = Path(__file__).parent.parent / f'debug_{request_id}.docx'
                with open(debug_docx_path, 'wb') as f:
                    f.write(output.getvalue())
                
                logger.info(f"Debug DOCX saved to {debug_docx_path}")
                
                # Reset the output buffer position
                output.seek(0)
            except Exception as e:
                logger.error(f"Error generating debug report: {e}")
        
        # O3: Clean up engine resources
        if o3_engine is not None:
            try:
                from utils.o3_bullet_core_engine import cleanup_o3_engine
                cleanup_o3_engine(request_id)
                logger.info("🚀 O3: Engine resources cleaned up")
            except Exception as e:
                logger.warning(f"🚀 O3: Engine cleanup failed: {e}")
        
        logger.info(f"Successfully built DOCX for request ID: {request_id}")
        return output
    
    except Exception as e:
        # Enhanced error reporting
        logger.error(f"Error building DOCX for request ID {request_id}: {e}")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error details: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise 

def _create_robust_company_style(doc):
    """O3's recommended robust solution for company style creation."""
    from docx.oxml.shared import qn
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE
    
    logger.info("🛠️ Creating robust MR_Company style using O3's method...")
    
    # Get or create the style
    try:
        st = doc.styles['MR_Company']
        logger.info("📝 MR_Company style already exists, updating it...")
    except KeyError:
        st = doc.styles.add_style('MR_Company', WD_STYLE_TYPE.PARAGRAPH)
        logger.info("📝 Created new MR_Company style...")
    
    # Set base style to No Spacing (inherits 0pt before/after)
    try:
        st.base_style = doc.styles['No Spacing']
        logger.info("✅ Set base style to 'No Spacing'")
    except KeyError:
        st.base_style = None
        logger.info("⚠️ 'No Spacing' style not found, set base_style to None")
    
    # Set paragraph format
    pf = st.paragraph_format
    pf.space_before, pf.space_after = Pt(0), Pt(0)
    logger.info("✅ Set space_before and space_after to 0pt")
    
    # Kill any line-based spacing using XML manipulation
    sp = st._element.get_or_add_pPr().get_or_add_spacing()
    sp.set(qn('w:after'), '0')
    sp.set(qn('w:afterLines'), '0')
    sp.set(qn('w:afterAutospacing'), '0')
    sp.set(qn('w:contextualSpacing'), '1')
    logger.info("✅ Applied XML-level spacing controls (w:after=0, w:afterLines=0, w:afterAutospacing=0, w:contextualSpacing=1)")
    
    return st 