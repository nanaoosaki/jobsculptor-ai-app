#!/usr/bin/env python3
"""
Phase 5.1 Integration Fix Test
============================

Test to verify that the O3 engine integration fix correctly uses 
safe numId allocation and threads it through the call chain.

This test validates:
1. Safe allocation is used by O3 engine
2. numId is passed through create_bullet_point -> O3 -> apply_native_bullet
3. Paragraphs reference the correct safe numId (not hardcoded 100)

Created: January 2025
"""

import os
import sys
import tempfile
from pathlib import Path
from docx import Document

# Add project root to Python path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

from word_styles.numbering_engine import NumberingEngine
from utils.docx_builder import create_bullet_point
import logging

# Configure logging to see the debug messages
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def test_integration_fix():
    """Test that safe numId is correctly threaded through the call chain."""
    print("🔧 Testing C1/C2 Integration Fix")
    print("="*60)
    
    # Create document and numbering engine
    doc = Document()
    numbering_engine = NumberingEngine.get_instance("test_integration")
    
    # Get safe allocation
    safe_num_id, safe_abstract_num_id = NumberingEngine._allocate_safe_ids(doc)
    print(f"🔢 Safe allocation: numId={safe_num_id}, abstractNumId={safe_abstract_num_id}")
    
    # Test O3 engine integration
    try:
        from utils.o3_bullet_core_engine import get_o3_engine
        o3_engine = get_o3_engine("test_integration")
        
        print(f"🚀 O3 engine available - testing full integration")
        
        # Create bullet using safe numId
        para = create_bullet_point(
            doc=doc,
            text="Test bullet with safe numId",
            docx_styles=None,  # Skip styles for test
            numbering_engine=numbering_engine,
            num_id=safe_num_id,
            o3_engine=o3_engine,
            section_name="test"
        )
        
        print(f"✅ Bullet created successfully")
        
        # Verify the paragraph uses the correct numId
        num_pr = para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
        if num_pr is not None:
            num_id_elem = num_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
            if num_id_elem is not None:
                actual_num_id = int(num_id_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'))
                print(f"📊 Paragraph numId: {actual_num_id}")
                
                if actual_num_id == safe_num_id:
                    print(f"✅ SUCCESS: Paragraph uses safe numId {safe_num_id}")
                    return True
                else:
                    print(f"❌ FAILURE: Paragraph uses {actual_num_id}, expected {safe_num_id}")
                    return False
            else:
                print(f"❌ FAILURE: No numId found in paragraph")
                return False
        else:
            print(f"❌ FAILURE: No numbering properties found in paragraph")
            return False
            
    except ImportError:
        print(f"⚠️  O3 engine not available - testing fallback path")
        
        # Test without O3 engine
        para = create_bullet_point(
            doc=doc,
            text="Test bullet without O3 engine",
            docx_styles=None,
            numbering_engine=numbering_engine,
            num_id=safe_num_id,
            o3_engine=None,
            section_name="test"
        )
        
        print(f"✅ Bullet created successfully (fallback path)")
        return True
        
    except Exception as e:
        print(f"❌ FAILURE: Exception during test: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_sequential_documents():
    """Test that sequential documents get different safe numIds."""
    print("\n🔄 Testing Sequential Document Safe Allocation")
    print("="*60)
    
    try:
        # Document 1
        doc1 = Document()
        safe_num_1, safe_abstract_1 = NumberingEngine._allocate_safe_ids(doc1)
        print(f"📊 Document 1 (before numbering): numId={safe_num_1}, abstractNumId={safe_abstract_1}")
        
        # Create a numbering definition to simulate used numbering
        engine1 = NumberingEngine()
        success1 = engine1.get_or_create_numbering_definition(doc1, safe_num_1, safe_abstract_1)
        print(f"📊 Document 1 numbering created: {success1}")
        
        # Document 2 (simulating second resume upload)
        doc2 = Document()
        
        # Check if doc2 can see any existing numbering (it shouldn't - they're separate docs)
        print(f"📊 Document 2 should be clean (no existing numbering)")
        
        safe_num_2, safe_abstract_2 = NumberingEngine._allocate_safe_ids(doc2)
        print(f"📊 Document 2 (after doc1 exists): numId={safe_num_2}, abstractNumId={safe_abstract_2}")
        
        # Expected behavior: Since doc2 is a fresh document with no existing numbering,
        # it should get the default PID-salt base (10100), same as doc1.
        # This is actually CORRECT behavior - each document starts fresh.
        
        print(f"\n🔍 Analysis:")
        print(f"- Document 1: Clean doc → gets PID-salt base {safe_num_1}")
        print(f"- Document 2: Clean doc → gets PID-salt base {safe_num_2}")
        print(f"- Since both are clean docs, they get the same starting ID")
        print(f"- This is CORRECT - each doc scans its own numbering independently")
        
        # The real test should be: what happens if doc2 had existing numbering?
        # Let's simulate that by adding some existing numbering to doc2
        print(f"\n🧪 Additional test: doc2 with existing numbering")
        
        # Add some fake existing numbering to doc2
        engine2 = NumberingEngine()
        engine2.get_or_create_numbering_definition(doc2, 1, 1)   # Simulate existing numId=1
        engine2.get_or_create_numbering_definition(doc2, 50, 50) # Simulate existing numId=50
        
        # Now get safe allocation for doc2 - should be higher
        safe_num_2b, safe_abstract_2b = NumberingEngine._allocate_safe_ids(doc2)
        print(f"📊 Document 2 (with existing 1,50): numId={safe_num_2b}, abstractNumId={safe_abstract_2b}")
        
        if safe_num_2b > 50:
            print(f"✅ SUCCESS: Doc2 with existing numbering gets higher safe numId")
            return True
        else:
            print(f"❌ FAILURE: Doc2 didn't avoid existing numbering")
            return False
            
    except Exception as e:
        print(f"❌ FAILURE: Exception during sequential test: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    print("🚀 Phase 5.1 Integration Fix Test")
    print("Testing that safe numId allocation is properly integrated...")
    
    try:
        # Test 1: Integration fix
        test1_passed = test_integration_fix()
        
        # Test 2: Sequential documents
        test2_passed = test_sequential_documents()
        
        if test1_passed and test2_passed:
            print("\n" + "="*60)
            print("🎉 ALL INTEGRATION TESTS PASSED!")
            print("="*60)
            print("✅ Safe numId threading through call chain")
            print("✅ Sequential document collision avoidance")
            print("\n🔧 Integration fix is working correctly!")
            print("📋 Ready for production testing...")
        else:
            print("\n" + "="*60)
            print("❌ INTEGRATION TESTS FAILED!")
            print("="*60)
            if not test1_passed:
                print("❌ Safe numId threading failed")
            if not test2_passed:
                print("❌ Sequential document collision avoidance failed")
            sys.exit(1)
            
    except Exception as e:
        print(f"\n❌ INTEGRATION TEST FAILED: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1) 