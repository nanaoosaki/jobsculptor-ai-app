#!/usr/bin/env python3
"""
Simple Document Test - Baseline Validation

Create a minimal DOCX document to ensure our basic document creation works
before adding complex numbering XML.
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt

# Add the project root to Python path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))


def test_basic_document():
    """Create a basic document without numbering to verify baseline functionality."""
    print("🧪 Testing basic document creation...")
    
    try:
        doc = Document()
        
        # Add title
        title = doc.add_paragraph()
        title.add_run("Basic Document Test")
        title.runs[0].font.size = Pt(14)
        title.runs[0].bold = True
        
        # Add some paragraphs
        para1 = doc.add_paragraph()
        para1.add_run("This is a basic paragraph without any special formatting.")
        
        para2 = doc.add_paragraph()
        para2.add_run("This is another paragraph to test basic functionality.")
        
        # Save document
        output_path = "test_basic.docx"
        doc.save(output_path)
        print(f"✅ Basic document saved: {output_path}")
        print("   Please test: Open this file in Word to verify it works")
        
        return True
    except Exception as e:
        print(f"❌ FAIL: Could not create basic document: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_manual_bullet_legacy():
    """Create document with manual bullets (legacy approach) to verify that works."""
    print("\n🧪 Testing legacy manual bullets...")
    
    try:
        doc = Document()
        
        # Add title
        title = doc.add_paragraph()
        title.add_run("Legacy Manual Bullets Test")
        title.runs[0].font.size = Pt(14)
        title.runs[0].bold = True
        
        # Add manual bullets (the way we know works)
        bullets = [
            "Manual bullet one with basic formatting",
            "Manual bullet two with basic formatting",
            "Manual bullet three with basic formatting"
        ]
        
        for bullet_text in bullets:
            para = doc.add_paragraph()
            # Add bullet character manually
            para.add_run("• " + bullet_text)
        
        # Save document
        output_path = "test_manual_bullets.docx"
        doc.save(output_path)
        print(f"✅ Manual bullets document saved: {output_path}")
        print("   Please test: Open this file in Word to verify it works")
        
        return True
    except Exception as e:
        print(f"❌ FAIL: Could not create manual bullets document: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """Run simple validation tests."""
    print("🔧 SIMPLE DOCUMENT VALIDATION")
    print("=" * 40)
    
    tests = [
        ("Basic document creation", test_basic_document),
        ("Legacy manual bullets", test_manual_bullet_legacy),
    ]
    
    results = []
    for test_name, test_func in tests:
        try:
            result = test_func()
            results.append((test_name, result))
        except Exception as e:
            print(f"❌ CRITICAL ERROR in {test_name}: {e}")
            results.append((test_name, False))
    
    # Summary
    print("\n" + "=" * 40)
    print("📊 SIMPLE VALIDATION RESULTS")
    print("=" * 40)
    
    passed = 0
    for test_name, result in results:
        status = "✅ PASS" if result else "❌ FAIL"
        print(f"{status}: {test_name}")
        if result:
            passed += 1
    
    if passed == len(results):
        print("\n✅ BASELINE WORKING - Ready to debug numbering issue")
        return True
    else:
        print("\n❌ BASELINE BROKEN - Fix basic issues first")
        return False


if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1) 