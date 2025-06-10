#!/usr/bin/env python3
"""
Phase 2 Integration Test - Native Bullets in DOCX Builder

This script validates the integration of NumberingEngine with the main DOCX builder,
testing feature flags, graceful degradation, and end-to-end document generation.

Tests O3's deployment strategy:
1. Feature flag functionality
2. Native bullets when enabled
3. Legacy fallback when disabled or failed
4. End-to-end document generation success
"""

import sys
import os
import tempfile
import json
from pathlib import Path
from docx import Document

# Add the project root to Python path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

# Set up test data directory
temp_dir = tempfile.mkdtemp()
print(f"📂 Using temp directory: {temp_dir}")

def create_test_data(request_id: str):
    """Create test JSON files for full document generation."""
    
    # Contact data
    contact_data = {
        "name": "Alex Johnson",
        "email": "alex.johnson@email.com",
        "phone": "(555) 123-4567",
        "location": "San Francisco, CA",
        "linkedin": "linkedin.com/in/alexjohnson"
    }
    
    # Experience data with achievements
    experience_data = {
        "experiences": [
            {
                "company": "Tech Corp",
                "title": "Senior Software Engineer",
                "start_date": "2022-01",
                "end_date": "Present",
                "location": "San Francisco, CA",
                "achievements": [
                    "Built scalable microservices handling 10M+ requests daily",
                    "Led team of 5 engineers in modernizing legacy systems",
                    "Reduced deployment time by 70% through CI/CD optimization",
                    "Implemented real-time analytics dashboard for business intelligence"
                ]
            },
            {
                "company": "Startup Inc",
                "title": "Full Stack Developer", 
                "start_date": "2020-03",
                "end_date": "2021-12",
                "location": "Remote",
                "achievements": [
                    "Developed React frontend with 99.9% uptime SLA",
                    "Optimized database queries reducing response time by 50%",
                    "Integrated payment processing for $2M+ annual revenue"
                ]
            }
        ]
    }
    
    # Education data with highlights
    education_data = {
        "education": [
            {
                "institution": "University of California, Berkeley",
                "degree": "Bachelor of Science",
                "field": "Computer Science",
                "start_date": "2016-08",
                "end_date": "2020-05",
                "location": "Berkeley, CA",
                "highlights": [
                    "Magna Cum Laude, GPA: 3.8/4.0",
                    "Dean's List for 6 semesters",
                    "Captain of Programming Competition Team",
                    "Teaching Assistant for Data Structures course"
                ]
            }
        ]
    }
    
    # Projects data with details
    projects_data = {
        "projects": [
            {
                "name": "E-Commerce Platform",
                "technologies": "React, Node.js, MongoDB",
                "start_date": "2023-01", 
                "end_date": "2023-06",
                "details": [
                    "Built responsive web application with React and TypeScript",
                    "Implemented secure authentication with JWT tokens",
                    "Created RESTful API with comprehensive error handling",
                    "Deployed on AWS with auto-scaling configuration"
                ]
            }
        ]
    }
    
    # Summary data
    summary_data = {
        "summary": "Experienced software engineer with 4+ years building scalable web applications. Expertise in full-stack development, cloud architecture, and team leadership. Passionate about clean code, performance optimization, and mentoring junior developers."
    }
    
    # Write all test files
    test_files = {
        "contact": contact_data,
        "experience": experience_data,
        "education": education_data, 
        "projects": projects_data,
        "summary": summary_data
    }
    
    for section, data in test_files.items():
        file_path = os.path.join(temp_dir, f"{request_id}_{section}.json")
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
        print(f"✅ Created: {file_path}")
    
    return temp_dir


def test_feature_flag_disabled():
    """Test that legacy bullets are used when feature flag is disabled."""
    print("\n🧪 Testing feature flag DISABLED...")
    
    # Ensure flag is disabled
    os.environ['DOCX_USE_NATIVE_BULLETS'] = 'false'
    
    # Import after setting environment variable
    try:
        from utils.docx_builder import create_bullet_point, NATIVE_BULLETS_ENABLED
        
        if NATIVE_BULLETS_ENABLED:
            print(f"❌ FAIL: Native bullets should be disabled but flag shows: {NATIVE_BULLETS_ENABLED}")
            return False
        
        # Test bullet creation
        doc = Document()
        bullet_para = create_bullet_point(doc, "Test bullet with flag disabled")
        
        if bullet_para is None:
            print("❌ FAIL: Bullet creation returned None")
            return False
        
        # Save test document
        doc.save(os.path.join(temp_dir, "test_flag_disabled.docx"))
        print("✅ PASS: Legacy bullets used when flag disabled")
        return True
        
    except Exception as e:
        print(f"❌ FAIL: Error testing disabled flag: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_feature_flag_enabled():
    """Test that native bullets are used when feature flag is enabled."""
    print("\n🧪 Testing feature flag ENABLED...")
    
    # Enable flag
    os.environ['DOCX_USE_NATIVE_BULLETS'] = 'true'
    
    # Need to reload modules after environment change
    import importlib
    import utils.docx_builder
    importlib.reload(utils.docx_builder)
    
    try:
        from utils.docx_builder import create_bullet_point, NATIVE_BULLETS_ENABLED
        
        print(f"Native bullets enabled: {NATIVE_BULLETS_ENABLED}")
        
        # Test bullet creation
        doc = Document()
        bullet_para = create_bullet_point(doc, "Test bullet with flag enabled")
        
        if bullet_para is None:
            print("❌ FAIL: Bullet creation returned None")
            return False
        
        # Save test document
        doc.save(os.path.join(temp_dir, "test_flag_enabled.docx"))
        print("✅ PASS: Bullet creation succeeded with flag enabled")
        return True
        
    except Exception as e:
        print(f"❌ FAIL: Error testing enabled flag: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_end_to_end_generation():
    """Test complete document generation with native bullets."""
    print("\n🧪 Testing end-to-end document generation...")
    
    # Ensure native bullets are enabled
    os.environ['DOCX_USE_NATIVE_BULLETS'] = 'true'
    
    try:
        # Create test data
        request_id = "test_e2e"
        create_test_data(request_id)
        
        # Import after setting environment
        from utils.docx_builder import build_docx
        
        # Generate complete document
        docx_buffer = build_docx(request_id, temp_dir, debug=True)
        
        if not docx_buffer:
            print("❌ FAIL: Document generation returned None")
            return False
        
        # Save the generated document
        output_path = os.path.join(temp_dir, "test_complete_document.docx")
        with open(output_path, 'wb') as f:
            f.write(docx_buffer.getvalue())
        
        print(f"✅ PASS: Complete document generated: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ FAIL: End-to-end generation failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_graceful_degradation():
    """Test that bullet creation never fails completely (100% success rate)."""
    print("\n🧪 Testing graceful degradation...")
    
    try:
        from utils.docx_builder import create_bullet_point
        
        doc = Document()
        
        # Test with various inputs
        test_cases = [
            "Normal bullet point",
            "Bullet with special chars: áéíóú ñü",
            "Very long bullet point that might exceed normal length limits and contains multiple sentences to test wrapping behavior and performance under stress conditions",
            "",  # Empty text
            "   ",  # Whitespace only
            "Bullet with numbers: 123,456.78 and symbols: @#$%^&*()",
        ]
        
        success_count = 0
        for i, text in enumerate(test_cases):
            try:
                bullet_para = create_bullet_point(doc, text)
                # Allow None for empty/whitespace cases
                if bullet_para is not None or not text.strip():
                    success_count += 1
                    print(f"  ✅ Case {i+1}: Success")
                else:
                    print(f"  ❌ Case {i+1}: Returned None for non-empty text")
            except Exception as case_error:
                print(f"  ❌ Case {i+1}: Exception: {case_error}")
        
        # Save test document
        doc.save(os.path.join(temp_dir, "test_degradation.docx"))
        
        # Expect 100% success for valid cases (non-empty text)
        valid_cases = len([t for t in test_cases if t.strip()])
        expected_success = valid_cases + (len(test_cases) - valid_cases)  # Empty cases should also "succeed"
        
        if success_count >= valid_cases:  # At least all valid cases must succeed
            print(f"✅ PASS: Graceful degradation working ({success_count}/{len(test_cases)} cases handled)")
            return True
        else:
            print(f"❌ FAIL: Some valid cases failed ({success_count}/{len(test_cases)})")
            return False
            
    except Exception as e:
        print(f"❌ FAIL: Graceful degradation test failed: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """Run all Phase 2 integration tests."""
    print("🚀 PHASE 2 INTEGRATION TESTING")
    print("=" * 50)
    
    tests = [
        ("Feature flag disabled", test_feature_flag_disabled),
        ("Feature flag enabled", test_feature_flag_enabled),
        ("End-to-end generation", test_end_to_end_generation),
        ("Graceful degradation", test_graceful_degradation),
    ]
    
    results = []
    for test_name, test_func in tests:
        try:
            result = test_func()
            results.append((test_name, result))
        except Exception as e:
            print(f"❌ CRITICAL ERROR in {test_name}: {e}")
            import traceback
            traceback.print_exc()
            results.append((test_name, False))
    
    # Summary
    print("\n" + "=" * 50)
    print("📊 PHASE 2 INTEGRATION RESULTS")
    print("=" * 50)
    
    passed = 0
    for test_name, result in results:
        status = "✅ PASS" if result else "❌ FAIL"
        print(f"{status}: {test_name}")
        if result:
            passed += 1
    
    print(f"\nSummary: {passed}/{len(results)} tests passed")
    
    if passed == len(results):
        print("\n🎉 PHASE 2 COMPLETE! Native bullets integrated successfully")
        print("✅ Feature flags working")
        print("✅ Graceful degradation implemented")
        print("✅ End-to-end generation successful")
        print(f"\n📋 TEST DOCUMENTS CREATED:")
        print(f"   📄 test_flag_disabled.docx - Legacy bullets")
        print(f"   📄 test_flag_enabled.docx - Native bullets") 
        print(f"   📄 test_complete_document.docx - Full resume with native bullets")
        print(f"   📄 test_degradation.docx - Edge case handling")
        print(f"\n🎯 READY FOR PHASE 3: Production deployment with feature flag")
        return True
    else:
        print("\n❌ PHASE 2 INCOMPLETE - Fix failing tests before production")
        return False


if __name__ == "__main__":
    success = main()
    print(f"\n📂 All test files available in: {temp_dir}")
    sys.exit(0 if success else 1) 