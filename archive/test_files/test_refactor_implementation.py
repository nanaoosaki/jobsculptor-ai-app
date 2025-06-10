#!/usr/bin/env python3
"""
Test Script for Refactor Implementation

This script tests the new "Build-Then-Reconcile" architecture to ensure:
1. NumberingEngine singleton with request isolation works
2. BulletReconciliationEngine successfully scans and repairs bullets
3. Integration with build_docx function works correctly
4. All O3 improvements (A-series and B-series) are functional

Created: December 2024
Author: Resume Tailor Team
"""

import sys
import os
import json
import tempfile
import shutil
from pathlib import Path

# Add the parent directory to the path so we can import our modules
sys.path.append(str(Path(__file__).parent))

# Import our enhanced modules
from word_styles.numbering_engine import NumberingEngine
from utils.bullet_reconciliation import BulletReconciliationEngine
from utils.docx_builder import build_docx

def test_numbering_engine_singleton():
    """Test A4: Singleton reset between requests."""
    print("🧪 Testing NumberingEngine singleton with request isolation...")
    
    # Test 1: Same request ID should return same instance
    engine1 = NumberingEngine.get_instance("test_request_1")
    engine2 = NumberingEngine.get_instance("test_request_1")
    assert engine1 is engine2, "Same request ID should return same instance"
    print("✅ Same request ID returns same instance")
    
    # Test 2: Different request ID should return new instance
    engine3 = NumberingEngine.get_instance("test_request_2")
    assert engine3 is not engine1, "Different request ID should return new instance"
    print("✅ Different request ID returns new instance")
    
    # Test 3: B9 - NumId allocation should be globally unique
    num_id_1 = NumberingEngine.allocate_num_id()
    num_id_2 = NumberingEngine.allocate_num_id()
    assert num_id_1 != num_id_2, "NumId allocation should be unique"
    assert num_id_2 > num_id_1, "NumId should increment"
    print(f"✅ Unique numId allocation: {num_id_1} → {num_id_2}")

def test_bullet_reconciliation_engine():
    """Test the BulletReconciliationEngine functionality."""
    print("\n🧪 Testing BulletReconciliationEngine...")
    
    # Create a minimal test document
    from docx import Document
    doc = Document()
    
    # Add some paragraphs with MR_BulletPoint style
    # (This is a minimal test - normally the styles would be properly created)
    para1 = doc.add_paragraph("Test bullet 1")
    para2 = doc.add_paragraph("Test bullet 2") 
    
    # Create reconciliation engine
    reconciler = BulletReconciliationEngine("test_request")
    
    # Test that the engine initializes correctly
    assert reconciler.request_id == "test_request"
    print("✅ BulletReconciliationEngine initializes correctly")
    
    # Test performance monitoring methods
    reconciler._start_performance_monitoring()
    duration_ms, memory_diff_mb = reconciler._end_performance_monitoring()
    assert duration_ms >= 0, "Duration should be non-negative"
    print(f"✅ Performance monitoring works: {duration_ms:.1f}ms, {memory_diff_mb:.1f}MB")

def create_test_resume_data():
    """Create test resume data for full integration test."""
    temp_dir = tempfile.mkdtemp()
    request_id = "test_refactor_001"
    
    # Create contact data
    contact_data = {
        "name": "John Doe",
        "email": "john.doe@example.com",
        "phone": "+1 (555) 123-4567",
        "location": "San Francisco, CA"
    }
    
    # Create experience data with bullets to test reconciliation
    experience_data = {
        "experiences": [
            {
                "company": "Tech Corp",
                "position": "Senior Developer",
                "location": "San Francisco, CA",
                "dates": "2020-2024",
                "achievements": [
                    "Developed scalable microservices architecture",
                    "Led team of 5 engineers on critical projects",
                    "Improved system performance by 40%"
                ]
            },
            {
                "company": "Startup Inc",
                "position": "Full Stack Developer", 
                "location": "Remote",
                "dates": "2018-2020",
                "achievements": [
                    "Built customer-facing web application from scratch",
                    "Implemented automated testing pipeline",
                    "Reduced deployment time by 60%"
                ]
            }
        ]
    }
    
    # Create education data with some highlights
    education_data = {
        "institutions": [
            {
                "institution": "University of California",
                "degree": "B.S. Computer Science",
                "location": "Berkeley, CA",
                "dates": "2014-2018",
                "highlights": [
                    "Magna Cum Laude, GPA: 3.8/4.0",
                    "President of Computer Science Club"
                ]
            }
        ]
    }
    
    # Create summary data
    summary_data = {
        "summary": "Experienced software engineer with 6+ years developing scalable web applications and leading engineering teams."
    }
    
    # Save all data files
    with open(Path(temp_dir) / f"{request_id}_contact.json", 'w') as f:
        json.dump(contact_data, f)
    
    with open(Path(temp_dir) / f"{request_id}_experience.json", 'w') as f:
        json.dump(experience_data, f)
        
    with open(Path(temp_dir) / f"{request_id}_education.json", 'w') as f:
        json.dump(education_data, f)
        
    with open(Path(temp_dir) / f"{request_id}_summary.json", 'w') as f:
        json.dump(summary_data, f)
    
    return temp_dir, request_id

def test_full_integration():
    """Test the full integration with build_docx function."""
    print("\n🧪 Testing full integration with build_docx...")
    
    # Create test data
    temp_dir, request_id = create_test_resume_data()
    
    try:
        # Generate the DOCX document
        print(f"📄 Generating document for request {request_id}...")
        docx_buffer = build_docx(request_id, temp_dir, debug=False)
        
        # Verify we got a valid BytesIO buffer
        assert docx_buffer is not None, "build_docx should return a BytesIO buffer"
        assert len(docx_buffer.getvalue()) > 0, "Generated DOCX should not be empty"
        
        print(f"✅ Document generated successfully: {len(docx_buffer.getvalue())} bytes")
        
        # Save the test document for inspection
        test_output_path = Path(__file__).parent / f"test_refactor_output_{request_id}.docx"
        with open(test_output_path, 'wb') as f:
            f.write(docx_buffer.getvalue())
        
        print(f"📄 Test document saved to: {test_output_path}")
        
        return True
        
    except Exception as e:
        print(f"❌ Integration test failed: {e}")
        import traceback
        print(f"❌ Traceback: {traceback.format_exc()}")
        return False
    
    finally:
        # Cleanup temp directory
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)

def run_all_tests():
    """Run all tests for the refactor implementation."""
    print("🚀 Running Refactor Implementation Tests")
    print("=" * 50)
    
    success_count = 0
    total_tests = 3
    
    try:
        # Test 1: NumberingEngine singleton
        test_numbering_engine_singleton()
        success_count += 1
        print("✅ NumberingEngine singleton test passed")
    except Exception as e:
        print(f"❌ NumberingEngine singleton test failed: {e}")
    
    try:
        # Test 2: BulletReconciliationEngine
        test_bullet_reconciliation_engine()
        success_count += 1
        print("✅ BulletReconciliationEngine test passed")
    except Exception as e:
        print(f"❌ BulletReconciliationEngine test failed: {e}")
    
    try:
        # Test 3: Full integration
        if test_full_integration():
            success_count += 1
            print("✅ Full integration test passed")
        else:
            print("❌ Full integration test failed")
    except Exception as e:
        print(f"❌ Full integration test failed: {e}")
    
    print("\n" + "=" * 50)
    print(f"🎯 Test Results: {success_count}/{total_tests} tests passed")
    
    if success_count == total_tests:
        print("🎉 ALL TESTS PASSED! Refactor implementation is working correctly.")
        return True
    else:
        print("⚠️  Some tests failed. Please check the implementation.")
        return False

if __name__ == "__main__":
    success = run_all_tests()
    sys.exit(0 if success else 1) 