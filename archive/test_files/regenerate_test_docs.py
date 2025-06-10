#!/usr/bin/env python3
"""
Regenerate Test Documents - Phase 2 Native Bullets

This script regenerates the test documents from Phase 2 integration testing
and saves them in the current directory for easy access and validation.
"""

import os
import sys
import tempfile
import json
from pathlib import Path
from docx import Document

# Add the project root to Python path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

def create_test_data_files(temp_dir: str, request_id: str):
    """Create test JSON files for complete document generation."""
    
    # Contact data
    contact_data = {
        "name": "Alex Johnson",
        "email": "alex.johnson@email.com",
        "phone": "(555) 123-4567",
        "location": "San Francisco, CA",
        "linkedin": "linkedin.com/in/alexjohnson"
    }
    
    # Experience data with achievements
    experience_data = {
        "experiences": [
            {
                "company": "Tech Corp",
                "title": "Senior Software Engineer",
                "start_date": "2022-01",
                "end_date": "Present",
                "location": "San Francisco, CA",
                "achievements": [
                    "Built scalable microservices handling 10M+ requests daily",
                    "Led team of 5 engineers in modernizing legacy systems", 
                    "Reduced deployment time by 70% through CI/CD optimization",
                    "Implemented real-time analytics dashboard for business intelligence"
                ]
            },
            {
                "company": "Startup Inc",
                "title": "Full Stack Developer",
                "start_date": "2020-03", 
                "end_date": "2021-12",
                "location": "Remote",
                "achievements": [
                    "Developed React frontend with 99.9% uptime SLA",
                    "Optimized database queries reducing response time by 50%",
                    "Integrated payment processing for $2M+ annual revenue"
                ]
            }
        ]
    }
    
    # Education data
    education_data = {
        "education": [
            {
                "institution": "University of California, Berkeley",
                "degree": "Bachelor of Science",
                "field": "Computer Science",
                "start_date": "2016-08",
                "end_date": "2020-05", 
                "location": "Berkeley, CA",
                "highlights": [
                    "Magna Cum Laude, GPA: 3.8/4.0",
                    "Dean's List for 6 semesters",
                    "Captain of Programming Competition Team",
                    "Teaching Assistant for Data Structures course"
                ]
            }
        ]
    }
    
    # Projects data
    projects_data = {
        "projects": [
            {
                "name": "E-Commerce Platform",
                "technologies": "React, Node.js, MongoDB",
                "start_date": "2023-01",
                "end_date": "2023-06",
                "details": [
                    "Built responsive web application with React and TypeScript",
                    "Implemented secure authentication with JWT tokens", 
                    "Created RESTful API with comprehensive error handling",
                    "Deployed on AWS with auto-scaling configuration"
                ]
            }
        ]
    }
    
    # Summary data
    summary_data = {
        "summary": "Experienced software engineer with 4+ years building scalable web applications. Expertise in full-stack development, cloud architecture, and team leadership. Passionate about clean code, performance optimization, and mentoring junior developers."
    }
    
    # Write all test files
    test_files = {
        "contact": contact_data,
        "experience": experience_data,
        "education": education_data,
        "projects": projects_data,
        "summary": summary_data
    }
    
    for section, data in test_files.items():
        file_path = os.path.join(temp_dir, f"{request_id}_{section}.json")
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
        print(f"✅ Created: {file_path}")
    
    return temp_dir

def generate_flag_disabled_test():
    """Generate test document with native bullets disabled."""
    print("\n📄 Generating test_flag_disabled.docx...")
    
    # Ensure flag is disabled
    os.environ['DOCX_USE_NATIVE_BULLETS'] = 'false'
    
    try:
        from utils.docx_builder import create_bullet_point, NATIVE_BULLETS_ENABLED
        
        print(f"   Native bullets enabled: {NATIVE_BULLETS_ENABLED}")
        
        # Create test document
        doc = Document()
        
        # Add title
        title = doc.add_paragraph()
        title.add_run("Native Bullets Feature Flag Test - DISABLED")
        title.runs[0].font.size = Pt(14)
        title.runs[0].bold = True
        
        # Add explanation
        doc.add_paragraph("This document tests bullet creation when the DOCX_USE_NATIVE_BULLETS flag is disabled. Should use legacy bullet approach.")
        
        # Add test bullets
        bullets = [
            "Legacy bullet point one - manual formatting",
            "Legacy bullet point two - style-based only", 
            "Legacy bullet point three - no native numbering"
        ]
        
        for bullet_text in bullets:
            bullet_para = create_bullet_point(doc, bullet_text)
        
        # Save document
        output_path = "test_flag_disabled.docx"
        doc.save(output_path)
        print(f"   ✅ Saved: {output_path}")
        return True
        
    except Exception as e:
        print(f"   ❌ Error: {e}")
        return False

def generate_flag_enabled_test():
    """Generate test document with native bullets enabled."""
    print("\n📄 Generating test_flag_enabled.docx...")
    
    # Enable flag
    os.environ['DOCX_USE_NATIVE_BULLETS'] = 'true'
    
    # Reload module to pick up environment change
    import importlib
    import utils.docx_builder
    importlib.reload(utils.docx_builder)
    
    try:
        from utils.docx_builder import create_bullet_point, NATIVE_BULLETS_ENABLED
        
        print(f"   Native bullets enabled: {NATIVE_BULLETS_ENABLED}")
        
        # Create test document
        doc = Document()
        
        # Add title
        title = doc.add_paragraph()
        title.add_run("Native Bullets Feature Flag Test - ENABLED")
        title.runs[0].font.size = Pt(14)
        title.runs[0].bold = True
        
        # Add explanation
        doc.add_paragraph("This document tests bullet creation when the DOCX_USE_NATIVE_BULLETS flag is enabled. Should use Word's native numbering system.")
        
        # Add test bullets
        bullets = [
            "Native bullet point one - uses Word numbering engine",
            "Native bullet point two - proper bullet continuation", 
            "Native bullet point three - professional Word behavior"
        ]
        
        for bullet_text in bullets:
            bullet_para = create_bullet_point(doc, bullet_text)
        
        # Save document
        output_path = "test_flag_enabled.docx" 
        doc.save(output_path)
        print(f"   ✅ Saved: {output_path}")
        return True
        
    except Exception as e:
        print(f"   ❌ Error: {e}")
        return False

def generate_complete_document_test():
    """Generate complete resume document with native bullets."""
    print("\n📄 Generating test_complete_document.docx...")
    
    # Ensure native bullets are enabled
    os.environ['DOCX_USE_NATIVE_BULLETS'] = 'true'
    
    try:
        # Create temporary directory and test data
        temp_dir = tempfile.mkdtemp()
        request_id = "test_regen"
        create_test_data_files(temp_dir, request_id)
        
        # Import and generate document
        from utils.docx_builder import build_docx
        
        # Generate complete document
        docx_buffer = build_docx(request_id, temp_dir, debug=False)
        
        # Save the generated document
        output_path = "test_complete_document.docx"
        with open(output_path, 'wb') as f:
            f.write(docx_buffer.getvalue())
        
        print(f"   ✅ Saved: {output_path}")
        
        # Cleanup temp directory
        import shutil
        shutil.rmtree(temp_dir)
        
        return True
        
    except Exception as e:
        print(f"   ❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False

def generate_degradation_test():
    """Generate test document for graceful degradation testing."""
    print("\n📄 Generating test_degradation.docx...")
    
    try:
        from utils.docx_builder import create_bullet_point
        
        # Create test document
        doc = Document()
        
        # Add title
        title = doc.add_paragraph()
        title.add_run("Graceful Degradation Test")
        title.runs[0].font.size = Pt(14)
        title.runs[0].bold = True
        
        # Add explanation
        doc.add_paragraph("This document tests bullet creation with various edge cases to ensure graceful degradation.")
        
        # Test with various edge cases
        test_cases = [
            "Normal bullet point",
            "Bullet with special chars: áéíóú ñü", 
            "Very long bullet point that might exceed normal length limits and contains multiple sentences to test wrapping behavior and performance under stress conditions",
            "Bullet with numbers: 123,456.78 and symbols: @#$%^&*()",
            "Bullet with emoji: 🚀 Native bullets are awesome! 🎉",
            "Mixed content: Code snippets `console.log('hello')` and **bold** text"
        ]
        
        for i, text in enumerate(test_cases):
            doc.add_paragraph(f"Test Case {i+1}:")
            bullet_para = create_bullet_point(doc, text)
            doc.add_paragraph("")  # Add spacing
        
        # Save document
        output_path = "test_degradation.docx"
        doc.save(output_path)
        print(f"   ✅ Saved: {output_path}")
        return True
        
    except Exception as e:
        print(f"   ❌ Error: {e}")
        return False

def main():
    """Regenerate all Phase 2 test documents."""
    print("🚀 REGENERATING PHASE 2 TEST DOCUMENTS")
    print("=" * 50)
    
    # Import Pt for document formatting
    from docx.shared import Pt
    
    tests = [
        ("Flag Disabled Test", generate_flag_disabled_test),
        ("Flag Enabled Test", generate_flag_enabled_test),
        ("Complete Document Test", generate_complete_document_test),
        ("Graceful Degradation Test", generate_degradation_test),
    ]
    
    results = []
    for test_name, test_func in tests:
        try:
            result = test_func()
            results.append((test_name, result))
        except Exception as e:
            print(f"❌ CRITICAL ERROR in {test_name}: {e}")
            results.append((test_name, False))
    
    # Summary
    print("\n" + "=" * 50)
    print("📊 DOCUMENT GENERATION RESULTS")
    print("=" * 50)
    
    passed = 0
    for test_name, result in results:
        status = "✅ SUCCESS" if result else "❌ FAILED"
        print(f"{status}: {test_name}")
        if result:
            passed += 1
    
    print(f"\nSummary: {passed}/{len(results)} documents generated successfully")
    
    if passed == len(results):
        print("\n🎉 ALL TEST DOCUMENTS REGENERATED!")
        print("📋 Available in current directory:")
        print("   📄 test_flag_disabled.docx - Legacy bullets")
        print("   📄 test_flag_enabled.docx - Native bullets")
        print("   📄 test_complete_document.docx - Full resume with native bullets")
        print("   📄 test_degradation.docx - Edge case handling")
        print("\n✨ Ready for testing and validation!")
        return True
    else:
        print("\n❌ Some documents failed to generate")
        return False

if __name__ == "__main__":
    main() 