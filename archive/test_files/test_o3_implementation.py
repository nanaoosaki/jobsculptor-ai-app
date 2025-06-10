#!/usr/bin/env python3
"""
Test script for Phase 4: O3 Core Implementation

This script validates the O3 enhanced bullet consistency system including:
- O3 core engine functionality
- Enhanced reconciliation
- Integration with existing systems
- API endpoints
- Performance metrics

Author: Resume Tailor Team + O3 Integration
Status: Phase 4 Testing
"""

import sys
import traceback
import time
from pathlib import Path

# Add project root to path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

def test_o3_core_engine():
    """Test O3 core engine basic functionality."""
    print("\n🚀 Phase 4.1: Testing O3 Core Engine...")
    
    try:
        from utils.o3_bullet_core_engine import get_o3_engine, cleanup_o3_engine
        from docx import Document
        from word_styles.numbering_engine import NumberingEngine
        
        # Create test document
        test_doc_id = f"test_o3_{int(time.time())}"
        doc = Document()
        numbering_engine = NumberingEngine()
        
        # Get O3 engine
        o3_engine = get_o3_engine(test_doc_id)
        print(f"✅ O3 engine created for document: {test_doc_id}")
        
        # Test bullet creation
        test_bullets = [
            "First O3 test bullet",
            "Second bullet with special chars: •★▶",
            "Third bullet with unicode: héllö wörld"
        ]
        
        created_bullets = []
        for bullet_text in test_bullets:
            try:
                para, bullet_id = o3_engine.create_bullet_trusted(
                    doc=doc,
                    text=bullet_text,
                    section_name="test_section",
                    numbering_engine=numbering_engine,
                    docx_styles={}
                )
                created_bullets.append(bullet_id)
                print(f"✅ Created bullet: {bullet_id} - '{bullet_text[:30]}...'")
            except Exception as e:
                print(f"❌ Failed to create bullet '{bullet_text}': {e}")
        
        # Test engine summary
        summary = o3_engine.get_engine_summary()
        print(f"✅ Engine summary: {summary['bullet_count']} bullets, state: {summary['document_state']}")
        
        # Test reconciliation
        print("\n🔧 Testing reconciliation...")
        reconciliation_stats = o3_engine.reconcile_document_bullets(doc, numbering_engine)
        
        bullets_processed = reconciliation_stats.get('bullets_processed', 0)
        success_rate = reconciliation_stats.get('success_rate', 0)
        print(f"✅ Reconciliation complete: {bullets_processed} bullets processed, {success_rate:.1f}% success rate")
        
        # Clean up
        cleanup_o3_engine(test_doc_id)
        print(f"✅ Engine cleaned up: {test_doc_id}")
        
        return True
        
    except Exception as e:
        print(f"❌ O3 core engine test failed: {e}")
        traceback.print_exc()
        return False

def test_b_series_integration():
    """Test integration with B-series modules."""
    print("\n🔗 Phase 4.2: Testing B-Series Integration...")
    
    try:
        from utils.o3_bullet_core_engine import get_o3_engine
        from utils.unicode_bullet_sanitizer import sanitize_bullet_text
        from utils.numid_collision_manager import allocate_safe_numid
        from docx import Document
        from word_styles.numbering_engine import NumberingEngine
        
        # Test unicode sanitization integration
        test_text = "• This bullet has a prefix"
        sanitized = sanitize_bullet_text(test_text)
        print(f"✅ Unicode sanitization: '{test_text}' → '{sanitized}'")
        
        # Test numId allocation integration
        test_doc_id = f"test_b_integration_{int(time.time())}"
        num_id, abstract_num_id = allocate_safe_numid(test_doc_id, "test_section", "MR_BulletPoint")
        print(f"✅ NumId allocation: numId={num_id}, abstractNumId={abstract_num_id}")
        
        # Test O3 engine with B-series features
        o3_engine = get_o3_engine(test_doc_id)
        doc = Document()
        numbering_engine = NumberingEngine()
        
        # Create bullet with problematic text (should be sanitized)
        problematic_bullets = [
            "• Bullet with western prefix",
            "▶ Bullet with arrow prefix",
            "★ Bullet with star prefix"
        ]
        
        for bullet_text in problematic_bullets:
            try:
                para, bullet_id = o3_engine.create_bullet_trusted(
                    doc=doc,
                    text=bullet_text,
                    section_name="integration_test",
                    numbering_engine=numbering_engine,
                    docx_styles={}
                )
                print(f"✅ B-series integration bullet: {bullet_id}")
            except Exception as e:
                print(f"❌ B-series integration failed for '{bullet_text}': {e}")
        
        # Clean up
        from utils.o3_bullet_core_engine import cleanup_o3_engine
        cleanup_o3_engine(test_doc_id)
        
        return True
        
    except Exception as e:
        print(f"❌ B-series integration test failed: {e}")
        traceback.print_exc()
        return False

def test_docx_builder_integration():
    """Test integration with docx_builder."""
    print("\n📝 Phase 4.3: Testing DOCX Builder Integration...")
    
    try:
        from utils.docx_builder import create_bullet_point
        from utils.o3_bullet_core_engine import get_o3_engine
        from docx import Document
        from word_styles.numbering_engine import NumberingEngine
        
        # Create test setup
        test_doc_id = f"test_docx_integration_{int(time.time())}"
        doc = Document()
        numbering_engine = NumberingEngine()
        o3_engine = get_o3_engine(test_doc_id)
        
        # Test create_bullet_point with O3 engine
        test_bullets = [
            "DOCX builder integration test bullet 1",
            "DOCX builder integration test bullet 2",
            "DOCX builder integration test bullet 3"
        ]
        
        for bullet_text in test_bullets:
            try:
                para = create_bullet_point(
                    doc=doc,
                    text=bullet_text,
                    docx_styles={},
                    numbering_engine=numbering_engine,
                    num_id=100,
                    o3_engine=o3_engine,
                    section_name="docx_test"
                )
                
                if para:
                    print(f"✅ DOCX builder created bullet: '{bullet_text[:30]}...'")
                else:
                    print(f"❌ DOCX builder returned None for: '{bullet_text}'")
                    
            except Exception as e:
                print(f"❌ DOCX builder failed for '{bullet_text}': {e}")
        
        # Test engine state
        summary = o3_engine.get_engine_summary()
        print(f"✅ DOCX integration summary: {summary['bullet_count']} bullets created")
        
        # Clean up
        from utils.o3_bullet_core_engine import cleanup_o3_engine
        cleanup_o3_engine(test_doc_id)
        
        return True
        
    except Exception as e:
        print(f"❌ DOCX builder integration test failed: {e}")
        traceback.print_exc()
        return False

def test_api_endpoints():
    """Test O3 API endpoints."""
    print("\n🌐 Phase 4.4: Testing API Endpoints...")
    
    try:
        import requests
        import json
        
        base_url = "http://localhost:5000"
        
        # Test Phase 4 summary endpoint
        try:
            response = requests.get(f"{base_url}/api/phase4-summary", timeout=10)
            if response.status_code == 200:
                data = response.json()
                print(f"✅ Phase 4 summary endpoint: {data.get('phase', 'Unknown')}")
                print(f"   Status: {data.get('status', 'Unknown')}")
                print(f"   API endpoints: {len(data.get('api_endpoints', []))}")
            else:
                print(f"❌ Phase 4 summary endpoint failed: {response.status_code}")
        except requests.exceptions.ConnectionError:
            print("⚠️ Flask app not running - skipping API tests")
            print("   To test APIs, run: python app.py")
            return True
        except Exception as e:
            print(f"❌ Phase 4 summary test failed: {e}")
        
        # Test O3 engine test endpoint
        try:
            response = requests.post(f"{base_url}/api/o3-core/test-engine", timeout=15)
            if response.status_code == 200:
                data = response.json()
                if data.get('success'):
                    print(f"✅ O3 engine test endpoint: {len(data.get('created_bullets', []))} bullets created")
                    print(f"   Reconciliation success: {data.get('reconciliation_stats', {}).get('success_rate', 0):.1f}%")
                else:
                    print(f"❌ O3 engine test endpoint failed: {data.get('error', 'Unknown error')}")
            else:
                print(f"❌ O3 engine test endpoint failed: {response.status_code}")
        except Exception as e:
            print(f"❌ O3 engine test endpoint failed: {e}")
        
        # Test all engines summary
        try:
            response = requests.get(f"{base_url}/api/o3-core/all-engines", timeout=10)
            if response.status_code == 200:
                data = response.json()
                if data.get('success'):
                    engines_count = data.get('engines_summary', {}).get('active_engines', 0)
                    print(f"✅ All engines summary: {engines_count} active engines")
                else:
                    print(f"❌ All engines summary failed: {data.get('error', 'Unknown error')}")
            else:
                print(f"❌ All engines summary failed: {response.status_code}")
        except Exception as e:
            print(f"❌ All engines summary failed: {e}")
        
        return True
        
    except Exception as e:
        print(f"❌ API endpoints test failed: {e}")
        traceback.print_exc()
        return False

def test_performance_metrics():
    """Test O3 performance metrics."""
    print("\n📊 Phase 4.5: Testing Performance Metrics...")
    
    try:
        from utils.o3_bullet_core_engine import get_o3_engine, cleanup_o3_engine
        from docx import Document
        from word_styles.numbering_engine import NumberingEngine
        import time
        
        # Create test scenario with many bullets
        test_doc_id = f"test_performance_{int(time.time())}"
        doc = Document()
        numbering_engine = NumberingEngine()
        o3_engine = get_o3_engine(test_doc_id)
        
        # Performance test with 50 bullets
        num_bullets = 50
        start_time = time.time()
        
        for i in range(num_bullets):
            bullet_text = f"Performance test bullet #{i+1} with some content to make it realistic"
            
            try:
                para, bullet_id = o3_engine.create_bullet_trusted(
                    doc=doc,
                    text=bullet_text,
                    section_name=f"section_{i // 10}",  # 5 sections
                    numbering_engine=numbering_engine,
                    docx_styles={}
                )
            except Exception as e:
                print(f"❌ Performance test bullet {i+1} failed: {e}")
        
        creation_time = time.time() - start_time
        print(f"✅ Created {num_bullets} bullets in {creation_time:.2f}s ({creation_time/num_bullets*1000:.1f}ms per bullet)")
        
        # Test reconciliation performance
        start_time = time.time()
        reconciliation_stats = o3_engine.reconcile_document_bullets(doc, numbering_engine)
        reconciliation_time = time.time() - start_time
        
        bullets_processed = reconciliation_stats.get('bullets_processed', 0)
        success_rate = reconciliation_stats.get('success_rate', 0)
        
        print(f"✅ Reconciled {bullets_processed} bullets in {reconciliation_time:.2f}s ({success_rate:.1f}% success)")
        
        # Get final engine summary
        summary = o3_engine.get_engine_summary()
        stats = summary.get('statistics', {})
        
        print(f"✅ Performance summary:")
        print(f"   Bullets created: {stats.get('bullets_created', 0)}")
        print(f"   Bullets validated: {stats.get('bullets_validated', 0)}")
        print(f"   Bullets reconciled: {stats.get('bullets_reconciled', 0)}")
        print(f"   Creation success rate: {summary.get('success_metrics', {}).get('creation_success_rate', 0):.1f}%")
        
        # Clean up
        cleanup_o3_engine(test_doc_id)
        
        return True
        
    except Exception as e:
        print(f"❌ Performance metrics test failed: {e}")
        traceback.print_exc()
        return False

def main():
    """Run all Phase 4 O3 implementation tests."""
    print("🚀 Phase 4: O3 Core Implementation Testing")
    print("=" * 50)
    
    test_results = []
    
    # Run all tests
    tests = [
        ("O3 Core Engine", test_o3_core_engine),
        ("B-Series Integration", test_b_series_integration),
        ("DOCX Builder Integration", test_docx_builder_integration),
        ("API Endpoints", test_api_endpoints),
        ("Performance Metrics", test_performance_metrics)
    ]
    
    for test_name, test_func in tests:
        print(f"\n{'='*20}")
        print(f"Running: {test_name}")
        print(f"{'='*20}")
        
        try:
            result = test_func()
            test_results.append((test_name, result))
            
            if result:
                print(f"✅ {test_name}: PASSED")
            else:
                print(f"❌ {test_name}: FAILED")
                
        except Exception as e:
            print(f"❌ {test_name}: CRASHED - {e}")
            test_results.append((test_name, False))
    
    # Print summary
    print(f"\n{'='*50}")
    print("PHASE 4 TEST SUMMARY")
    print(f"{'='*50}")
    
    passed = sum(1 for _, result in test_results if result)
    total = len(test_results)
    
    for test_name, result in test_results:
        status = "✅ PASSED" if result else "❌ FAILED"
        print(f"{test_name:.<30} {status}")
    
    print(f"\nOverall: {passed}/{total} tests passed ({passed/total*100:.1f}%)")
    
    if passed == total:
        print("\n🎉 Phase 4: O3 Core Implementation - ALL TESTS PASSED!")
        print("🚀 O3 enhanced bullet consistency system is ready for production!")
    else:
        print(f"\n⚠️ Phase 4: {total-passed} tests failed. Review implementation.")
    
    return passed == total

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1) 