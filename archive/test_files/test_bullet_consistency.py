#!/usr/bin/env python3
"""
Test bullet consistency issue - isolate which bullets fail
"""

import os
import tempfile
import json
from docx import Document

# Set environment for native bullets
os.environ['DOCX_USE_NATIVE_BULLETS'] = 'true'

from utils.docx_builder import build_docx, create_bullet_point
from style_manager import StyleManager
from style_engine import StyleEngine


def create_test_data():
    """Create test data with multiple bullets to test consistency."""
    temp_dir = tempfile.mkdtemp(prefix="bullet_test_")
    request_id = "consistency_test"
    
    # Create complete test data that matches production complexity
    experience_data = {
        "experiences": [
            {
                "company": "Integration Test Corp",
                "position": "Senior Test Engineer",
                "startDate": "2023",
                "endDate": "2024",
                "location": "Remote",
                "achievements": [
                    "First achievement that should trigger native bullets with proper 0.13 inch indentation",
                    "Second achievement to test multiple bullet creation and potential legacy fallbacks",
                    "Third achievement with complex formatting to stress-test the numbering system",
                    "Fourth achievement that might trigger different code paths in production",
                    "Fifth achievement to test style repair loops and second-pass formatting"
                ]
            }
        ]
    }
    
    # Add education data with bullets
    education_data = {
        "institutions": [
            {
                "institution": "Test University",
                "degree": "BS in Testing",
                "dates": "2018-2022",
                "location": "Test City",
                "highlights": [
                    "Education bullet that tests numbering in different sections",
                    "Another education bullet to verify consistent formatting"
                ]
            }
        ]
    }
    
    # Add contact data to match production
    contact_data = {
        "name": "Test Person",
        "location": "Test City, TC",
        "phone": "555-1234",
        "email": "test@example.com"
    }
    
    # Write all the test files
    with open(os.path.join(temp_dir, f"{request_id}_experience.json"), 'w') as f:
        json.dump(experience_data, f)
    
    with open(os.path.join(temp_dir, f"{request_id}_education.json"), 'w') as f:
        json.dump(education_data, f)
        
    with open(os.path.join(temp_dir, f"{request_id}_contact.json"), 'w') as f:
        json.dump(contact_data, f)
    
    return temp_dir, request_id


def test_individual_bullet_creation():
    """Test creating individual bullets to see which ones fail."""
    print("🧪 Testing individual bullet creation...")
    
    # Create a fresh document
    doc = Document()
    docx_styles = StyleManager.load_docx_styles()
    
    # Create styles
    StyleEngine.create_docx_custom_styles(doc)
    
    # Test bullet texts
    bullet_texts = [
        "First bullet point to test consistency",
        "Second bullet point to test consistency", 
        "Third bullet point to test consistency",
        "Fourth bullet point to test consistency",
        "Fifth bullet point to test consistency"
    ]
    
    print(f"📝 Testing {len(bullet_texts)} bullets individually...")
    
    success_count = 0
    failure_count = 0
    
    for i, text in enumerate(bullet_texts, 1):
        try:
            print(f"\n🔫 Creating bullet {i}: '{text[:30]}...'")
            para = create_bullet_point(doc, text, docx_styles)
            
            if para and para._element.xpath('.//w:numPr'):
                print(f"✅ Bullet {i} SUCCESS: Has numbering")
                success_count += 1
            else:
                print(f"❌ Bullet {i} FAILED: No numbering found")
                failure_count += 1
                
        except Exception as e:
            print(f"❌ Bullet {i} ERROR: {e}")
            failure_count += 1
    
    print(f"\n📊 Results:")
    print(f"   ✅ Successful bullets: {success_count}")
    print(f"   ❌ Failed bullets: {failure_count}")
    print(f"   📈 Success rate: {success_count/(success_count+failure_count)*100:.1f}%")
    
    # Save document for inspection
    doc.save("bullet_consistency_test.docx")
    print(f"💾 Saved test document: bullet_consistency_test.docx")


def test_full_document_generation():
    """Test full document generation to replicate the production issue."""
    print("\n🧪 Testing full document generation...")
    
    temp_dir, request_id = create_test_data()
    print(f"📁 Created test data in: {temp_dir}")
    
    try:
        # Generate document using production build_docx
        doc_bytes = build_docx(request_id, temp_dir)
        
        # Save the document  
        output_file = "bullet_consistency_full_test.docx"
        with open(output_file, 'wb') as f:
            f.write(doc_bytes.getvalue())
        
        print(f"💾 Saved full test document: {output_file}")
        
        # Analyze the document
        from docx import Document
        doc = Document(output_file)
        
        bullet_count = 0
        numbered_count = 0
        
        for para in doc.paragraphs:
            if para.text and ("bullet" in para.text.lower() or "achievement" in para.text.lower()):
                bullet_count += 1
                has_numbering = bool(para._element.xpath('.//w:numPr'))
                print(f"📋 Bullet {bullet_count}: '{para.text[:40]}...' - Numbering: {has_numbering}")
                if has_numbering:
                    numbered_count += 1
        
        print(f"\n📊 Full Document Results:")
        print(f"   🔫 Total bullets found: {bullet_count}")
        print(f"   ✅ Bullets with numbering: {numbered_count}")
        print(f"   ❌ Bullets missing numbering: {bullet_count - numbered_count}")
        
    except Exception as e:
        print(f"❌ Full document generation failed: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    print("🎯 BULLET CONSISTENCY DIAGNOSTIC TEST")
    print("=" * 50)
    
    # Test 1: Individual bullet creation
    test_individual_bullet_creation()
    
    # Test 2: Full document generation
    test_full_document_generation()
    
    print("\n🎉 Testing complete!") 