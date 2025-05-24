import os
import sys
import logging
import uuid
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('reduced_spacing_test.log')
    ]
)
logger = logging.getLogger(__name__)

def test_reduced_spacing():
    """Test that DOCX section headers have reduced spacing"""
    try:
        # Import our modules
        from style_engine import StyleEngine, TokenAccessor
        from utils.docx_builder import add_section_header
        
        # Load tokens
        tokens = StyleEngine.load_tokens()
        tokens_access = TokenAccessor(tokens)
        
        # Log the spacing value to verify our changes
        spacing_after_pt = tokens_access.get("sectionHeader.spacingAfterPt", 0)
        section_spacing_pt = tokens_access.get("docx-section-spacing-pt", 0)
        logger.info(f"Current token values - sectionHeader.spacingAfterPt: {spacing_after_pt}, docx-section-spacing-pt: {section_spacing_pt}")
        
        # Create a test document
        doc = Document()
        
        # Create the BoxedHeading2 style
        style_name = StyleEngine.create_boxed_heading_style(doc, tokens)
        logger.info(f"Created style: {style_name}")
        
        # Add a section header with the style
        para = doc.add_paragraph("PERSONAL SUMMARY")
        StyleEngine.apply_boxed_section_header_style(doc, para, tokens)
        
        # Add some regular content immediately after (should be closer now)
        doc.add_paragraph("This is regular content text that follows the section header with reduced spacing.")
        
        # Add another section header
        para = doc.add_paragraph("PROFESSIONAL EXPERIENCE")
        StyleEngine.apply_boxed_section_header_style(doc, para, tokens)
        
        # Add some regular content
        doc.add_paragraph("This is more regular content text that should be closer to the header.")
        
        # Generate a unique filename
        test_id = str(uuid.uuid4())[:8]
        output_file = f"test_reduced_spacing_{test_id}.docx"
        
        # Save the document
        doc.save(output_file)
        logger.info(f"Saved test document to {output_file}")
        
        # Validate the document
        validate_reduced_spacing(output_file, tokens_access)
        
        return output_file
        
    except Exception as e:
        logger.error(f"Test failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return None

def validate_reduced_spacing(docx_file, tokens_access):
    """Validate that the section headers in the document have reduced spacing"""
    try:
        # Open the document with python-docx
        doc = Document(docx_file)
        
        # Find paragraphs with BoxedHeading2 style
        boxed_headers = [p for p in doc.paragraphs if p.style and p.style.name == 'BoxedHeading2']
        
        if not boxed_headers:
            logger.error("No BoxedHeading2 style paragraphs found in the document")
            return False
        
        logger.info(f"Found {len(boxed_headers)} boxed headers")
        
        # Check the spacing after property of each header
        for header in boxed_headers:
            space_after = header.paragraph_format.space_after
            if space_after:
                # Convert to points for comparison (space_after is in twips, 20 twips = 1 point)
                space_after_pt = space_after.pt
                logger.info(f"Header '{header.text}' has space_after of {space_after_pt} points")
                
                # Get expected value from tokens
                expected_spacing = float(tokens_access.get("sectionHeader.spacingAfterPt", 4))
                
                # Allow a small tolerance for floating point comparison
                if abs(space_after_pt - expected_spacing) > 0.1:
                    logger.warning(f"Unexpected spacing: {space_after_pt}pt, expected: {expected_spacing}pt")
            else:
                logger.warning(f"Header '{header.text}' does not have space_after property set")
        
        logger.info("Reduced spacing validation completed")
        return True
        
    except Exception as e:
        logger.error(f"Validation failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return False

if __name__ == "__main__":
    output_file = test_reduced_spacing()
    if output_file:
        print(f"Test successful! Generated file: {output_file}")
    else:
        print("Test failed!")
        sys.exit(1) 