#!/usr/bin/env python
"""
Test to ensure there are no empty paragraphs between section headers in DOCX files.

This test verifies that the spacing between section headers is controlled properly through
paragraph properties rather than empty paragraphs, as recommended in the code review.
"""

import os
import sys
import logging
import uuid
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('no_empty_paragraphs_test.log')
    ]
)
logger = logging.getLogger(__name__)

def verify_no_empty_paragraphs_between_sections(docx_path):
    """
    Verify that there are no empty paragraphs between section headers.
    
    Args:
        docx_path: Path to the DOCX document to check
        
    Returns:
        True if there are no empty paragraphs between sections, False otherwise
    """
    doc = Document(docx_path)
    section_headers = []
    empty_paragraphs_between_sections = 0
    
    # Find all section headers and their positions
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == 'BoxedHeading2':
            section_headers.append((i, para))
    
    # No section headers found
    if not section_headers:
        logger.warning("No BoxedHeading2 section headers found in document")
        return True
        
    # Only one section header, so no "between sections" to check
    if len(section_headers) < 2:
        logger.info("Only one section header found, no 'between sections' to check")
        return True
        
    # Check for empty paragraphs between consecutive section headers
    for i in range(len(section_headers) - 1):
        current_header_idx = section_headers[i][0]
        next_header_idx = section_headers[i+1][0]
        
        # Check each paragraph between these two section headers
        for j in range(current_header_idx + 1, next_header_idx):
            if j >= len(doc.paragraphs):
                break
                
            para = doc.paragraphs[j]
            if not para.text.strip():
                empty_paragraphs_between_sections += 1
                logger.error(f"Empty paragraph found at index {j} between section headers at {current_header_idx} and {next_header_idx}")
    
    # Log the result
    if empty_paragraphs_between_sections == 0:
        logger.info("No empty paragraphs found between section headers")
        return True
    else:
        logger.error(f"Found {empty_paragraphs_between_sections} empty paragraphs between section headers")
        return False

def measure_header_spacing(docx_path):
    """
    Measure the spacing between section headers in twips.
    
    Args:
        docx_path: Path to the DOCX document to check
        
    Returns:
        List of spacing measurements between consecutive headers
    """
    doc = Document(docx_path)
    headers = []
    spacing_measurements = []
    
    # Find all BoxedHeading2 paragraphs
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == 'BoxedHeading2':
            headers.append((i, para))
    
    if len(headers) < 2:
        logger.warning("Not enough headers to measure spacing")
        return spacing_measurements
    
    # For each consecutive pair of headers, measure the distance
    for i in range(len(headers) - 1):
        header1_idx, header1 = headers[i]
        header2_idx, header2 = headers[i+1]
        
        # Calculate actual spacing - this is approximate and would need real 
        # rendering measurements for perfect accuracy
        paragraphs_between = header2_idx - header1_idx - 1
        
        # Get spacing after first header
        space_after_header1 = header1.paragraph_format.space_after.twips if header1.paragraph_format.space_after else 0
        
        # Get spacing before second header
        space_before_header2 = header2.paragraph_format.space_before.twips if header2.paragraph_format.space_before else 0
        
        # Calculate total spacing between sections
        total_spacing = space_after_header1 + space_before_header2
        
        spacing_measurements.append(total_spacing)
        
        logger.info(f"Spacing between headers {i+1} and {i+2}:")
        logger.info(f"  - Space after header {i+1}: {space_after_header1} twips")
        logger.info(f"  - Space before header {i+2}: {space_before_header2} twips")
        logger.info(f"  - Total spacing: {total_spacing} twips")
        logger.info(f"  - Paragraphs between: {paragraphs_between}")
        
        # Check if spacing is within acceptable limits (less than 200 twips / ~0.14 inches)
        if total_spacing < 200:
            logger.info(f"  ✓ Spacing is acceptable (< 200 twips)")
        else:
            logger.warning(f"  ✗ Spacing is too large (>= 200 twips)")
    
    return spacing_measurements

def test_sections_spacing():
    """
    Generate a test document and verify section spacing.
    """
    try:
        # Import our modules
        from style_engine import StyleEngine
        from utils.docx_builder import add_section_header
        
        # Load tokens
        tokens = StyleEngine.load_tokens()
        
        # Create a test document
        doc = Document()
        
        # Section 1
        header1 = add_section_header(doc, 'PERSONAL SUMMARY', {})
        para1 = doc.add_paragraph('This is a personal summary section with text content.')
        para1.paragraph_format.space_before = Pt(0)
        para1.paragraph_format.space_after = Pt(6)
        
        # More content
        para2 = doc.add_paragraph('Additional paragraph in the summary section.')
        para2.paragraph_format.space_before = Pt(0)
        para2.paragraph_format.space_after = Pt(6)
        
        # Section 2 - should be close to preceding content with no empty paragraphs
        header2 = add_section_header(doc, 'EXPERIENCE', {})
        para3 = doc.add_paragraph('Experience section with work history.')
        para3.paragraph_format.space_before = Pt(0)
        para3.paragraph_format.space_after = Pt(6)
        
        # Section 3
        header3 = add_section_header(doc, 'EDUCATION', {})
        para4 = doc.add_paragraph('Education section with academic credentials.')
        para4.paragraph_format.space_before = Pt(0)
        para4.paragraph_format.space_after = Pt(6)
        
        # Generate a unique filename
        test_id = str(uuid.uuid4())[:8]
        output_file = f"test_sections_spacing_{test_id}.docx"
        
        # Save the document
        doc.save(output_file)
        logger.info(f"Saved test document to {output_file}")
        
        # Verify no empty paragraphs between sections
        no_empty = verify_no_empty_paragraphs_between_sections(output_file)
        
        # Measure header spacing
        spacings = measure_header_spacing(output_file)
        
        # All spacings should be less than 200 twips
        spacing_ok = all(spacing < 200 for spacing in spacings)
        
        # Both checks should pass
        return no_empty and spacing_ok
        
    except Exception as e:
        logger.error(f"Test failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return False

if __name__ == "__main__":
    success = test_sections_spacing()
    if success:
        print("✅ SUCCESS: No empty paragraphs between sections and spacing is acceptable")
        sys.exit(0)
    else:
        print("❌ FAILURE: Found issues with empty paragraphs or spacing")
        sys.exit(1) 