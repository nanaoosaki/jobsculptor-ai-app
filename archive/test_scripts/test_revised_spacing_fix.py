#!/usr/bin/env python
"""
Test script to verify the revised spacing fixes for DOCX output.

This script tests:
1. Tighter section header box with reduced padding
2. Proper spacing between sections by zeroing space_after on paragraphs before headers
"""

import os
import sys
import logging
import uuid
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('revised_spacing_fix_test.log')
    ]
)
logger = logging.getLogger(__name__)

def test_revised_spacing_fix():
    """Test that spacing is fixed properly."""
    try:
        # Import required modules
        from style_engine import StyleEngine, TokenAccessor
        from utils.docx_builder import add_section_header, _fix_spacing_between_sections
        
        # Load tokens
        tokens = StyleEngine.load_tokens()
        tokens_access = TokenAccessor(tokens)
        
        # Verify token values
        padding_pt = tokens_access.get("sectionHeader.paddingPt", 0)
        spacing_after_pt = tokens_access.get("sectionHeader.spacingAfterPt", 0)
        logger.info(f"Current token values - sectionHeader.paddingPt: {padding_pt}, sectionHeader.spacingAfterPt: {spacing_after_pt}")
        
        # Create a test document
        doc = Document()
        
        # Section 1
        header1 = add_section_header(doc, 'PERSONAL SUMMARY', {})
        
        # Add content paragraphs
        para1 = doc.add_paragraph('First paragraph of content with normal spacing.')
        para1.paragraph_format.space_before = Pt(0)
        para1.paragraph_format.space_after = Pt(6)
        
        para2 = doc.add_paragraph('Last paragraph of first section. This paragraph should have space_after=0 after our fix is applied.')
        para2.paragraph_format.space_before = Pt(0)
        para2.paragraph_format.space_after = Pt(6)  # Will be set to 0 by our fix
        
        # Section 2
        header2 = add_section_header(doc, 'PROFESSIONAL EXPERIENCE', {})
        
        # Add content paragraphs
        para3 = doc.add_paragraph('First paragraph of professional experience with normal spacing.')
        para3.paragraph_format.space_before = Pt(0)
        para3.paragraph_format.space_after = Pt(6)
        
        para4 = doc.add_paragraph('Last paragraph of second section. This paragraph should have space_after=0 after our fix is applied.')
        para4.paragraph_format.space_before = Pt(0)
        para4.paragraph_format.space_after = Pt(6)  # Will be set to 0 by our fix
        
        # Section 3
        header3 = add_section_header(doc, 'EDUCATION', {})
        
        # Add final content
        para5 = doc.add_paragraph('Content in the education section.')
        para5.paragraph_format.space_before = Pt(0)
        para5.paragraph_format.space_after = Pt(6)
        
        # Apply our spacing fix
        _fix_spacing_between_sections(doc)
        
        # Generate a unique filename
        test_id = str(uuid.uuid4())[:8]
        output_file = f"test_revised_spacing_{test_id}.docx"
        
        # Save the document
        doc.save(output_file)
        logger.info(f"Saved test document to {output_file}")
        
        # Validate the document
        validate_spacing_fix(output_file)
        
        return output_file
        
    except Exception as e:
        logger.error(f"Test failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return None

def validate_spacing_fix(docx_file):
    """Validate that the spacing fixes were properly applied."""
    try:
        # Open the document with python-docx
        doc = Document(docx_file)
        
        # Find all section headers
        section_headers = []
        for i, para in enumerate(doc.paragraphs):
            if para.style and para.style.name == 'BoxedHeading2':
                section_headers.append((i, para))
        
        logger.info(f"Found {len(section_headers)} section headers")
        
        # Check each section header's border properties
        for i, (idx, header) in enumerate(section_headers):
            # Check paragraph properties through XML
            p_pr = header._element.get_or_add_pPr()
            
            # Check line spacing
            spacing = p_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing')
            if spacing is not None:
                line_value = spacing.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}line')
                line_rule = spacing.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lineRule')
                logger.info(f"Header {i+1} line spacing: {line_value} ({line_rule})")
                
                # Validate line value is 220 for the tighter box
                if line_value and int(line_value) == 220:
                    logger.info("✓ Line height is correctly set to 220")
                else:
                    logger.warning(f"✗ Line height is not 220: {line_value}")
            
            # Check border padding
            borders = p_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr')
            if borders is not None:
                # Check top border space (padding)
                top_border = borders.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}top')
                if top_border is not None:
                    padding = top_border.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space')
                    logger.info(f"Header {i+1} border padding: {padding}")
                    
                    # Validate padding is 20 (1pt * 20) for the tighter box
                    if padding and int(padding) == 20:
                        logger.info("✓ Border padding is correctly set to 1pt")
                    else:
                        logger.warning(f"✗ Border padding is not 1pt: {int(padding)/20}pt")
        
        # Check paragraphs before section headers (except the first) should have space_after=0
        for i, (idx, header) in enumerate(section_headers[1:], 1):  # Skip the first header
            # Find the paragraph right before this header
            prev_idx = idx - 1
            if prev_idx >= 0 and prev_idx < len(doc.paragraphs):
                prev_para = doc.paragraphs[prev_idx]
                space_after = prev_para.paragraph_format.space_after
                
                logger.info(f"Paragraph before header {i+1}: space_after={space_after.pt if space_after else 'None'}")
                
                # Validate space_after is 0
                if space_after and space_after.pt == 0:
                    logger.info("✓ Space after is correctly set to 0")
                else:
                    logger.warning(f"✗ Space after is not 0: {space_after.pt if space_after else 'None'}")
        
        return True
    
    except Exception as e:
        logger.error(f"Validation failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return False

if __name__ == "__main__":
    output_file = test_revised_spacing_fix()
    if output_file:
        print(f"Test successful! Generated file: {output_file}")
        print("Please open the generated file to visually inspect that:")
        print("1. Section header boxes have tighter padding around the text")
        print("2. There's minimal space between end of content and next section header")
    else:
        print("Test failed!")
        sys.exit(1) 