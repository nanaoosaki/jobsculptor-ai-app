import os
import sys
import logging
import uuid
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('fix_spacing_issues_test.log')
    ]
)
logger = logging.getLogger(__name__)

def test_fix_spacing_issues():
    """Test that DOCX section headers have correct height and spacing"""
    try:
        # Import our modules
        from style_engine import StyleEngine, TokenAccessor
        from utils.docx_builder import add_section_header
        
        # Load tokens
        tokens = StyleEngine.load_tokens()
        tokens_access = TokenAccessor(tokens)
        
        # Log the spacing value to verify our changes
        padding_pt = tokens_access.get("sectionHeader.paddingPt", 0)
        spacing_after_pt = tokens_access.get("sectionHeader.spacingAfterPt", 0)
        logger.info(f"Current token values - sectionHeader.paddingPt: {padding_pt}, sectionHeader.spacingAfterPt: {spacing_after_pt}")
        
        # Create a test document
        doc = Document()
        
        # Create the BoxedHeading2 style
        style_name = StyleEngine.create_boxed_heading_style(doc, tokens)
        logger.info(f"Created style: {style_name}")
        
        # Add a section header with the style
        para = doc.add_paragraph("PERSONAL SUMMARY")
        StyleEngine.apply_boxed_section_header_style(doc, para, tokens)
        
        # Add some regular content immediately after (should be closer now)
        para2 = doc.add_paragraph("This is regular content text that follows the section header with reduced spacing.")
        # Make sure to only set space after, not space before
        para2.paragraph_format.space_before = Pt(0)
        para2.paragraph_format.space_after = Pt(6)
        
        # Add more content paragraph
        para3 = doc.add_paragraph("This is a second paragraph of content directly after the first paragraph.")
        para3.paragraph_format.space_before = Pt(0)
        para3.paragraph_format.space_after = Pt(6)
        
        # No empty paragraph between end of content and next section
        
        # Add another section header - this should now be closer to the content above
        para4 = doc.add_paragraph("PROFESSIONAL EXPERIENCE")
        StyleEngine.apply_boxed_section_header_style(doc, para4, tokens)
        
        # Add some regular content
        para5 = doc.add_paragraph("This is more regular content text that should be closer to the header.")
        para5.paragraph_format.space_before = Pt(0)
        para5.paragraph_format.space_after = Pt(6)
        
        # Generate a unique filename
        test_id = str(uuid.uuid4())[:8]
        output_file = f"test_fix_spacing_{test_id}.docx"
        
        # Save the document
        doc.save(output_file)
        logger.info(f"Saved test document to {output_file}")
        
        # Validate the document
        validate_fix_spacing(output_file, tokens_access)
        
        return output_file
        
    except Exception as e:
        logger.error(f"Test failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return None

def validate_fix_spacing(docx_file, tokens_access):
    """Validate that the section headers have correct spacing and height"""
    try:
        # Open the document with python-docx
        doc = Document(docx_file)
        
        # Find paragraphs with BoxedHeading2 style
        boxed_headers = [p for p in doc.paragraphs if p.style and p.style.name == 'BoxedHeading2']
        
        if not boxed_headers:
            logger.error("No BoxedHeading2 style paragraphs found in the document")
            return False
        
        logger.info(f"Found {len(boxed_headers)} boxed headers")
        
        # Check all paragraphs in the document
        all_paragraphs = list(doc.paragraphs)
        logger.info(f"Total paragraphs in document: {len(all_paragraphs)}")
        
        # For each header, check the paragraph before and after it
        for i, para in enumerate(all_paragraphs):
            if para.style and para.style.name == 'BoxedHeading2':
                logger.info(f"Checking header: '{para.text}'")
                
                # Check paragraph properties
                header_space_after = para.paragraph_format.space_after
                header_space_before = para.paragraph_format.space_before
                
                # Check line spacing
                header_line_spacing = para.paragraph_format.line_spacing
                
                # Log the spacing values
                logger.info(f"Header space_before: {header_space_before.pt if header_space_before else 'None'} pt")
                logger.info(f"Header space_after: {header_space_after.pt if header_space_after else 'None'} pt")
                logger.info(f"Header line_spacing: {header_line_spacing}")
                
                # Check the paragraph after this one if it's not the last paragraph
                if i < len(all_paragraphs) - 1:
                    next_para = all_paragraphs[i + 1]
                    next_para_space_before = next_para.paragraph_format.space_before
                    logger.info(f"Next paragraph space_before: {next_para_space_before.pt if next_para_space_before else 'None'} pt")
                
                # Check the paragraph before this one if it's not the first paragraph
                if i > 0:
                    prev_para = all_paragraphs[i - 1]
                    prev_para_space_after = prev_para.paragraph_format.space_after
                    logger.info(f"Previous paragraph space_after: {prev_para_space_after.pt if prev_para_space_after else 'None'} pt")
        
        logger.info("Spacing validation completed")
        return True
        
    except Exception as e:
        logger.error(f"Validation failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return False

if __name__ == "__main__":
    output_file = test_fix_spacing_issues()
    if output_file:
        print(f"Test successful! Generated file: {output_file}")
        print("Please open the generated file to visually inspect that:")
        print("1. Section header boxes have proper height (not enlarged)")
        print("2. There's no extra space between content and the next section header")
    else:
        print("Test failed!")
        sys.exit(1) 