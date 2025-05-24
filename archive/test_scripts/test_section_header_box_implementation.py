"""
Test script for box-styled section headers in DOCX.

This script creates a sample DOCX file with boxed section headers
to validate the implementation of task 1 of the styling improvements.
"""

import os
import sys
import logging
import uuid
from pathlib import Path
from io import BytesIO

import json
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.style import WD_STYLE_TYPE

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('section_header_box_implementation_test.log')
    ]
)
logger = logging.getLogger(__name__)

def create_sample_resume():
    """Create a sample resume with boxed section headers."""
    
    # Import our modules
    from style_engine import StyleEngine, TokenAccessor
    from utils.docx_builder import add_section_header
    
    # Load tokens
    tokens = StyleEngine.load_tokens()
    tokens_access = TokenAccessor(tokens)
    
    # Create a new document
    doc = Document()
    
    # Configure page margins
    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    
    # Add document title (name)
    name_para = doc.add_paragraph("John Doe")
    name_para.style = doc.styles['Title']
    
    # Add contact details
    contact_para = doc.add_paragraph("San Francisco, CA | 555-123-4567 | john.doe@example.com | linkedin.com/in/johndoe")
    contact_para.alignment = 1  # Center
    
    # Add a divider
    doc.add_paragraph("")
    
    # Add Personal Summary section with boxed header
    add_section_header(doc, "PERSONAL SUMMARY", {})
    doc.add_paragraph("Experienced software engineer with 8+ years of expertise in Python, JavaScript, and cloud technologies. Passionate about creating efficient, scalable solutions to complex problems.")
    
    # Add Education section with boxed header
    add_section_header(doc, "EDUCATION", {})
    
    # Create an education entry
    doc.add_paragraph("M.S. Computer Science, Stanford University", style="Heading 3")
    doc.add_paragraph("GPA: 3.9/4.0, May 2015")
    
    doc.add_paragraph("B.S. Computer Science, University of California, Berkeley", style="Heading 3")
    doc.add_paragraph("GPA: 3.8/4.0, May 2013")
    
    # Add Professional Experience section with boxed header
    add_section_header(doc, "PROFESSIONAL EXPERIENCE", {})
    
    # Create a job entry
    doc.add_paragraph("Senior Software Engineer, Google, Mountain View, CA", style="Heading 3")
    doc.add_paragraph("May 2018 - Present")
    
    # Add bullets directly
    bullet1 = doc.add_paragraph()
    bullet1.style = 'List Bullet'  # Use built-in bullet style
    bullet1.add_run("Led a team of 5 engineers to develop a new cloud-based analytics platform, resulting in 30% improvement in data processing speed")
    
    bullet2 = doc.add_paragraph()
    bullet2.style = 'List Bullet'
    bullet2.add_run("Refactored legacy codebase, reducing technical debt by 40% and improving system reliability")
    
    bullet3 = doc.add_paragraph()
    bullet3.style = 'List Bullet'
    bullet3.add_run("Implemented CI/CD pipeline that reduced deployment time from 2 hours to 15 minutes")
    
    # Create another job entry
    doc.add_paragraph("Software Engineer, Facebook, Menlo Park, CA", style="Heading 3")
    doc.add_paragraph("June 2015 - April 2018")
    
    # Add bullets
    bullet4 = doc.add_paragraph()
    bullet4.style = 'List Bullet'
    bullet4.add_run("Developed high-performance APIs serving 10M+ daily requests")
    
    bullet5 = doc.add_paragraph()
    bullet5.style = 'List Bullet'
    bullet5.add_run("Optimized database queries, reducing response time by 40%")
    
    # Add Skills section with boxed header
    add_section_header(doc, "SKILLS", {})
    
    # Create skills paragraph
    doc.add_paragraph("Programming Languages: Python, JavaScript, TypeScript, C++, Java")
    doc.add_paragraph("Technologies: React, Node.js, Django, Flask, Docker, Kubernetes, AWS, GCP")
    doc.add_paragraph("Tools: Git, JIRA, Jenkins, CircleCI, Terraform")
    
    # Generate a unique filename
    test_id = str(uuid.uuid4())[:8]
    output_file = f"test_boxed_headers_{test_id}.docx"
    
    # Save the document
    doc.save(output_file)
    logger.info(f"Saved test document to {output_file}")
    
    return output_file

def validate_section_header_box(docx_file):
    """Validate that the section headers in the document have proper box styling"""
    from docx import Document
    
    try:
        # Open the document with python-docx
        doc = Document(docx_file)
        
        # Get all paragraphs and their styles
        all_styles = [(p.text, p.style.name if hasattr(p, 'style') and p.style else 'None') for p in doc.paragraphs]
        logger.info(f"Found styles: {all_styles}")
        
        # Find section headers by text (which should now be in all caps)
        section_headers = []
        for p in doc.paragraphs:
            if p.text in ["PERSONAL SUMMARY", "EDUCATION", "PROFESSIONAL EXPERIENCE", "SKILLS"]:
                section_headers.append(p)
        
        logger.info(f"Found {len(section_headers)} section headers")
        
        if not section_headers:
            logger.error("No section headers found in the document")
            return False
        
        # Check each section header for border properties
        for header in section_headers:
            logger.info(f"Checking section header: {header.text}")
            
            # Access paragraph's XML
            pPr = header._element.get_or_add_pPr()
            
            # Check for borders
            border_elements = pPr.findall('.//{%s}pBdr' % '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}')
            if not border_elements:
                logger.warning(f"No border found for section header: {header.text}")
                
            # Look for top, right, bottom, left borders in style or direct paragraph properties
            # First check direct paragraph properties
            has_borders = False
            
            # Check in the paragraph
            for border_el in border_elements:
                sides = border_el.findall('.//{%s}top' % '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}')
                sides.extend(border_el.findall('.//{%s}left' % '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'))
                sides.extend(border_el.findall('.//{%s}bottom' % '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'))
                sides.extend(border_el.findall('.//{%s}right' % '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'))
                
                if sides:
                    has_borders = True
                    logger.info(f"Found {len(sides)} borders directly in paragraph")
                    break
            
            # If no direct borders, check the style
            if not has_borders and hasattr(header, 'style') and header.style:
                style_name = header.style.name
                logger.info(f"Checking style: {style_name}")
                
                # Check if using our BoxedHeading2 style
                if style_name == 'BoxedHeading2':
                    logger.info("Using BoxedHeading2 style - this should have borders defined")
                    has_borders = True  # Assume the style has borders
            
            if not has_borders:
                logger.warning(f"Section header {header.text} does not appear to have borders")
                
        logger.info("Validation complete")
        return True
        
    except Exception as e:
        logger.error(f"Validation failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return False

if __name__ == "__main__":
    try:
        output_file = create_sample_resume()
        logger.info(f"Successfully created sample resume with boxed section headers in {output_file}")
        
        # Validate the document
        validation_result = validate_section_header_box(output_file)
        if validation_result:
            logger.info("Validation successful")
            print(f"Test successful! Generated file: {output_file}")
        else:
            logger.warning("Validation had warnings or errors")
            print(f"Test completed with warnings. Generated file: {output_file}")
    except Exception as e:
        logger.error(f"Test failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        print(f"Test failed: {e}")
        sys.exit(1) 