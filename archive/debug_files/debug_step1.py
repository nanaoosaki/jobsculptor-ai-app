import docx
import re

# Load the NEW DOCX file after our O3 fixes
doc = docx.Document('downloadedYCinspiredResume_v2.docx')

print("=== O3's Step 1 Diagnostic - POST-FIX ANALYSIS ===")
print("Checking if our fixes worked...")
print()

# Find paragraphs containing EXPERIENCE
experience_found = False
for i, p in enumerate(doc.paragraphs):
    if re.search(r'EXPERIENCE', p.text, re.I):
        experience_found = True
        print(f"Paragraph {i}:")
        print(f"  Text: '{p.text[:50]}...'")
        print(f"  Style: {p.style.name}")
        print(f"  space_after: {p.paragraph_format.space_after}")
        print(f"  space_before: {p.paragraph_format.space_before}")
        print()

if not experience_found:
    print("No paragraphs containing 'EXPERIENCE' found. Let's check all section headers:")
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == 'MR_SectionHeader':
            print(f"Paragraph {i} (MR_SectionHeader):")
            print(f"  Text: '{p.text[:50]}...'")
            print(f"  space_after: {p.paragraph_format.space_after}")
            print()

print("=== CRITICAL TEST: Company Paragraphs Analysis ===")
print("🎯 This is the key test - are company paragraphs now using MR_Company style?")
company_found = False
for i, p in enumerate(doc.paragraphs):
    if 'Global Cloud Inc' in p.text or 'Cloud Inc' in p.text:
        company_found = True
        print(f"Company Paragraph {i}:")
        print(f"  Text: '{p.text[:50]}...'")
        print(f"  Style: {p.style.name}")
        print(f"  space_after: {p.paragraph_format.space_after}")
        print(f"  space_before: {p.paragraph_format.space_before}")
        
        # O3's success criteria
        if p.style.name == 'MR_Company':
            print(f"  ✅ SUCCESS: Using MR_Company style!")
        else:
            print(f"  ❌ FAILED: Using '{p.style.name}' instead of MR_Company")
        
        if p.paragraph_format.space_after == None or p.paragraph_format.space_after.pt == 0:
            print(f"  ✅ SUCCESS: space_after is 0pt!")
        else:
            print(f"  ❌ FAILED: space_after is {p.paragraph_format.space_after}")
        
        print()

if not company_found:
    print("❌ No company paragraphs found! This is unexpected.")

print("=== STYLE VERIFICATION ===")
print("Checking if MR_Company style exists and its properties...")
try:
    mr_company_style = doc.styles['MR_Company']
    print(f"✅ MR_Company style exists!")
    print(f"  Base style: {mr_company_style.base_style.name if mr_company_style.base_style else 'None'}")
    print(f"  space_after: {mr_company_style.paragraph_format.space_after}")
    print(f"  space_before: {mr_company_style.paragraph_format.space_before}")
except KeyError:
    print("❌ MR_Company style does NOT exist in document!") 