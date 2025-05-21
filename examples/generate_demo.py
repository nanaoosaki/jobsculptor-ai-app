"""
Demo Resume Generator

This script demonstrates the new style registry and section builder
by generating a sample resume with proper section header boxes and spacing.
"""

import os
import sys
import logging
from pathlib import Path
from io import BytesIO

# Add parent directory to path to allow imports
parent_dir = Path(__file__).resolve().parent.parent
sys.path.append(str(parent_dir))

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

from word_styles.registry import StyleRegistry, get_or_create_style
from word_styles.section_builder import add_section_header, add_content_paragraph, add_bullet_point

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('demo_resume.log')
    ]
)
logger = logging.getLogger(__name__)

def create_demo_resume():
    """Generate a sample resume document with proper styling."""
    # Create output directory if it doesn't exist
    output_dir = Path(__file__).parent / "output"
    output_dir.mkdir(exist_ok=True)
    
    # Create output document
    doc = Document()
    
    # Initialize style registry and apply compatibility settings
    registry = StyleRegistry()
    registry.apply_compatibility_settings(doc)
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
    
    # Add contact information
    name_para = doc.add_paragraph("John Doe")
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.style = "Title"
    
    contact_para = doc.add_paragraph("P: +1(000) 000-0000 | 1234@gmail.com | LinkedIn Link | Github Link")
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Professional Summary section
    add_section_header(doc, "PROFESSIONAL SUMMARY")
    
    summary_text = (
        "Experienced AI and Machine Learning leader with a Master's degree in Information Science "
        "and a proven track record in executing large-scale AI/ML projects. Expertise in leveraging "
        "machine learning frameworks and cloud-based ML platforms to drive innovation and align with "
        "corporate goals. Successfully designed scalable, secure data pipelines and enhanced model "
        "accuracy through advanced techniques, showcasing proficiency in Python, Java, and Scala. "
        "Adept at leading cross-functional teams to deliver AI/ML innovations on time and within budget, "
        "while translating complex technical concepts into actionable business strategies. Passionate "
        "about evolving AI/ML trends, fostering industry relationships, and advancing energy systems "
        "knowledge to benefit organizational growth."
    )
    add_content_paragraph(doc, summary_text)
    
    # Add Experience section
    add_section_header(doc, "EXPERIENCE")
    
    # Function to add experience entry with tab stops for right alignment
    def add_experience_entry(doc, company, location, position, dates, role_desc, achievements):
        # Create paragraph with tab stops for right alignment
        para = doc.add_paragraph()
        para.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT)
        
        # Apply content paragraph style
        para.style = get_or_create_style("ContentParagraph", doc)
        
        # Add company and location with tab alignment
        company_run = para.add_run(f"{company} ")
        company_run.bold = True
        para.add_run("\t")  # Tab
        para.add_run(f"{location}").italic = True
        
        # Add position and dates
        pos_para = doc.add_paragraph()
        pos_para.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT)
        pos_para.style = get_or_create_style("ContentParagraph", doc)
        
        pos_run = pos_para.add_run(f"{position} ")
        pos_run.bold = True
        pos_para.add_run("\t")  # Tab
        pos_para.add_run(f"{dates}").italic = True
        
        # Add role description if provided
        if role_desc:
            role_para = doc.add_paragraph()
            role_para.style = get_or_create_style("ContentParagraph", doc)
            role_para.paragraph_format.left_indent = Pt(18)
            role_run = role_para.add_run(role_desc)
            role_run.italic = True
        
        # Add achievements
        for achievement in achievements:
            add_bullet_point(doc, achievement)
    
    # Experience entries
    add_experience_entry(
        doc,
        "Tech Company A",
        "New York, NY",
        "Senior Software Engineer",
        "Jan 2020 - Present",
        "Lead developer for cloud-based applications and services.",
        [
            "Implemented CI/CD pipeline reducing deployment time by 50%",
            "Developed microservice architecture for scalable backend systems",
            "Led team of 5 engineers in successful product launches"
        ]
    )
    
    add_experience_entry(
        doc,
        "Tech Company B",
        "Boston, MA",
        "Software Engineer",
        "Jun 2016 - Dec 2019",
        "Full-stack developer for web applications.",
        [
            "Redesigned user interface improving user satisfaction by 35%",
            "Optimized database queries reducing load times by 40%",
            "Implemented automated testing framework with 85% code coverage"
        ]
    )
    
    # Add Education section
    add_section_header(doc, "EDUCATION")
    
    # Function to add education entry
    def add_education_entry(doc, institution, location, degree, dates, highlights):
        # Create paragraph with tab stops for right alignment
        para = doc.add_paragraph()
        para.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT)
        
        # Apply content paragraph style
        para.style = get_or_create_style("ContentParagraph", doc)
        
        # Add institution and location with tab alignment
        inst_run = para.add_run(f"{institution} ")
        inst_run.bold = True
        para.add_run("\t")  # Tab
        para.add_run(f"{location}").italic = True
        
        # Add degree and dates
        deg_para = doc.add_paragraph()
        deg_para.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT)
        deg_para.style = get_or_create_style("ContentParagraph", doc)
        
        deg_run = deg_para.add_run(f"{degree} ")
        deg_run.bold = True
        deg_para.add_run("\t")  # Tab
        deg_para.add_run(f"{dates}").italic = True
        
        # Add highlights
        for highlight in highlights:
            add_bullet_point(doc, highlight)
    
    # Education entries
    add_education_entry(
        doc,
        "University of Technology",
        "New York, NY",
        "Master of Science in Computer Science",
        "2014 - 2016",
        [
            "GPA: 3.8/4.0",
            "Specialization in Artificial Intelligence"
        ]
    )
    
    # Add Skills section
    add_section_header(doc, "SKILLS")
    
    # Function to add skill category
    def add_skill_category(doc, category, skills):
        # Add category heading
        cat_para = doc.add_paragraph()
        cat_para.style = get_or_create_style("ContentParagraph", doc)
        cat_run = cat_para.add_run(f"{category}: ")
        cat_run.bold = True
        
        # Add skills as comma-separated list
        skills_text = ", ".join(skills)
        cat_para.add_run(skills_text)
    
    # Skill categories
    add_skill_category(
        doc,
        "Technical Skills",
        ["Python", "JavaScript", "React", "Node.js", "AWS", "Docker", "Kubernetes"]
    )
    
    add_skill_category(
        doc,
        "Soft Skills",
        ["Leadership", "Communication", "Problem Solving", "Team Collaboration"]
    )
    
    # Save the document
    output_path = output_dir / "demo_resume.docx"
    doc.save(output_path)
    
    logger.info(f"Generated demo resume: {output_path}")
    print(f"Demo resume generated at: {output_path}")
    print("Please open in Word to visually inspect the styling.")
    
    return output_path

if __name__ == "__main__":
    # Create examples directory structure if it doesn't exist
    Path("examples/output").mkdir(parents=True, exist_ok=True)
    create_demo_resume() 