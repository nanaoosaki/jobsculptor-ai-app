"""
Section Builder for DOCX Documents

This module provides functions to build sections (headers and content)
with proper spacing in DOCX documents.
"""

import logging
from typing import Optional

from docx import Document
from docx.shared import Pt, Twips
from docx.table import Table, _Cell
from docx.oxml.ns import qn

from .registry import get_or_create_style, apply_direct_paragraph_formatting

# Simple stub for rendering tracer to avoid import errors
def trace(name):
    """Simple stub decorator for tracing function calls"""
    def decorator(func):
        return func
    return decorator

logger = logging.getLogger(__name__)

@trace("docx.section_header")
def add_section_header(doc: Document, text: str, style_name: str = "BoxedHeading2Table") -> None:
    """
    Add a section header with proper styling and spacing.
    
    This function adds a section header with the specified style and ensures
    that there's no unwanted spacing before it. If the previous paragraph is not
    empty, it adds an empty paragraph with zero spacing after.
    
    Args:
        doc: The document to add the section header to
        text: The section header text
        style_name: The name of the style to apply (must be in registry)
    
    Returns:
        The added paragraph or table, depending on the style wrapper type
    """
    # Ensure the style exists
    style_name = get_or_create_style(style_name, doc)
    
    # Check if we need to add an empty paragraph first to control spacing
    if len(doc.paragraphs) > 0:
        last_para = doc.paragraphs[-1]
        
        # Only add a spacing control paragraph if the last paragraph is not empty
        if last_para.text.strip():
            # Create an empty paragraph with zero spacing after
            empty_para = doc.add_paragraph()
            empty_para.text = ""
            
            # Apply EmptyParagraph style - very small, no spacing
            empty_style = get_or_create_style("EmptyParagraph", doc)
            empty_para.style = empty_style
            
            # Double-ensure zero spacing after via direct XML
            apply_direct_paragraph_formatting(empty_para, "EmptyParagraph")
            logger.info("Added empty control paragraph before section header")
    
    # Get the style details to determine if we should use table or paragraph
    from .registry import _registry
    style_def = _registry.get(style_name)
    
    if style_def and style_def.wrapper == "table":
        # Use table-based approach
        return _add_table_section_header(doc, text, style_def)
    else:
        # Use paragraph-based approach
        return _add_paragraph_section_header(doc, text, style_name)

@trace("docx.section_header")
def _add_paragraph_section_header(doc: Document, text: str, style_name: str):
    """
    Add a section header using a paragraph with border.
    
    Args:
        doc: The document to add the section header to
        text: The section header text
        style_name: The name of the style to apply
        
    Returns:
        The added paragraph
    """
    # Now add the section header
    header_para = doc.add_paragraph()
    header_para.text = text
    
    # Apply the style
    header_para.style = style_name
    
    # Also apply direct formatting to ensure all properties are set
    # This helps with cross-platform compatibility
    apply_direct_paragraph_formatting(header_para, style_name)
    
    logger.info(f"Added section header: {text} with style {style_name}")
    return header_para

@trace("docx.section_header")
def _add_table_section_header(doc: Document, text: str, style_def):
    """
    Add a section header using a 1x1 table with borders.
    
    This approach avoids Word's paragraph spacing issues by using a table,
    which has more reliable height control.
    
    Args:
        doc: The document to add the section header to
        text: The section header text
        style_def: The style definition from the registry
        
    Returns:
        The added table
    """
    # Create a 1x1 table
    tbl = doc.add_table(rows=1, cols=1)
    
    # Disable autofit to prevent Word from resizing
    tbl.autofit = False
    
    # Get the cell
    cell = tbl.rows[0].cells[0]
    
    # Apply border to the cell
    _apply_cell_border(cell, style_def)
    
    # Set cell vertical alignment to top
    _set_cell_vertical_alignment(cell, 'top')
    
    # Set asymmetric cell margins (less on top)
    margins = {
        'top': 10,     # Half the side margins for minimal top padding
        'left': 20,
        'bottom': 20,
        'right': 20
    }
    _set_cell_margins(cell, margins)
    
    # Get the existing paragraph in the cell
    para = cell.paragraphs[0]
    
    # Apply our purpose-built HeaderBoxH2 style instead of Heading 2
    # This avoids inheriting unwanted spacing from Heading 2
    header_style = get_or_create_style("HeaderBoxH2", doc)
    para.style = header_style
    
    # Add the text
    para.text = text
    
    # Promote outline level to maintain document structure
    _promote_outline_level(para, style_def.outline_level)
    
    logger.info(f"Added table section header: {text} with style {style_def.name}")
    return tbl

def _apply_cell_border(cell: _Cell, style_def):
    """
    Apply border to a table cell based on style definition.
    
    Args:
        cell: The cell to apply border to
        style_def: The style definition with border properties
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    # Convert style's border width to 1/8th points for Word
    width_8th_pt = int(style_def.border_width_pt * 8)
    
    # Get cell properties
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Remove any existing borders
    tcBorders_elements = tcPr.xpath('./w:tcBorders')
    for element in tcBorders_elements:
        tcPr.remove(element)
    
    # Create border XML
    tcBorders_xml = f'''
    <w:tcBorders {nsdecls("w")}>
        <w:top w:val="{style_def.border_style}" w:sz="{width_8th_pt}" w:color="{style_def.border_color}"/>
        <w:left w:val="{style_def.border_style}" w:sz="{width_8th_pt}" w:color="{style_def.border_color}"/>
        <w:bottom w:val="{style_def.border_style}" w:sz="{width_8th_pt}" w:color="{style_def.border_color}"/>
        <w:right w:val="{style_def.border_style}" w:sz="{width_8th_pt}" w:color="{style_def.border_color}"/>
    </w:tcBorders>
    '''
    
    # Add borders to cell
    tcBorders = parse_xml(tcBorders_xml)
    tcPr.append(tcBorders)

def _set_cell_margins(cell: _Cell, margins):
    """
    Set custom margins for each side of a table cell.
    
    Args:
        cell: The cell to set margins for
        margins: Either an int (all sides equal) or dict with 'top', 'left', etc.
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    # Handle different margin formats
    if isinstance(margins, int):
        top = margins
        left = margins
        bottom = margins
        right = margins
    else:
        # Use dict with defaults
        top = margins.get('top', 10)  # Less on top by default
        left = margins.get('left', 20)
        bottom = margins.get('bottom', 20)
        right = margins.get('right', 20)
    
    # Get cell properties
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Remove any existing margins
    tcMar_elements = tcPr.xpath('./w:tcMar')
    for element in tcMar_elements:
        tcPr.remove(element)
    
    # Create margins XML
    tcMar_xml = f'''
    <w:tcMar {nsdecls("w")}>
        <w:top w:w="{top}" w:type="dxa"/>
        <w:left w:w="{left}" w:type="dxa"/>
        <w:bottom w:w="{bottom}" w:type="dxa"/>
        <w:right w:w="{right}" w:type="dxa"/>
    </w:tcMar>
    '''
    
    # Add margins to cell
    tcMar = parse_xml(tcMar_xml)
    tcPr.append(tcMar)

def _set_cell_vertical_alignment(cell: _Cell, alignment='top'):
    """
    Set vertical alignment for cell content.
    
    Args:
        cell: The cell to set alignment for
        alignment: The vertical alignment ('top', 'center', 'bottom')
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    # Get cell properties
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Remove any existing vertical alignment
    for va in tcPr.xpath('./w:vAlign'):
        tcPr.remove(va)
    
    # Add vertical alignment
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{alignment}"/>')
    tcPr.append(vAlign)

def _promote_outline_level(para, level: int):
    """
    Set the outline level for a paragraph to maintain document structure.
    
    This is important for tables, as paragraphs in tables don't inherit
    outline level from styles automatically.
    
    Args:
        para: The paragraph to set outline level for
        level: The outline level (0 for Heading 1, 1 for Heading 2, etc.)
    """
    from docx.oxml.ns import qn
    
    # Get paragraph properties
    p_pr = para._element.get_or_add_pPr()
    
    # Remove any existing outline level
    for existing in p_pr.xpath('./w:outlineLvl'):
        p_pr.remove(existing)
    
    # Create new outline level
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    # Create outline level XML
    outlineLvl_xml = f'<w:outlineLvl {nsdecls("w")} w:val="{level}"/>'
    
    # Add outline level to paragraph
    outlineLvl = parse_xml(outlineLvl_xml)
    p_pr.append(outlineLvl)

def add_content_paragraph(doc: Document, text: str, style_name: str = "ContentParagraph") -> None:
    """
    Add a content paragraph with proper styling and spacing.
    
    Args:
        doc: The document to add the paragraph to
        text: The paragraph text
        style_name: The name of the style to apply (must be in registry)
    
    Returns:
        The added paragraph
    """
    # Ensure the style exists
    style_name = get_or_create_style(style_name, doc)
    
    # Add the paragraph
    para = doc.add_paragraph()
    para.text = text
    
    # Apply the style
    para.style = style_name
    
    # Apply direct formatting to ensure all properties are set
    apply_direct_paragraph_formatting(para, style_name)
    
    logger.info(f"Added content paragraph with style {style_name}")
    return para

def add_bullet_point(doc: Document, text: str, level: int = 0, style_name: str = "ContentParagraph") -> None:
    """
    Add a bullet point with proper styling and spacing.
    
    Args:
        doc: The document to add the bullet point to
        text: The bullet point text
        level: The indentation level (0 = first level)
        style_name: The base style name to apply (will be formatted as bullet)
    
    Returns:
        The added paragraph
    """
    # Ensure the style exists
    style_name = get_or_create_style(style_name, doc)
    
    # Add the paragraph
    para = doc.add_paragraph(style=style_name)
    
    # Apply bullet styling
    para.style = style_name
    para.paragraph_format.left_indent = Pt(18 * (level + 1))  # 18pt per level
    para.paragraph_format.first_line_indent = Pt(-18)  # Hanging indent for bullet
    
    # Add text as a run
    run = para.add_run(text)
    
    # Apply bullet character
    para._p.get_or_add_pPr().get_or_add_numPr()
    
    # Apply direct formatting to ensure all properties are set
    apply_direct_paragraph_formatting(para, style_name)
    
    logger.info(f"Added bullet point with style {style_name}")
    return para

def remove_empty_paragraphs(doc: Document) -> int:
    """
    Remove all empty paragraphs from the document.
    
    Args:
        doc: The document to clean
        
    Returns:
        Number of paragraphs removed
    """
    removed_count = 0
    
    # We need to work backwards since removing elements changes indices
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        para = doc.paragraphs[i]
        
        # Check if paragraph is empty (no text or only whitespace)
        if not para.text.strip():
            # Don't remove if it's the last paragraph (Word needs at least one)
            if i < len(doc.paragraphs) - 1:
                p_element = para._element
                parent = p_element.getparent()
                if parent is not None:
                    parent.remove(p_element)
                    removed_count += 1
    
    logger.info(f"Removed {removed_count} empty paragraphs")
    return removed_count

def add_role_box(doc: Document, role: str, dates: Optional[str] = None) -> Table:
    """
    Add a role box using a table with light grey-blue background and no borders.
    Creates a single unified box containing both role and dates, similar to section headers.
    
    Args:
        doc: The document to add the role box to
        role: The role/position text
        dates: Optional dates text (if provided, displays right-aligned within the same box)
        
    Returns:
        The added table
    """
    # Always create a single-column table for unified styling
    tbl = doc.add_table(rows=1, cols=1)
    
    # Disable autofit to prevent Word from resizing
    tbl.autofit = False
    
    # Apply no borders to the table (remove outline)
    _apply_role_box_table_border(tbl)
    
    # CRITICAL: Apply tight spacing after role box table to match design tokens
    # This controls the gap between role boxes and role descriptions
    _apply_role_box_table_spacing(tbl)
    
    # Get the single cell
    cell = tbl.rows[0].cells[0]
    
    # Apply light grey-blue background to the cell
    _apply_role_box_cell_shading(cell)
    
    # Set cell vertical alignment to top
    _set_cell_vertical_alignment(cell, 'top')
    
    # Set cell margins for role box (consistent with design tokens)
    cell_margins = {
        'top': 5,      # Minimal top padding
        'left': 15,    # Moderate side padding
        'bottom': 15,
        'right': 15
    }
    _set_cell_margins(cell, cell_margins)
    
    # Get the existing paragraph in the cell
    para = cell.paragraphs[0]
    
    # Apply role box style
    from .registry import get_or_create_style
    role_style = get_or_create_style("RoleBoxText", doc)
    para.style = role_style
    
    # Create the content with role on left and dates on right (like HTML flex layout)
    if dates:
        # Calculate the correct tab position to align with location text (same logic as format_right_aligned_pair)
        from docx.shared import Cm
        from docx.enum.text import WD_TAB_ALIGNMENT
        
        # Get document margins and page width (same calculation as in format_right_aligned_pair)
        sections = doc.sections
        section = sections[0] if sections else None
        
        if section:
            # Page width and margins in cm
            page_width_cm = section.page_width.cm
            global_left_margin_cm = section.left_margin.cm
            global_right_margin_cm = section.right_margin.cm
            
            # Calculate content width in cm using global margins
            content_width_cm = page_width_cm - global_left_margin_cm - global_right_margin_cm
            
            # Set tab position at the right edge of content area (same as format_right_aligned_pair)
            if content_width_cm <= 0:  # Fallback if dimensions are invalid
                tab_position_cm = 16.0  # A reasonable default
            else:
                tab_position_cm = content_width_cm
        else:
            # Fallback if no section found
            tab_position_cm = 16.0
        
        # Clear any existing tab stops
        para.paragraph_format.tab_stops.clear_all()
        
        # Add tab stop for right alignment at the calculated position (matches location alignment)
        para.paragraph_format.tab_stops.add_tab_stop(Cm(tab_position_cm), WD_TAB_ALIGNMENT.RIGHT)
        
        # Add role text (left-aligned)
        role_run = para.add_run(role)
        role_run.bold = True
        
        # Add tab and dates text (right-aligned)
        para.add_run('\t')  # Tab to push dates to the right
        dates_run = para.add_run(dates)
        dates_run.italic = True
    else:
        # Just add the role text
        role_run = para.add_run(role)
        role_run.bold = True
    
    # O3: Prevent border merging in LibreOffice
    tbl.allow_overlap = False  # Word ignores, LibreOffice respects
    
    logger.info(f"Added role box: {role}" + (f" with dates: {dates}" if dates else ""))
    return tbl

def _apply_role_box_table_spacing(tbl: Table):
    """
    Apply spacing after role box table - SIMPLIFIED VERSION.
    
    The previous implementation used floating table positioning (tblpPr) which
    caused role boxes to float outside page margins. This version uses normal
    table properties or relies on following paragraph spacing control.
    
    Args:
        tbl: The table to apply spacing to
    """
    # REMOVED: Floating table positioning (tblpPr) that caused the layout issues
    # The tight spacing between role boxes and role descriptions should be 
    # controlled by the MR_RoleDescription paragraph spacing (space_before: 0pt)
    # rather than making the table float outside the document flow.
    
    logger.info("Applied normal table spacing (no floating positioning)")
    
    # Note: If we need tighter control over table spacing in the future,
    # we should use table margins or the spacing properties of following paragraphs,
    # NOT table positioning properties (tblpPr) which break document flow.

def _apply_role_box_table_border(tbl: Table):
    """
    Remove all borders from the role box table for a clean, borderless appearance.
    
    Args:
        tbl: The table to remove borders from
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    # Get table element and its properties
    tbl_element = tbl._tbl
    
    # Get or create tblPr element using proper XML approach
    tblPr = tbl_element.find('.//w:tblPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    if tblPr is None:
        # Create tblPr element if it doesn't exist
        tblPr_xml = f'<w:tblPr {nsdecls("w")}></w:tblPr>'
        tblPr = parse_xml(tblPr_xml)
        # Insert tblPr as the first child of tbl (after any range markup elements)
        tbl_element.insert(0, tblPr)
    
    # Remove any existing borders
    tblBorders_elements = tblPr.xpath('./w:tblBorders')
    for element in tblBorders_elements:
        tblPr.remove(element)
    
    # Create table border XML with no borders
    tblBorders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="none"/>
        <w:left w:val="none"/>
        <w:bottom w:val="none"/>
        <w:right w:val="none"/>
        <w:insideH w:val="none"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>
    '''
    
    # Add borderless styling to table
    tblBorders = parse_xml(tblBorders_xml)
    tblPr.append(tblBorders)

def _apply_role_box_cell_shading(cell: _Cell):
    """
    Apply light grey-blue background color to a table cell for modern, subtle appearance.
    
    Args:
        cell: The cell to apply background color to
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    # Light grey-blue color - subtle and professional
    # Similar to #E8F2FF or #F0F6FF for a very light blue-grey
    background_color = "E8F2FF"  # Light grey-blue without # prefix for XML
    
    # Get cell properties
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Remove any existing shading
    shading_elements = tcPr.xpath('./w:shd')
    for element in shading_elements:
        tcPr.remove(element)
    
    # Create shading XML with correct format
    shading_xml = f'''
    <w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{background_color}"/>
    '''
    
    # Add shading to cell
    shading = parse_xml(shading_xml)
    tcPr.append(shading) 