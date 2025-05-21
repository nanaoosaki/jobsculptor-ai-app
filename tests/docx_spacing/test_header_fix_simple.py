"""
Simple Test for Header Box Height Fix

This script generates a test document to visually verify the header box styling.
"""

import sys
import logging
from pathlib import Path

# Add parent directory to path to allow imports
parent_dir = Path(__file__).resolve().parent.parent.parent
sys.path.append(str(parent_dir))

from docx import Document
from docx.shared import Pt

from word_styles.registry import StyleRegistry
from word_styles.section_builder import add_section_header, add_content_paragraph

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def create_test_document():
    """Generate a test document with the improved header styling."""
    # Create output directory if it doesn't exist
    output_dir = Path(__file__).parent / "output"
    output_dir.mkdir(exist_ok=True)
    
    # Create output document
    doc = Document()
    
    # Get the StyleRegistry and apply compatibility settings
    registry = StyleRegistry()
    try:
        registry.apply_compatibility_settings(doc)
    except Exception as e:
        logger.warning(f"Could not apply compatibility settings: {e}")
    
    # Add title
    title = doc.add_paragraph("Header Style Fix Test")
    title.style = "Title"
    
    # Add explanation
    explanation = doc.add_paragraph(
        "This document tests the improved header styling with: "
        "1) Table-based wrapper, 2) HeaderBoxH2 style based on Normal, "
        "3) Asymmetric cell margins, and 4) Top vertical alignment."
    )
    
    # Add a section header
    add_section_header(doc, "HEADER WITH MINIMAL TOP SPACE")
    add_content_paragraph(doc, "This header should have minimal space above the text.")
    
    # Add another section header with longer text to test wrapping
    add_section_header(doc, "HEADER WITH LONGER TEXT THAT SHOULD WRAP TO THE NEXT LINE")
    add_content_paragraph(doc, "This header should wrap properly and still have minimal top space.")
    
    # Add a few more different length headers to test consistency
    add_section_header(doc, "SHORT HEADER")
    add_content_paragraph(doc, "This tests a short header.")
    
    add_section_header(doc, "EDUCATION")
    add_content_paragraph(doc, "This tests a typical resume section header.")
    
    add_section_header(doc, "SKILLS")
    add_content_paragraph(doc, "This tests another typical resume section header.")
    
    # Save the document
    output_path = output_dir / "header_style_fix_simple.docx"
    doc.save(output_path)
    
    return output_path

def main():
    """Run the test."""
    # Create test document
    docx_path = create_test_document()
    
    print("\nTest DOCX generated successfully!")
    print(f"File location: {docx_path}")
    print("\nPlease examine in Word to verify that:")
    print("1. Header boxes have minimal space above the text (similar to PDF output)")
    print("2. The total height is approximately font size plus border padding")
    print("3. Spacing is consistent across all headers")

if __name__ == "__main__":
    # Create tests directory structure if it doesn't exist
    Path("tests/docx_spacing/output").mkdir(parents=True, exist_ok=True)
    main() 