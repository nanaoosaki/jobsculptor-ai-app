"""
Test Exact Line Height for DOCX Section Headers

This script generates a test document with section headers using the 'exact' line rule
to verify that header boxes have the correct compact height.
"""

import os
import sys
import logging
from pathlib import Path

# Add parent directory to path to allow imports
parent_dir = Path(__file__).resolve().parent.parent.parent
sys.path.append(str(parent_dir))

from docx import Document
from docx.shared import Pt

from word_styles.registry import StyleRegistry, ParagraphBoxStyle
from word_styles.xml_utils import make_spacing_node, make_border_node, twips_from_pt
from word_styles.section_builder import add_section_header, add_content_paragraph

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('exact_line_height.log')
    ]
)
logger = logging.getLogger(__name__)

def create_test_document():
    """Generate a test document with section headers using exact line height."""
    # Create output directory if it doesn't exist
    output_dir = Path(__file__).parent / "output"
    output_dir.mkdir(exist_ok=True)
    
    # Create output document
    doc = Document()
    
    # Get the StyleRegistry and apply compatibility settings
    registry = StyleRegistry()
    registry.apply_compatibility_settings(doc)
    
    # Verify BoxedHeading2 style has exact line height
    style = registry.get("BoxedHeading2")
    logger.info(f"BoxedHeading2 line rule: {style.line_rule}")
    logger.info(f"BoxedHeading2 line height: {style.line_height_pt}pt ({twips_from_pt(style.line_height_pt)} twips)")
    
    # Add title
    title = doc.add_paragraph("Exact Line Height Test")
    title.style = "Title"
    
    # Add explanation
    explanation = doc.add_paragraph(
        "This document tests section headers with 'exact' line height to verify "
        "that header boxes have a compact height with minimal extra space."
    )
    
    # Add a section header
    add_section_header(doc, "REGULAR SECTION HEADER")
    add_content_paragraph(doc, "This is content under a regular section header.")
    
    # Add a section header with longer text to test wrapping
    add_section_header(doc, "SECTION HEADER WITH LONGER TEXT THAT SHOULD WRAP TO THE NEXT LINE")
    add_content_paragraph(doc, "This is content under a section header with wrapped text.")
    
    # Add a section header with mixed case
    add_section_header(doc, "Section Header with Mixed Case")
    add_content_paragraph(doc, "This is content under a section header with mixed case.")
    
    # Save the document
    output_path = output_dir / "exact_line_height.docx"
    doc.save(output_path)
    
    logger.info(f"Generated test document: {output_path}")
    logger.info(f"Please open in Word to verify the following:")
    logger.info(f"1. Single-line header boxes should be compact (font size + ~2pt)")
    logger.info(f"2. Multi-line headers should still be readable (not clipped)")
    
    return output_path

def main():
    """Run the test."""
    # Create test document
    docx_path = create_test_document()
    print(f"\nTest DOCX generated at: {docx_path}")
    print("Please examine in Word to verify that:")
    print("1. Header boxes are now compact with minimal extra space")
    print("2. The total height is approximately 'font size + 2pt border padding'")
    print("3. No extra spacing appears between content and headers")

if __name__ == "__main__":
    # Create tests directory structure if it doesn't exist
    Path("tests/docx_spacing/output").mkdir(parents=True, exist_ok=True)
    main() 