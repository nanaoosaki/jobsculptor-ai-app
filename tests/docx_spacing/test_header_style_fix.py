"""
Test Header Box Height Fix

This script tests the implementation of the fix for header box height using:
1. Table-based wrapper
2. Custom HeaderBoxH2 style based on Normal (not Heading 2) 
3. Asymmetric cell margins
4. Vertical alignment

The goal is to eliminate excess space above the header text.
"""

import os
import sys
import logging
import zipfile
from pathlib import Path
import xml.etree.ElementTree as ET

# Add parent directory to path to allow imports
parent_dir = Path(__file__).resolve().parent.parent.parent
sys.path.append(str(parent_dir))

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from word_styles.registry import StyleRegistry, ParagraphBoxStyle
from word_styles.section_builder import add_section_header, add_content_paragraph

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('header_style_fix.log')
    ]
)
logger = logging.getLogger(__name__)

def create_test_document():
    """Generate a test document with the improved header styling."""
    # Create output directory if it doesn't exist
    output_dir = Path(__file__).parent / "output"
    output_dir.mkdir(exist_ok=True)
    
    # Create output document
    doc = Document()
    
    # Get the StyleRegistry and apply compatibility settings
    registry = StyleRegistry()
    registry.apply_compatibility_settings(doc)
    
    # Add title
    title = doc.add_paragraph("Header Style Fix Test")
    title.style = "Title"
    
    # Add explanation
    explanation = doc.add_paragraph(
        "This document tests the improved header styling with: "
        "1) Table-based wrapper, 2) HeaderBoxH2 style based on Normal, "
        "3) Asymmetric cell margins, and 4) Top vertical alignment."
    )
    
    # Add a section header
    add_section_header(doc, "HEADER WITH MINIMAL TOP SPACE")
    add_content_paragraph(doc, "This header should have minimal space above the text.")
    
    # Add another section header with longer text to test wrapping
    add_section_header(doc, "HEADER WITH LONGER TEXT THAT SHOULD WRAP TO THE NEXT LINE")
    add_content_paragraph(doc, "This header should wrap properly and still have minimal top space.")
    
    # Save the document
    output_path = output_dir / "header_style_fix.docx"
    doc.save(output_path)
    
    logger.info(f"Generated test document: {output_path}")
    return output_path

def validate_docx_xml(docx_path):
    """Verify the XML structure matches our expectations."""
    with zipfile.ZipFile(docx_path) as z:
        with z.open('word/document.xml') as f:
            xml_content = f.read()
            
    # Parse XML
    ET.register_namespace('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
    root = ET.fromstring(xml_content)
    
    # Count spacing nodes in header paragraphs
    # The namespace handling is a bit tricky with ElementTree
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # Find all paragraphs
    all_paras = root.findall('.//w:p', ns)
    logger.info(f"Found {len(all_paras)} paragraphs total")
    
    # Find paragraphs in table cells
    cell_paras = root.findall('.//w:tc//w:p', ns)
    logger.info(f"Found {len(cell_paras)} paragraphs in table cells")
    
    # Check each paragraph in cells
    header_count = 0
    for i, para in enumerate(cell_paras):
        # Check style
        style_nodes = para.findall('.//w:pStyle', ns)
        for style in style_nodes:
            style_val = style.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            logger.info(f"Cell paragraph {i+1} has style: {style_val}")
            if style_val == "HeaderBoxH2":
                header_count += 1
        
        # Check spacing nodes
        spacing_nodes = para.findall('.//w:spacing', ns)
        spacing_count = len(spacing_nodes)
        logger.info(f"Cell paragraph {i+1} has {spacing_count} spacing nodes")
        
        # Check for vertical alignment in parent cell
        cell = para.find('./ancestor::w:tc', ns)
        if cell is not None:
            v_align = cell.find('.//w:vAlign', ns)
            if v_align is not None:
                logger.info(f"Cell {i+1} vertical alignment: {v_align.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')}")
            else:
                logger.warning(f"Cell {i+1} has no vertical alignment set")
    
    logger.info(f"Found {header_count} paragraphs with HeaderBoxH2 style")
    return header_count > 0

def main():
    """Run the test."""
    # Create test document
    docx_path = create_test_document()
    
    # Validate XML structure
    valid = validate_docx_xml(docx_path)
    
    if valid:
        print("\nTest DOCX generated successfully!")
        print(f"File location: {docx_path}")
        print("\nPlease examine in Word to verify that:")
        print("1. Header boxes have minimal space above the text")
        print("2. The total height is approximately font size plus border padding")
        print("3. Headers appear in document navigation/outline")
    else:
        print("\nTest failed: XML structure validation issues")

if __name__ == "__main__":
    # Create tests directory structure if it doesn't exist
    Path("tests/docx_spacing/output").mkdir(parents=True, exist_ok=True)
    main() 