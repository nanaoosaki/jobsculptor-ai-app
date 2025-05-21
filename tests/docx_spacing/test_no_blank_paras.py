"""
Test No Blank Paragraphs

This script validates that there are no unwanted blank paragraphs
between sections in the generated document.
"""

import os
import sys
import logging
from pathlib import Path

# Add parent directory to path to allow imports
parent_dir = Path(__file__).resolve().parent.parent.parent
sys.path.append(str(parent_dir))

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from word_styles.registry import StyleRegistry, ParagraphBoxStyle
from word_styles.section_builder import add_section_header, add_content_paragraph, remove_empty_paragraphs

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('no_blank_paras.log')
    ]
)
logger = logging.getLogger(__name__)

def generate_test_document():
    """Generate test document with sections and content."""
    # Create output directory if it doesn't exist
    output_dir = Path(__file__).parent / "output"
    output_dir.mkdir(exist_ok=True)
    
    # Create output document
    doc = Document()
    
    # Add title
    title = doc.add_paragraph("No Blank Paragraphs Test")
    title.style = "Title"
    
    # Initialize a style registry
    registry = StyleRegistry()
    
    # Add section headers and content
    add_section_header(doc, "SECTION 1")
    add_content_paragraph(doc, "This is content paragraph 1 for Section 1.")
    add_content_paragraph(doc, "This is content paragraph 2 for Section 1.")
    
    add_section_header(doc, "SECTION 2")
    add_content_paragraph(doc, "This is content paragraph 1 for Section 2.")
    add_content_paragraph(doc, "This is content paragraph 2 for Section 2.")
    
    add_section_header(doc, "SECTION 3")
    add_content_paragraph(doc, "This is content paragraph 1 for Section 3.")
    add_content_paragraph(doc, "This is content paragraph 2 for Section 3.")
    
    add_section_header(doc, "SECTION 4")
    add_content_paragraph(doc, "This is content paragraph 1 for Section 4.")
    add_content_paragraph(doc, "This is content paragraph 2 for Section 4.")
    
    # Save the document
    output_path = output_dir / "no_blank_paras.docx"
    doc.save(output_path)
    
    logger.info(f"Generated test document: {output_path}")
    return output_path

def validate_no_blank_paragraphs(docx_path):
    """Validate that there are no unwanted blank paragraphs in the document."""
    # Open the document
    doc = Document(docx_path)
    
    # Check for blank paragraphs between sections
    blank_paras_count = 0
    header_indices = []
    
    # Find all section headers
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == "BoxedHeading2":
            header_indices.append(i)
    
    # Check for blank paragraphs between content and section headers
    for i in range(len(header_indices)):
        header_idx = header_indices[i]
        
        # Skip the first header (no preceding content)
        if i == 0:
            continue
        
        # Check paragraphs between previous content and this header
        prev_header_idx = header_indices[i-1]
        
        for j in range(prev_header_idx + 1, header_idx):
            para = doc.paragraphs[j]
            
            # Skip non-empty content paragraphs
            if para.text.strip():
                continue
            
            # Skip the intentional control paragraph (should be only one)
            if para.style.name == "EmptyParagraph":
                continue
                
            # Count any other blank paragraphs
            blank_paras_count += 1
            logger.warning(f"Found unwanted blank paragraph at index {j}, between headers {prev_header_idx} and {header_idx}")
    
    # Check XML structure for empty <w:p></w:p> elements
    xml_empty_paras = 0
    
    # Get document XML
    doc_xml = doc._body._element.xml
    
    # Count empty paragraph tags (simplified approach)
    empty_p_tags = doc_xml.count("<w:p></w:p>")
    xml_empty_paras += empty_p_tags
    
    logger.info(f"Found {blank_paras_count} unwanted blank paragraphs between sections")
    logger.info(f"Found {xml_empty_paras} empty <w:p></w:p> elements in XML")
    
    return blank_paras_count == 0 and xml_empty_paras == 0

def main():
    """Run the test."""
    # Generate test document
    docx_path = generate_test_document()
    
    # Validate no blank paragraphs
    result = validate_no_blank_paragraphs(docx_path)
    
    if result:
        logger.info("✅ TEST PASSED: No unwanted blank paragraphs found")
        return 0
    else:
        logger.error("❌ TEST FAILED: Unwanted blank paragraphs found")
        return 1

if __name__ == "__main__":
    # Create tests directory structure if it doesn't exist
    Path("tests/docx_spacing/output").mkdir(parents=True, exist_ok=True)
    sys.exit(main()) 