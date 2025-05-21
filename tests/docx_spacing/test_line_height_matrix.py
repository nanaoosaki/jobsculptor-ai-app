"""
Test Line Height Matrix for DOCX Section Headers

This script generates a test document with a matrix of different lineRule and
line values to find the optimal settings for section header boxes.
"""

import os
import sys
import logging
from pathlib import Path

# Add parent directory to path to allow imports
parent_dir = Path(__file__).resolve().parent.parent.parent
sys.path.append(str(parent_dir))

from docx import Document
from docx.shared import Pt

from word_styles.registry import StyleRegistry, ParagraphBoxStyle
from word_styles.xml_utils import make_spacing_node, make_border_node, twips_from_pt
from word_styles.section_builder import add_section_header, add_content_paragraph, remove_empty_paragraphs

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('line_height_matrix.log')
    ]
)
logger = logging.getLogger(__name__)

def main():
    """Generate test document with line height matrix."""
    # Create output directory if it doesn't exist
    output_dir = Path(__file__).parent / "output"
    output_dir.mkdir(exist_ok=True)
    
    # Create output document
    doc = Document()
    
    # Add title
    title = doc.add_paragraph("Line Height Matrix Test")
    title.style = "Title"
    
    # Add explanation
    explanation = doc.add_paragraph(
        "This document shows section headers with different line height settings "
        "to find the optimal configuration for minimal box height."
    )
    
    # Define matrix of line rules and line values
    line_rules = ["auto", "exact", "atLeast"]
    
    # Values in points, will be converted to twips
    line_values_pt = [13.0, 13.8, 14.0, 14.5, 15.0, 16.0]
    
    # Register test styles
    registry = StyleRegistry()
    
    # Create test styles for each combination
    for rule in line_rules:
        for value_pt in line_values_pt:
            style_name = f"Test_{rule}_{int(value_pt * 10)}"
            
            # Register with registry
            registry.register(ParagraphBoxStyle(
                name=style_name,
                base_style_name="Heading 2",
                outline_level=1,
                font_name="Calibri",
                font_size_pt=14.0,
                font_bold=True,
                font_color="0D2B7E",
                border_width_pt=1.0,
                border_color="0D2B7E",
                border_style="single",
                border_padding_pt=1.0,
                space_before_pt=0.0,
                space_after_pt=4.0,
                line_rule=rule,
                line_height_pt=value_pt,
                keep_with_next=True,
                has_border=True
            ))
    
    # Add a section for each line rule
    for rule in line_rules:
        # Add rule section header
        rule_header = doc.add_paragraph(f"Line Rule: {rule}")
        rule_header.style = "Heading 1"
        
        # Add section for each value
        for value_pt in line_values_pt:
            style_name = f"Test_{rule}_{int(value_pt * 10)}"
            twips_value = twips_from_pt(value_pt)
            
            # Add section header with test style
            add_section_header(
                doc, 
                f"Header with {rule}/{value_pt}pt ({twips_value} twips)",
                style_name
            )
            
            # Add content paragraph
            add_content_paragraph(
                doc,
                f"This is a test paragraph following the section header with {rule}/{value_pt}pt."
            )
            
            # Add another section header with the same style and different content
            add_section_header(
                doc,
                f"Another header with {rule}/{value_pt}pt - with longer text to test wrapping behavior for multi-line headers",
                style_name
            )
            
            # Add content paragraph
            add_content_paragraph(
                doc,
                f"This is another test paragraph following the section header with {rule}/{value_pt}pt."
            )
    
    # Clean up any unintended empty paragraphs
    remove_empty_paragraphs(doc)
    
    # Save the document
    output_path = output_dir / "line_height_matrix.docx"
    doc.save(output_path)
    
    logger.info(f"Generated line height matrix test document: {output_path}")
    logger.info(f"Please open in Word to visually inspect the results.")
    logger.info(f"Look for the combination that produces the most compact box height.")

if __name__ == "__main__":
    # Create tests directory structure if it doesn't exist
    Path("tests/docx_spacing/output").mkdir(parents=True, exist_ok=True)
    main() 