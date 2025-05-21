"""
Test Table-Based Section Headers

This script generates a test document with section headers using the table-based approach
to verify that header boxes have the correct compact height.
"""

import os
import sys
import logging
from pathlib import Path

# Add parent directory to path to allow imports
parent_dir = Path(__file__).resolve().parent.parent.parent
sys.path.append(str(parent_dir))

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from word_styles.registry import StyleRegistry, ParagraphBoxStyle
from word_styles.section_builder import add_section_header, add_content_paragraph

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('table_section_headers.log')
    ]
)
logger = logging.getLogger(__name__)

def create_test_document():
    """Generate a test document with section headers using the table-based approach."""
    # Create output directory if it doesn't exist
    output_dir = Path(__file__).parent / "output"
    output_dir.mkdir(exist_ok=True)
    
    # Create output document
    doc = Document()
    
    # Get the StyleRegistry and apply compatibility settings
    registry = StyleRegistry()
    registry.apply_compatibility_settings(doc)
    
    # Verify BoxedHeading2Table style is properly registered
    style = registry.get("BoxedHeading2Table")
    logger.info(f"BoxedHeading2Table wrapper type: {style.wrapper}")
    logger.info(f"BoxedHeading2Table border_twips: {style.border_twips}")
    logger.info(f"BoxedHeading2Table padding_twips: {style.padding_twips}")
    
    # Add title
    title = doc.add_paragraph("Table-Based Section Headers Test")
    title.style = "Title"
    
    # Add explanation
    explanation = doc.add_paragraph(
        "This document tests section headers using the table-based approach "
        "to verify that header boxes have a compact height with no extra space."
    )
    
    # Add a section header
    add_section_header(doc, "REGULAR SECTION HEADER")
    add_content_paragraph(doc, "This is content under a regular section header.")
    
    # Add a section header with longer text to test wrapping
    add_section_header(doc, "SECTION HEADER WITH LONGER TEXT THAT SHOULD WRAP TO THE NEXT LINE")
    add_content_paragraph(doc, "This is content under a section header with wrapped text.")
    
    # Add a section header with mixed case
    add_section_header(doc, "Section Header with Mixed Case")
    add_content_paragraph(doc, "This is content under a section header with mixed case.")
    
    # Add multiple section headers right after each other
    add_section_header(doc, "FIRST SECTION")
    add_section_header(doc, "SECOND SECTION")
    add_content_paragraph(doc, "This tests spacing between consecutive headers.")
    
    # Save the document
    output_path = output_dir / "table_section_headers.docx"
    doc.save(output_path)
    
    logger.info(f"Generated test document: {output_path}")
    logger.info(f"Please open in Word to verify the following:")
    logger.info(f"1. Section header boxes should be compact (exact content height)")
    logger.info(f"2. There should be no extra space above or below the text in boxes")
    logger.info(f"3. Multi-line headers should display correctly")
    
    return output_path

def main():
    """Run the test."""
    # Create test document
    docx_path = create_test_document()
    print(f"\nTest DOCX generated at: {docx_path}")
    print("Please examine in Word to verify that:")
    print("1. Header boxes have exactly the right height (no extra space)")
    print("2. The total height is approximately font size plus border padding")
    print("3. Section headers show up in the document navigation pane/outline")
    print("4. Spacing between sections is consistent and minimal")

if __name__ == "__main__":
    # Create tests directory structure if it doesn't exist
    Path("tests/docx_spacing/output").mkdir(parents=True, exist_ok=True)
    main() 