"""
Migration Guard Tests

These tests ensure that deprecated rendering paths cannot be accidentally re-enabled,
preventing architectural regression.
"""

import pytest
from docx import Document
import sys
import os

# Add project root to path for imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))

import config
from utils.docx_builder import add_section_header

@pytest.mark.architecture
def test_paragraph_writer_removed():
    """Ensure legacy paragraph writer cannot be re-enabled"""
    
    # Save original config
    original_config = config.RENDERING_CONFIG.copy()
    
    try:
        # Attempt to enable legacy paragraph writer
        config.RENDERING_CONFIG["enable_legacy_paragraph_path"] = True
        config.RENDERING_CONFIG["section_header_writer"] = "paragraph"
        
        # This should raise an error
        with pytest.raises(ValueError, match="Legacy paragraph writer is deprecated"):
            doc = Document()
            add_section_header(doc, "TEST_SECTION")
    
    finally:
        # Restore original config
        config.RENDERING_CONFIG.update(original_config)

@pytest.mark.architecture
def test_invalid_section_writer_config():
    """Ensure invalid section writer configurations are rejected"""
    
    # Save original config
    original_config = config.RENDERING_CONFIG.copy()
    
    try:
        # Test invalid writer type
        config.RENDERING_CONFIG["section_header_writer"] = "invalid_writer"
        config.RENDERING_CONFIG["enable_legacy_paragraph_path"] = False
        
        # This should raise an error
        with pytest.raises(ValueError, match="No valid section header writer configured"):
            doc = Document()
            add_section_header(doc, "TEST_SECTION")
    
    finally:
        # Restore original config
        config.RENDERING_CONFIG.update(original_config)

@pytest.mark.architecture
def test_table_writer_works():
    """Ensure the table writer is the default and works correctly"""
    
    # Save original config
    original_config = config.RENDERING_CONFIG.copy()
    
    try:
        # Ensure table writer is configured (should be default)
        config.RENDERING_CONFIG["section_header_writer"] = "table"
        config.RENDERING_CONFIG["enable_legacy_paragraph_path"] = False
        
        # This should work without error
        doc = Document()
        result = add_section_header(doc, "TEST_SECTION")
        
        # Verify something was added to the document
        assert len(doc.paragraphs) > 0 or len(doc.tables) > 0, "No content was added to document"
    
    finally:
        # Restore original config
        config.RENDERING_CONFIG.update(original_config) 