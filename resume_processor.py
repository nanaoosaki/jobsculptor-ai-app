import os
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid
import logging

# Import the PDF parser module
from pdf_parser import read_pdf_file

# Configure logging
logger = logging.getLogger(__name__)

def create_upload_directory(upload_folder):
    """Ensure upload directory exists"""
    os.makedirs(upload_folder, exist_ok=True)
    return upload_folder

def save_uploaded_file(file, upload_folder):
    """Save uploaded file with unique filename and return the path"""
    # Determine file extension from original file
    original_extension = os.path.splitext(file.filename)[1].lower()
    
    # Check if it's a supported file type
    if original_extension not in ['.docx', '.pdf']:
        raise ValueError(f"Unsupported file type: {original_extension}")
    
    # Generate a unique filename with the correct extension
    filename = str(uuid.uuid4()) + original_extension
    filepath = os.path.join(upload_folder, filename)
    file.save(filepath)
    return filename, filepath

def read_docx_file(filepath):
    """Read and extract content from a DOCX file"""
    doc = docx.Document(filepath)
    content = {
        'paragraphs': [],
        'tables': []
    }
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            content['paragraphs'].append({
                'text': para.text,
                'style': para.style.name,
                'alignment': para.alignment,
                'bold': any(run.bold for run in para.runs),
                'italic': any(run.italic for run in para.runs),
                'font_size': para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else None
            })
    
    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data.append(row_data)
        content['tables'].append(table_data)
    
    return content

def extract_resume_sections(content):
    """Extract and categorize resume sections from content"""
    sections = {
        'contact_info': [],
        'summary': [],
        'education': [],
        'experience': [],
        'skills': [],
        'projects': [],
        'other': []
    }
    
    current_section = 'other'
    
    for para in content['paragraphs']:
        text = para['text'].lower()
        
        # Determine section based on heading text
        if para['bold'] or para['style'].startswith('Heading'):
            if any(term in text for term in ['contact', 'email', 'phone', 'address']):
                current_section = 'contact_info'
            elif any(term in text for term in ['summary', 'objective', 'profile']):
                current_section = 'summary'
            elif any(term in text for term in ['education', 'academic']):
                current_section = 'education'
            elif any(term in text for term in ['experience', 'employment', 'work']):
                current_section = 'experience'
            elif any(term in text for term in ['skill', 'technology', 'competenc']):
                current_section = 'skills'
            elif any(term in text for term in ['project', 'portfolio']):
                current_section = 'projects'
            else:
                current_section = 'other'
        
        # Add paragraph to current section
        sections[current_section].append(para)
    
    return sections

def analyze_resume(filepath):
    """Analyze resume content and structure"""
    # Determine file type based on extension
    file_ext = os.path.splitext(filepath)[1].lower()
    
    if file_ext == '.docx':
        logger.info(f"Processing DOCX file: {filepath}")
        content = read_docx_file(filepath)
    elif file_ext == '.pdf':
        logger.info(f"Processing PDF file: {filepath}")
        content = read_pdf_file(filepath)
    else:
        logger.error(f"Unsupported file format: {file_ext}")
        raise ValueError(f"Unsupported file format: {file_ext}")
    
    sections = extract_resume_sections(content)
    
    # Basic analysis
    analysis = {
        'has_contact_info': len(sections['contact_info']) > 0,
        'has_summary': len(sections['summary']) > 0,
        'has_education': len(sections['education']) > 0,
        'has_experience': len(sections['experience']) > 0,
        'has_skills': len(sections['skills']) > 0,
        'has_projects': len(sections['projects']) > 0,
        'total_paragraphs': len(content['paragraphs']),
        'total_tables': len(content['tables'])
    }
    
    return {
        'content': content,
        'sections': sections,
        'analysis': analysis
    }

def generate_resume_preview_html(analysis):
    """Generate HTML preview of resume content for display in the UI"""
    sections = analysis['sections']
    
    html = '<div class="resume-preview-container">'
    html += '<h3 class="preview-title">User Resume Parsed</h3>'
    html += '<div class="resume-preview-content">'
    
    # Contact Information
    if sections['contact_info']:
        html += '<div class="preview-section">'
        html += '<h4 class="preview-section-title">Contact Information</h4>'
        for para in sections['contact_info']:
            html += f'<p class="preview-text">{para["text"]}</p>'
        html += '</div>'
    
    # Summary
    if sections['summary']:
        html += '<div class="preview-section">'
        html += '<h4 class="preview-section-title">Summary</h4>'
        for para in sections['summary']:
            html += f'<p class="preview-text">{para["text"]}</p>'
        html += '</div>'
    
    # Experience
    if sections['experience']:
        html += '<div class="preview-section">'
        html += '<h4 class="preview-section-title">Experience</h4>'
        for para in sections['experience']:
            if para.get('bold', False):
                html += f'<p class="preview-text preview-bold">{para["text"]}</p>'
            else:
                html += f'<p class="preview-text">{para["text"]}</p>'
        html += '</div>'
    
    # Education
    if sections['education']:
        html += '<div class="preview-section">'
        html += '<h4 class="preview-section-title">Education</h4>'
        for para in sections['education']:
            if para.get('bold', False):
                html += f'<p class="preview-text preview-bold">{para["text"]}</p>'
            else:
                html += f'<p class="preview-text">{para["text"]}</p>'
        html += '</div>'
    
    # Skills
    if sections['skills']:
        html += '<div class="preview-section">'
        html += '<h4 class="preview-section-title">Skills</h4>'
        for para in sections['skills']:
            html += f'<p class="preview-text">{para["text"]}</p>'
        html += '</div>'
    
    # Projects
    if sections['projects']:
        html += '<div class="preview-section">'
        html += '<h4 class="preview-section-title">Projects</h4>'
        for para in sections['projects']:
            if para.get('bold', False):
                html += f'<p class="preview-text preview-bold">{para["text"]}</p>'
            else:
                html += f'<p class="preview-text">{para["text"]}</p>'
        html += '</div>'
    
    # Other sections
    if sections['other']:
        html += '<div class="preview-section">'
        html += '<h4 class="preview-section-title">Additional Information</h4>'
        for para in sections['other']:
            html += f'<p class="preview-text">{para["text"]}</p>'
        html += '</div>'
    
    html += '</div></div>'
    
    return html
