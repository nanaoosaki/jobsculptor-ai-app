#!/usr/bin/env python3
"""
XML Table Comparison Tool

Generates documents with both legacy and universal table approaches,
then extracts and compares the raw XML to identify positioning differences.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from word_styles.section_builder import add_section_header as legacy_add_section_header
from rendering.components import SectionHeader
from style_engine import StyleEngine
import xml.etree.ElementTree as ET
import xml.dom.minidom as minidom

def create_legacy_table_document():
    """Create document using legacy table approach"""
    doc = Document()
    doc.add_paragraph("Reference paragraph for spacing")
    
    # Use legacy approach
    legacy_add_section_header(doc, "Legacy Test Header", "BoxedHeading2Table")
    
    doc.add_paragraph("Another reference paragraph")
    doc.save("legacy_table_test.docx")
    return doc

def create_universal_table_document():
    """Create document using universal renderer approach"""
    doc = Document()
    doc.add_paragraph("Reference paragraph for spacing")
    
    # Use universal renderer
    tokens = StyleEngine.load_tokens()
    universal_header = SectionHeader(tokens, "Universal Test Header")
    universal_header.to_docx(doc)
    
    doc.add_paragraph("Another reference paragraph")
    doc.save("universal_table_test.docx")
    return doc

def extract_table_xml(doc, approach_name):
    """Extract XML for the first table in document"""
    tables = doc.tables
    if not tables:
        print(f"❌ No tables found in {approach_name} document")
        return None
    
    table = tables[0]
    # Get the underlying XML element
    table_xml = table._tbl
    
    # Convert to pretty-printed string
    rough_string = ET.tostring(table_xml, encoding='unicode')
    reparsed = minidom.parseString(f"<root>{rough_string}</root>")
    pretty = reparsed.toprettyxml(indent="  ")
    
    # Remove the wrapper and XML declaration
    lines = pretty.split('\n')
    lines = [line for line in lines if line.strip() and not line.startswith('<?xml') and not line.strip() in ['<root>', '</root>']]
    
    return '\n'.join(lines)

def compare_table_properties():
    """Compare key table properties between approaches"""
    print("🔍 Creating comparison documents...")
    
    # Create both documents
    legacy_doc = create_legacy_table_document()
    universal_doc = create_universal_table_document()
    
    print("✅ Documents created")
    print("📄 Legacy: legacy_table_test.docx")
    print("📄 Universal: universal_table_test.docx")
    
    # Extract XML
    print("\n🔍 Extracting table XML...")
    legacy_xml = extract_table_xml(legacy_doc, "legacy")
    universal_xml = extract_table_xml(universal_doc, "universal")
    
    if not legacy_xml or not universal_xml:
        print("❌ Failed to extract XML from one or both documents")
        return
    
    # Save XML to files for comparison
    with open("legacy_table.xml", "w", encoding='utf-8') as f:
        f.write(legacy_xml)
    
    with open("universal_table.xml", "w", encoding='utf-8') as f:
        f.write(universal_xml)
    
    print("✅ XML extracted and saved:")
    print("📄 Legacy XML: legacy_table.xml")
    print("📄 Universal XML: universal_table.xml")
    
    # Basic comparison
    print("\n🔍 Quick Comparison:")
    legacy_lines = legacy_xml.split('\n')
    universal_lines = universal_xml.split('\n')
    
    print(f"Legacy XML lines: {len(legacy_lines)}")
    print(f"Universal XML lines: {len(universal_lines)}")
    
    # Find differences
    print("\n🔍 Line-by-line differences:")
    max_lines = max(len(legacy_lines), len(universal_lines))
    differences_found = 0
    
    for i in range(max_lines):
        legacy_line = legacy_lines[i].strip() if i < len(legacy_lines) else "<<MISSING>>"
        universal_line = universal_lines[i].strip() if i < len(universal_lines) else "<<MISSING>>"
        
        if legacy_line != universal_line:
            differences_found += 1
            if differences_found <= 10:  # Show first 10 differences
                print(f"  Line {i+1}:")
                print(f"    Legacy:    {legacy_line}")
                print(f"    Universal: {universal_line}")
    
    if differences_found > 10:
        print(f"    ... and {differences_found - 10} more differences")
    elif differences_found == 0:
        print("  ✅ No differences found in XML structure")
    else:
        print(f"  Found {differences_found} differences")

def analyze_table_positioning():
    """Analyze specific positioning-related XML elements"""
    print("\n🔍 Analyzing positioning-specific XML elements...")
    
    # Load the XML files
    try:
        with open("legacy_table.xml", "r", encoding='utf-8') as f:
            legacy_xml = f.read()
        with open("universal_table.xml", "r", encoding='utf-8') as f:
            universal_xml = f.read()
    except FileNotFoundError:
        print("❌ XML files not found. Run compare_table_properties() first.")
        return
    
    # Check for positioning-related elements
    positioning_elements = [
        'w:tblPr',      # Table properties
        'w:tblW',       # Table width
        'w:tblInd',     # Table indentation
        'w:jc',         # Justification
        'w:tblpPr',     # Table positioning properties
        'w:tcPr',       # Cell properties
        'w:tcW',        # Cell width
        'w:tcMar',      # Cell margins
        'w:gridCol',    # Grid column
    ]
    
    print("Checking for positioning elements:")
    for element in positioning_elements:
        legacy_has = element in legacy_xml
        universal_has = element in universal_xml
        
        if legacy_has and universal_has:
            print(f"  ✅ {element}: Both have it")
        elif legacy_has and not universal_has:
            print(f"  ❌ {element}: Only legacy has it")
        elif not legacy_has and universal_has:
            print(f"  ❌ {element}: Only universal has it")
        else:
            print(f"  ⚪ {element}: Neither has it")

def main():
    """Main function to run XML comparison"""
    print("🔍 XML Table Comparison Tool")
    print("=" * 50)
    
    try:
        compare_table_properties()
        analyze_table_positioning()
        
        print("\n" + "=" * 50)
        print("✅ XML comparison completed!")
        print("\n📋 Next steps:")
        print("1. Review legacy_table.xml vs universal_table.xml")
        print("2. Look for positioning differences in table properties")
        print("3. Test specific property modifications in universal renderer")
        print("4. Focus on w:tblPr, w:tblW, w:tblInd elements")
        
    except Exception as e:
        print(f"❌ Comparison failed: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 