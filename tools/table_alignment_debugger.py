#!/usr/bin/env python3
"""
Table Alignment Debugger

Generates a test document with alternating tables and paragraphs 
to visually compare alignment between legacy and universal renderers.
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from word_styles.section_builder import add_section_header as legacy_add_section_header
from rendering.components import SectionHeader
from style_engine import StyleEngine
import logging

logger = logging.getLogger(__name__)

def create_alignment_test_document():
    """Create test document comparing legacy vs universal table alignment"""
    doc = Document()
    
    # Add title
    title = doc.add_paragraph("Table Alignment Comparison Test")
    title.runs[0].bold = True
    title.runs[0].font.size = 16
    
    # Add reference paragraph
    ref1 = doc.add_paragraph("Reference paragraph #1 - This shows normal content alignment")
    
    # Add legacy table section header
    print("Adding legacy section header...")
    try:
        legacy_add_section_header(doc, "Legacy Section Header", "BoxedHeading2Table")
        print("✅ Legacy section header added successfully")
    except Exception as e:
        print(f"❌ Legacy section header failed: {e}")
        doc.add_paragraph(f"LEGACY FAILED: {e}")
    
    # Add another reference paragraph
    ref2 = doc.add_paragraph("Reference paragraph #2 - This should align with the section headers")
    
    # Add universal renderer section header
    print("Adding universal section header...")
    try:
        tokens = StyleEngine.load_tokens()
        universal_header = SectionHeader(tokens, "Universal Section Header (Simplified)")
        universal_header.to_docx(doc)
        print("✅ Universal section header added successfully")
    except Exception as e:
        print(f"❌ Universal section header failed: {e}")
        doc.add_paragraph(f"UNIVERSAL FAILED: {e}")
    
    # Add final reference paragraph
    ref3 = doc.add_paragraph("Reference paragraph #3 - Compare alignment with both headers above")
    
    # Add comparison notes
    notes = doc.add_paragraph("\nNOTES:")
    notes.runs[0].bold = True
    doc.add_paragraph("• All three reference paragraphs should align at the same left position")
    doc.add_paragraph("• Both section headers should align with the reference paragraphs")
    doc.add_paragraph("• If universal header is shifted left, the issue persists")
    doc.add_paragraph("• If both headers align with paragraphs, the issue is fixed")
    
    # Save test document
    output_path = "table_alignment_comparison.docx"
    doc.save(output_path)
    print(f"📄 Test document saved: {output_path}")
    print("👀 Please open this file and visually inspect the alignment")
    
    return output_path

def main():
    """Main function to run alignment test"""
    print("🔍 Creating table alignment test document...")
    print("=" * 50)
    
    try:
        output_file = create_alignment_test_document()
        print("=" * 50)
        print("✅ Test document created successfully!")
        print(f"📂 Open: {output_file}")
        print("🔍 Visual inspection needed:")
        print("   1. Check if all paragraphs align at same left position")
        print("   2. Check if both section headers align with paragraphs")
        print("   3. Report if universal header is still shifted left")
        
    except Exception as e:
        print(f"❌ Test failed: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 