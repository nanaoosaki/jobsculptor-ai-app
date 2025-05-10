#!/usr/bin/env python3
"""
DOCX Debug Tool

This script is used to debug DOCX formatting by creating test documents and inspecting
existing DOCX files.
"""

import os
import sys
import json
import logging
import argparse
from pathlib import Path
from io import BytesIO

# Add parent directory to path for imports
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

from utils.docx_debug import inspect_docx_paragraphs, inspect_docx_styles, generate_debug_report
from utils.docx_builder import create_bullet_point, add_section_header, add_role_description, format_right_aligned_pair
from style_engine import StyleEngine

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def create_test_docx(output_path=None):
    """Create a test DOCX file with various formatting examples."""
    logger.info("Creating test DOCX file...")
    
    # Create a new document
    doc = Document()
    
    # Set up the document with our custom styles
    custom_styles = StyleEngine.create_docx_custom_styles(doc)
    
    # Add a title
    doc.add_paragraph("TEST RESUME", style='Title')
    
    # Add contact information
    doc.add_paragraph("email@example.com | (123) 456-7890 | LinkedIn | GitHub", style='Subtitle')
    
    # Add section headers and content with our custom styles
    
    # Professional Summary section
    section_para = add_section_header(doc, "Professional Summary", {})
    content_para = doc.add_paragraph("Professional with experience in software development and cloud engineering. Strong skills in Python, JavaScript, and AWS.", style='MR_Content')
    
    # Experience section
    section_para = add_section_header(doc, "Experience", {})
    
    # Company 1
    company_para = format_right_aligned_pair(
        doc, 
        "Company ABC - Software Developer", 
        "2020 - Present",
        "company",
        "date",
        {}
    )
    
    # Role description
    role_para = add_role_description(
        doc,
        "Backend developer specializing in Python and Django frameworks. Lead a team of 3 developers.",
        {}
    )
    
    # Bullet points
    bullet_para = create_bullet_point(doc, "Developed RESTful APIs using Django Rest Framework, serving over 1 million requests per day", {})
    bullet_para = create_bullet_point(doc, "Optimized database queries resulting in 30% performance improvement", {})
    bullet_para = create_bullet_point(doc, "Implemented CI/CD pipeline using GitHub Actions", {})
    
    # Company 2
    company_para = format_right_aligned_pair(
        doc, 
        "Company XYZ - Data Analyst", 
        "2018 - 2020",
        "company",
        "date",
        {}
    )
    
    # Role description
    role_para = add_role_description(
        doc,
        "Data analyst focusing on business intelligence",
        {}
    )
    
    # Bullet points
    bullet_para = create_bullet_point(doc, "Created interactive dashboards using Power BI", {})
    bullet_para = create_bullet_point(doc, "Conducted data mining and statistical analysis", {})
    bullet_para = create_bullet_point(doc, "Automated monthly reporting process saving 10 hours per month", {})
    
    # Education section
    section_para = add_section_header(doc, "Education", {})
    
    # University
    company_para = format_right_aligned_pair(
        doc, 
        "University Name - Bachelor of Science, Computer Science", 
        "2014 - 2018",
        "company",
        "date",
        {}
    )
    
    # Bullet points
    bullet_para = create_bullet_point(doc, "GPA: 3.8/4.0", {})
    bullet_para = create_bullet_point(doc, "Dean's List: All semesters", {})
    bullet_para = create_bullet_point(doc, "Relevant coursework: Data Structures, Algorithms, Database Systems", {})
    
    # Determine output path
    if not output_path:
        output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "debug_test.docx")
    
    # Save the document
    doc.save(output_path)
    logger.info(f"Test DOCX saved to {output_path}")
    
    return output_path

def inspect_docx(docx_path=None):
    """Inspect a DOCX file and generate a debug report."""
    if not docx_path:
        # Default: Create a test document first
        docx_path = create_test_docx()
    
    logger.info(f"Inspecting DOCX file: {docx_path}")
    
    # Open the document
    doc = Document(docx_path)
    
    # Inspect paragraphs
    paragraphs_info = inspect_docx_paragraphs(doc)
    
    # Inspect styles
    styles_info = inspect_docx_styles(doc)
    
    # Generate a debug report
    debug_report = generate_debug_report(doc)
    
    # Save the debug report to a file
    debug_output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "docx_debug_report.json")
    with open(debug_output_path, 'w') as f:
        json.dump(debug_report, f, indent=2)
    
    logger.info(f"Debug report saved to {debug_output_path}")
    
    # Log a summary of the document
    logger.info(f"Summary: {len(paragraphs_info)} paragraphs, {len(styles_info)} paragraph styles")
    
    # Log specific element types for easier debugging
    for p in paragraphs_info:
        if p.get('style', '') == 'MR_SectionHeader':
            logger.info(f"Section Header: '{p.get('text')}' - Left indent: {p.get('left_indent')}")
        elif 'Company' in p.get('text', '') or 'University' in p.get('text', ''):
            logger.info(f"Company: '{p.get('text')}' - Left indent: {p.get('left_indent')}")
        elif p.get('style', '') == 'MR_RoleDescription':
            logger.info(f"Role: '{p.get('text')}' - Left indent: {p.get('left_indent')}")
        elif p.get('text', '').startswith('•'):
            logger.info(f"Bullet: '{p.get('text')}' - Left indent: {p.get('left_indent')}, First line: {p.get('first_line_indent')}")
    
    return debug_report

def main():
    """Main function to run the debugger."""
    parser = argparse.ArgumentParser(description='Debug DOCX files')
    parser.add_argument('file_path', nargs='?', help='Path to DOCX file to debug, or "create" to create a test file')
    parser.add_argument('--debug-output', help='Path to save the debug report JSON file')
    parser.add_argument('--output-path', help='Path to save the created DOCX file (only with "create" mode)')
    args = parser.parse_args()
    
    if not args.file_path or args.file_path.lower() == 'create':
        # Create a test document
        docx_path = create_test_docx(args.output_path)
        
        # If debug output is requested, also inspect the file
        if args.debug_output:
            debug_report = inspect_docx(docx_path)
            with open(args.debug_output, 'w') as f:
                json.dump(debug_report, f, indent=2)
            logger.info(f"Debug report saved to {args.debug_output}")
    else:
        # Inspect an existing file
        debug_report = inspect_docx(args.file_path)
        
        # Save to custom output path if specified
        if args.debug_output:
            with open(args.debug_output, 'w') as f:
                json.dump(debug_report, f, indent=2)
            logger.info(f"Debug report saved to {args.debug_output}")

if __name__ == "__main__":
    main() 